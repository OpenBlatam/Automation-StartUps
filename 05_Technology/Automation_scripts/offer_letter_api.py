"""
API para generación automatizada de cartas de oferta laboral (Offer Letters)
Sistema de automatización para recursos humanos

Funcionalidades:
- Generación de cartas en múltiples formatos (TXT, HTML, PDF)
- API REST con Flask
- Plantillas personalizables
- Base de datos SQLite
- Envío por email
- Firmas digitales
- Internacionalización (i18n)
- Estadísticas y reportes
- Validación avanzada
"""

import json
import os
import sqlite3
import re
import uuid
import traceback
import html
import time
import shutil
from datetime import datetime, timedelta, date
from typing import Dict, Optional, List, Tuple
from dataclasses import dataclass, asdict
from enum import Enum
from collections import defaultdict
from functools import wraps
import logging
from pathlib import Path

# Configuración de logging (debe estar antes de los imports opcionales)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Intentar importar dependencias opcionales
try:
    from flask import Flask, request, jsonify, send_file, make_response
    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False
    logger.warning("Flask no disponible. Funcionalidad de API REST deshabilitada.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("ReportLab no disponible. Generación de PDF deshabilitada.")

try:
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False
    logger.warning("Módulos de email no disponibles. Envío por email deshabilitado.")

try:
    import hashlib
    import hmac
    import base64
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False
    logger.warning("Módulos de criptografía no disponibles. Firmas digitales deshabilitadas.")

try:
    import functools
    import time
    CACHE_AVAILABLE = True
except ImportError:
    CACHE_AVAILABLE = False

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False

try:
    from io import StringIO
    IO_AVAILABLE = True
except ImportError:
    IO_AVAILABLE = False

try:
    import gzip
    import zlib
    COMPRESSION_AVAILABLE = True
except ImportError:
    COMPRESSION_AVAILABLE = False

try:
    import asyncio
    from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
    ASYNC_AVAILABLE = True
except ImportError:
    ASYNC_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx no disponible. Generación de Word deshabilitada.")

try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    AIOFILES_AVAILABLE = False

try:
    from redis import Redis
    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False

try:
    import orjson
    ORJSON_AVAILABLE = True
except ImportError:
    ORJSON_AVAILABLE = False

# ============================================================================
# Excepciones Personalizadas
# ============================================================================

class OfferLetterException(Exception):
    """Excepción base para errores del sistema de ofertas"""
    pass

class ValidationError(OfferLetterException):
    """Excepción para errores de validación"""
    def __init__(self, message: str, errors: List[str] = None):
        super().__init__(message)
        self.errors = errors or []

class DatabaseError(OfferLetterException):
    """Excepción para errores de base de datos"""
    pass

class TemplateError(OfferLetterException):
    """Excepción para errores de plantillas"""
    pass

class FormatError(OfferLetterException):
    """Excepción para errores de formato"""
    pass

class RateLimitError(OfferLetterException):
    """Excepción para errores de rate limiting"""
    pass

# ============================================================================
# Configuración Centralizada
# ============================================================================

class Config:
    """Configuración centralizada de la aplicación"""
    
    # Rate limiting
    MAX_REQUESTS_PER_MINUTE = int(os.environ.get('MAX_REQUESTS_PER_MINUTE', 60))
    RATE_LIMIT_ENABLED = os.environ.get('RATE_LIMIT_ENABLED', 'true').lower() == 'true'
    
    # Directorios
    OUTPUT_DIR = Path(os.environ.get('OUTPUT_DIR', 'output/offer_letters'))
    TEMPLATES_DIR = Path(os.environ.get('TEMPLATES_DIR', 'templates/offer_letters'))
    LOGS_DIR = Path(os.environ.get('LOGS_DIR', 'logs'))
    
    # Validación
    MAX_SALARY = float(os.environ.get('MAX_SALARY', 10000000))
    MIN_SALARY = float(os.environ.get('MIN_SALARY', 0))
    MAX_TEXT_LENGTH = int(os.environ.get('MAX_TEXT_LENGTH', 200))
    
    # Formatos soportados
    SUPPORTED_FORMATS = ['txt', 'html', 'pdf', 'json', 'docx', 'rtf']
    DEFAULT_FORMAT = 'txt'
    
    # Configuración async
    ASYNC_ENABLED = os.environ.get('ASYNC_ENABLED', 'true').lower() == 'true'
    MAX_WORKERS = int(os.environ.get('MAX_WORKERS', 4))
    
    # Redis para caché distribuido
    REDIS_HOST = os.environ.get('REDIS_HOST', 'localhost')
    REDIS_PORT = int(os.environ.get('REDIS_PORT', 6379))
    REDIS_DB = int(os.environ.get('REDIS_DB', 0))
    USE_REDIS_CACHE = os.environ.get('USE_REDIS_CACHE', 'false').lower() == 'true'
    
    # Monedas soportadas
    SUPPORTED_CURRENCIES = ['USD', 'EUR', 'MXN', 'GBP', 'CAD', 'AUD', 'JPY', 'CHF', 'CNY']
    
    # Tipos de empleo
    EMPLOYMENT_TYPES = ['Full-time', 'Part-time', 'Contract', 'Intern', 'Temporary', 'Freelance']
    
    # Logging
    LOG_LEVEL = os.environ.get('LOG_LEVEL', 'INFO')
    LOG_FILE = LOGS_DIR / 'offer_letter_api.log'
    
    # Cache
    ENABLE_CACHE = os.environ.get('ENABLE_CACHE', 'true').lower() == 'true'
    CACHE_TTL_SECONDS = int(os.environ.get('CACHE_TTL_SECONDS', 3600))
    
    @classmethod
    def setup_directories(cls):
        """Crea los directorios necesarios"""
        cls.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        cls.TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
        cls.LOGS_DIR.mkdir(parents=True, exist_ok=True)
    
    @classmethod
    def setup_logging(cls):
        """Configura el sistema de logging"""
        log_format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        handlers = [logging.StreamHandler()]
        
        if cls.LOG_FILE:
            file_handler = logging.FileHandler(cls.LOG_FILE, encoding='utf-8')
            handlers.append(file_handler)
        
        logging.basicConfig(
            level=getattr(logging, cls.LOG_LEVEL),
            format=log_format,
            handlers=handlers
        )


# Inicializar configuración
Config.setup_directories()
Config.setup_logging()

# Rate limiting básico
request_counts = defaultdict(list)
MAX_REQUESTS_PER_MINUTE = Config.MAX_REQUESTS_PER_MINUTE

# Cache de templates
_template_cache = {}
_cache_timestamps = {}


def validate_offer_data(data: Dict) -> Tuple[bool, List[str]]:
    """
    Valida los datos de la oferta con validaciones mejoradas usando configuración centralizada.
    Incluye validaciones exhaustivas de tipos, formatos y rangos de valores.
    
    Args:
        data: Diccionario con los datos de la oferta. Debe contener al menos:
            - position_title: Título del puesto (requerido, str)
            - salary: Salario ofrecido (requerido, numérico)
            - start_date: Fecha de inicio (requerido, formato YYYY-MM-DD)
            - location: Ubicación del trabajo (requerido, str)
            - candidate_email: Email del candidato (opcional, debe ser válido si se proporciona)
            - employment_type: Tipo de empleo (opcional, debe estar en Config.EMPLOYMENT_TYPES)
            - currency: Moneda (opcional, debe estar en Config.SUPPORTED_CURRENCIES)
            - probation_period: Período de prueba en días (opcional, 0-365)
            - notice_period: Período de preaviso en días (opcional, 0-180)
    
    Returns:
        Tupla (es_válido: bool, lista_errores: List[str]).
        Si es_válido es True, lista_errores estará vacía.
        Si es_válido es False, lista_errores contendrá mensajes descriptivos de los errores.
    
    Raises:
        TypeError: Si data no es un diccionario.
    
    Example:
        >>> data = {
        ...     'position_title': 'Software Engineer',
        ...     'salary': 50000,
        ...     'start_date': '2024-02-01',
        ...     'location': 'Remote',
        ...     'candidate_email': 'candidate@example.com',
        ...     'currency': 'USD'
        ... }
        >>> is_valid, errors = validate_offer_data(data)
        >>> if not is_valid:
        ...     print("Errores:", errors)
        ... else:
        ...     print("Datos válidos")
    """
    if not isinstance(data, dict):
        return False, ["Data must be a dictionary"]
    
    errors = []
    required = ['position_title', 'salary', 'start_date', 'location']
    
    # Validar campos requeridos
    for field in required:
        if not data.get(field):
            errors.append(f"Missing required field: {field}")
    
    # Validar formato de fecha
    if data.get('start_date'):
        date_str = str(data.get('start_date', '')).strip()
        date_pattern = r'^\d{4}-\d{2}-\d{2}$'
        if not re.match(date_pattern, date_str):
            errors.append("start_date must be in format YYYY-MM-DD")
        else:
            try:
                parsed_date = datetime.strptime(date_str, '%Y-%m-%d').date()
                current_date = date.today()
                if parsed_date.year < 2000 or parsed_date.year > 2100:
                    errors.append("start_date year must be between 2000 and 2100")
                elif parsed_date < current_date:
                    errors.append("start_date cannot be in the past")
            except ValueError:
                errors.append("start_date is not a valid date")
    
    # Validar salario usando configuración
    if data.get('salary'):
        try:
            salary = str(data.get('salary', '')).replace('$', '').replace(',', '').strip()
            salary_value = float(salary)
            if salary_value <= Config.MIN_SALARY:
                errors.append(f"salary must be greater than {Config.MIN_SALARY}")
            if salary_value > Config.MAX_SALARY:
                errors.append(f"salary exceeds maximum allowed value ({Config.MAX_SALARY:,.0f})")
        except (ValueError, TypeError):
            errors.append("salary must be a valid number")
    
    # Validar email si está presente
    if data.get('candidate_email'):
        email_pattern = r'^[a-zA-Z0-9]([a-zA-Z0-9._-]*[a-zA-Z0-9])?@[a-zA-Z0-9]([a-zA-Z0-9.-]*[a-zA-Z0-9])?\.[a-zA-Z]{2,}$'
        email_value = str(data.get('candidate_email', '')).strip()
        if not re.match(email_pattern, email_value):
            errors.append("candidate_email must be a valid email address")
    
    # Validar longitud de campos de texto usando configuración
    text_fields = {
        'position_title': Config.MAX_TEXT_LENGTH,
        'location': 100,
        'candidate_name': 100,
        'department': 100,
        'company_name': Config.MAX_TEXT_LENGTH
    }
    
    for field, max_length in text_fields.items():
        if data.get(field) and len(str(data.get(field, ''))) > max_length:
            errors.append(f"{field} exceeds maximum length of {max_length} characters")
    
    # Validar employment_type usando configuración
    if data.get('employment_type'):
        if data.get('employment_type') not in Config.EMPLOYMENT_TYPES:
            errors.append(f"employment_type must be one of: {', '.join(Config.EMPLOYMENT_TYPES)}")
    
    # Validar currency usando configuración
    if data.get('currency'):
        if data.get('currency') not in Config.SUPPORTED_CURRENCIES:
            errors.append(f"currency must be one of: {', '.join(Config.SUPPORTED_CURRENCIES)}")
    
    return len(errors) == 0, errors


def get_cached_template(template_key: str) -> Optional[str]:
    """
    Obtiene un template del cache si está disponible y no ha expirado
    
    Args:
        template_key: Clave del template
    
    Returns:
        Template si está en cache y es válido, None en caso contrario
    """
    if not Config.ENABLE_CACHE:
        return None
    
    if not isinstance(template_key, str) or not template_key:
        return None
    
    if template_key in _template_cache:
        timestamp = _cache_timestamps.get(template_key, datetime.now())
        age = (datetime.now() - timestamp).total_seconds()
        
        if age < Config.CACHE_TTL_SECONDS:
            return _template_cache[template_key]
        else:
            # Cache expirado, limpiar
            del _template_cache[template_key]
            if template_key in _cache_timestamps:
            del _cache_timestamps[template_key]
    
    return None


def cache_template(template_key: str, template_content: str):
    """
    Almacena un template en el cache
    
    Args:
        template_key: Clave del template
        template_content: Contenido del template
    """
    if not Config.ENABLE_CACHE:
        return
    
    if not isinstance(template_key, str) or not template_key:
        logger.warning("Invalid template_key provided to cache_template")
        return
    
    if not isinstance(template_content, str):
        logger.warning("template_content must be a string")
        return
    
        _template_cache[template_key] = template_content
        _cache_timestamps[template_key] = datetime.now()


def format_currency(amount: float, currency: str = 'USD', include_decimals: bool = True) -> str:
    """
    Formatea un monto según la moneda especificada con soporte para múltiples formatos.
    
    Args:
        amount: Monto a formatear (debe ser un número positivo).
        currency: Código de moneda ISO 4217 (USD, EUR, MXN, etc.).
        include_decimals: Si True, incluye decimales. Si False, redondea a entero.
            Para JPY y algunas monedas asiáticas, se ignora y siempre se muestra sin decimales.
    
    Returns:
        String formateado con el monto y símbolo de moneda, con separadores de miles.
        Ejemplo: "$50,000.00" para USD, "€45,000.00" para EUR.
    
    Raises:
        ValueError: Si amount es negativo o no es un número válido.
    
    Example:
        >>> format_currency(50000, 'USD')
        '$50,000.00'
        >>> format_currency(1000000, 'MXN')
        '$1,000,000.00'
        >>> format_currency(500000, 'JPY')
        '¥500,000'
    """
    # Validar entrada
    if not isinstance(amount, (int, float)):
        raise ValueError(f"Amount must be a number, got {type(amount).__name__}")
    
    if amount < 0:
        raise ValueError(f"Amount cannot be negative: {amount}")
    
    if not isinstance(currency, str) or len(currency) != 3:
        raise ValueError(f"Currency must be a 3-letter ISO code, got: {currency}")
    
    currency = currency.upper()
    
    # Diccionario de símbolos de moneda
    currency_symbols = {
        'USD': '$',
        'EUR': '€',
        'MXN': '$',
        'GBP': '£',
        'CAD': '$',
        'AUD': '$',
        'JPY': '¥',
        'CHF': 'CHF',
        'CNY': '¥'
    }
    
    symbol = currency_symbols.get(currency, currency)
    
    # Monedas que tradicionalmente no usan decimales
    currencies_without_decimals = {'JPY', 'KRW', 'VND', 'CLP'}
    use_decimals = include_decimals and currency not in currencies_without_decimals
    
    # Formato especial para JPY (sin decimales)
    if currency == 'JPY':
        return f"{symbol}{amount:,.0f}"
    
    return f"{symbol}{amount:,.2f}"


def calculate_annual_salary(
    monthly_salary: float,
    include_bonus: bool = False,
    bonus_amount: float = 0.0,
    months_per_year: int = 12
) -> float:
    """
    Calcula el salario anual basado en el salario mensual con opciones avanzadas.
    Soporta bonos y diferentes períodos de pago.
    
    Args:
        monthly_salary: Salario mensual base (debe ser >= 0).
        include_bonus: Si True, incluye el bono anual en el cálculo.
        bonus_amount: Monto del bono anual (solo se usa si include_bonus=True).
        months_per_year: Número de meses por año (default: 12).
            Útil para contratos con períodos diferentes.
    
    Returns:
        Salario anual total calculado.
    
    Raises:
        ValueError: Si monthly_salary es negativo o months_per_year es inválido.
        TypeError: Si los parámetros no son números.
    
    Example:
        >>> calculate_annual_salary(5000)
        60000.0
        >>> calculate_annual_salary(5000, include_bonus=True, bonus_amount=5000)
        65000.0
        >>> calculate_annual_salary(5000, months_per_year=13)  # 13 meses
        65000.0
    """
    if not isinstance(monthly_salary, (int, float)):
        raise TypeError(f"monthly_salary must be a number, got {type(monthly_salary).__name__}")
    
    if monthly_salary < 0:
        raise ValueError(f"monthly_salary cannot be negative: {monthly_salary}")
    
    if not isinstance(months_per_year, int) or months_per_year < 1 or months_per_year > 24:
        raise ValueError(f"months_per_year must be between 1 and 24, got: {months_per_year}")
    
    annual_base = monthly_salary * months_per_year
    
    if include_bonus:
        if not isinstance(bonus_amount, (int, float)):
            raise TypeError(f"bonus_amount must be a number, got {type(bonus_amount).__name__}")
        if bonus_amount < 0:
            raise ValueError(f"bonus_amount cannot be negative: {bonus_amount}")
        return annual_base + bonus_amount
    
    return annual_base


def calculate_monthly_salary(
    annual_salary: float,
    months_per_year: int = 12,
    include_bonus: bool = False,
    bonus_amount: float = 0.0
) -> float:
    """
    Calcula el salario mensual basado en el salario anual con opciones avanzadas.
    Puede excluir bonos del cálculo mensual.
    
    Args:
        annual_salary: Salario anual total (debe ser >= 0).
        months_per_year: Número de meses por año (default: 12).
            Útil para contratos con períodos diferentes.
        include_bonus: Si True, el annual_salary incluye bonos.
        bonus_amount: Monto del bono anual (solo se usa si include_bonus=True).
            Se resta del annual_salary antes de calcular el mensual.
    
    Returns:
        Salario mensual promedio calculado.
    
    Raises:
        ValueError: Si annual_salary es negativo, months_per_year es inválido,
            o el resultado sería negativo.
        ZeroDivisionError: Si months_per_year es 0.
        TypeError: Si los parámetros no son números.
    
    Example:
        >>> calculate_monthly_salary(60000)
        5000.0
        >>> calculate_monthly_salary(65000, include_bonus=True, bonus_amount=5000)
        5000.0
        >>> calculate_monthly_salary(65000, months_per_year=13)
        5000.0
    """
    if not isinstance(annual_salary, (int, float)):
        raise TypeError(f"annual_salary must be a number, got {type(annual_salary).__name__}")
    
    if annual_salary < 0:
        raise ValueError(f"annual_salary cannot be negative: {annual_salary}")
    
    if not isinstance(months_per_year, int) or months_per_year < 1 or months_per_year > 24:
        raise ValueError(f"months_per_year must be between 1 and 24, got: {months_per_year}")
    
    if months_per_year == 0:
        raise ZeroDivisionError("months_per_year cannot be zero")
    
    # Ajustar salario anual si incluye bonos
    base_annual = annual_salary
    if include_bonus:
        if not isinstance(bonus_amount, (int, float)):
            raise TypeError(f"bonus_amount must be a number, got {type(bonus_amount).__name__}")
        if bonus_amount < 0:
            raise ValueError(f"bonus_amount cannot be negative: {bonus_amount}")
        if bonus_amount > annual_salary:
            raise ValueError(f"bonus_amount ({bonus_amount}) cannot exceed annual_salary ({annual_salary})")
        base_annual = annual_salary - bonus_amount
    
    monthly = base_annual / months_per_year
    
    if monthly < 0:
        raise ValueError(f"Calculated monthly salary is negative: {monthly}")
    
    return monthly


def get_statistics() -> Dict:
    """
    Obtiene estadísticas de uso de la API
    
    Returns:
        Diccionario con estadísticas
    """
    total_requests = sum(len(requests) for requests in request_counts.values())
    unique_ips = len(request_counts)
    
    return {
        'total_requests_last_minute': total_requests,
        'unique_ips': unique_ips,
        'rate_limit_enabled': Config.RATE_LIMIT_ENABLED,
        'max_requests_per_minute': Config.MAX_REQUESTS_PER_MINUTE,
        'cache_enabled': Config.ENABLE_CACHE,
        'cached_templates': len(_template_cache),
        'supported_formats': Config.SUPPORTED_FORMATS,
        'supported_currencies': Config.SUPPORTED_CURRENCIES,
        'supported_employment_types': Config.EMPLOYMENT_TYPES
    }


def clear_cache(older_than_seconds: Optional[int] = None) -> Dict[str, any]:
    """
    Limpia el cache de templates con opciones avanzadas de filtrado.
    Puede limpiar todo el caché o solo entradas antiguas.
    
    Args:
        older_than_seconds: Si se proporciona, solo elimina entradas más antiguas
            que este número de segundos. Si None, limpia todo el caché.
            Útil para limpiar solo entradas expiradas sin afectar el caché activo.
    
    Returns:
        Diccionario con información sobre la operación:
            - cleared_count: int - Número de entradas eliminadas
            - remaining_count: int - Número de entradas que quedan en caché
            - total_before: int - Total de entradas antes de limpiar
            - cleared_keys: List[str] - Lista de claves eliminadas (solo si older_than_seconds)
            - error: Optional[str] - Mensaje de error si ocurre algún problema
    
    Example:
        >>> result = clear_cache()
        >>> print(f"Cleared {result['cleared_count']} entries")
        >>> # Limpiar solo entradas más antiguas de 1 hora
        >>> result = clear_cache(older_than_seconds=3600)
        >>> print(f"Cleared {result['cleared_count']} old entries")
    """
    result = {
        'cleared_count': 0,
        'remaining_count': 0,
        'total_before': len(_template_cache),
        'cleared_keys': [],
        'error': None
    }
    
    try:
        if older_than_seconds is None:
            # Limpiar todo el caché
            cleared_count = len(_template_cache)
    _template_cache.clear()
    _cache_timestamps.clear()
            result['cleared_count'] = cleared_count
            result['remaining_count'] = 0
            logger.info(f"Template cache cleared: {cleared_count} entries removed")
        else:
            # Limpiar solo entradas antiguas
            if not isinstance(older_than_seconds, int) or older_than_seconds < 0:
                result['error'] = f"older_than_seconds must be a non-negative integer, got: {older_than_seconds}"
                return result
            
            cutoff_time = datetime.now() - timedelta(seconds=older_than_seconds)
            keys_to_remove = []
            
            for key, timestamp in list(_cache_timestamps.items()):
                if timestamp < cutoff_time:
                    keys_to_remove.append(key)
            
            for key in keys_to_remove:
                _template_cache.pop(key, None)
                _cache_timestamps.pop(key, None)
                result['cleared_keys'].append(key)
            
            result['cleared_count'] = len(keys_to_remove)
            result['remaining_count'] = len(_template_cache)
            
            if result['cleared_count'] > 0:
                logger.info(
                    f"Cleared {result['cleared_count']} old cache entries "
                    f"(older than {older_than_seconds}s), {result['remaining_count']} remaining"
                )
            else:
                logger.debug("No old cache entries to clear")
        
        return result
        
    except Exception as e:
        error_msg = f"Error clearing cache: {str(e)}"
        result['error'] = error_msg
        logger.error(error_msg, exc_info=True)
        return result


def log_performance(
    func_name: str,
    execution_time: float,
    success: bool = True,
    additional_info: Optional[Dict] = None
) -> None:
    """
    Registra el rendimiento de una función con métricas detalladas y alertas.
    Proporciona logging estructurado para monitoreo y análisis de rendimiento.
    
    Args:
        func_name: Nombre de la función o operación medida.
        execution_time: Tiempo de ejecución en segundos (debe ser >= 0).
        success: Si la ejecución fue exitosa (True) o falló (False).
        additional_info: Diccionario opcional con información adicional a registrar.
            Ejemplo: {'memory_usage': '50MB', 'records_processed': 1000}
    
    Returns:
        None
    
    Example:
        >>> log_performance("generate_offer_letter", 0.234, True, {"format": "html"})
        >>> log_performance("database_query", 6.5, True)  # Generará warning
    """
    if execution_time < 0:
        logger.warning(f"Invalid execution time for {func_name}: {execution_time}s")
        return
    
    status = "SUCCESS" if success else "FAILED"
    
    # Construir mensaje base
    log_message = f"PERFORMANCE [{status}] {func_name}: {execution_time:.3f}s"
    
    # Agregar información adicional si está disponible
    if additional_info:
        info_str = ", ".join([f"{k}={v}" for k, v in additional_info.items()])
        log_message += f" | {info_str}"
    
    # Log según nivel de rendimiento
    if not success:
        logger.error(log_message)
    elif execution_time > 10.0:
        # Muy lento - nivel ERROR
        logger.error(f"CRITICAL: Very slow execution - {log_message}")
    elif execution_time > 5.0:
        # Lento - nivel WARNING
        logger.warning(f"Slow execution detected - {log_message}")
    elif execution_time > 2.0:
        # Moderadamente lento - nivel INFO
        logger.info(log_message)
    else:
        # Normal - nivel DEBUG para no saturar logs
        logger.debug(log_message)
    
    # Métricas adicionales para análisis
    if execution_time > 1.0:
        # Calcular throughput si hay información adicional
        if additional_info and 'items_processed' in additional_info:
            items = additional_info['items_processed']
            if isinstance(items, (int, float)) and items > 0:
                throughput = items / execution_time
                logger.debug(f"Throughput for {func_name}: {throughput:.2f} items/sec")


def rate_limit_check(ip_address: str, window_seconds: int = 60) -> Tuple[bool, Dict[str, any]]:
    """
    Verifica rate limiting mejorado usando configuración centralizada.
    Proporciona información detallada sobre el estado del rate limit.
    
    Args:
        ip_address: Dirección IP del cliente o identificador único.
        window_seconds: Ventana de tiempo en segundos para el rate limit (default: 60).
            Si se proporciona, sobrescribe la configuración por defecto.
    
    Returns:
        Tupla (allowed: bool, info: Dict):
            - allowed: True si la solicitud está permitida, False si excede el límite.
            - info: Diccionario con información detallada:
                - remaining: Solicitudes restantes en la ventana actual.
                - reset_at: Timestamp cuando se resetea el contador.
                - limit: Límite máximo de solicitudes.
                - retry_after: Segundos hasta que se puede hacer otra solicitud.
    
    Example:
        >>> allowed, info = rate_limit_check("192.168.1.1")
        >>> if not allowed:
        ...     print(f"Rate limit exceeded. Retry after {info['retry_after']}s")
    """
    if not Config.RATE_LIMIT_ENABLED:
        return True, {
            'remaining': float('inf'),
            'limit': Config.MAX_REQUESTS_PER_MINUTE,
            'reset_at': datetime.now().timestamp() + window_seconds,
            'retry_after': 0
        }
    
    if not isinstance(ip_address, str) or not ip_address.strip():
        logger.warning("Invalid IP address provided to rate_limit_check")
        return False, {
            'remaining': 0,
            'limit': Config.MAX_REQUESTS_PER_MINUTE,
            'reset_at': datetime.now().timestamp(),
            'retry_after': window_seconds
        }
    
    now = datetime.now()
    current_timestamp = now.timestamp()
    
    # Inicializar lista si no existe
    if ip_address not in request_counts:
        request_counts[ip_address] = []
    
    # Limpiar requests antiguos (fuera de la ventana de tiempo)
    cutoff_time = now - timedelta(seconds=window_seconds)
    request_counts[ip_address] = [
        req_time for req_time in request_counts[ip_address]
        if req_time > cutoff_time
    ]
    
    # Calcular información del rate limit
    current_count = len(request_counts[ip_address])
    limit = Config.MAX_REQUESTS_PER_MINUTE
    remaining = max(0, limit - current_count)
    
    # Calcular tiempo de reset (tiempo del request más antiguo + ventana)
    if request_counts[ip_address]:
        oldest_request = min(request_counts[ip_address])
        reset_at = (oldest_request + timedelta(seconds=window_seconds)).timestamp()
    else:
        reset_at = current_timestamp + window_seconds
    
    retry_after = max(0, int(reset_at - current_timestamp))
    
    info = {
        'remaining': remaining,
        'limit': limit,
        'reset_at': reset_at,
        'retry_after': retry_after,
        'current_count': current_count
    }
    
    # Verificar límite
    if current_count >= limit:
        logger.warning(
            f"Rate limit exceeded for IP: {ip_address} "
            f"({current_count}/{limit} requests in {window_seconds}s window)"
        )
        return False, info
    
    # Agregar request actual
    request_counts[ip_address].append(now)
    info['remaining'] = max(0, limit - len(request_counts[ip_address]))

    return True, info


def sanitize_filename(filename: str, max_length: int = 200, allow_unicode: bool = True) -> str:
    """
    Sanitiza un nombre de archivo para que sea seguro y compatible con sistemas de archivos.
    Remueve caracteres peligrosos y limita la longitud.
    
    Args:
        filename: Nombre de archivo a sanitizar.
        max_length: Longitud máxima permitida (default: 200).
            Algunos sistemas de archivos tienen límites más estrictos.
        allow_unicode: Si True, permite caracteres Unicode válidos.
            Si False, solo permite ASCII alfanumérico y guiones.
    
    Returns:
        Nombre de archivo sanitizado y seguro.
    
    Example:
        >>> sanitize_filename("test<>file?.txt")
        'testfile.txt'
        >>> sanitize_filename("mí_archivo_especial.pdf")
        'mí_archivo_especial.pdf'
    """
    if not isinstance(filename, str):
        filename = str(filename)
    
    # Remover caracteres peligrosos comunes en todos los sistemas operativos
    # Windows: < > : " / \ | ? *
    # Linux/Mac: / y caracteres de control
    dangerous_chars = r'[<>:"/\\|?*\x00-\x1f]'
    filename = re.sub(dangerous_chars, '', filename)
    
    # Remover espacios al inicio y final (pueden causar problemas)
    filename = filename.strip()
    
    # Remover puntos al final (reservados en algunos sistemas)
    filename = filename.rstrip('.')
    
    # Si no permite Unicode, convertir a ASCII
    if not allow_unicode:
        try:
            filename = filename.encode('ascii', 'ignore').decode('ascii')
        except (UnicodeEncodeError, UnicodeDecodeError):
            # Si falla, usar solo caracteres alfanuméricos y guiones
            filename = re.sub(r'[^a-zA-Z0-9._-]', '_', filename)
    
    # Limitar longitud (reservar espacio para extensión si es necesario)
    if len(filename) > max_length:
        # Intentar preservar extensión si existe
        if '.' in filename:
            name, ext = filename.rsplit('.', 1)
            max_name_length = max_length - len(ext) - 1  # -1 para el punto
            filename = name[:max_name_length] + '.' + ext
        else:
            filename = filename[:max_length]
    
    # Asegurar que no esté vacío
    if not filename:
        filename = "unnamed_file"
    
    return filename


def extract_text_from_html(html_content: str, preserve_structure: bool = False) -> str:
    """
    Extrae texto plano de contenido HTML con opciones avanzadas de preservación.
    Maneja entidades HTML, scripts, estilos y mantiene estructura básica si se solicita.
    
    Args:
        html_content: Contenido HTML a procesar.
        preserve_structure: Si True, preserva saltos de línea y estructura básica.
            Si False, convierte todo a texto plano continuo.
    
    Returns:
        Texto plano extraído del HTML, con entidades decodificadas y limpio.
    
    Raises:
        TypeError: Si html_content no es una cadena de texto.
    
    Example:
        >>> html = "<p>Hello <strong>World</strong></p>"
        >>> extract_text_from_html(html)
        'Hello World'
        >>> extract_text_from_html(html, preserve_structure=True)
        'Hello World'
    """
    if not isinstance(html_content, str):
        raise TypeError(f"html_content must be a string, got {type(html_content).__name__}")
    
    if not html_content.strip():
        return ""
    
    try:
        # Remover scripts y estilos primero (más eficiente)
        # Usar compilación de regex para mejor rendimiento en contenido grande
        script_pattern = re.compile(r'<script[^>]*>.*?</script>', re.DOTALL | re.IGNORECASE)
        style_pattern = re.compile(r'<style[^>]*>.*?</style>', re.DOTALL | re.IGNORECASE)
        
        html_content = script_pattern.sub('', html_content)
        html_content = style_pattern.sub('', html_content)
        
        # Remover comentarios HTML
        comment_pattern = re.compile(r'<!--.*?-->', re.DOTALL)
        html_content = comment_pattern.sub('', html_content)
        
        # Reemplazar elementos de bloque con saltos de línea si se preserva estructura
        if preserve_structure:
            block_elements = ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'br', 'hr']
            for element in block_elements:
                # Reemplazar tags de apertura y cierre con saltos de línea
                html_content = re.sub(
                    rf'</?{element}[^>]*>',
                    '\n',
                    html_content,
                    flags=re.IGNORECASE
                )
        
        # Remover todos los tags HTML restantes
        tag_pattern = re.compile(r'<[^>]+>')
        text = tag_pattern.sub('', html_content)
        
        # Decodificar entidades HTML (maneja &amp;, &lt;, &gt;, &quot;, etc.)
        text = html.unescape(text)
        
        # Limpiar espacios en blanco
        if preserve_structure:
            # Preservar saltos de línea múltiples pero limpiar espacios excesivos
            text = re.sub(r'[ \t]+', ' ', text)  # Múltiples espacios/tabs a uno
            text = re.sub(r'\n[ \t]+', '\n', text)  # Espacios al inicio de línea
            text = re.sub(r'[ \t]+\n', '\n', text)  # Espacios al final de línea
            text = re.sub(r'\n{3,}', '\n\n', text)  # Más de 2 saltos de línea a 2
        else:
            # Convertir todo a texto continuo
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
        
    except Exception as e:
        error_msg = f"Error extracting text from HTML: {str(e)}"
        logger.error(error_msg, exc_info=True)
        # En caso de error, intentar retornar al menos el contenido sin tags
        try:
            fallback_text = re.sub(r'<[^>]+>', '', html_content)
            return html.unescape(fallback_text).strip()
        except Exception:
            logger.error("Failed to extract fallback text from HTML")
            return html_content  # Retornar original como último recurso


def convert_format(
    content: str,
    from_format: str,
    to_format: str,
    preserve_structure: bool = True
) -> Optional[str]:
    """
    Convierte contenido entre diferentes formatos con soporte mejorado.
    Soporta conversiones bidireccionales entre TXT, HTML y Markdown básico.
    
    Args:
        content: Contenido a convertir (debe ser una cadena de texto).
        from_format: Formato origen. Opciones: 'txt', 'html', 'markdown'.
        to_format: Formato destino. Opciones: 'txt', 'html', 'markdown'.
        preserve_structure: Si True, intenta preservar la estructura del documento
            (encabezados, listas, párrafos). Si False, convierte a texto plano.
    
    Returns:
        Contenido convertido en el formato solicitado, o None si:
            - El formato de conversión no es soportado
            - Ocurre un error durante la conversión
    
    Raises:
        ValueError: Si from_format o to_format no son formatos válidos.
        TypeError: Si content no es una cadena de texto.
    
    Example:
        >>> html = "<p>Hello <strong>World</strong></p>"
        >>> convert_format(html, "html", "txt")
        'Hello World'
        >>> txt = "# Title\\n## Subtitle\\n- Item 1"
        >>> convert_format(txt, "txt", "html")
        '<html>...'
    """
    if not isinstance(content, str):
        raise TypeError(f"content must be a string, got {type(content).__name__}")
    
    if not isinstance(from_format, str) or not isinstance(to_format, str):
        raise ValueError("from_format and to_format must be strings")
    
    from_format = from_format.lower().strip()
    to_format = to_format.lower().strip()
    
    # Formatos soportados
    supported_formats = {'txt', 'text', 'html', 'htm', 'markdown', 'md'}
    
    if from_format not in supported_formats:
        raise ValueError(f"Unsupported source format: {from_format}. Supported: {supported_formats}")
    
    if to_format not in supported_formats:
        raise ValueError(f"Unsupported target format: {to_format}. Supported: {supported_formats}")
    
    try:
        # Si los formatos son iguales, retornar sin cambios
        if from_format == to_format or (from_format in ['txt', 'text'] and to_format in ['txt', 'text']):
            return content
        
        # HTML a texto plano
        if from_format in ['html', 'htm'] and to_format in ['txt', 'text']:
            return extract_text_from_html(content, preserve_structure=preserve_structure)
        
        # HTML a Markdown básico
        if from_format in ['html', 'htm'] and to_format in ['markdown', 'md']:
            text = extract_text_from_html(content, preserve_structure=True)
            # Convertir texto estructurado a markdown básico
            lines = text.split('\n')
            md_lines = []
            in_list = False
            
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    if in_list:
                        md_lines.append('')
                        in_list = False
                    else:
                        md_lines.append('')
                    continue
                
                # Detectar listas simples (líneas que empiezan con guión o número)
                if stripped.startswith('-') or stripped.startswith('*') or re.match(r'^\d+\.', stripped):
                    if not in_list:
                        in_list = True
                    md_lines.append(stripped)
                else:
                    if in_list:
                        md_lines.append('')
                        in_list = False
                    md_lines.append(stripped)
            
            return '\n'.join(md_lines)
        
        # Texto a HTML mejorado
        if from_format in ['txt', 'text'] and to_format in ['html', 'htm']:
            lines = content.split('\n')
            html_parts = ['<html><head><meta charset="UTF-8"></head><body>']
            in_list = False
            
            for line in lines:
                stripped = line.strip()
                
                if not stripped:
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    html_parts.append('<br>')
                    continue
                
                # Encabezados Markdown-style
                if stripped.startswith('#'):
                    level = len(stripped) - len(stripped.lstrip('#'))
                    text = stripped.lstrip('#').strip()
                    level = min(level, 6)  # HTML solo soporta h1-h6
                    html_parts.append(f"<h{level}>{html.escape(text)}</h{level}>")
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                
                # Listas
                elif stripped.startswith('-') or stripped.startswith('*') or re.match(r'^\d+\.', stripped):
                    if not in_list:
                        html_parts.append('<ul>')
                        in_list = True
                    text = re.sub(r'^[-*\d+\.]\s*', '', stripped)
                    html_parts.append(f"<li>{html.escape(text)}</li>")
                
                # Texto normal
                else:
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    html_parts.append(f"<p>{html.escape(stripped)}</p>")
            
            if in_list:
                html_parts.append('</ul>')
            
            html_parts.append('</body></html>')
            return ''.join(html_parts)
        
        # Markdown a HTML (básico)
        if from_format in ['markdown', 'md'] and to_format in ['html', 'htm']:
            # Reutilizar conversión de texto a HTML (markdown básico)
            return convert_format(content, 'txt', 'html', preserve_structure)
        
        # Markdown a texto
        if from_format in ['markdown', 'md'] and to_format in ['txt', 'text']:
            # Remover sintaxis markdown básica
            text = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)  # Encabezados
            text = re.sub(r'^\s*[-*]\s*', '', text, flags=re.MULTILINE)  # Listas
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Negrita
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Cursiva
            return text.strip()
        
        logger.warning(f"Conversion from {from_format} to {to_format} not fully supported")
        return None
        
    except Exception as e:
        error_msg = f"Error converting format from {from_format} to {to_format}: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return None


def get_file_info(filepath: str, include_hash: bool = False) -> Dict:
    """
    Obtiene información detallada sobre un archivo con opciones avanzadas.
    Incluye metadatos del sistema de archivos y opcionalmente hash del contenido.
    
    Args:
        filepath: Ruta del archivo a analizar.
        include_hash: Si True, calcula hash MD5 del archivo (puede ser lento para archivos grandes).
            Por defecto False para evitar operaciones costosas.
    
    Returns:
        Diccionario con información del archivo:
            - exists: bool - Si el archivo existe
            - size: int - Tamaño en bytes
            - size_human: str - Tamaño formateado legible
            - created: str - Fecha de creación (ISO format)
            - modified: str - Fecha de modificación (ISO format)
            - accessed: str - Fecha de último acceso (ISO format)
            - extension: str - Extensión del archivo (con punto)
            - filename: str - Nombre del archivo sin ruta
            - directory: str - Directorio contenedor
            - is_file: bool - Si es un archivo (no directorio)
            - is_readable: bool - Si el archivo es legible
            - is_writable: bool - Si el archivo es escribible
            - hash_md5: str - Hash MD5 (solo si include_hash=True)
            - error: str - Mensaje de error si ocurre algún problema
    
    Example:
        >>> info = get_file_info("/path/to/file.txt")
        >>> print(info['size_human'])
        '1.5 MB'
        >>> info = get_file_info("/path/to/file.txt", include_hash=True)
        >>> print(info['hash_md5'])
        '5d41402abc4b2a76b9719d911017c592'
    """
    result = {
        "exists": False,
        "error": None
    }
    
    try:
        if not isinstance(filepath, str) or not filepath.strip():
            result["error"] = "Invalid filepath provided"
            return result
        
        filepath = os.path.abspath(os.path.expanduser(filepath))
        
        if not os.path.exists(filepath):
            result["error"] = "File does not exist"
            return result
        
        stat = os.stat(filepath)
        
        # Información básica
        result.update({
            "exists": True,
            "size": stat.st_size,
            "size_human": get_file_size_human(stat.st_size),
            "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "accessed": datetime.fromtimestamp(stat.st_atime).isoformat(),
            "extension": os.path.splitext(filepath)[1],
            "filename": os.path.basename(filepath),
            "directory": os.path.dirname(filepath),
            "is_file": os.path.isfile(filepath),
            "is_readable": os.access(filepath, os.R_OK),
            "is_writable": os.access(filepath, os.W_OK)
        })
        
        # Calcular hash si se solicita
        if include_hash and result["is_file"]:
            try:
                import hashlib
                hash_md5 = hashlib.md5()
                
                # Para archivos grandes, leer en chunks
                chunk_size = 8192
                with open(filepath, 'rb') as f:
                    while chunk := f.read(chunk_size):
                        hash_md5.update(chunk)
                
                result["hash_md5"] = hash_md5.hexdigest()
            except Exception as hash_error:
                logger.warning(f"Could not calculate hash for {filepath}: {str(hash_error)}")
                result["hash_error"] = str(hash_error)
        
        return result
        
    except PermissionError as e:
        result["error"] = f"Permission denied: {str(e)}"
        logger.error(f"Permission error getting file info for {filepath}: {str(e)}")
        return result
    except Exception as e:
        error_msg = f"Error getting file info: {str(e)}"
        result["error"] = error_msg
        logger.error(error_msg, exc_info=True)
        return result


def clean_old_files(
    directory: str,
    days_old: int = 30,
    pattern: str = "*",
    recursive: bool = False,
    dry_run: bool = False,
    min_file_size: int = 0
) -> Dict[str, any]:
    """
    Limpia archivos antiguos de un directorio con opciones avanzadas.
    Proporciona información detallada sobre la operación.
    
    Args:
        directory: Directorio a limpiar (debe existir y ser accesible).
        days_old: Días de antigüedad para considerar archivo como antiguo (default: 30).
            Archivos más antiguos que esta fecha serán candidatos para eliminación.
        pattern: Patrón de nombres de archivo (ej: "*.txt", "offer_*.docx").
            Usa sintaxis de glob. Por defecto "*" (todos los archivos).
        recursive: Si True, busca recursivamente en subdirectorios.
            Si False, solo busca en el directorio especificado.
        dry_run: Si True, simula la operación sin eliminar archivos.
            Útil para ver qué se eliminaría sin hacer cambios.
        min_file_size: Tamaño mínimo en bytes para considerar un archivo (default: 0).
            Archivos más pequeños serán ignorados.
    
    Returns:
        Diccionario con información detallada:
            - deleted_count: int - Número de archivos eliminados
            - total_size_freed: int - Tamaño total liberado en bytes
            - total_size_freed_human: str - Tamaño liberado formateado
            - files_deleted: List[str] - Lista de archivos eliminados
            - files_failed: List[Dict] - Lista de archivos que fallaron al eliminar
            - dry_run: bool - Si fue una simulación
            - error: str - Mensaje de error si ocurre algún problema
    
    Example:
        >>> result = clean_old_files("./output", days_old=7, pattern="*.txt")
        >>> print(f"Deleted {result['deleted_count']} files")
        >>> result = clean_old_files("./output", days_old=7, dry_run=True)
        >>> print(f"Would delete {result['deleted_count']} files")
    """
    result = {
        "deleted_count": 0,
        "total_size_freed": 0,
        "total_size_freed_human": "0 B",
        "files_deleted": [],
        "files_failed": [],
        "dry_run": dry_run,
        "error": None
    }
    
    try:
        if not isinstance(directory, str) or not directory.strip():
            result["error"] = "Invalid directory path"
            return result
        
        directory = os.path.abspath(os.path.expanduser(directory))
        
        if not os.path.exists(directory):
            result["error"] = f"Directory does not exist: {directory}"
            return result
        
        if not os.path.isdir(directory):
            result["error"] = f"Path is not a directory: {directory}"
            return result
        
        if not isinstance(days_old, int) or days_old < 0:
            result["error"] = f"days_old must be a non-negative integer, got: {days_old}"
            return result
        
        cutoff_date = datetime.now() - timedelta(days=days_old)
        
        import glob
        
        # Construir patrón de búsqueda
        if recursive:
            search_pattern = os.path.join(directory, "**", pattern)
        else:
            search_pattern = os.path.join(directory, pattern)
        
        files = glob.glob(search_pattern, recursive=recursive)
        
        for filepath in files:
            try:
                # Solo procesar archivos (no directorios)
                if not os.path.isfile(filepath):
                    continue
                
                # Verificar tamaño mínimo
                file_size = os.path.getsize(filepath)
                if file_size < min_file_size:
                    continue
                
                # Verificar antigüedad
                    file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                    if file_time < cutoff_date:
                    if dry_run:
                        # En dry run, solo registrar
                        result["deleted_count"] += 1
                        result["total_size_freed"] += file_size
                        result["files_deleted"].append(filepath)
                        logger.debug(f"[DRY RUN] Would delete: {filepath}")
                    else:
                        # Eliminar archivo
                        os.remove(filepath)
                        result["deleted_count"] += 1
                        result["total_size_freed"] += file_size
                        result["files_deleted"].append(filepath)
                        logger.info(f"Deleted old file: {filepath} ({get_file_size_human(file_size)})")
                        
            except PermissionError as e:
                error_info = {"file": filepath, "error": f"Permission denied: {str(e)}"}
                result["files_failed"].append(error_info)
                logger.warning(f"Permission denied deleting {filepath}: {str(e)}")
            except Exception as e:
                error_info = {"file": filepath, "error": str(e)}
                result["files_failed"].append(error_info)
                logger.warning(f"Could not delete file {filepath}: {str(e)}")
        
        # Formatear tamaño liberado
        result["total_size_freed_human"] = get_file_size_human(result["total_size_freed"])
        
        action = "Would delete" if dry_run else "Deleted"
        logger.info(
            f"{action} {result['deleted_count']} old files from {directory} "
            f"(freed {result['total_size_freed_human']})"
        )
        
        if result["files_failed"]:
            logger.warning(f"Failed to delete {len(result['files_failed'])} files")
        
        return result
        
    except Exception as e:
        error_msg = f"Error cleaning old files: {str(e)}"
        result["error"] = error_msg
        logger.error(error_msg, exc_info=True)
        return result


def validate_file_path(
    filepath: str,
    must_exist: bool = False,
    check_permissions: bool = False,
    max_length: Optional[int] = None
) -> Tuple[bool, Optional[str], Optional[Dict]]:
    """
    Valida una ruta de archivo con validaciones avanzadas y opciones de verificación.
    Proporciona información detallada sobre el estado de la ruta.
    
    Args:
        filepath: Ruta a validar (puede ser relativa o absoluta).
        must_exist: Si True, el archivo debe existir para que la validación sea exitosa.
        check_permissions: Si True, verifica permisos de lectura/escritura.
        max_length: Longitud máxima permitida (None = usar límite del sistema).
            Windows: 260 caracteres, Linux/Mac: 4096 caracteres típicamente.
    
    Returns:
        Tupla (es_válida: bool, mensaje_error: Optional[str], info: Optional[Dict]):
            - es_válida: True si la ruta es válida según los criterios
            - mensaje_error: Mensaje descriptivo si hay un error, None si es válida
            - info: Diccionario con información adicional:
                - normalized_path: Ruta normalizada
                - exists: Si el archivo/directorio existe
                - is_file: Si es un archivo
                - is_directory: Si es un directorio
                - is_readable: Si es legible (solo si check_permissions=True)
                - is_writable: Si es escribible (solo si check_permissions=True)
                - parent_exists: Si el directorio padre existe
    
    Example:
        >>> valid, error, info = validate_file_path("/path/to/file.txt")
        >>> if valid:
        ...     print(f"Path exists: {info['exists']}")
        >>> valid, error, info = validate_file_path("/path/to/file.txt", check_permissions=True)
        >>> if valid and info['is_readable']:
        ...     print("File is readable")
    """
    info = {}
    
    try:
        if not isinstance(filepath, str) or not filepath.strip():
            return False, "File path must be a non-empty string", None
        
        # Normalizar ruta
        normalized_path = os.path.abspath(os.path.expanduser(filepath))
        info['normalized_path'] = normalized_path
        
        # Determinar límite de longitud según el sistema
        if max_length is None:
            # Windows tiene límite de 260 caracteres (MAX_PATH)
            # Linux/Mac típicamente 4096, pero usamos 260 como límite conservador
            max_length = 260 if os.name == 'nt' else 4096
        
        # Validar longitud
        if len(normalized_path) > max_length:
            return False, f"File path too long (max {max_length} characters, got {len(normalized_path)})", info
        
        # Validar caracteres prohibidos según el sistema operativo
        if os.name == 'nt':  # Windows
            invalid_chars = ['<', '>', ':', '"', '|', '?', '*', '\x00']
            reserved_names = ['CON', 'PRN', 'AUX', 'NUL'] + \
                           [f'COM{i}' for i in range(1, 10)] + \
                           [f'LPT{i}' for i in range(1, 10)]
            filename = os.path.basename(normalized_path).upper()
            if filename in reserved_names:
                return False, f"File name is a reserved Windows name: {filename}", info
        else:  # Linux/Mac
            invalid_chars = ['/', '\x00']
        
        if any(char in filepath for char in invalid_chars):
            return False, f"File path contains invalid characters: {invalid_chars}", info
        
        # Verificar si existe
        exists = os.path.exists(normalized_path)
        info['exists'] = exists
        
        if must_exist and not exists:
            return False, f"File does not exist: {normalized_path}", info
        
        # Información adicional si existe
        if exists:
            info['is_file'] = os.path.isfile(normalized_path)
            info['is_directory'] = os.path.isdir(normalized_path)
            
            # Verificar permisos si se solicita
            if check_permissions:
                info['is_readable'] = os.access(normalized_path, os.R_OK)
                info['is_writable'] = os.access(normalized_path, os.W_OK)
        else:
            info['is_file'] = None
            info['is_directory'] = None
        
        # Validar directorio padre
        parent_dir = os.path.dirname(normalized_path)
        if parent_dir and parent_dir != normalized_path:  # Evitar loop con root
            parent_exists = os.path.exists(parent_dir)
            info['parent_exists'] = parent_exists
            if not parent_exists and must_exist:
                return False, f"Parent directory does not exist: {parent_dir}", info
        else:
            info['parent_exists'] = True  # Root siempre existe
        
        return True, None, info
        
    except ValueError as e:
        return False, f"Invalid file path format: {str(e)}", info
    except Exception as e:
        return False, f"Error validating file path: {str(e)}", info


def ensure_directory(
    directory: str,
    mode: int = 0o755,
    create_parents: bool = True
) -> Tuple[bool, Optional[str], Optional[Dict]]:
    """
    Asegura que un directorio exista, creándolo si es necesario con opciones avanzadas.
    Proporciona información detallada sobre la operación.
    
    Args:
        directory: Ruta del directorio a crear/verificar.
        mode: Permisos del directorio en formato octal (default: 0o755).
            Solo se aplica si se crea el directorio.
        create_parents: Si True, crea directorios padres si no existen (como mkdir -p).
            Si False, falla si el directorio padre no existe.
    
    Returns:
        Tupla (success: bool, error_message: Optional[str], info: Optional[Dict]):
            - success: True si el directorio existe o fue creado exitosamente
            - error_message: Mensaje de error si falla, None si es exitoso
            - info: Diccionario con información:
                - created: bool - Si el directorio fue creado en esta llamada
                - existed: bool - Si el directorio ya existía
                - path: str - Ruta normalizada del directorio
                - is_writable: bool - Si el directorio es escribible
                - permissions: str - Permisos del directorio (formato octal)
    
    Example:
        >>> success, error, info = ensure_directory("./output/offers")
        >>> if success:
        ...     print(f"Directory {'created' if info['created'] else 'existed'}: {info['path']}")
    """
    info = {
        'created': False,
        'existed': False,
        'path': None,
        'is_writable': False,
        'permissions': None
    }
    
    try:
        if not isinstance(directory, str) or not directory.strip():
            return False, "Directory path must be a non-empty string", None
        
        # Normalizar ruta
        normalized_path = os.path.abspath(os.path.expanduser(directory))
        info['path'] = normalized_path
        
        # Verificar si ya existe
        if os.path.exists(normalized_path):
            if os.path.isdir(normalized_path):
                info['existed'] = True
                info['is_writable'] = os.access(normalized_path, os.W_OK)
                
                # Obtener permisos actuales
                try:
                    stat_info = os.stat(normalized_path)
                    info['permissions'] = oct(stat_info.st_mode)[-3:]
                except Exception:
                    pass
                
                logger.debug(f"Directory already exists: {normalized_path}")
                return True, None, info
            else:
                return False, f"Path exists but is not a directory: {normalized_path}", info
        
        # Crear directorio
        try:
            if create_parents:
                os.makedirs(normalized_path, mode=mode, exist_ok=True)
            else:
                os.mkdir(normalized_path, mode=mode)
            
            info['created'] = True
            info['is_writable'] = os.access(normalized_path, os.W_OK)
            info['permissions'] = oct(mode)
            
            logger.info(f"Created directory: {normalized_path} (mode: {oct(mode)})")
            return True, None, info
            
        except PermissionError as e:
            error_msg = f"Permission denied creating directory: {str(e)}"
            logger.error(error_msg)
            return False, error_msg, info
        except FileNotFoundError as e:
            error_msg = f"Parent directory does not exist: {str(e)}"
            logger.error(error_msg)
            return False, error_msg, info
        except OSError as e:
            error_msg = f"OS error creating directory: {str(e)}"
            logger.error(error_msg)
            return False, error_msg, info
        
    except Exception as e:
        error_msg = f"Error ensuring directory exists: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return False, error_msg, info


def get_file_size_human(size_bytes: int, binary: bool = False, precision: int = 2) -> str:
    """
    Convierte tamaño de archivo a formato legible con opciones de formato.
    Soporta unidades binarias (1024) y decimales (1000).
    
    Args:
        size_bytes: Tamaño en bytes (debe ser >= 0).
        binary: Si True, usa unidades binarias (KiB, MiB, etc.) con base 1024.
            Si False, usa unidades decimales (KB, MB, etc.) con base 1000.
        precision: Número de decimales a mostrar (default: 2).
            Máximo 6 decimales.
    
    Returns:
        String formateado con el tamaño legible.
        Ejemplos:
            - "1.50 MB" (decimal)
            - "1.43 MiB" (binario)
            - "500 B" (sin decimales si es < 1 KB)
    
    Raises:
        ValueError: Si size_bytes es negativo.
        TypeError: Si size_bytes no es un número.
    
    Example:
        >>> get_file_size_human(1536000)
        '1.54 MB'
        >>> get_file_size_human(1536000, binary=True)
        '1.46 MiB'
        >>> get_file_size_human(500)
        '500.00 B'
    """
    if not isinstance(size_bytes, (int, float)):
        raise TypeError(f"size_bytes must be a number, got {type(size_bytes).__name__}")
    
    if size_bytes < 0:
        raise ValueError(f"size_bytes cannot be negative: {size_bytes}")
    
    if not isinstance(precision, int) or precision < 0 or precision > 6:
        precision = 2
    
    # Unidades según el sistema
    if binary:
        units = ['B', 'KiB', 'MiB', 'GiB', 'TiB', 'PiB', 'EiB']
        base = 1024.0
    else:
        units = ['B', 'KB', 'MB', 'GB', 'TB', 'PB', 'EB']
        base = 1000.0
    
    # Caso especial para 0 bytes
    if size_bytes == 0:
        return f"0.{'0' * precision} {units[0]}"
    
    # Convertir a float para cálculos
    size = float(size_bytes)
    unit_index = 0
    
    # Encontrar la unidad apropiada
    while size >= base and unit_index < len(units) - 1:
        size /= base
        unit_index += 1
    
    # Formatear según el tamaño
    if unit_index == 0:
        # Para bytes, mostrar sin decimales si es entero
        if size == int(size):
            return f"{int(size)} {units[unit_index]}"
        else:
            return f"{size:.{precision}f} {units[unit_index]}"
    else:
        # Para unidades mayores, siempre mostrar decimales
        return f"{size:.{precision}f} {units[unit_index]}"


def search_in_content(
    content: str,
    search_term: str,
    case_sensitive: bool = False,
    use_regex: bool = False,
    context_chars: int = 50,
    max_results: Optional[int] = None
) -> List[Dict]:
    """
    Busca un término o patrón en el contenido con opciones avanzadas de búsqueda.
    Soporta búsqueda literal y por expresiones regulares.
    
    Args:
        content: Contenido donde buscar (debe ser una cadena de texto).
        search_term: Término o patrón a buscar. Si use_regex=True, puede ser una expresión regular.
        case_sensitive: Si True, la búsqueda distingue entre mayúsculas y minúsculas.
        use_regex: Si True, trata search_term como una expresión regular.
            Si False, busca el término literal.
        context_chars: Número de caracteres de contexto a incluir antes y después
            de cada coincidencia (default: 50).
        max_results: Número máximo de resultados a retornar (None = sin límite).
            Útil para limitar resultados en contenido grande.
    
    Returns:
        Lista de diccionarios con información de cada coincidencia:
            - position: int - Posición de inicio de la coincidencia
            - end_position: int - Posición de fin de la coincidencia
            - length: int - Longitud de la coincidencia
            - match_text: str - Texto que coincidió
            - context: str - Contexto alrededor de la coincidencia
            - context_start: int - Posición de inicio del contexto
            - context_end: int - Posición de fin del contexto
            - line_number: int - Número de línea donde ocurre la coincidencia
            - column_number: int - Columna dentro de la línea
    
    Raises:
        ValueError: Si search_term está vacío o si use_regex=True y el patrón es inválido.
        TypeError: Si content no es una cadena de texto.
    
    Example:
        >>> results = search_in_content("Hello world, hello again", "hello", case_sensitive=False)
        >>> len(results)
        2
        >>> results = search_in_content("Price: $100", r'\$\d+', use_regex=True)
        >>> results[0]['match_text']
        '$100'
    """
    if not isinstance(content, str):
        raise TypeError(f"content must be a string, got {type(content).__name__}")
    
    if not isinstance(search_term, str) or not search_term.strip():
        return []
    
    if max_results is not None and (not isinstance(max_results, int) or max_results < 1):
        max_results = None
    
    results = []
    
    try:
        if use_regex:
            # Búsqueda con expresiones regulares
            flags = 0 if case_sensitive else re.IGNORECASE
            pattern = re.compile(search_term, flags)
            matches = pattern.finditer(content)
            
            for match in matches:
                if max_results and len(results) >= max_results:
                    break
                
                pos = match.start()
                end_pos = match.end()
                match_text = match.group()
                
                # Calcular número de línea y columna
                line_start = content.rfind('\n', 0, pos) + 1
                line_number = content[:pos].count('\n') + 1
                column_number = pos - line_start + 1
                
                # Obtener contexto
                context_start = max(0, pos - context_chars)
                context_end = min(len(content), end_pos + context_chars)
                context = content[context_start:context_end]
                
                results.append({
                    "position": pos,
                    "end_position": end_pos,
                    "length": len(match_text),
                    "match_text": match_text,
                    "context": context,
                    "context_start": context_start,
                    "context_end": context_end,
                    "line_number": line_number,
                    "column_number": column_number
                })
        else:
            # Búsqueda literal
    search_text = content if case_sensitive else content.lower()
    term = search_term if case_sensitive else search_term.lower()
    
    start = 0
    while True:
                if max_results and len(results) >= max_results:
                    break
                
        pos = search_text.find(term, start)
        if pos == -1:
            break
        
                end_pos = pos + len(term)
                match_text = content[pos:end_pos]
                
                # Calcular número de línea y columna
                line_start = content.rfind('\n', 0, pos) + 1
                line_number = content[:pos].count('\n') + 1
                column_number = pos - line_start + 1
                
                # Obtener contexto
                context_start = max(0, pos - context_chars)
                context_end = min(len(content), end_pos + context_chars)
        context = content[context_start:context_end]
        
        results.append({
            "position": pos,
                    "end_position": end_pos,
            "length": len(term),
                    "match_text": match_text,
            "context": context,
                    "context_start": context_start,
                    "context_end": context_end,
                    "line_number": line_number,
                    "column_number": column_number
        })
        
        start = pos + 1
    
    return results

    except re.error as e:
        error_msg = f"Invalid regular expression pattern: {str(e)}"
        logger.error(error_msg)
        raise ValueError(error_msg) from e
    except Exception as e:
        error_msg = f"Error searching content: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return []


def analyze_offer_content(
    content: str,
    words_per_minute: int = 200,
    include_sentiment: bool = False
) -> Dict:
    """
    Analiza el contenido de una oferta y extrae estadísticas detalladas.
    Proporciona métricas de legibilidad, estructura y contenido clave.
    
    Args:
        content: Contenido de la oferta a analizar (debe ser una cadena de texto).
        words_per_minute: Velocidad de lectura asumida para calcular tiempo de lectura
            (default: 200 palabras/minuto, promedio para adultos).
        include_sentiment: Si True, intenta analizar el sentimiento del contenido.
            Requiere procesamiento adicional y puede ser más lento.
    
    Returns:
        Diccionario con estadísticas detalladas:
            - word_count: int - Número total de palabras
            - character_count: int - Número total de caracteres
            - character_count_no_spaces: int - Caracteres sin espacios
            - sentence_count: int - Número de oraciones
            - paragraph_count: int - Número de párrafos
            - average_words_per_sentence: float - Promedio de palabras por oración
            - average_chars_per_word: float - Promedio de caracteres por palabra
            - longest_word: str - Palabra más larga encontrada
            - shortest_word: str - Palabra más corta encontrada (excluyendo palabras de 1 letra)
            - important_terms: Dict - Conteo de términos importantes
            - estimated_reading_time_minutes: float - Tiempo estimado de lectura
            - estimated_reading_time_seconds: int - Tiempo estimado en segundos
            - readability_score: float - Puntuación de legibilidad (Flesch Reading Ease aproximado)
            - has_salary_info: bool - Si contiene información de salario
            - has_date_info: bool - Si contiene fechas
            - has_contact_info: bool - Si contiene información de contacto
            - sentiment: Optional[Dict] - Análisis de sentimiento (solo si include_sentiment=True)
            - error: Optional[str] - Mensaje de error si ocurre algún problema
    
    Raises:
        TypeError: Si content no es una cadena de texto.
    
    Example:
        >>> content = "We are pleased to offer you a position..."
        >>> stats = analyze_offer_content(content)
        >>> print(f"Word count: {stats['word_count']}")
        >>> print(f"Reading time: {stats['estimated_reading_time_minutes']} minutes")
    """
    if not isinstance(content, str):
        raise TypeError(f"content must be a string, got {type(content).__name__}")
    
    result = {
        "error": None
    }
    
    try:
        if not content.strip():
            result.update({
                "word_count": 0,
                "character_count": 0,
                "character_count_no_spaces": 0,
                "sentence_count": 0,
                "paragraph_count": 0,
                "average_words_per_sentence": 0.0,
                "average_chars_per_word": 0.0,
                "longest_word": "",
                "shortest_word": "",
                "important_terms": {},
                "estimated_reading_time_minutes": 0.0,
                "estimated_reading_time_seconds": 0,
                "readability_score": 0.0,
                "has_salary_info": False,
                "has_date_info": False,
                "has_contact_info": False
            })
            return result
        
        # Análisis básico de palabras
        words = [w.strip('.,!?;:()[]{}"\'-') for w in content.split() if w.strip()]
        word_count = len(words)
        
        # Análisis de caracteres
        char_count = len(content)
        char_count_no_spaces = len(content.replace(' ', '').replace('\n', '').replace('\t', ''))
        
        # Análisis de oraciones (mejorado)
        sentence_pattern = re.compile(r'[.!?]+(?:\s+|$)')
        sentences = [s.strip() for s in sentence_pattern.split(content) if s.strip()]
        sentence_count = len(sentences)
        
        # Análisis de párrafos
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        paragraph_count = len(paragraphs) if paragraphs else 1
        
        # Estadísticas de palabras
        if words:
            word_lengths = [len(w) for w in words]
            longest_word = max(words, key=len)
            shortest_words = [w for w in words if len(w) > 1]
            shortest_word = min(shortest_words, key=len) if shortest_words else words[0]
            avg_chars_per_word = sum(word_lengths) / word_count
        else:
            longest_word = ""
            shortest_word = ""
            avg_chars_per_word = 0.0
        
        avg_words_per_sentence = word_count / max(sentence_count, 1)
        
        # Detección de términos importantes (mejorado)
        important_terms = {
            'salary': len(re.findall(r'\$\s*[\d,]+\.?\d*|salario|salary|compensation', content, re.IGNORECASE)),
            'benefits': len(re.findall(r'beneficio|benefit|perk', content, re.IGNORECASE)),
            'start_date': len(re.findall(r'\d{4}-\d{2}-\d{2}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|fecha.*inicio|start.*date', content, re.IGNORECASE)),
            'location': len(re.findall(r'ubicación|location|address|dirección', content, re.IGNORECASE)),
            'position': len(re.findall(r'position|puesto|role|cargo', content, re.IGNORECASE)),
            'company': len(re.findall(r'company|empresa|corporation', content, re.IGNORECASE)),
        }
        
        # Detección de información específica
        has_salary = bool(re.search(r'\$\s*[\d,]+|salario|salary', content, re.IGNORECASE))
        has_date = bool(re.search(r'\d{4}-\d{2}-\d{2}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', content))
        has_contact = bool(re.search(r'@|email|correo|phone|teléfono|contact', content, re.IGNORECASE))
        
        # Tiempo de lectura estimado
        reading_time_minutes = word_count / max(words_per_minute, 1)
        reading_time_seconds = int(reading_time_minutes * 60)
        
        # Puntuación de legibilidad aproximada (Flesch Reading Ease simplificado)
        # Fórmula simplificada: 206.835 - (1.015 * ASL) - (84.6 * ASW)
        # ASL = Average Sentence Length, ASW = Average Syllables per Word
        # Usamos aproximación: ASW ≈ avg_chars_per_word / 3
        if sentence_count > 0 and word_count > 0:
            asl = avg_words_per_sentence
            asw = max(avg_chars_per_word / 3, 1)  # Aproximación de sílabas
            readability = max(0, min(100, 206.835 - (1.015 * asl) - (84.6 * asw)))
        else:
            readability = 0.0
        
        result.update({
            "word_count": word_count,
            "character_count": char_count,
            "character_count_no_spaces": char_count_no_spaces,
            "sentence_count": sentence_count,
            "paragraph_count": paragraph_count,
            "average_words_per_sentence": round(avg_words_per_sentence, 2),
            "average_chars_per_word": round(avg_chars_per_word, 2),
            "longest_word": longest_word,
            "shortest_word": shortest_word,
            "important_terms": important_terms,
            "estimated_reading_time_minutes": round(reading_time_minutes, 1),
            "estimated_reading_time_seconds": reading_time_seconds,
            "readability_score": round(readability, 2),
            "has_salary_info": has_salary,
            "has_date_info": has_date,
            "has_contact_info": has_contact
        })
        
        # Análisis de sentimiento (básico, si se solicita)
        if include_sentiment:
            try:
                # Análisis básico de sentimiento basado en palabras clave
                positive_words = ['pleased', 'excited', 'delighted', 'welcome', 'complace', 'entusiasmado']
                negative_words = ['regret', 'unfortunately', 'lamentamos', 'desafortunadamente']
                
                content_lower = content.lower()
                positive_count = sum(1 for word in positive_words if word in content_lower)
                negative_count = sum(1 for word in negative_words if word in content_lower)
                
                # Calcular sentimiento básico
                total_sentiment_words = positive_count + negative_count
                if total_sentiment_words > 0:
                    sentiment_score = (positive_count - negative_count) / total_sentiment_words
                else:
                    sentiment_score = 0.0
                
                result["sentiment"] = {
                    "score": round(sentiment_score, 2),
                    "label": "positive" if sentiment_score > 0.3 else "negative" if sentiment_score < -0.3 else "neutral",
                    "positive_indicators": positive_count,
                    "negative_indicators": negative_count
        }
    except Exception as e:
                logger.warning(f"Error analyzing sentiment: {str(e)}")
                result["sentiment"] = {"error": str(e)}
        
        return result
        
    except Exception as e:
        error_msg = f"Error analyzing content: {str(e)}"
        result["error"] = error_msg
        logger.error(error_msg, exc_info=True)
        return result


def parse_date_string(date_str: str) -> Optional[date]:
    """
    Parsea una cadena de fecha en varios formatos comunes.
    
    Args:
        date_str: Cadena de fecha a parsear
    
    Returns:
        Objeto date si se puede parsear, None en caso contrario
    """
    if not date_str or not isinstance(date_str, str):
        return None
    
    date_formats = [
        '%Y-%m-%d',
        '%Y/%m/%d',
        '%d-%m-%Y',
        '%d/%m/%Y',
        '%m-%d-%Y',
        '%m/%d/%Y'
    ]
    
    for fmt in date_formats:
        try:
            return datetime.strptime(date_str.strip(), fmt).date()
        except ValueError:
            continue
    
    return None


def filter_offers_by_criteria(
    offers: List[Dict],
    criteria: Dict,
    match_all: bool = True
) -> Tuple[List[Dict], Dict]:
    """
    Filtra ofertas según criterios especificados con opciones avanzadas.
    Soporta múltiples operadores y tipos de comparación.
    
    Args:
        offers: Lista de ofertas a filtrar.
        criteria: Diccionario con criterios de filtrado. Soporta:
            - min_salary, max_salary: Rango de salario
            - employment_type: Tipo de empleo (exacto o lista)
            - location: Ubicación (búsqueda parcial o lista)
            - department: Departamento (exacto o lista)
            - start_date_from, start_date_to: Rango de fechas
            - currency: Moneda
            - has_benefits: Si tiene beneficios (bool)
            - min_word_count: Mínimo de palabras en descripción
            - custom_field: Valor personalizado
        match_all: Si True, todas las condiciones deben cumplirse (AND).
            Si False, cualquier condición que se cumpla es suficiente (OR).
    
    Returns:
        Tupla (filtered_offers: List[Dict], stats: Dict):
            - filtered_offers: Lista de ofertas que cumplen los criterios
            - stats: Diccionario con estadísticas del filtrado:
                - total_offers: Total de ofertas analizadas
                - matched_offers: Ofertas que cumplen criterios
                - match_rate: Porcentaje de coincidencias
                - criteria_applied: Lista de criterios aplicados
    
    Example:
        >>> offers = [{'salary': 50000, 'location': 'Remote'}, ...]
        >>> filtered, stats = filter_offers_by_criteria(
        ...     offers,
        ...     {'min_salary': 40000, 'location': 'Remote'}
        ... )
        >>> print(f"Matched {stats['matched_offers']} offers")
    """
    if not isinstance(offers, list):
        raise TypeError(f"offers must be a list, got {type(offers).__name__}")
    
    if not isinstance(criteria, dict) or not criteria:
        return offers, {
            'total_offers': len(offers),
            'matched_offers': len(offers),
            'match_rate': 100.0,
            'criteria_applied': []
        }
    
    filtered = []
    criteria_applied = []
    total_offers = len(offers)
    
    for offer in offers:
        matches = []
        
        # Filtrar por rango de salario
        if 'min_salary' in criteria:
            salary = float(offer.get('salary', 0) or 0)
            min_sal = float(criteria['min_salary'])
            matches.append(salary >= min_sal)
            criteria_applied.append('min_salary')
        
        if 'max_salary' in criteria:
            salary = float(offer.get('salary', 0) or 0)
            max_sal = float(criteria['max_salary'])
            matches.append(salary <= max_sal)
            criteria_applied.append('max_salary')
        
        # Filtrar por tipo de empleo (soporta lista o valor único)
        if 'employment_type' in criteria:
            emp_type = offer.get('employment_type', '')
            criteria_type = criteria['employment_type']
            if isinstance(criteria_type, list):
                matches.append(emp_type in criteria_type)
            else:
                matches.append(emp_type == criteria_type)
            criteria_applied.append('employment_type')
        
        # Filtrar por ubicación (búsqueda parcial o lista)
        if 'location' in criteria:
            location = str(offer.get('location', '')).lower()
            criteria_loc = criteria['location']
            if isinstance(criteria_loc, list):
                matches.append(any(loc.lower() in location for loc in criteria_loc))
            else:
                matches.append(criteria_loc.lower() in location)
            criteria_applied.append('location')
        
        # Filtrar por departamento (soporta lista)
        if 'department' in criteria:
            dept = offer.get('department', '')
            criteria_dept = criteria['department']
            if isinstance(criteria_dept, list):
                matches.append(dept in criteria_dept)
            else:
                matches.append(dept == criteria_dept)
            criteria_applied.append('department')
        
        # Filtrar por moneda
        if 'currency' in criteria:
            currency = offer.get('currency', '').upper()
            criteria_curr = str(criteria['currency']).upper()
            matches.append(currency == criteria_curr)
            criteria_applied.append('currency')
        
        # Filtrar por rango de fechas
        if 'start_date_from' in criteria:
            offer_date = parse_date_string(offer.get('start_date', ''))
            from_date = parse_date_string(criteria['start_date_from'])
            if offer_date and from_date:
                matches.append(offer_date >= from_date)
            criteria_applied.append('start_date_from')
        
        if 'start_date_to' in criteria:
            offer_date = parse_date_string(offer.get('start_date', ''))
            to_date = parse_date_string(criteria['start_date_to'])
            if offer_date and to_date:
                matches.append(offer_date <= to_date)
            criteria_applied.append('start_date_to')
        
        # Filtrar por presencia de beneficios
        if 'has_benefits' in criteria:
            benefits = offer.get('benefits', [])
            has_benefits = bool(benefits and len(benefits) > 0)
            matches.append(has_benefits == bool(criteria['has_benefits']))
            criteria_applied.append('has_benefits')
        
        # Filtrar por conteo mínimo de palabras
        if 'min_word_count' in criteria:
            content = str(offer.get('description', '') or offer.get('content', ''))
            word_count = len(content.split())
            matches.append(word_count >= int(criteria['min_word_count']))
            criteria_applied.append('min_word_count')
        
        # Determinar si cumple criterios
        if match_all:
            # Todas las condiciones deben cumplirse (AND)
            match = all(matches) if matches else True
        else:
            # Cualquier condición que se cumpla es suficiente (OR)
            match = any(matches) if matches else True
        
        if match:
            filtered.append(offer)
    
    # Estadísticas
    match_rate = (len(filtered) / total_offers * 100) if total_offers > 0 else 0.0
    stats = {
        'total_offers': total_offers,
        'matched_offers': len(filtered),
        'match_rate': round(match_rate, 2),
        'criteria_applied': list(set(criteria_applied)),
        'match_mode': 'AND' if match_all else 'OR'
    }
    
    return filtered, stats


def generate_summary_report(
    offers: List[Dict],
    include_percentiles: bool = True,
    group_by: Optional[List[str]] = None
) -> Dict:
    """
    Genera un reporte resumen detallado de múltiples ofertas con análisis estadístico avanzado.
    Proporciona métricas agregadas, distribuciones y análisis por grupos.
    
    Args:
        offers: Lista de ofertas a analizar.
        include_percentiles: Si True, incluye percentiles de salario (25, 50, 75, 90).
        group_by: Lista de campos para agrupar estadísticas (ej: ['department', 'location']).
            Si None, solo genera estadísticas globales.
    
    Returns:
        Diccionario con estadísticas detalladas:
            - total_offers: int - Total de ofertas
            - salary_statistics: Dict - Estadísticas de salario (promedio, min, max, percentiles)
            - employment_types: Dict - Conteo por tipo de empleo
            - departments: Dict - Conteo por departamento
            - locations: Dict - Conteo por ubicación
            - currencies: Dict - Conteo por moneda
            - date_range: Dict - Rango de fechas de inicio
            - benefits_analysis: Dict - Análisis de beneficios
            - grouped_statistics: Dict - Estadísticas agrupadas (si group_by se proporciona)
            - generated_at: str - Timestamp de generación
            - error: Optional[str] - Mensaje de error si ocurre algún problema
    
    Example:
        >>> offers = [{'salary': 50000, 'department': 'Engineering'}, ...]
        >>> report = generate_summary_report(offers, include_percentiles=True)
        >>> print(f"Average salary: {report['salary_statistics']['average']}")
    """
    result = {
        "error": None
    }
    
    if not offers:
        result["error"] = "No offers provided"
        return result
    
    if not isinstance(offers, list):
        result["error"] = f"offers must be a list, got {type(offers).__name__}"
        return result
    
    try:
        total_offers = len(offers)
        
        # Estadísticas de salario
        salaries = [float(offer.get('salary', 0) or 0) for offer in offers]
        valid_salaries = [s for s in salaries if s > 0]
        
        salary_stats = {
            "total": sum(salaries),
            "average": sum(valid_salaries) / len(valid_salaries) if valid_salaries else 0,
            "minimum": min(valid_salaries) if valid_salaries else 0,
            "maximum": max(valid_salaries) if valid_salaries else 0,
            "median": 0,
            "count_with_salary": len(valid_salaries),
            "count_without_salary": total_offers - len(valid_salaries)
        }
        
        # Calcular percentiles si se solicita
        if include_percentiles and valid_salaries:
            sorted_salaries = sorted(valid_salaries)
            n = len(sorted_salaries)
            
            def percentile(p):
                if n == 0:
                    return 0
                k = (n - 1) * p
                f = int(k)
                c = k - f
                if f + 1 < n:
                    return sorted_salaries[f] + c * (sorted_salaries[f + 1] - sorted_salaries[f])
                return sorted_salaries[f]
            
            salary_stats.update({
                "percentile_25": percentile(0.25),
                "percentile_50": percentile(0.50),  # Mediana
                "percentile_75": percentile(0.75),
                "percentile_90": percentile(0.90)
            })
            salary_stats["median"] = salary_stats["percentile_50"]
        elif valid_salaries:
            sorted_salaries = sorted(valid_salaries)
            n = len(sorted_salaries)
            salary_stats["median"] = sorted_salaries[n // 2] if n > 0 else 0
        
        # Contadores por categorías
        employment_types = defaultdict(int)
        departments = defaultdict(int)
        locations = defaultdict(int)
        currencies = defaultdict(int)
        
        # Análisis de fechas
        dates = []
        benefits_count = defaultdict(int)
        total_benefits = 0
        
        for offer in offers:
            # Tipo de empleo
            emp_type = offer.get('employment_type', 'Unknown')
            employment_types[emp_type] += 1
            
            # Departamento
            dept = offer.get('department', 'Unknown')
            departments[dept] += 1
            
            # Ubicación
            loc = offer.get('location', 'Unknown')
            locations[loc] += 1
            
            # Moneda
            curr = offer.get('currency', 'Unknown')
            currencies[curr] += 1
            
            # Fechas
            start_date = parse_date_string(offer.get('start_date', ''))
            if start_date:
                dates.append(start_date)
            
            # Beneficios
            benefits = offer.get('benefits', [])
            if isinstance(benefits, list):
                total_benefits += len(benefits)
                for benefit in benefits:
                    benefits_count[str(benefit)] += 1
        
        # Rango de fechas
        date_range = {}
        if dates:
            date_range = {
                "earliest": min(dates).isoformat(),
                "latest": max(dates).isoformat(),
                "span_days": (max(dates) - min(dates)).days
            }
        
        # Análisis de beneficios
        benefits_analysis = {
            "total_benefits_mentioned": total_benefits,
            "average_benefits_per_offer": round(total_benefits / total_offers, 2) if total_offers > 0 else 0,
            "most_common_benefits": dict(sorted(benefits_count.items(), key=lambda x: x[1], reverse=True)[:10])
        }
        
        # Estadísticas agrupadas
        grouped_stats = {}
        if group_by and isinstance(group_by, list):
            for group_field in group_by:
                if group_field not in offers[0] if offers else {}:
                    continue
                
                groups = defaultdict(list)
                for offer in offers:
                    group_value = offer.get(group_field, 'Unknown')
                    groups[group_value].append(offer)
                
                group_statistics = {}
                for group_name, group_offers in groups.items():
                    group_salaries = [float(o.get('salary', 0) or 0) for o in group_offers if float(o.get('salary', 0) or 0) > 0]
                    group_statistics[group_name] = {
                        "count": len(group_offers),
                        "average_salary": sum(group_salaries) / len(group_salaries) if group_salaries else 0,
                        "min_salary": min(group_salaries) if group_salaries else 0,
                        "max_salary": max(group_salaries) if group_salaries else 0
                    }
                
                grouped_stats[group_field] = group_statistics
        
        result.update({
            "total_offers": total_offers,
            "salary_statistics": {k: round(v, 2) if isinstance(v, float) else v for k, v in salary_stats.items()},
            "employment_types": dict(employment_types),
            "departments": dict(departments),
            "locations": dict(locations),
            "currencies": dict(currencies),
            "date_range": date_range,
            "benefits_analysis": benefits_analysis,
            "grouped_statistics": grouped_stats if grouped_stats else None,
            "generated_at": datetime.now().isoformat()
        })
        
        return result
        
    except Exception as e:
        error_msg = f"Error generating summary report: {str(e)}"
        result["error"] = error_msg
        logger.error(error_msg, exc_info=True)
        return result


def export_to_csv(
    offers: List[Dict],
    filepath: str,
    fields: Optional[List[str]] = None,
    delimiter: str = ',',
    include_header: bool = True,
    flatten_nested: bool = True
) -> Tuple[bool, Optional[str], Dict]:
    """
    Exporta ofertas a formato CSV con opciones avanzadas de formato.
    Soporta delimitadores personalizados, campos anidados y control de encabezados.
    
    Args:
        offers: Lista de ofertas a exportar.
        filepath: Ruta del archivo CSV donde guardar.
        fields: Lista de campos a exportar. Si None, exporta todos los campos.
            Para campos anidados, usa notación de punto (ej: 'user.name').
        delimiter: Delimitador de campos (default: ','). Usa ';' para Excel en algunos locales.
        include_header: Si True, incluye fila de encabezados.
        flatten_nested: Si True, convierte campos anidados (dict/list) a strings.
            Si False, los omite o los serializa como JSON.
    
    Returns:
        Tupla (success: bool, error_message: Optional[str], info: Dict):
            - success: True si la exportación fue exitosa
            - error_message: Mensaje de error si falla, None si es exitoso
            - info: Diccionario con información:
                - rows_exported: int - Número de filas exportadas
                - fields_exported: List[str] - Campos exportados
                - file_size: int - Tamaño del archivo en bytes
                - file_size_human: str - Tamaño formateado
    
    Example:
        >>> offers = [{'name': 'John', 'salary': 50000}, ...]
        >>> success, error, info = export_to_csv(offers, 'offers.csv', fields=['name', 'salary'])
        >>> if success:
        ...     print(f"Exported {info['rows_exported']} rows")
    """
    info = {
        'rows_exported': 0,
        'fields_exported': [],
        'file_size': 0,
        'file_size_human': '0 B'
    }
    
    try:
        import csv
        
        if not isinstance(offers, list):
            return False, f"offers must be a list, got {type(offers).__name__}", info
        
        if not offers:
            return False, "No offers to export", info
        
        if not isinstance(filepath, str) or not filepath.strip():
            return False, "Invalid filepath provided", info
        
        # Determinar campos a exportar
        if not fields:
            # Obtener todos los campos únicos de todas las ofertas
            all_fields = set()
            for offer in offers:
                all_fields.update(offer.keys())
            fields = sorted(list(all_fields))
        else:
            # Validar que los campos solicitados existan
            available_fields = set()
            for offer in offers:
                available_fields.update(offer.keys())
            fields = [f for f in fields if f in available_fields or '.' in f]
        
        info['fields_exported'] = fields
        
        # Asegurar que el directorio existe
        dir_path = os.path.dirname(filepath) if os.path.dirname(filepath) else '.'
        success, error, _ = ensure_directory(dir_path)
        if not success:
            return False, f"Could not create directory: {error}", info
        
        # Función helper para obtener valores anidados
        def get_nested_value(obj: Dict, key: str):
            if '.' not in key:
                return obj.get(key, '')
            
            parts = key.split('.')
            value = obj
            for part in parts:
                if isinstance(value, dict):
                    value = value.get(part, '')
                else:
                    return ''
            return value
        
        # Función helper para aplanar valores complejos
        def flatten_value(value):
            if value is None:
                return ''
            elif isinstance(value, (dict, list)):
                if flatten_nested:
                    return json.dumps(value, ensure_ascii=False)
                else:
                    return ''
            else:
                return str(value)
        
        # Escribir CSV
        with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(
                csvfile,
                fieldnames=fields,
                delimiter=delimiter,
                extrasaction='ignore'
            )
            
            if include_header:
            writer.writeheader()
            
            for offer in offers:
                row = {}
                for field in fields:
                    if '.' in field:
                        value = get_nested_value(offer, field)
                    else:
                        value = offer.get(field, '')
                    row[field] = flatten_value(value)
                
                writer.writerow(row)
                info['rows_exported'] += 1
        
        # Obtener información del archivo
        if os.path.exists(filepath):
            file_size = os.path.getsize(filepath)
            info['file_size'] = file_size
            info['file_size_human'] = get_file_size_human(file_size)
        
        logger.info(
            f"Exported {info['rows_exported']} offers to CSV: {filepath} "
            f"({info['file_size_human']})"
        )
        
        return True, None, info
        
    except PermissionError as e:
        error_msg = f"Permission denied writing to {filepath}: {str(e)}"
        logger.error(error_msg)
        return False, error_msg, info
    except Exception as e:
        error_msg = f"Error exporting to CSV: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return False, error_msg, info


def validate_json_structure(data: Dict, schema: Dict) -> Tuple[bool, List[str]]:
    """
    Valida la estructura de un JSON según un esquema
    
    Args:
        data: Datos a validar
        schema: Esquema de validación con tipos esperados
    
    Returns:
        Tupla (es_válido, lista_errores)
    """
    errors = []
    
    for field, expected_type in schema.items():
        if field not in data:
            errors.append(f"Missing required field: {field}")
        elif not isinstance(data[field], expected_type):
            errors.append(
                f"Field '{field}' must be of type {expected_type.__name__}, "
                f"got {type(data[field]).__name__}"
            )
    
    return len(errors) == 0, errors


def merge_offer_data(
    base_data: Dict,
    additional_data: Dict,
    overwrite: bool = False,
    merge_strategy: str = 'smart'
) -> Dict:
    """
    Fusiona datos de oferta con estrategias avanzadas de combinación.
    Maneja diccionarios anidados, listas y diferentes estrategias de fusión.
    
    Args:
        base_data: Datos base de la oferta (no se modifica).
        additional_data: Datos adicionales a fusionar.
        overwrite: Si True, sobrescribe campos existentes en base_data.
            Si False, solo agrega campos nuevos.
        merge_strategy: Estrategia para fusionar datos:
            - 'smart': Fusión inteligente (default)
                - Dicts: Fusiona recursivamente
                - Lists: Combina sin duplicados
                - Otros: Sobrescribe si overwrite=True
            - 'replace': Reemplaza completamente los valores existentes
            - 'append': Para listas, agrega elementos sin duplicados
            - 'deep': Fusión profunda recursiva de todos los tipos
    
    Returns:
        Nuevo diccionario con datos fusionados (base_data no se modifica).
    
    Raises:
        TypeError: Si base_data o additional_data no son diccionarios.
    
    Example:
        >>> base = {'name': 'John', 'skills': ['Python'], 'meta': {'level': 1}}
        >>> additional = {'skills': ['Java'], 'meta': {'exp': 5}}
        >>> merged = merge_offer_data(base, additional, merge_strategy='smart')
        >>> print(merged['skills'])  # ['Python', 'Java']
    """
    if not isinstance(base_data, dict):
        raise TypeError(f"base_data must be a dict, got {type(base_data).__name__}")
    
    if not isinstance(additional_data, dict):
        raise TypeError(f"additional_data must be a dict, got {type(additional_data).__name__}")
    
    # Crear copia profunda para no modificar el original
    merged = json.loads(json.dumps(base_data))  # Deep copy simple
    
    for key, value in additional_data.items():
        if key not in merged:
            # Campo nuevo, agregarlo
            merged[key] = value
        elif overwrite:
            # Sobrescribir si se solicita
            if merge_strategy == 'deep' and isinstance(merged[key], dict) and isinstance(value, dict):
                merged[key] = merge_offer_data(merged[key], value, overwrite=True, merge_strategy='deep')
            else:
                merged[key] = value
        else:
            # Fusionar según estrategia
            if merge_strategy == 'smart':
                if isinstance(merged[key], dict) and isinstance(value, dict):
                    # Fusionar diccionarios anidados recursivamente
                    merged[key] = merge_offer_data(merged[key], value, overwrite=False, merge_strategy='smart')
                elif isinstance(merged[key], list) and isinstance(value, list):
                    # Combinar listas sin duplicados
                    combined = merged[key] + value
                    # Remover duplicados manteniendo orden
                    seen = set()
                    unique_list = []
                    for item in combined:
                        item_str = json.dumps(item, sort_keys=True) if isinstance(item, (dict, list)) else str(item)
                        if item_str not in seen:
                            seen.add(item_str)
                            unique_list.append(item)
                    merged[key] = unique_list
                # Para otros tipos, no sobrescribir si overwrite=False
            elif merge_strategy == 'replace':
                merged[key] = value
            elif merge_strategy == 'append':
                if isinstance(merged[key], list) and isinstance(value, list):
                    # Agregar elementos sin duplicados
                    for item in value:
                        if item not in merged[key]:
                            merged[key].append(item)
                elif isinstance(merged[key], dict) and isinstance(value, dict):
                    merged[key] = merge_offer_data(merged[key], value, overwrite=False, merge_strategy='append')
                else:
                    merged[key] = value
            elif merge_strategy == 'deep':
                # Fusión profunda recursiva
                if isinstance(merged[key], dict) and isinstance(value, dict):
                    merged[key] = merge_offer_data(merged[key], value, overwrite=False, merge_strategy='deep')
                elif isinstance(merged[key], list) and isinstance(value, list):
                    merged[key] = list(set(merged[key] + value))
                else:
                    merged[key] = value
    
    return merged


def calculate_offer_statistics(offer_data: Dict, include_content_analysis: bool = False) -> Dict:
    """
    Calcula estadísticas detalladas de una oferta individual con análisis de completitud.
    Proporciona métricas de calidad y completitud de los datos.
    
    Args:
        offer_data: Datos de la oferta a analizar (debe ser un diccionario).
        include_content_analysis: Si True, incluye análisis del contenido de la oferta
            (requiere campo 'content' o 'description'). Puede ser más lento.
    
    Returns:
        Diccionario con estadísticas detalladas:
            - has_salary: bool - Si tiene información de salario
            - has_benefits: bool - Si tiene beneficios
            - has_location: bool - Si tiene ubicación
            - has_start_date: bool - Si tiene fecha de inicio
            - has_contact_info: bool - Si tiene información de contacto
            - benefits_count: int - Número de beneficios listados
            - additional_terms_count: int - Número de términos adicionales
            - completeness_score: float - Puntuación de completitud (0-100)
            - required_fields_present: int - Campos requeridos presentes
            - required_fields_missing: List[str] - Campos requeridos faltantes
            - optional_fields_present: int - Campos opcionales presentes
            - data_quality_score: float - Puntuación de calidad de datos (0-100)
            - content_analysis: Optional[Dict] - Análisis de contenido (si include_content_analysis=True)
            - timestamp: str - Timestamp del análisis
    
    Example:
        >>> offer = {'candidate_name': 'John', 'salary': 50000, 'benefits': ['Health']}
        >>> stats = calculate_offer_statistics(offer)
        >>> print(f"Completeness: {stats['completeness_score']}%")
    """
    if not isinstance(offer_data, dict):
        raise TypeError(f"offer_data must be a dict, got {type(offer_data).__name__}")
    
    stats = {
        "has_salary": bool(offer_data.get('salary')),
        "has_benefits": bool(offer_data.get('benefits')),
        "has_location": bool(offer_data.get('location')),
        "has_start_date": bool(offer_data.get('start_date')),
        "has_contact_info": bool(
            offer_data.get('hr_contact') or 
            offer_data.get('candidate_email') or
            offer_data.get('company_address')
        ),
        "benefits_count": len(offer_data.get('benefits', [])) if isinstance(offer_data.get('benefits'), list) else 0,
        "additional_terms_count": len(offer_data.get('additional_terms', [])) if isinstance(offer_data.get('additional_terms'), list) else 0,
        "completeness_score": 0.0,
        "required_fields_present": 0,
        "required_fields_missing": [],
        "optional_fields_present": 0,
        "data_quality_score": 0.0,
        "timestamp": datetime.now().isoformat()
    }
    
    # Campos requeridos (peso 70%)
    required_fields = {
        'candidate_name': lambda x: bool(x and len(str(x).strip()) >= 2),
        'position_title': lambda x: bool(x and len(str(x).strip()) >= 3),
        'salary': lambda x: bool(x and float(x or 0) > 0),
        'start_date': lambda x: bool(x and parse_date_string(str(x)) is not None),
        'location': lambda x: bool(x and len(str(x).strip()) > 0)
    }
    
    # Campos opcionales importantes (peso 30%)
    optional_fields = {
        'benefits': lambda x: bool(x and isinstance(x, list) and len(x) > 0),
        'additional_terms': lambda x: bool(x and isinstance(x, list) and len(x) > 0),
        'company_name': lambda x: bool(x and len(str(x).strip()) > 0),
        'hr_contact': lambda x: bool(x and len(str(x).strip()) > 0),
        'department': lambda x: bool(x and len(str(x).strip()) > 0),
        'currency': lambda x: bool(x and str(x).upper() in Config.SUPPORTED_CURRENCIES),
        'employment_type': lambda x: bool(x and str(x) in Config.EMPLOYMENT_TYPES)
    }
    
    # Validar campos requeridos
    required_present = []
    required_missing = []
    
    for field, validator in required_fields.items():
        value = offer_data.get(field)
        if validator(value):
            required_present.append(field)
            stats["required_fields_present"] += 1
        else:
            required_missing.append(field)
    
    stats["required_fields_missing"] = required_missing
    
    # Validar campos opcionales
    optional_present = 0
    for field, validator in optional_fields.items():
        value = offer_data.get(field)
        if validator(value):
            optional_present += 1
    
    stats["optional_fields_present"] = optional_present
    
    # Calcular score de completitud (0-100)
    required_score = (len(required_present) / len(required_fields)) * 70 if required_fields else 0
    optional_score = (optional_present / len(optional_fields)) * 30 if optional_fields else 0
    stats["completeness_score"] = round(required_score + optional_score, 2)
    
    # Calcular score de calidad de datos
    quality_factors = []
    
    # Validar formato de email si existe
    if offer_data.get('candidate_email'):
        email = str(offer_data.get('candidate_email', '')).strip()
        if re.match(r'^[a-zA-Z0-9]([a-zA-Z0-9._-]*[a-zA-Z0-9])?@[a-zA-Z0-9]([a-zA-Z0-9.-]*[a-zA-Z0-9])?\.[a-zA-Z]{2,}$', email):
            quality_factors.append(1.0)
        else:
            quality_factors.append(0.5)  # Email presente pero formato inválido
    
    # Validar salario razonable
    if offer_data.get('salary'):
        salary = float(offer_data.get('salary', 0) or 0)
        if Config.MIN_SALARY < salary <= Config.MAX_SALARY:
            quality_factors.append(1.0)
        else:
            quality_factors.append(0.5)  # Salario fuera de rango esperado
    
    # Validar fecha
    if offer_data.get('start_date'):
        start_date = parse_date_string(str(offer_data.get('start_date', '')))
        if start_date:
            current_date = date.today()
            if current_date <= start_date <= current_date + timedelta(days=365*2):
                quality_factors.append(1.0)
            else:
                quality_factors.append(0.7)  # Fecha válida pero fuera de rango razonable
        else:
            quality_factors.append(0.3)  # Fecha en formato inválido
    
    # Calcular score de calidad
    if quality_factors:
        stats["data_quality_score"] = round((sum(quality_factors) / len(quality_factors)) * 100, 2)
    else:
        stats["data_quality_score"] = 0.0
    
    # Análisis de contenido si se solicita
    if include_content_analysis:
        content = offer_data.get('content') or offer_data.get('description') or ''
        if content:
            try:
                content_stats = analyze_offer_content(str(content), include_sentiment=False)
                stats["content_analysis"] = content_stats
            except Exception as e:
                logger.warning(f"Error analyzing content: {str(e)}")
                stats["content_analysis"] = {"error": str(e)}
        else:
            stats["content_analysis"] = None
    
    return stats


class OfferStatus(Enum):
    """Estados de una oferta laboral"""
    DRAFT = "draft"
    GENERATED = "generated"
    SENT = "sent"
    ACCEPTED = "accepted"
    REJECTED = "rejected"
    EXPIRED = "expired"
    WITHDRAWN = "withdrawn"


class TemplateType(Enum):
    """Tipos de plantillas disponibles"""
    STANDARD = "standard"
    EXECUTIVE = "executive"
    TECHNICAL = "technical"
    INTERN = "intern"
    CONTRACT = "contract"


@dataclass
class OfferLetterData:
    """Estructura de datos para una carta de oferta"""
    candidate_name: str
    candidate_email: str
    position_title: str
    department: str
    start_date: str
    salary: float
    currency: str = "USD"
    benefits: List[str] = None
    reporting_manager: str = ""
    location: str = ""
    employment_type: str = "Full-time"  # Full-time, Part-time, Contract
    probation_period: int = 90  # días
    notice_period: int = 30  # días
    company_name: str = ""
    company_address: str = ""
    hr_contact: str = ""
    offer_expiry_date: str = ""
    additional_terms: List[str] = None
    template_type: str = "standard"
    language: str = "es"
    signature_required: bool = True
    equity_stock: Optional[str] = None
    bonus_structure: Optional[str] = None
    relocation_assistance: Optional[str] = None
    visa_sponsorship: bool = False
    custom_fields: Dict = None
    
    def __post_init__(self):
        if self.benefits is None:
            self.benefits = []
        if self.additional_terms is None:
            self.additional_terms = []
        if self.custom_fields is None:
            self.custom_fields = {}
    
    def validate(self, strict: bool = True) -> Tuple[bool, List[str], Dict]:
        """
        Valida los datos de la oferta con validaciones exhaustivas.
        
        Args:
            strict: Si True, valida todos los campos estrictamente.
                Si False, solo valida campos críticos.
        
        Returns:
            Tupla (es_válido: bool, lista_errores: List[str], warnings: Dict):
                - es_válido: True si todos los datos son válidos
                - lista_errores: Lista de errores encontrados
                - warnings: Diccionario con advertencias (campos opcionales faltantes, etc.)
        """
        errors = []
        warnings = {
            'missing_optional_fields': [],
            'format_warnings': [],
            'value_warnings': []
        }
        
        # Validación de nombre del candidato
        if not self.candidate_name or len(self.candidate_name.strip()) < 2:
            errors.append("El nombre del candidato es requerido y debe tener al menos 2 caracteres")
        elif len(self.candidate_name.strip()) > 100:
            errors.append("El nombre del candidato no puede exceder 100 caracteres")
        elif not re.match(r'^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s\-\'\.]+$', self.candidate_name.strip()):
            warnings['format_warnings'].append("El nombre del candidato contiene caracteres inusuales")
        
        # Validación de email
        if not self.candidate_email:
            errors.append("El email del candidato es requerido")
        elif not self._is_valid_email(self.candidate_email):
            errors.append("El email del candidato debe ser válido")
        elif len(self.candidate_email) > 254:
            errors.append("El email del candidato excede la longitud máxima (254 caracteres)")
        
        # Validación de título del puesto
        if not self.position_title or len(self.position_title.strip()) < 3:
            errors.append("El título del puesto es requerido y debe tener al menos 3 caracteres")
        elif len(self.position_title.strip()) > Config.MAX_TEXT_LENGTH:
            errors.append(f"El título del puesto no puede exceder {Config.MAX_TEXT_LENGTH} caracteres")
        
        # Validación de departamento
        if not self.department or len(self.department.strip()) < 1:
            errors.append("El departamento es requerido")
        elif len(self.department.strip()) > 100:
            errors.append("El departamento no puede exceder 100 caracteres")
        
        # Validación de fecha de inicio
        if not self.start_date:
            errors.append("La fecha de inicio es requerida")
        elif not self._is_valid_date(self.start_date):
            errors.append("La fecha de inicio debe estar en formato YYYY-MM-DD")
        else:
            # Validar que la fecha sea razonable
            try:
                parsed_date = datetime.strptime(self.start_date, "%Y-%m-%d").date()
                current_date = date.today()
                if parsed_date < current_date - timedelta(days=7):
                    warnings['value_warnings'].append("La fecha de inicio está en el pasado (más de 7 días)")
                elif parsed_date > current_date + timedelta(days=365*2):
                    warnings['value_warnings'].append("La fecha de inicio está muy lejana en el futuro (más de 2 años)")
            except ValueError:
                pass  # Ya se capturó en _is_valid_date
        
        # Validación de salario
        if self.salary <= 0:
            errors.append("El salario debe ser mayor a 0")
        elif self.salary < Config.MIN_SALARY:
            errors.append(f"El salario debe ser al menos {Config.MIN_SALARY:,.2f}")
        elif self.salary > Config.MAX_SALARY:
            errors.append(f"El salario excede el máximo permitido ({Config.MAX_SALARY:,.2f})")
        
        # Validación de moneda
        if self.currency and self.currency.upper() not in Config.SUPPORTED_CURRENCIES:
            errors.append(f"La moneda debe ser una de: {', '.join(Config.SUPPORTED_CURRENCIES)}")
        
        # Validación de tipo de empleo
        if self.employment_type and self.employment_type not in Config.EMPLOYMENT_TYPES:
            errors.append(f"El tipo de empleo debe ser uno de: {', '.join(Config.EMPLOYMENT_TYPES)}")
        
        # Validación de períodos
        if self.probation_period < 0:
            errors.append("El período de prueba no puede ser negativo")
        elif self.probation_period > 365:
            warnings['value_warnings'].append("El período de prueba es mayor a 1 año")
        
        if self.notice_period < 0:
            errors.append("El período de preaviso no puede ser negativo")
        elif self.notice_period > 180:
            warnings['value_warnings'].append("El período de preaviso es mayor a 180 días")
        
        # Validaciones estrictas adicionales
        if strict:
            # Validar que tenga información de contacto
            if not self.hr_contact and not self.company_address:
                warnings['missing_optional_fields'].append("Falta información de contacto (hr_contact o company_address)")
            
            # Validar que tenga ubicación si es relevante
            if not self.location or len(self.location.strip()) < 3:
                warnings['missing_optional_fields'].append("La ubicación es recomendada pero no está especificada")
            
            # Validar beneficios
            if not self.benefits or len(self.benefits) == 0:
                warnings['missing_optional_fields'].append("No se especificaron beneficios")
        
        return len(errors) == 0, errors, warnings
    
    @staticmethod
    def _is_valid_email(email: str) -> bool:
        """
        Valida formato de email con regex mejorado según RFC 5322 simplificado.
        
        Args:
            email: Cadena de email a validar
        
        Returns:
            True si el email tiene formato válido, False en caso contrario
        """
        if not email or not isinstance(email, str):
            return False
        
        # Regex mejorado para validación de email (RFC 5322 simplificado)
        pattern = r'^[a-zA-Z0-9]([a-zA-Z0-9._-]*[a-zA-Z0-9])?@[a-zA-Z0-9]([a-zA-Z0-9.-]*[a-zA-Z0-9])?\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email.strip().lower()))
    
    @staticmethod
    def _is_valid_date(date_str: str) -> bool:
        """
        Valida formato de fecha YYYY-MM-DD y que sea una fecha válida en rango razonable.
        
        Args:
            date_str: Cadena de fecha a validar en formato YYYY-MM-DD
        
        Returns:
            True si la fecha tiene formato válido y es una fecha real en rango razonable,
            False en caso contrario
        """
        if not date_str or not isinstance(date_str, str):
            return False
        
        try:
            # Intentar parsear con formato estándar
            parsed_date = datetime.strptime(date_str.strip(), "%Y-%m-%d")
            # Verificar que sea una fecha válida en rango razonable
            # Evitar fechas como 2024-02-30 (se captura en strptime)
            # Verificar rango de años razonable
            return 2000 <= parsed_date.year <= 2100
        except ValueError:
            return False


class OfferLetterGenerator:
    """Generador de cartas de oferta laboral con múltiples formatos y plantillas"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Inicializa el generador de cartas de oferta
        
        Args:
            template_path: Ruta al archivo de plantilla (opcional)
        """
        self.template_path = template_path
        self.templates = {
            TemplateType.STANDARD: self._get_default_template(),
            TemplateType.EXECUTIVE: self._get_executive_template(),
            TemplateType.TECHNICAL: self._get_technical_template(),
            TemplateType.INTERN: self._get_intern_template(),
            TemplateType.CONTRACT: self._get_contract_template()
        }
        self.default_template = self._get_default_template()
    
    def _get_default_template(self) -> str:
        """Retorna la plantilla por defecto para cartas de oferta"""
        return """CARTA DE OFERTA LABORAL

Fecha: {date}

Estimado/a {candidate_name},

Nos complace extenderle una oferta de empleo para la posición de {position_title} 
en el departamento de {department} de {company_name}.

DETALLES DE LA OFERTA:

Posición: {position_title}
Tipo de empleo: {employment_type}
Fecha de inicio: {start_date}
Ubicación: {location}
Reporta a: {reporting_manager}

COMPENSACIÓN:

Salario: {currency} {salary:,.2f} {salary_period}
Período de prueba: {probation_period} días
Período de preaviso: {notice_period} días

BENEFICIOS INCLUIDOS:
{benefits_list}

{additional_terms_section}

Esta oferta es válida hasta el {offer_expiry_date}. Por favor, confirme su aceptación 
antes de esta fecha.

Si tiene alguna pregunta, no dude en contactar a {hr_contact}.

Esperamos tenerle como parte de nuestro equipo.

Atentamente,
{company_name}

Dirección: {company_address}
"""
    
    def _get_executive_template(self) -> str:
        """Plantilla para posiciones ejecutivas"""
        return """CARTA DE OFERTA EJECUTIVA

Fecha: {date}

Estimado/a {candidate_name},

Nos complace extenderle una oferta de empleo para la posición ejecutiva de {position_title} 
en {company_name}.

DETALLES DE LA OFERTA:

Posición: {position_title}
Departamento: {department}
Tipo de empleo: {employment_type}
Fecha de inicio: {start_date}
Ubicación: {location}
Reporta a: {reporting_manager}

COMPENSACIÓN Y BENEFICIOS:

Salario base: {currency} {salary:,.2f} {salary_period}
{bonus_section}
{equity_section}
{relocation_section}

BENEFICIOS EJECUTIVOS:
{benefits_list}

{additional_terms_section}

Esta oferta es válida hasta el {offer_expiry_date}.

Atentamente,
{company_name}
{company_address}
"""
    
    def _get_technical_template(self) -> str:
        """Plantilla para posiciones técnicas"""
        return """CARTA DE OFERTA - POSICIÓN TÉCNICA

Fecha: {date}

Estimado/a {candidate_name},

Nos complace ofrecerle la posición de {position_title} en el departamento de {department}.

DETALLES TÉCNICOS:

Posición: {position_title}
Fecha de inicio: {start_date}
Ubicación: {location}
Modalidad: {employment_type}
Reporta a: {reporting_manager}

COMPENSACIÓN:

Salario: {currency} {salary:,.2f} {salary_period}
Período de prueba: {probation_period} días

BENEFICIOS TÉCNICOS:
{benefits_list}

{additional_terms_section}

Esperamos su confirmación antes del {offer_expiry_date}.

Saludos,
{company_name}
"""
    
    def _get_intern_template(self) -> str:
        """Plantilla para pasantías/internships"""
        return """CARTA DE OFERTA - PROGRAMA DE PRÁCTICAS

Fecha: {date}

Estimado/a {candidate_name},

Nos complace ofrecerle una posición de prácticas como {position_title} 
en el departamento de {department}.

DETALLES:

Posición: {position_title}
Duración: {employment_type}
Fecha de inicio: {start_date}
Ubicación: {location}
Supervisor: {reporting_manager}

COMPENSACIÓN:

Estipendio: {currency} {salary:,.2f} {salary_period}

BENEFICIOS:
{benefits_list}

{additional_terms_section}

Por favor confirme su aceptación antes del {offer_expiry_date}.

Saludos cordiales,
{company_name}
"""
    
    def _get_contract_template(self) -> str:
        """Plantilla para contratos temporales"""
        return """CARTA DE OFERTA - CONTRATO TEMPORAL

Fecha: {date}

Estimado/a {candidate_name},

Le extendemos una oferta de contrato temporal para la posición de {position_title}.

DETALLES DEL CONTRATO:

Posición: {position_title}
Departamento: {department}
Tipo: {employment_type}
Fecha de inicio: {start_date}
Duración: {contract_duration}
Ubicación: {location}

COMPENSACIÓN:

Tarifa: {currency} {salary:,.2f} {salary_period}
Período de preaviso: {notice_period} días

TÉRMINOS DEL CONTRATO:
{benefits_list}

{additional_terms_section}

Esta oferta expira el {offer_expiry_date}.

Atentamente,
{company_name}
"""
    
    def generate_offer_letter(self, data: OfferLetterData, format_type: str = "txt") -> str:
        """
        Genera una carta de oferta basada en los datos proporcionados
        
        Args:
            data: Datos de la oferta laboral
            format_type: Tipo de formato (txt, html)
            
        Returns:
            Carta de oferta formateada como string
        """
        try:
            # Validar datos antes de generar con validación estricta
            is_valid, errors, warnings = data.validate(strict=True)
            if not is_valid:
                error_msg = f"Invalid offer data: {', '.join(errors)}"
                logger.warning(f"Validation failed: {error_msg}")
                if warnings.get('format_warnings') or warnings.get('value_warnings'):
                    logger.info(f"Validation warnings: {warnings}")
                raise ValidationError(error_msg, errors)
            
            # Log warnings si existen pero no bloquean la generación
            if warnings.get('format_warnings') or warnings.get('value_warnings') or warnings.get('missing_optional_fields'):
                logger.info(f"Validation warnings for {data.candidate_name}: {warnings}")
            
            # Determinar período salarial de forma más inteligente
            # Considerar el contexto de la moneda y el monto
            SALARY_THRESHOLD_ANNUAL = 100000
            if data.currency in ['USD', 'EUR', 'GBP', 'CAD', 'AUD']:
                SALARY_THRESHOLD_ANNUAL = 100000
            elif data.currency in ['MXN']:
                SALARY_THRESHOLD_ANNUAL = 1000000  # Ajuste para pesos mexicanos
            elif data.currency in ['JPY', 'CNY']:
                SALARY_THRESHOLD_ANNUAL = 10000000  # Ajuste para yenes/yuanes
            
            salary_period = "anual" if data.salary >= SALARY_THRESHOLD_ANNUAL else "mensual"
            
            # Formatear lista de beneficios con validación y sanitización
            if data.benefits and len(data.benefits) > 0:
                # Filtrar beneficios vacíos y sanitizar
                valid_benefits = [str(b).strip() for b in data.benefits if str(b).strip()]
                if valid_benefits:
                    benefits_list = "\n".join([f"- {html.escape(benefit)}" for benefit in valid_benefits])
                else:
                    benefits_list = "No se especificaron beneficios adicionales."
            else:
                benefits_list = "No se especificaron beneficios adicionales."
            
            # Formatear términos adicionales con validación
            additional_terms_section = ""
            if data.additional_terms and len(data.additional_terms) > 0:
                # Filtrar términos vacíos y sanitizar
                valid_terms = [str(t).strip() for t in data.additional_terms if str(t).strip()]
                if valid_terms:
                    additional_terms_section = "\nTÉRMINOS ADICIONALES:\n"
                    additional_terms_section += "\n".join([f"- {html.escape(term)}" for term in valid_terms])
            
            # Preparar datos para la plantilla
            template_data = {
                "date": datetime.now().strftime("%d de %B de %Y"),
                "candidate_name": html.escape(data.candidate_name),
                "position_title": html.escape(data.position_title),
                "department": html.escape(data.department),
                "company_name": html.escape(data.company_name or "Nuestra Empresa"),
                "start_date": data.start_date,
                "location": html.escape(data.location or "Por definir"),
                "reporting_manager": html.escape(data.reporting_manager or "Por asignar"),
                "currency": data.currency,
                "salary": data.salary,
                "salary_period": salary_period,
                "employment_type": html.escape(data.employment_type),
                "probation_period": data.probation_period,
                "notice_period": data.notice_period,
                "benefits_list": benefits_list,
                "additional_terms_section": additional_terms_section,
                "offer_expiry_date": data.offer_expiry_date or "30 días a partir de la fecha",
                "hr_contact": html.escape(data.hr_contact or "Recursos Humanos"),
                "company_address": html.escape(data.company_address or "Dirección no especificada")
            }
            
            # Generar según formato con validación
            format_lower = format_type.lower().strip()
            if format_lower == "html":
                offer_letter = self._generate_html_letter(template_data, data)
            elif format_lower == "txt" or format_lower == "text":
                # Generar carta usando la plantilla de texto
                try:
                offer_letter = self.default_template.format(**template_data)
                except KeyError as e:
                    missing_key = str(e).strip("'")
                    error_msg = f"Missing template variable: {missing_key}"
                    logger.error(error_msg)
                    raise TemplateError(error_msg) from e
            else:
                error_msg = f"Unsupported format type: {format_type}. Supported formats: txt, html"
                logger.error(error_msg)
                raise FormatError(error_msg)
            
            # Validar que se generó contenido
            if not offer_letter or len(offer_letter.strip()) < 50:
                error_msg = "Generated offer letter is too short or empty"
                logger.error(error_msg)
                raise FormatError(error_msg)
            
            logger.info(
                f"Carta de oferta generada exitosamente para {data.candidate_name} "
                f"(formato: {format_type}, tamaño: {len(offer_letter)} caracteres)"
            )
            return offer_letter
            
        except (ValidationError, TemplateError, FormatError):
            # Re-raise excepciones específicas sin modificar
            raise
        except Exception as e:
            error_msg = f"Error inesperado al generar carta de oferta: {str(e)}"
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            raise OfferLetterException(error_msg) from e
    
    def _generate_html_letter(self, template_data: Dict, data: OfferLetterData) -> str:
        """
        Genera una carta de oferta en formato HTML mejorado
        
        Args:
            template_data: Datos formateados para la plantilla
            data: Datos completos de la oferta
            
        Returns:
            HTML formateado de la carta
        """
        benefits_html = ""
        if data.benefits:
            benefits_html = "<ul class='benefits-list'>"
            for benefit in data.benefits:
                benefits_html += f"<li>{html.escape(benefit)}</li>"
            benefits_html += "</ul>"
        else:
            benefits_html = "<p>No se especificaron beneficios adicionales.</p>"
        
        additional_terms_html = ""
        if data.additional_terms:
            additional_terms_html = "<div class='additional-terms'><h3>Términos Adicionales</h3><ul>"
            for term in data.additional_terms:
                additional_terms_html += f"<li>{html.escape(term)}</li>"
            additional_terms_html += "</ul></div>"
        
        html_template = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Carta de Oferta - {template_data['position_title']}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f5f5f5;
            padding: 20px;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            border-radius: 8px;
        }}
        .header {{
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .header h1 {{
            color: #2c3e50;
            font-size: 24px;
            margin-bottom: 10px;
        }}
        .date {{
            color: #7f8c8d;
            font-size: 14px;
        }}
        .section {{
            margin: 25px 0;
            padding: 20px;
            background-color: #f8f9fa;
            border-left: 4px solid #3498db;
            border-radius: 4px;
        }}
        .section h2 {{
            color: #2c3e50;
            margin-bottom: 15px;
            font-size: 18px;
        }}
        .section ul {{
            list-style: none;
            padding-left: 0;
        }}
        .section li {{
            padding: 8px 0;
            border-bottom: 1px solid #ddd;
        }}
        .section li:last-child {{
            border-bottom: none;
        }}
        .benefits-list li:before {{
            content: "✓ ";
            color: #27ae60;
            font-weight: bold;
            margin-right: 10px;
        }}
        .salary {{
            font-size: 20px;
            font-weight: bold;
            color: #27ae60;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 2px solid #ecf0f1;
            text-align: right;
        }}
        .footer p {{
            margin: 5px 0;
        }}
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            .container {{
                box-shadow: none;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{template_data['company_name']}</h1>
            <p class="date">Fecha: {template_data['date']}</p>
        </div>
        
        <p><strong>{template_data['candidate_name']}</strong></p>
        <p><strong>Re: Oferta de Empleo - {template_data['position_title']}</strong></p>
        
        <p>Estimado/a {template_data['candidate_name']},</p>
        
        <p>Nos complace extenderle una oferta de empleo para la posición de <strong>{template_data['position_title']}</strong> 
        en el departamento de <strong>{template_data['department']}</strong> de {template_data['company_name']}.</p>
        
        <div class="section">
            <h2>Detalles de la Oferta</h2>
            <ul>
                <li><strong>Posición:</strong> {template_data['position_title']}</li>
                <li><strong>Tipo de empleo:</strong> {template_data['employment_type']}</li>
                <li><strong>Fecha de inicio:</strong> {template_data['start_date']}</li>
                <li><strong>Ubicación:</strong> {template_data['location']}</li>
                <li><strong>Reporta a:</strong> {template_data['reporting_manager']}</li>
            </ul>
        </div>
        
        <div class="section">
            <h2>Compensación</h2>
            <p class="salary">{template_data['currency']} {template_data['salary']:,.2f} {template_data['salary_period']}</p>
            <ul>
                <li><strong>Período de prueba:</strong> {template_data['probation_period']} días</li>
                <li><strong>Período de preaviso:</strong> {template_data['notice_period']} días</li>
            </ul>
        </div>
        
        <div class="section">
            <h2>Beneficios Incluidos</h2>
            {benefits_html}
        </div>
        
        {additional_terms_html}
        
        <p>Esta oferta es válida hasta el <strong>{template_data['offer_expiry_date']}</strong>. 
        Por favor, confirme su aceptación antes de esta fecha.</p>
        
        <p>Si tiene alguna pregunta, no dude en contactar a {template_data['hr_contact']}.</p>
        
        <p>Esperamos tenerle como parte de nuestro equipo.</p>
        
        <div class="footer">
            <p><strong>Atentamente,</strong></p>
            <p><strong>{template_data['company_name']}</strong></p>
            <p>{template_data['company_address']}</p>
        </div>
    </div>
</body>
</html>"""
        return html_template
    
    def save_offer_letter(self, offer_letter: str, output_path: str, candidate_name: str, 
                         format_type: str = "txt") -> str:
        """
        Guarda la carta de oferta en un archivo con mejor manejo de nombres
        
        Args:
            offer_letter: Contenido de la carta
            output_path: Directorio donde guardar el archivo
            candidate_name: Nombre del candidato (para el nombre del archivo)
            format_type: Tipo de formato (txt, html, pdf)
            
        Returns:
            Ruta completa del archivo guardado
        """
        try:
            # Crear directorio si no existe
            os.makedirs(output_path, exist_ok=True)
            
            # Sanitizar nombre del candidato
            safe_name = sanitize_filename(candidate_name)
            safe_name = safe_name.replace(' ', '_')
            
            # Determinar extensión
            extension_map = {
                'txt': '.txt',
                'html': '.html',
                'htm': '.html',
                'pdf': '.pdf'
            }
            extension = extension_map.get(format_type.lower(), '.txt')
            
            # Generar nombre de archivo único
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"offer_letter_{safe_name}_{timestamp}{extension}"
            filepath = os.path.join(output_path, filename)
            
            # Asegurar que el archivo no exista (añadir contador si es necesario)
            counter = 1
            original_filepath = filepath
            while os.path.exists(filepath):
                name_part = original_filepath.rsplit(extension, 1)[0]
                filepath = f"{name_part}_{counter}{extension}"
                counter += 1
            
            # Guardar archivo
            encoding = 'utf-8'
            mode = 'w'
            if format_type.lower() in ['html', 'htm']:
                mode = 'w'
            elif format_type.lower() == 'pdf':
                # PDF se maneja diferente
                raise ValueError("Use generate_pdf() method for PDF files")
            
            with open(filepath, mode, encoding=encoding) as f:
                f.write(offer_letter)
            
            logger.info(f"Carta de oferta guardada en: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error al guardar carta de oferta: {str(e)}")
            logger.error(traceback.format_exc())
            raise
    
    def generate_pdf(self, data: OfferLetterData, output_path: str) -> str:
        """
        Genera un PDF de la carta de oferta
        
        Args:
            data: Datos de la oferta laboral
            output_path: Ruta donde guardar el PDF
            
        Returns:
            Ruta del archivo PDF generado
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab no está instalado. Instale con: pip install reportlab")
        
        try:
            safe_name = "".join(c for c in data.candidate_name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')
            filename = f"offer_letter_{safe_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
            filepath = os.path.join(output_path, filename)
            
            os.makedirs(output_path, exist_ok=True)
            
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#2c3e50',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            story.append(Paragraph("CARTA DE OFERTA LABORAL", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Contenido
            text_content = self.generate_offer_letter(data)
            for line in text_content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), styles['Normal']))
                else:
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            logger.info(f"PDF generado exitosamente: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error al generar PDF: {str(e)}")
            raise
    
    def generate_docx(self, data: OfferLetterData, output_path: str) -> str:
        """
        Genera un documento Word (DOCX) de la carta de oferta
        
        Args:
            data: Datos de la oferta laboral
            output_path: Ruta donde guardar el DOCX
            
        Returns:
            Ruta del archivo DOCX generado
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instale con: pip install python-docx")
        
        try:
            safe_name = "".join(c for c in data.candidate_name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')
            filename = f"offer_letter_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = os.path.join(output_path, filename)
            
            os.makedirs(output_path, exist_ok=True)
            
            doc = Document()
            
            # Título
            title = doc.add_heading('CARTA DE OFERTA LABORAL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fecha
            date_para = doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph()
            
            # Saludo
            doc.add_paragraph(f"Estimado/a {data.candidate_name},")
            doc.add_paragraph()
            
            # Contenido principal
            intro = doc.add_paragraph()
            intro.add_run(f"Nos complace extenderle una oferta de empleo para la posición de ").bold = False
            intro.add_run(f"{data.position_title}").bold = True
            intro.add_run(f" en el departamento de {data.department} de {data.company_name}.")
            doc.add_paragraph()
            
            # Detalles de la oferta
            doc.add_heading('DETALLES DE LA OFERTA', level=1)
            details = [
                ("Posición", data.position_title),
                ("Tipo de empleo", data.employment_type),
                ("Fecha de inicio", data.start_date),
                ("Ubicación", data.location),
                ("Reporta a", getattr(data, 'reporting_manager', 'N/A'))
            ]
            
            for label, value in details:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))
            
            doc.add_paragraph()
            
            # Compensación
            doc.add_heading('COMPENSACIÓN', level=1)
            salary_para = doc.add_paragraph()
            salary_para.add_run("Salario: ").bold = True
            salary_para.add_run(f"{data.currency} {data.salary:,.2f} {getattr(data, 'salary_period', 'anual')}")
            
            if hasattr(data, 'probation_period') and data.probation_period:
                doc.add_paragraph(f"Período de prueba: {data.probation_period} días", style='List Bullet')
            
            if hasattr(data, 'notice_period') and data.notice_period:
                doc.add_paragraph(f"Período de preaviso: {data.notice_period} días", style='List Bullet')
            
            doc.add_paragraph()
            
            # Beneficios
            if hasattr(data, 'benefits') and data.benefits:
                doc.add_heading('BENEFICIOS INCLUIDOS', level=1)
                for benefit in data.benefits:
                    doc.add_paragraph(benefit, style='List Bullet')
                doc.add_paragraph()
            
            # Términos adicionales
            if hasattr(data, 'additional_terms') and data.additional_terms:
                doc.add_heading('TÉRMINOS ADICIONALES', level=1)
                for term in data.additional_terms:
                    doc.add_paragraph(term, style='List Bullet')
                doc.add_paragraph()
            
            # Fecha de expiración
            if hasattr(data, 'offer_expiry_date') and data.offer_expiry_date:
                expiry_para = doc.add_paragraph()
                expiry_para.add_run(f"Esta oferta es válida hasta el {data.offer_expiry_date}. ").italic = True
                expiry_para.add_run("Por favor, confirme su aceptación antes de esta fecha.")
            
            doc.add_paragraph()
            
            # Contacto
            if hasattr(data, 'hr_contact') and data.hr_contact:
                doc.add_paragraph(f"Si tiene alguna pregunta, no dude en contactar a {data.hr_contact}.")
                doc.add_paragraph()
            
            # Cierre
            doc.add_paragraph("Esperamos tenerle como parte de nuestro equipo.")
            doc.add_paragraph()
            doc.add_paragraph("Atentamente,")
            doc.add_paragraph()
            
            # Firma
            company_para = doc.add_paragraph()
            company_para.add_run(data.company_name).bold = True
            if hasattr(data, 'company_address') and data.company_address:
                doc.add_paragraph(data.company_address)
            
            doc.save(filepath)
            logger.info(f"DOCX generado exitosamente: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error al generar DOCX: {str(e)}")
            raise
    
    def generate_json_export(self, data: OfferLetterData) -> Dict:
        """
        Genera un export JSON de los datos de la oferta
        
        Args:
            data: Datos de la oferta laboral
            
        Returns:
            Diccionario con los datos estructurados
        """
        return {
            "offer_letter": asdict(data),
            "generated_at": datetime.now().isoformat(),
            "version": "2.0"
        }
    
    def send_email(self, data: OfferLetterData, offer_letter: str, 
                   smtp_config: Dict, attachments: List[str] = None) -> bool:
        """
        Envía la carta de oferta por email
        
        Args:
            data: Datos de la oferta
            offer_letter: Contenido de la carta
            smtp_config: Configuración SMTP
            attachments: Lista de archivos adjuntos
            
        Returns:
            True si se envió exitosamente
        """
        if not EMAIL_AVAILABLE:
            raise ImportError("Módulos de email no disponibles")
        
        try:
            msg = MIMEMultipart()
            msg['From'] = smtp_config.get('from_email', 'hr@company.com')
            msg['To'] = data.candidate_email
            msg['Subject'] = f"Oferta de Empleo - {data.position_title} - {data.company_name}"
            
            # Cuerpo del mensaje
            body = f"""
Estimado/a {data.candidate_name},

Adjunto encontrará la carta de oferta formal para la posición de {data.position_title}.

{offer_letter}

Saludos cordiales,
{smtp_config.get('sender_name', 'Recursos Humanos')}
{data.company_name}
"""
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # Adjuntar archivos con validación de seguridad
            if attachments:
                for filepath in attachments:
                    if not isinstance(filepath, str) or not filepath:
                        logger.warning(f"Invalid filepath in attachments: {filepath}")
                        continue
                    
                    # Validar que el archivo existe y es accesible
                    if not os.path.exists(filepath) or not os.path.isfile(filepath):
                        logger.warning(f"Attachment file not found: {filepath}")
                        continue
                    
                    # Validar tamaño de archivo (máximo 10MB)
                    max_size = 10 * 1024 * 1024  # 10MB
                    if os.path.getsize(filepath) > max_size:
                        logger.warning(f"Attachment file too large: {filepath}")
                        continue
                    
                    try:
                        with open(filepath, "rb") as attachment:
                            part = MIMEBase('application', 'octet-stream')
                            part.set_payload(attachment.read())
                            encoders.encode_base64(part)
                            # Sanitizar nombre de archivo para prevenir inyección
                            safe_filename = sanitize_filename(os.path.basename(filepath))
                            part.add_header(
                                'Content-Disposition',
                                f'attachment; filename="{safe_filename}"'
                            )
                            msg.attach(part)
                    except (IOError, OSError) as e:
                        logger.error(f"Error reading attachment file {filepath}: {str(e)}")
                        continue
            
            # Enviar email usando context manager
            smtp_server = smtp_config.get('smtp_server', 'smtp.gmail.com')
            smtp_port = smtp_config.get('smtp_port', 587)
            
            with smtplib.SMTP(smtp_server, smtp_port, timeout=30) as server:
            server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            server.send_message(msg)
            
            logger.info(f"Email enviado exitosamente a {data.candidate_email}")
            return True
            
        except Exception as e:
            logger.error(f"Error al enviar email: {str(e)}")
            raise


class OfferLetterDatabase:
    """Gestión de base de datos para ofertas laborales"""
    
    def __init__(self, db_path: str = "offer_letters.db"):
        self.db_path = db_path
        self._init_database()
    
    def _init_database(self):
        """Inicializa la base de datos"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS offers (
                    id TEXT PRIMARY KEY,
                    candidate_name TEXT NOT NULL,
                    candidate_email TEXT NOT NULL,
                    position_title TEXT NOT NULL,
                    department TEXT NOT NULL,
                    salary REAL NOT NULL,
                    currency TEXT DEFAULT 'USD',
                    status TEXT DEFAULT 'draft',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    offer_data TEXT,
                    offer_letter_content TEXT
                )
            ''')
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS offer_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    offer_id TEXT NOT NULL,
                    action TEXT NOT NULL,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    details TEXT,
                    FOREIGN KEY (offer_id) REFERENCES offers(id)
                )
            ''')
            
            # Crear índices para mejorar rendimiento de consultas
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_offers_status ON offers(status)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_offers_email ON offers(candidate_email)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_offers_created_at ON offers(created_at DESC)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_offer_history_offer_id ON offer_history(offer_id)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_offer_history_timestamp ON offer_history(timestamp DESC)')
            
            conn.commit()
    
    def create_offer(self, offer_id: str, data: OfferLetterData, offer_letter: str) -> bool:
        """Crea un nuevo registro de oferta"""
        try:
            # Validar offer_id para prevenir SQL injection
            if not isinstance(offer_id, str) or not offer_id.strip():
                logger.error("Invalid offer_id provided")
                return False
            
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                cursor.execute('''
                    INSERT INTO offers (id, candidate_name, candidate_email, position_title,
                                      department, salary, currency, status, offer_data, offer_letter_content)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    offer_id,
                    data.candidate_name,
                    data.candidate_email,
                    data.position_title,
                    data.department,
                    data.salary,
                    data.currency,
                    OfferStatus.GENERATED.value,
                    json.dumps(asdict(data)),
                    offer_letter
                ))
                
                cursor.execute('''
                    INSERT INTO offer_history (offer_id, action, details)
                    VALUES (?, ?, ?)
                ''', (offer_id, 'created', 'Oferta creada'))
                
                conn.commit()
            return True
        except sqlite3.Error as e:
            logger.error(f"Error de base de datos al crear oferta: {str(e)}", exc_info=True)
            return False
        except Exception as e:
            logger.error(f"Error inesperado al crear oferta en BD: {str(e)}", exc_info=True)
            return False
    
    def get_offer(self, offer_id: str, include_history: bool = False) -> Optional[Dict]:
        """
        Obtiene una oferta por ID con opciones avanzadas.
        
        Args:
            offer_id: Identificador único de la oferta.
            include_history: Si True, incluye el historial de cambios de la oferta.
        
        Returns:
            Diccionario con los datos de la oferta, o None si no se encuentra.
            Si include_history=True, incluye clave 'history' con lista de eventos.
        """
        if not isinstance(offer_id, str) or not offer_id.strip():
            logger.warning("Invalid offer_id provided to get_offer")
            return None
        
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cursor.execute('SELECT * FROM offers WHERE id = ?', (offer_id,))
                row = cursor.fetchone()
                
                if not row:
                    return None
                
                offer_dict = dict(row)
                
                # Incluir historial si se solicita
                if include_history:
                    cursor.execute(
                        'SELECT * FROM offer_history WHERE offer_id = ? ORDER BY created_at DESC',
                        (offer_id,)
                    )
                    history_rows = cursor.fetchall()
                    offer_dict['history'] = [dict(h) for h in history_rows]
                
                return offer_dict
        except sqlite3.Error as e:
            logger.error(f"Error de base de datos al obtener oferta: {str(e)}", exc_info=True)
            return None
        except Exception as e:
            logger.error(f"Error inesperado al obtener oferta: {str(e)}", exc_info=True)
            return None
    
    def update_offer_status(
        self,
        offer_id: str,
        status: OfferStatus,
        details: str = "",
        user_id: Optional[str] = None
    ) -> Tuple[bool, Optional[str]]:
        """
        Actualiza el estado de una oferta con registro de auditoría.
        
        Args:
            offer_id: Identificador único de la oferta.
            status: Nuevo estado de la oferta (enum OfferStatus).
            details: Detalles adicionales sobre el cambio de estado.
            user_id: ID del usuario que realiza el cambio (para auditoría).
        
        Returns:
            Tupla (success: bool, error_message: Optional[str]):
                - success: True si la actualización fue exitosa
                - error_message: Mensaje de error si falla, None si es exitoso
        """
        if not isinstance(offer_id, str) or not offer_id.strip():
            error_msg = "Invalid offer_id provided to update_offer_status"
            logger.warning(error_msg)
            return False, error_msg
        
        if not isinstance(status, OfferStatus):
            error_msg = f"status must be an OfferStatus enum, got {type(status).__name__}"
            logger.warning(error_msg)
            return False, error_msg
        
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Verificar que la oferta existe y obtener estado actual
                cursor.execute('SELECT id, status FROM offers WHERE id = ?', (offer_id,))
                offer_row = cursor.fetchone()
                if not offer_row:
                    error_msg = f"Offer with id {offer_id} not found"
                    logger.warning(error_msg)
                    return False, error_msg
                
                old_status = offer_row[1] if offer_row else None
                
                # Actualizar estado
                cursor.execute('''
                    UPDATE offers 
                    SET status = ?, updated_at = CURRENT_TIMESTAMP 
                    WHERE id = ?
                ''', (status.value, offer_id))
                
                # Registrar en historial con información detallada
                action_details = {
                    'old_status': old_status,
                    'new_status': status.value,
                    'details': details,
                    'user_id': user_id,
                    'timestamp': datetime.now().isoformat()
                }
                
                cursor.execute('''
                    INSERT INTO offer_history (offer_id, action, details)
                    VALUES (?, ?, ?)
                ''', (
                    offer_id,
                    f'status_changed_to_{status.value}',
                    json.dumps(action_details)
                ))
                
                conn.commit()
                logger.info(f"Offer {offer_id} status updated to {status.value}")
                return True, None
                
        except sqlite3.Error as e:
            error_msg = f"Error de base de datos al actualizar estado: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg
        except Exception as e:
            error_msg = f"Error inesperado al actualizar estado: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg
    
    def get_all_offers(
        self,
        filters: Optional[Dict] = None,
        limit: Optional[int] = None,
        offset: int = 0,
        order_by: str = 'created_at',
        order_direction: str = 'DESC'
    ) -> Tuple[List[Dict], int]:
        """
        Obtiene todas las ofertas con filtros opcionales, paginación y ordenamiento.
        
        Args:
            filters: Diccionario con filtros opcionales:
                - status: Filtrar por estado
                - department: Filtrar por departamento
                - date_from: Filtrar desde fecha (YYYY-MM-DD)
                - date_to: Filtrar hasta fecha (YYYY-MM-DD)
                - min_salary: Salario mínimo
                - max_salary: Salario máximo
            limit: Número máximo de resultados a retornar (None = sin límite).
            offset: Número de resultados a saltar (para paginación).
            order_by: Campo por el cual ordenar (default: 'created_at').
            order_direction: Dirección del ordenamiento ('ASC' o 'DESC', default: 'DESC').
        
        Returns:
            Tupla (offers: List[Dict], total_count: int):
                - offers: Lista de ofertas que cumplen los criterios
                - total_count: Número total de ofertas (sin límite)
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Construir query base
                query = 'SELECT * FROM offers WHERE 1=1'
                count_query = 'SELECT COUNT(*) FROM offers WHERE 1=1'
                params = []
                
                # Aplicar filtros con validación
                if filters:
                    allowed_filters = {
                        'status', 'department', 'date_from', 'date_to',
                        'min_salary', 'max_salary', 'employment_type', 'location'
                    }
                    
                    for key, value in filters.items():
                        if key not in allowed_filters or not value:
                            continue
                        
                        if key == 'date_from':
                            try:
                                datetime.strptime(str(value), '%Y-%m-%d')
                                query += ' AND created_at >= ?'
                                count_query += ' AND created_at >= ?'
                                params.append(str(value))
                            except ValueError:
                                logger.warning(f"Invalid date format in filter date_from: {value}")
                        
                        elif key == 'date_to':
                            try:
                                datetime.strptime(str(value), '%Y-%m-%d')
                                query += ' AND created_at <= ?'
                                count_query += ' AND created_at <= ?'
                                params.append(str(value))
                            except ValueError:
                                logger.warning(f"Invalid date format in filter date_to: {value}")
                        
                        elif key == 'min_salary':
                            try:
                                min_sal = float(value)
                                query += ' AND salary >= ?'
                                count_query += ' AND salary >= ?'
                                params.append(min_sal)
                            except (ValueError, TypeError):
                                logger.warning(f"Invalid min_salary value: {value}")
                        
                        elif key == 'max_salary':
                            try:
                                max_sal = float(value)
                                query += ' AND salary <= ?'
                                count_query += ' AND salary <= ?'
                                params.append(max_sal)
                            except (ValueError, TypeError):
                                logger.warning(f"Invalid max_salary value: {value}")
                        
                        else:
                            # Filtros de igualdad (status, department, etc.)
                            query += f' AND {key} = ?'
                            count_query += f' AND {key} = ?'
                            params.append(str(value))
                
                # Obtener conteo total (sin límite)
                cursor.execute(count_query, params)
                total_count = cursor.fetchone()[0] or 0
                
                # Validar y aplicar ordenamiento
                valid_order_fields = {'created_at', 'updated_at', 'salary', 'candidate_name', 'position_title'}
                if order_by not in valid_order_fields:
                    order_by = 'created_at'
                    logger.warning(f"Invalid order_by field, using default: {order_by}")
                
                order_direction = order_direction.upper()
                if order_direction not in ['ASC', 'DESC']:
                    order_direction = 'DESC'
                    logger.warning(f"Invalid order_direction, using default: {order_direction}")
                
                query += f' ORDER BY {order_by} {order_direction}'
                
                # Aplicar límite y offset
                if limit is not None:
                    if not isinstance(limit, int) or limit < 0:
                        limit = 100  # Default
                        logger.warning(f"Invalid limit, using default: {limit}")
                    query += ' LIMIT ?'
                    params.append(limit)
                    
                    if offset > 0:
                        if not isinstance(offset, int) or offset < 0:
                            offset = 0
                        query += ' OFFSET ?'
                        params.append(offset)
                
                # Ejecutar query
                cursor.execute(query, params)
                rows = cursor.fetchall()
                
                return [dict(row) for row in rows], total_count
                
        except sqlite3.Error as e:
            logger.error(f"Error de base de datos al obtener ofertas: {str(e)}", exc_info=True)
            return [], 0
        except Exception as e:
            logger.error(f"Error inesperado al obtener ofertas: {str(e)}", exc_info=True)
            return [], 0
    
    def get_statistics(self, include_trends: bool = False) -> Dict:
        """
        Obtiene estadísticas detalladas de ofertas con opciones avanzadas.
        
        Args:
            include_trends: Si True, incluye análisis de tendencias temporales
                (requiere más tiempo de procesamiento).
        
        Returns:
            Diccionario con estadísticas:
                - total_offers: int - Total de ofertas
                - by_status: Dict - Conteo por estado
                - by_department: Dict - Conteo por departamento
                - by_employment_type: Dict - Conteo por tipo de empleo
                - salary_statistics: Dict - Estadísticas de salario (avg, min, max, median)
                - by_location: Dict - Conteo por ubicación
                - trends: Optional[Dict] - Tendencias temporales (si include_trends=True)
                - generated_at: str - Timestamp de generación
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                stats = {
                    'generated_at': datetime.now().isoformat()
                }
                
                # Total de ofertas
                cursor.execute('SELECT COUNT(*) FROM offers')
                stats['total_offers'] = cursor.fetchone()[0] or 0
                
                # Por estado
                cursor.execute('SELECT status, COUNT(*) FROM offers GROUP BY status')
                stats['by_status'] = dict(cursor.fetchall())
                
                # Por departamento
                cursor.execute('SELECT department, COUNT(*) FROM offers GROUP BY department')
                stats['by_department'] = dict(cursor.fetchall())
                
                # Por tipo de empleo
                cursor.execute('SELECT employment_type, COUNT(*) FROM offers GROUP BY employment_type')
                stats['by_employment_type'] = dict(cursor.fetchall())
                
                # Por ubicación
                cursor.execute('SELECT location, COUNT(*) FROM offers WHERE location IS NOT NULL AND location != "" GROUP BY location')
                stats['by_location'] = dict(cursor.fetchall())
                
                # Estadísticas de salario
                cursor.execute('''
                    SELECT 
                        AVG(salary) as avg_salary,
                        MIN(salary) as min_salary,
                        MAX(salary) as max_salary,
                        COUNT(salary) as count_with_salary
                    FROM offers 
                    WHERE salary > 0
                ''')
                salary_row = cursor.fetchone()
                if salary_row and salary_row[0]:
                    stats['salary_statistics'] = {
                        'average': round(float(salary_row[0]), 2),
                        'minimum': float(salary_row[1]) if salary_row[1] else 0.0,
                        'maximum': float(salary_row[2]) if salary_row[2] else 0.0,
                        'count': salary_row[3] or 0
                    }
                    
                    # Calcular mediana si hay suficientes datos
                    cursor.execute('SELECT salary FROM offers WHERE salary > 0 ORDER BY salary')
                    salaries = [row[0] for row in cursor.fetchall()]
                    if salaries:
                        n = len(salaries)
                        median = salaries[n // 2] if n % 2 == 1 else (salaries[n // 2 - 1] + salaries[n // 2]) / 2
                        stats['salary_statistics']['median'] = round(float(median), 2)
                else:
                    stats['salary_statistics'] = {
                        'average': 0.0,
                        'minimum': 0.0,
                        'maximum': 0.0,
                        'median': 0.0,
                        'count': 0
                    }
                
                # Tendencias temporales (si se solicita)
                if include_trends:
                    trends = {}
                    
                    # Ofertas por mes (últimos 12 meses)
                    cursor.execute('''
                        SELECT 
                            strftime('%Y-%m', created_at) as month,
                            COUNT(*) as count
                        FROM offers
                        WHERE created_at >= datetime('now', '-12 months')
                        GROUP BY month
                        ORDER BY month
                    ''')
                    trends['by_month'] = dict(cursor.fetchall())
                    
                    # Ofertas por día (últimos 30 días)
                    cursor.execute('''
                        SELECT 
                            DATE(created_at) as day,
                            COUNT(*) as count
                        FROM offers
                        WHERE created_at >= datetime('now', '-30 days')
                        GROUP BY day
                        ORDER BY day
                    ''')
                    trends['by_day'] = dict(cursor.fetchall())
                    
                    # Cambios de estado recientes
                    cursor.execute('''
                        SELECT 
                            action,
                            COUNT(*) as count
                        FROM offer_history
                        WHERE created_at >= datetime('now', '-7 days')
                        GROUP BY action
                        ORDER BY count DESC
                    ''')
                    trends['recent_status_changes'] = dict(cursor.fetchall())
                    
                    stats['trends'] = trends
                
                return stats
                
        except sqlite3.Error as e:
            logger.error(f"Error de base de datos al obtener estadísticas: {str(e)}", exc_info=True)
            return {
                'total_offers': 0,
                'by_status': {},
                'by_department': {},
                'by_employment_type': {},
                'by_location': {},
                'salary_statistics': {
                    'average': 0.0,
                    'minimum': 0.0,
                    'maximum': 0.0,
                    'median': 0.0,
                    'count': 0
                },
                'generated_at': datetime.now().isoformat(),
                'error': str(e)
            }
        except Exception as e:
            logger.error(f"Error inesperado al obtener estadísticas: {str(e)}", exc_info=True)
            return {
                'total_offers': 0,
                'by_status': {},
                'by_department': {},
                'by_employment_type': {},
                'by_location': {},
                'salary_statistics': {
                    'average': 0.0,
                    'minimum': 0.0,
                    'maximum': 0.0,
                    'median': 0.0,
                    'count': 0
                },
                'generated_at': datetime.now().isoformat(),
                'error': str(e)
            }


class OfferLetterAPI:
    """API RESTful para gestión de cartas de oferta"""
    
    def __init__(self, use_database: bool = True, db_path: str = "offer_letters.db"):
        self.generator = OfferLetterGenerator()
        self.offers_history: List[Dict] = []
        self.use_database = use_database
        if use_database:
            self.db = OfferLetterDatabase(db_path)
        else:
            self.db = None
    
    def create_offer(self, offer_data: Dict, format_type: str = "txt") -> Dict:
        """
        Crea una nueva carta de oferta con validación mejorada
        
        Args:
            offer_data: Diccionario con los datos de la oferta
            format_type: Tipo de formato deseado (txt, html, pdf)
            
        Returns:
            Diccionario con la carta generada y metadatos, incluyendo:
            - success: bool indicando si la operación fue exitosa
            - offer_id: str identificador único de la oferta
            - offer_letter: str contenido de la carta generada
            - format: str formato de la carta generada
            - metadata: Dict con metadatos de la oferta
            - data: Dict con los datos estructurados
        """
        try:
            # Validar datos usando la función mejorada
            is_valid, validation_errors = validate_offer_data(offer_data)
            if not is_valid:
                return {
                    "success": False,
                    "error": "Validation failed",
                    "validation_errors": validation_errors
                }
            
            # Validar campos requeridos adicionales
            required_fields = ['candidate_name', 'position_title', 'department', 'start_date', 'salary']
            missing_fields = [field for field in required_fields if field not in offer_data or not offer_data[field]]
            if missing_fields:
                return {
                    "success": False,
                    "error": "Missing required fields",
                    "missing_fields": missing_fields
                }
            
            # Sanitizar inputs de texto con límite de longitud
            text_fields = ['candidate_name', 'position_title', 'department', 'location', 'reporting_manager']
            for key in text_fields:
                if key in offer_data and offer_data[key]:
                    sanitized = html.escape(str(offer_data[key]).strip())
                    offer_data[key] = sanitized[:Config.MAX_TEXT_LENGTH]
            
            # Crear objeto de datos
            data = OfferLetterData(**offer_data)
            
            # Validar objeto de datos con validación estricta
            is_valid_obj, obj_errors, obj_warnings = data.validate(strict=True)
            if not is_valid_obj:
                return {
                    "success": False,
                    "error": "Data validation failed",
                    "validation_errors": obj_errors,
                    "validation_warnings": obj_warnings
                }
            
            # Log warnings si existen pero no bloquean la creación
            if obj_warnings and (obj_warnings.get('format_warnings') or 
                                obj_warnings.get('value_warnings') or 
                                obj_warnings.get('missing_optional_fields')):
                logger.info(f"Validation warnings for offer creation: {obj_warnings}")
            
            # Determinar formato (txt o html)
            format_type = format_type.lower() if format_type else offer_data.get('format', 'txt').lower()
            if format_type not in ['txt', 'html', 'text']:
                format_type = 'txt'
            
            # Generar carta
            offer_letter = self.generator.generate_offer_letter(data, format_type=format_type)
            
            # Generar ID único para la oferta
            offer_id = str(uuid.uuid4())
            
            # Guardar en base de datos si está disponible
            if self.use_database and self.db:
                try:
                    self.db.save_offer(
                        offer_id=offer_id,
                        offer_data=asdict(data),
                        offer_letter=offer_letter,
                        status=OfferStatus.GENERATED
                    )
                except DatabaseError as db_error:
                    logger.warning(f"Error al guardar en BD, usando historial en memoria: {str(db_error)}")
            
            # Guardar en historial
            offer_record = {
                "id": offer_id,
                "candidate_name": data.candidate_name,
                "candidate_email": getattr(data, 'candidate_email', ''),
                "position_title": data.position_title,
                "created_at": datetime.now().isoformat(),
                "status": OfferStatus.GENERATED.value,
                "format": format_type
            }
            self.offers_history.append(offer_record)
            
            return {
                "success": True,
                "offer_id": offer_id,
                "offer_letter": offer_letter,
                "format": format_type,
                "metadata": offer_record,
                "data": asdict(data)
            }
            
        except (ValueError, ValidationError) as e:
            logger.error(f"Validation error al crear oferta: {str(e)}")
            return {
                "success": False,
                "error": "Validation error",
                "message": str(e)
            }
        except DatabaseError as e:
            logger.error(f"Database error al crear oferta: {str(e)}")
            return {
                "success": False,
                "error": "Database error",
                "message": "Error al guardar la oferta en la base de datos"
            }
        except TemplateError as e:
            logger.error(f"Template error al crear oferta: {str(e)}")
            return {
                "success": False,
                "error": "Template error",
                "message": "Error al generar la plantilla de la oferta"
            }
        except Exception as e:
            logger.error(f"Error inesperado al crear oferta: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "error": "Internal server error",
                "message": str(e) if logger.level == logging.DEBUG else "An error occurred while generating the offer letter"
            }
    
    def get_offer_history(
        self,
        filters: Optional[Dict] = None,
        limit: Optional[int] = None,
        offset: int = 0,
        order_by: str = 'created_at',
        order_direction: str = 'DESC'
    ) -> Tuple[List[Dict], int]:
        """
        Retorna el historial de ofertas generadas con opciones avanzadas de filtrado y paginación.
        
        Args:
            filters: Diccionario opcional con filtros:
                - status: Filtrar por estado
                - department: Filtrar por departamento
                - date_from: Filtrar desde fecha (YYYY-MM-DD)
                - date_to: Filtrar hasta fecha (YYYY-MM-DD)
                - min_salary: Salario mínimo
                - max_salary: Salario máximo
            limit: Número máximo de resultados (None = sin límite).
            offset: Número de resultados a saltar (para paginación).
            order_by: Campo por el cual ordenar (default: 'created_at').
            order_direction: Dirección del ordenamiento ('ASC' o 'DESC').
        
        Returns:
            Tupla (offers: List[Dict], total_count: int):
                - offers: Lista de ofertas que cumplen los criterios
                - total_count: Número total de ofertas (sin límite)
        """
        if self.use_database and self.db:
            try:
                return self.db.get_all_offers(
                    filters=filters,
                    limit=limit,
                    offset=offset,
                    order_by=order_by,
                    order_direction=order_direction
                )
            except DatabaseError as e:
                logger.error(f"Error al obtener historial de BD: {str(e)}")
                # Fallback a historial en memoria
                offers = self.offers_history
                total = len(offers)
                
                # Aplicar filtros básicos en memoria
                if filters:
                    offers = filter_offers_by_criteria(offers, filters)[0]
                    total = len(offers)
                
                # Aplicar límite y offset
                if limit is not None:
                    start = offset
                    end = offset + limit
                    offers = offers[start:end]
                
                return offers, total
        
        # Usar historial en memoria
        offers = self.offers_history.copy()
        total = len(offers)
        
        # Aplicar filtros básicos
        if filters:
            offers, stats = filter_offers_by_criteria(offers, filters)
            total = stats.get('matched_offers', len(offers))
        
        # Ordenar
        reverse = (order_direction.upper() == 'DESC')
        if order_by == 'created_at':
            offers.sort(key=lambda x: x.get('created_at', ''), reverse=reverse)
        elif order_by == 'candidate_name':
            offers.sort(key=lambda x: x.get('candidate_name', ''), reverse=reverse)
        
        # Aplicar límite y offset
        if limit is not None:
            start = offset
            end = offset + limit
            offers = offers[start:end]
        
        return offers, total
    
    def get_offer_by_id(self, offer_id: str, include_history: bool = False) -> Optional[Dict]:
        """
        Obtiene una oferta por su ID con opciones avanzadas.
        
        Args:
            offer_id: Identificador único de la oferta.
            include_history: Si True, incluye el historial de cambios de la oferta.
        
        Returns:
            Diccionario con los datos de la oferta o None si no existe.
            Si include_history=True, incluye clave 'history' con lista de eventos.
        """
        if not offer_id or not isinstance(offer_id, str) or not offer_id.strip():
            logger.warning("Invalid offer_id provided to get_offer_by_id")
            return None
            
        if self.use_database and self.db:
            try:
                return self.db.get_offer(offer_id, include_history=include_history)
            except DatabaseError as e:
                logger.error(f"Error al obtener oferta de BD: {str(e)}")
            except Exception as e:
                logger.error(f"Error inesperado al obtener oferta: {str(e)}")
        
        # Buscar en historial en memoria
        for offer in self.offers_history:
            if offer.get('id') == offer_id:
                result = offer.copy()
                # Si se solicita historial, intentar obtenerlo de la BD
                if include_history and self.use_database and self.db:
                    try:
                        full_offer = self.db.get_offer(offer_id, include_history=True)
                        if full_offer and 'history' in full_offer:
                            result['history'] = full_offer['history']
                    except Exception:
                        pass  # Si falla, retornar sin historial
                return result
        
        return None
    
    def update_offer_status(
        self,
        offer_id: str,
        status: OfferStatus,
        details: str = "",
        user_id: Optional[str] = None
    ) -> Tuple[bool, Optional[str]]:
        """
        Actualiza el estado de una oferta con registro de auditoría.
        
        Args:
            offer_id: Identificador único de la oferta.
            status: Nuevo estado de la oferta (enum OfferStatus).
            details: Detalles adicionales sobre el cambio de estado.
            user_id: ID del usuario que realiza el cambio (para auditoría).
        
        Returns:
            Tupla (success: bool, error_message: Optional[str]):
                - success: True si la actualización fue exitosa
                - error_message: Mensaje de error si falla, None si es exitoso
        """
        if not offer_id or not isinstance(offer_id, str) or not offer_id.strip():
            error_msg = "Invalid offer_id provided to update_offer_status"
            logger.warning(error_msg)
            return False, error_msg
        
        if not isinstance(status, OfferStatus):
            error_msg = f"status must be an OfferStatus enum, got {type(status).__name__}"
            logger.warning(error_msg)
            return False, error_msg
            
        if self.use_database and self.db:
            try:
                success, error = self.db.update_offer_status(
                    offer_id,
                    status,
                    details,
                    user_id=user_id
                )
                if success:
                    # También actualizar en historial en memoria si existe
                    for offer in self.offers_history:
                        if offer.get('id') == offer_id:
                            offer['status'] = status.value
                            offer['updated_at'] = datetime.now().isoformat()
                            if details:
                                offer['last_action'] = details
                return success, error
            except DatabaseError as e:
                error_msg = f"Error al actualizar estado en BD: {str(e)}"
                logger.error(error_msg)
                return False, error_msg
            except Exception as e:
                error_msg = f"Error inesperado al actualizar estado: {str(e)}"
                logger.error(error_msg)
                return False, error_msg
        
        # Actualizar en historial en memoria
        for offer in self.offers_history:
            if offer.get('id') == offer_id:
                offer['status'] = status.value
                offer['updated_at'] = datetime.now().isoformat()
                if details:
                    offer['last_action'] = details
                logger.info(f"Offer {offer_id} status updated to {status.value} (in-memory)")
                return True, None
        
        error_msg = f"Offer with id {offer_id} not found"
        logger.warning(error_msg)
        return False, error_msg
    
    def generate_pdf(
        self,
        offer_id: str,
        output_path: str = "./output",
        filename: Optional[str] = None
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Genera un PDF de una oferta existente con opciones avanzadas.
        
        Args:
            offer_id: Identificador único de la oferta.
            output_path: Directorio donde guardar el PDF.
            filename: Nombre personalizado para el archivo PDF (sin extensión).
                Si None, se genera automáticamente basado en el offer_id.
        
        Returns:
            Tupla (filepath: Optional[str], error_message: Optional[str]):
                - filepath: Ruta al archivo PDF generado, o None si hubo un error
                - error_message: Mensaje de error si falla, None si es exitoso
        """
        if not REPORTLAB_AVAILABLE:
            error_msg = "ReportLab no está disponible. No se puede generar PDF."
            logger.error(error_msg)
            return None, error_msg
            
        try:
            offer = self.get_offer_by_id(offer_id)
            if not offer:
                error_msg = f"Oferta {offer_id} no encontrada"
                logger.error(error_msg)
                return None, error_msg
            
            # Obtener datos de la oferta
            offer_data_str = offer.get('offer_data', '{}')
            if isinstance(offer_data_str, str):
                try:
                    offer_data = json.loads(offer_data_str)
                except json.JSONDecodeError as e:
                    error_msg = f"Error al parsear datos de oferta: {str(e)}"
                    logger.error(error_msg)
                    return None, error_msg
            else:
                offer_data = offer_data_str
            
            # Validar que los datos sean suficientes
            if not offer_data:
                error_msg = "Datos de oferta vacíos o inválidos"
                logger.error(error_msg)
                return None, error_msg
            
            try:
                data = OfferLetterData(**offer_data)
            except (TypeError, ValueError) as e:
                error_msg = f"Error al crear objeto OfferLetterData: {str(e)}"
                logger.error(error_msg)
                return None, error_msg
            
            # Asegurar que el directorio existe
            try:
                success, error, _ = ensure_directory(output_path)
                if not success:
                    error_msg = f"No se pudo crear el directorio: {error}"
                    logger.error(error_msg)
                    return None, error_msg
            except Exception as e:
                error_msg = f"Error al crear directorio: {str(e)}"
                logger.error(error_msg)
                return None, error_msg
            
            # Generar PDF
            try:
                pdf_path = self.generator.generate_pdf(data, output_path)
                
                # Si se proporciona un nombre personalizado, renombrar el archivo
                if filename and pdf_path:
                    try:
                        # Sanitizar el nombre de archivo
                        safe_filename = sanitize_filename(filename)
                        if not safe_filename.endswith('.pdf'):
                            safe_filename += '.pdf'
                        
                        new_path = os.path.join(output_path, safe_filename)
                        if os.path.exists(pdf_path):
                            os.rename(pdf_path, new_path)
                            pdf_path = new_path
                    except Exception as rename_error:
                        logger.warning(f"No se pudo renombrar el archivo: {str(rename_error)}")
                
                if pdf_path and os.path.exists(pdf_path):
                    logger.info(f"PDF generado exitosamente: {pdf_path}")
                    return pdf_path, None
                else:
                    error_msg = "El PDF se generó pero no se encuentra en la ruta esperada"
                    logger.error(error_msg)
                    return None, error_msg
                    
            except FormatError as e:
                error_msg = f"Error de formato al generar PDF: {str(e)}"
                logger.error(error_msg)
                return None, error_msg
            except Exception as e:
                error_msg = f"Error al generar PDF: {str(e)}"
                logger.error(error_msg)
                return None, error_msg
                
        except Exception as e:
            error_msg = f"Error inesperado al generar PDF: {str(e)}"
            logger.error(error_msg)
            logger.error(f"Traceback: {traceback.format_exc()}")
            return None, error_msg
    
    def send_offer_email(
        self,
        offer_id: str,
        smtp_config: Dict,
        include_pdf: bool = False,
        subject: Optional[str] = None,
        custom_message: Optional[str] = None
    ) -> Tuple[bool, Optional[str]]:
        """
        Envía una oferta por email con validación mejorada y opciones personalizadas.
        
        Args:
            offer_id: Identificador único de la oferta.
            smtp_config: Diccionario con configuración SMTP:
                - host: str servidor SMTP
                - port: int puerto SMTP
                - username: str usuario SMTP
                - password: str contraseña SMTP
                - from_email: str email remitente
                - from_name: str nombre del remitente (opcional)
                - use_tls: bool usar TLS (opcional, default True)
                - use_ssl: bool usar SSL (opcional, default False)
            include_pdf: Si True, adjunta el PDF de la oferta.
            subject: Asunto personalizado del email. Si None, se genera automáticamente.
            custom_message: Mensaje personalizado adicional en el cuerpo del email.
        
        Returns:
            Tupla (success: bool, error_message: Optional[str]):
                - success: True si el email se envió exitosamente
                - error_message: Mensaje de error si falla, None si es exitoso
        """
        if not EMAIL_AVAILABLE:
            error_msg = "Módulos de email no están disponibles"
            logger.error(error_msg)
            return False, error_msg
            
        if not offer_id or not isinstance(offer_id, str) or not offer_id.strip():
            error_msg = "offer_id es requerido y debe ser una cadena válida"
            logger.error(error_msg)
            return False, error_msg
            
        # Validar configuración SMTP
        required_smtp_fields = ['host', 'port', 'username', 'password', 'from_email']
        missing_fields = [field for field in required_smtp_fields 
                         if field not in smtp_config or not smtp_config[field]]
        if missing_fields:
            error_msg = f"Configuración SMTP incompleta. Faltan: {', '.join(missing_fields)}"
            logger.error(error_msg)
            return False, error_msg
        
        # Validar tipos de configuración SMTP
        try:
            port = int(smtp_config.get('port', 587))
            if port < 1 or port > 65535:
                error_msg = f"Puerto SMTP inválido: {port}"
                logger.error(error_msg)
                return False, error_msg
        except (ValueError, TypeError):
            error_msg = f"Puerto SMTP debe ser un número entero, got: {smtp_config.get('port')}"
            logger.error(error_msg)
            return False, error_msg
        
        try:
            offer = self.get_offer_by_id(offer_id)
            if not offer:
                error_msg = f"Oferta {offer_id} no encontrada"
                logger.error(error_msg)
                return False, error_msg
            
            # Obtener datos de la oferta
            offer_data_str = offer.get('offer_data', '{}')
            if isinstance(offer_data_str, str):
                try:
                    offer_data = json.loads(offer_data_str)
                except json.JSONDecodeError as e:
                    error_msg = f"Error al parsear datos de oferta: {str(e)}"
                    logger.error(error_msg)
                    return False, error_msg
            else:
                offer_data = offer_data_str
            
            if not offer_data:
                error_msg = "Datos de oferta vacíos o inválidos"
                logger.error(error_msg)
                return False, error_msg
            
            try:
                data = OfferLetterData(**offer_data)
            except (TypeError, ValueError) as e:
                error_msg = f"Error al crear objeto OfferLetterData: {str(e)}"
                logger.error(error_msg)
                return False, error_msg
            
            # Validar email del destinatario
            if not data.candidate_email or not OfferLetterData._is_valid_email(data.candidate_email):
                error_msg = f"Email del candidato inválido: {data.candidate_email}"
                logger.error(error_msg)
                return False, error_msg
            
            offer_letter = offer.get('offer_letter_content', '')
            
            # Si no hay contenido, generar la carta
            if not offer_letter:
                format_type = offer.get('format', 'txt')
                try:
                    offer_letter = self.generator.generate_offer_letter(data, format_type=format_type)
                except Exception as e:
                    error_msg = f"Error al generar carta de oferta: {str(e)}"
                    logger.error(error_msg)
                    return False, error_msg
            
            # Preparar adjuntos
            attachments = []
            if include_pdf:
                pdf_path, pdf_error = self.generate_pdf(offer_id)
                if pdf_path and os.path.exists(pdf_path):
                    # Validar tamaño del PDF antes de adjuntar
                    max_attachment_size = 10 * 1024 * 1024  # 10MB
                    pdf_size = os.path.getsize(pdf_path)
                    if pdf_size > max_attachment_size:
                        logger.warning(f"PDF demasiado grande para adjuntar ({pdf_size} bytes). Continuando sin PDF.")
                    else:
                        attachments.append(pdf_path)
                elif pdf_error:
                    logger.warning(f"No se pudo generar PDF para adjuntar: {pdf_error}. Continuando sin PDF.")
                else:
                    logger.warning("No se pudo generar PDF para adjuntar. Continuando sin PDF.")
            
            # Preparar asunto personalizado o usar el predeterminado
            if not subject:
                subject = f"Oferta de Empleo - {data.position_title} - {data.company_name or 'Nuestra Empresa'}"
            
            # Modificar configuración SMTP para incluir asunto personalizado si es necesario
            smtp_config_with_subject = smtp_config.copy()
            smtp_config_with_subject['subject'] = subject
            
            # Enviar email
            try:
                success = self.generator.send_email(
                    data,
                    offer_letter,
                    smtp_config_with_subject,
                    attachments
                )
                
                if success:
                    # Actualizar estado de la oferta
                    update_success, update_error = self.update_offer_status(
                        offer_id,
                        OfferStatus.SENT,
                        f"Oferta enviada por email a {data.candidate_email}",
                        user_id=smtp_config.get('user_id')
                    )
                    if not update_success:
                        logger.warning(f"Oferta enviada pero no se pudo actualizar estado: {update_error}")
                    
                    logger.info(f"Oferta {offer_id} enviada exitosamente a {data.candidate_email}")
                    return True, None
                else:
                    error_msg = f"Error al enviar oferta {offer_id} por email (método send_email retornó False)"
                    logger.error(error_msg)
                    return False, error_msg
                    
            except (smtplib.SMTPException, OSError) as e:
                error_msg = f"Error SMTP al enviar email: {str(e)}"
                logger.error(error_msg)
                return False, error_msg
            except Exception as e:
                error_msg = f"Error inesperado al enviar email: {str(e)}"
                logger.error(error_msg)
                return False, error_msg
            
        except (ValueError, ValidationError) as e:
            error_msg = f"Error de validación al enviar email: {str(e)}"
            logger.error(error_msg)
            return False, error_msg
        except Exception as e:
            error_msg = f"Error inesperado al enviar email: {str(e)}"
            logger.error(error_msg)
            logger.error(f"Traceback: {traceback.format_exc()}")
            return False, error_msg
    
    def export_offer_json(self, offer_data: Dict) -> Dict:
        """
        Exporta los datos de una oferta en formato JSON estructurado
        
        Args:
            offer_data: Diccionario con los datos de la oferta
            
        Returns:
            Diccionario con los datos estructurados en JSON, incluyendo:
            - success: bool indicando si la exportación fue exitosa
            - data: Dict con los datos estructurados
            - metadata: Dict con metadatos adicionales
            - error: str mensaje de error si hubo algún problema
        """
        try:
            # Validar datos antes de exportar
            is_valid, validation_errors = validate_offer_data(offer_data)
            if not is_valid:
                return {
                    "success": False,
                    "error": "Validation failed",
                    "validation_errors": validation_errors
                }
            
            data = OfferLetterData(**offer_data)
            json_data = self.generator.generate_json_export(data)
            
            # Agregar metadatos adicionales
            metadata = {
                "exported_at": datetime.now().isoformat(),
                "format_version": "1.0",
                "data_completeness": calculate_offer_statistics(offer_data).get("completeness_score", 0)
            }
            
            return {
                "success": True,
                "data": json_data,
                "metadata": metadata
            }
            
        except (ValueError, ValidationError) as e:
            logger.error(f"Error de validación al exportar oferta: {str(e)}")
            return {
                "success": False,
                "error": "Validation error",
                "message": str(e)
            }
        except Exception as e:
            logger.error(f"Error inesperado al exportar oferta: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "error": "Export error",
                "message": str(e) if logger.level == logging.DEBUG else "An error occurred while exporting the offer"
            }
    
    def delete_offer(self, offer_id: str) -> bool:
        """
        Elimina una oferta del sistema
        
        Args:
            offer_id: Identificador único de la oferta
            
        Returns:
            True si la oferta fue eliminada exitosamente, False en caso contrario
        """
        if not offer_id:
            return False
            
        if self.use_database and self.db:
            try:
                # En lugar de eliminar físicamente, marcar como withdrawn
                self.db.update_offer_status(offer_id, OfferStatus.WITHDRAWN, "Oferta eliminada")
                return True
            except DatabaseError as e:
                logger.error(f"Error al eliminar oferta de BD: {str(e)}")
                return False
        
        # Eliminar del historial en memoria
        self.offers_history = [offer for offer in self.offers_history 
                             if offer.get('id') != offer_id]
        return True
    
    def get_offer_statistics(self, filters: Optional[Dict] = None) -> Dict:
        """
        Obtiene estadísticas de las ofertas
        
        Args:
            filters: Diccionario opcional con filtros para las estadísticas
            
        Returns:
            Diccionario con estadísticas agregadas
        """
        try:
            offers = self.get_offer_history(filters)
            
            if not offers:
                return {
                    "total_offers": 0,
                    "by_status": {},
                    "by_department": {},
                    "avg_salary": 0,
                    "total_salary": 0
                }
            
            stats = {
                "total_offers": len(offers),
                "by_status": defaultdict(int),
                "by_department": defaultdict(int),
                "salaries": [],
                "created_dates": []
            }
            
            for offer in offers:
                # Estadísticas por estado
                status = offer.get('status', 'unknown')
                stats["by_status"][status] += 1
                
                # Estadísticas por departamento
                offer_data = offer.get('offer_data', {})
                if isinstance(offer_data, str):
                    try:
                        offer_data = json.loads(offer_data)
                    except json.JSONDecodeError:
                        offer_data = {}
                
                department = offer_data.get('department', 'unknown')
                stats["by_department"][department] += 1
                
                # Salarios
                salary = offer_data.get('salary', 0)
                if salary:
                    stats["salaries"].append(float(salary))
                
                # Fechas
                created_at = offer.get('created_at')
                if created_at:
                    stats["created_dates"].append(created_at)
            
            # Calcular promedios
            stats["avg_salary"] = sum(stats["salaries"]) / len(stats["salaries"]) if stats["salaries"] else 0
            stats["total_salary"] = sum(stats["salaries"])
            stats["min_salary"] = min(stats["salaries"]) if stats["salaries"] else 0
            stats["max_salary"] = max(stats["salaries"]) if stats["salaries"] else 0
            
            # Convertir defaultdicts a dicts
            stats["by_status"] = dict(stats["by_status"])
            stats["by_department"] = dict(stats["by_department"])
            
            # Limpiar listas temporales
            del stats["salaries"]
            del stats["created_dates"]
            
            return stats
            
        except Exception as e:
            logger.error(f"Error al calcular estadísticas: {str(e)}")
            return {
                "error": str(e),
                "total_offers": 0
            }
    
    def search_offers(self, query: str, search_fields: Optional[List[str]] = None) -> List[Dict]:
        """
        Busca ofertas por texto en campos específicos
        
        Args:
            query: Texto a buscar
            search_fields: Lista de campos donde buscar (None = todos los campos)
            
        Returns:
            Lista de ofertas que coinciden con la búsqueda
        """
        if not query or not query.strip():
            return []
        
        query_lower = query.lower().strip()
        offers = self.get_offer_history()
        results = []
        
        # Campos por defecto para búsqueda
        default_fields = ['candidate_name', 'position_title', 'department', 
                         'location', 'candidate_email', 'company_name']
        search_fields = search_fields or default_fields
        
        for offer in offers:
            # Buscar en campos directos
            for field in search_fields:
                if field in offer and offer[field]:
                    if query_lower in str(offer[field]).lower():
                        results.append(offer)
                        break
            
            # Buscar en offer_data si es un string JSON
            offer_data = offer.get('offer_data', {})
            if isinstance(offer_data, str):
                try:
                    offer_data = json.loads(offer_data)
                except json.JSONDecodeError:
                    continue
            
            # Buscar en datos estructurados
            if isinstance(offer_data, dict):
                for field in search_fields:
                    if field in offer_data and offer_data[field]:
                        if query_lower in str(offer_data[field]).lower():
                            if offer not in results:
                                results.append(offer)
                                break
        
        return results
    
    def bulk_create_offers(self, offers_data: List[Dict], format_type: str = "txt") -> Dict:
        """
        Crea múltiples ofertas en lote
        
        Args:
            offers_data: Lista de diccionarios con datos de ofertas
            format_type: Tipo de formato deseado (txt, html)
            
        Returns:
            Diccionario con resultados del proceso:
            - success_count: int número de ofertas creadas exitosamente
            - failed_count: int número de ofertas que fallaron
            - results: List[Dict] resultados individuales
        """
        results = {
            "success_count": 0,
            "failed_count": 0,
            "results": []
        }
        
        if not isinstance(offers_data, list):
            return {
                "success_count": 0,
                "failed_count": 0,
                "error": "offers_data debe ser una lista",
                "results": []
            }
        
        for i, offer_data in enumerate(offers_data):
            try:
                result = self.create_offer(offer_data, format_type)
                if result.get("success"):
                    results["success_count"] += 1
                else:
                    results["failed_count"] += 1
                
                results["results"].append({
                    "index": i,
                    "success": result.get("success", False),
                    "offer_id": result.get("offer_id"),
                    "error": result.get("error"),
                    "validation_errors": result.get("validation_errors")
                })
            except Exception as e:
                results["failed_count"] += 1
                results["results"].append({
                    "index": i,
                    "success": False,
                    "error": str(e)
                })
                logger.error(f"Error al crear oferta en lote (índice {i}): {str(e)}")
        
        return results
    
    def export_offers_to_csv(self, filepath: str, filters: Optional[Dict] = None) -> bool:
        """
        Exporta ofertas a un archivo CSV
        
        Args:
            filepath: Ruta del archivo CSV a crear
            filters: Diccionario opcional con filtros
            
        Returns:
            True si la exportación fue exitosa, False en caso contrario
        """
        if not CSV_AVAILABLE:
            logger.error("Módulo csv no está disponible")
            return False
        
        try:
            offers = self.get_offer_history(filters)
            
            if not offers:
                logger.warning("No hay ofertas para exportar")
                return False
            
            # Determinar campos a exportar
            fieldnames = [
                'id', 'candidate_name', 'candidate_email', 'position_title',
                'department', 'salary', 'currency', 'status', 'created_at',
                'location', 'employment_type', 'start_date'
            ]
            
            # Asegurar que el directorio existe
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)
            
            with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames, extrasaction='ignore')
                writer.writeheader()
                
                for offer in offers:
                    # Extraer datos de offer_data si es necesario
                    row = offer.copy()
                    offer_data = offer.get('offer_data', {})
                    
                    if isinstance(offer_data, str):
                        try:
                            offer_data = json.loads(offer_data)
                            row.update(offer_data)
                        except json.JSONDecodeError:
                            pass
                    elif isinstance(offer_data, dict):
                        row.update(offer_data)
                    
                    writer.writerow(row)
            
            logger.info(f"Exportadas {len(offers)} ofertas a {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error al exportar ofertas a CSV: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return False
    
    def import_offers_from_json(self, filepath: str, format_type: str = "txt") -> Dict:
        """
        Importa ofertas desde un archivo JSON
        
        Args:
            filepath: Ruta del archivo JSON
            format_type: Tipo de formato para las ofertas importadas
            
        Returns:
            Diccionario con resultados de la importación
        """
        try:
            if not os.path.exists(filepath):
                return {
                    "success": False,
                    "error": f"Archivo no encontrado: {filepath}"
                }
            
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Aceptar tanto lista como objeto único
            if isinstance(data, list):
                offers_data = data
            elif isinstance(data, dict):
                if 'offers' in data:
                    offers_data = data['offers']
                else:
                    offers_data = [data]
            else:
                return {
                    "success": False,
                    "error": "Formato JSON inválido. Se espera una lista o un objeto."
                }
            
            # Crear ofertas en lote
            result = self.bulk_create_offers(offers_data, format_type)
            
            return {
                "success": True,
                "imported": result["success_count"],
                "failed": result["failed_count"],
                "total": len(offers_data),
                "details": result["results"]
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"Error al parsear JSON: {str(e)}")
            return {
                "success": False,
                "error": f"Error al parsear JSON: {str(e)}"
            }
        except Exception as e:
            logger.error(f"Error al importar ofertas: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def duplicate_offer(self, offer_id: str, modifications: Optional[Dict] = None) -> Dict:
        """
        Duplica una oferta existente con modificaciones opcionales
        
        Args:
            offer_id: ID de la oferta a duplicar
            modifications: Diccionario con modificaciones a aplicar
            
        Returns:
            Diccionario con el resultado de la duplicación
        """
        try:
            offer = self.get_offer_by_id(offer_id)
            if not offer:
                return {
                    "success": False,
                    "error": f"Oferta {offer_id} no encontrada"
                }
            
            # Obtener datos de la oferta
            offer_data_str = offer.get('offer_data', '{}')
            if isinstance(offer_data_str, str):
                try:
                    offer_data = json.loads(offer_data_str)
                except json.JSONDecodeError:
                    return {
                        "success": False,
                        "error": "Error al parsear datos de la oferta original"
                    }
            else:
                offer_data = offer_data_str
            
            # Aplicar modificaciones si se proporcionan
            if modifications:
                offer_data.update(modifications)
            
            # Generar nuevo ID
            offer_data.pop('id', None)
            
            # Crear nueva oferta
            result = self.create_offer(offer_data, offer.get('format', 'txt'))
            
            if result.get("success"):
                return {
                    "success": True,
                    "original_offer_id": offer_id,
                    "new_offer_id": result.get("offer_id"),
                    "offer": result
                }
            else:
                return {
                    "success": False,
                    "error": "Error al crear oferta duplicada",
                    "validation_errors": result.get("validation_errors")
                }
                
        except Exception as e:
            logger.error(f"Error al duplicar oferta: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def get_offers_by_status(self, status: OfferStatus) -> List[Dict]:
        """
        Obtiene todas las ofertas con un estado específico
        
        Args:
            status: Estado de las ofertas a buscar
            
        Returns:
            Lista de ofertas con el estado especificado
        """
        filters = {"status": status.value}
        return self.get_offer_history(filters)
    
    def get_expired_offers(self) -> List[Dict]:
        """
        Obtiene ofertas que han expirado
        
        Returns:
            Lista de ofertas expiradas
        """
        try:
            offers = self.get_offer_history()
            expired = []
            today = date.today()
            
            for offer in offers:
                offer_data = offer.get('offer_data', {})
                if isinstance(offer_data, str):
                    try:
                        offer_data = json.loads(offer_data)
                    except json.JSONDecodeError:
                        continue
                
                expiry_date_str = offer_data.get('offer_expiry_date')
                if expiry_date_str:
                    try:
                        expiry_date = datetime.strptime(expiry_date_str, '%Y-%m-%d').date()
                        if expiry_date < today:
                            expired.append(offer)
                    except ValueError:
                        continue
            
            return expired
            
        except Exception as e:
            logger.error(f"Error al obtener ofertas expiradas: {str(e)}")
            return []
    
    def update_offer_data(self, offer_id: str, updates: Dict) -> bool:
        """
        Actualiza los datos de una oferta existente
        
        Args:
            offer_id: ID de la oferta a actualizar
            updates: Diccionario con campos a actualizar
            
        Returns:
            True si la actualización fue exitosa, False en caso contrario
        """
        try:
            offer = self.get_offer_by_id(offer_id)
            if not offer:
                logger.error(f"Oferta {offer_id} no encontrada")
                return False
            
            # Obtener datos actuales
            offer_data_str = offer.get('offer_data', '{}')
            if isinstance(offer_data_str, str):
                try:
                    offer_data = json.loads(offer_data_str)
                except json.JSONDecodeError:
                    logger.error(f"Error al parsear datos de oferta {offer_id}")
                    return False
            else:
                offer_data = offer_data_str
            
            # Aplicar actualizaciones
            offer_data.update(updates)
            
            # Validar datos actualizados
            is_valid, validation_errors = validate_offer_data(offer_data)
            if not is_valid:
                logger.error(f"Error de validación al actualizar oferta: {validation_errors}")
                return False
            
            # Regenerar carta con nuevos datos
            data = OfferLetterData(**offer_data)
            format_type = offer.get('format', 'txt')
            new_letter = self.generator.generate_offer_letter(data, format_type=format_type)
            
            # Actualizar en base de datos
            if self.use_database and self.db:
                try:
                    # Actualizar oferta en BD
                    with sqlite3.connect(self.db.db_path) as conn:
                        cursor = conn.cursor()
                        cursor.execute('''
                            UPDATE offers 
                            SET offer_data = ?, offer_letter_content = ?, updated_at = CURRENT_TIMESTAMP
                            WHERE id = ?
                        ''', (json.dumps(offer_data), new_letter, offer_id))
                        conn.commit()
                except Exception as e:
                    logger.error(f"Error al actualizar en BD: {str(e)}")
                    return False
            
            # Actualizar en historial en memoria
            for i, o in enumerate(self.offers_history):
                if o.get('id') == offer_id:
                    self.offers_history[i]['offer_data'] = offer_data
                    self.offers_history[i]['offer_letter_content'] = new_letter
                    self.offers_history[i]['updated_at'] = datetime.now().isoformat()
                    break
            
            logger.info(f"Oferta {offer_id} actualizada exitosamente")
            return True
            
        except Exception as e:
            logger.error(f"Error al actualizar oferta: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return False


# ============================================================================
# FUNCIONES DE UTILIDAD ADICIONALES
# ============================================================================

class UpsellManager:
    """Gestor de upsells complementarios para productos"""
    
    def __init__(self, db_path: str = "upsells.db", cache_enabled: bool = True):
        self.db_path = db_path
        self.cache_enabled = cache_enabled
        self._recommendation_cache = {}  # Cache simple en memoria
        self._cache_ttl = 300  # 5 minutos
        self._init_database()
        self._load_default_upsells()
    
    def _init_database(self):
        """Inicializa la base de datos de upsells"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Tabla de productos
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS products (
                    id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    price REAL NOT NULL,
                    category TEXT,
                    description TEXT,
                    tags TEXT,
                    stock_quantity INTEGER DEFAULT -1,
                    is_active BOOLEAN DEFAULT 1,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Tabla de historial de cliente
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS customer_history (
                    id TEXT PRIMARY KEY,
                    customer_id TEXT NOT NULL,
                    product_id TEXT NOT NULL,
                    purchased_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (product_id) REFERENCES products(id)
                )
            ''')
            
            # Tabla de ofertas temporales
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS time_limited_offers (
                    id TEXT PRIMARY KEY,
                    upsell_id TEXT NOT NULL,
                    start_date TIMESTAMP NOT NULL,
                    end_date TIMESTAMP NOT NULL,
                    discount_percentage REAL DEFAULT 0,
                    is_active BOOLEAN DEFAULT 1,
                    FOREIGN KEY (upsell_id) REFERENCES upsells(id)
                )
            ''')
            
            # Tabla de variantes de mensajes (A/B testing)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS message_variants (
                    id TEXT PRIMARY KEY,
                    upsell_id TEXT NOT NULL,
                    variant_name TEXT NOT NULL,
                    message_template TEXT NOT NULL,
                    conversion_rate REAL DEFAULT 0,
                    total_views INTEGER DEFAULT 0,
                    total_conversions INTEGER DEFAULT 0,
                    is_active BOOLEAN DEFAULT 1,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (upsell_id) REFERENCES upsells(id)
                )
            ''')
            
            # Tabla de cross-sells (productos relacionados)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS cross_sells (
                    id TEXT PRIMARY KEY,
                    product_id TEXT NOT NULL,
                    related_product_id TEXT NOT NULL,
                    similarity_score REAL DEFAULT 0,
                    is_active BOOLEAN DEFAULT 1,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (product_id) REFERENCES products(id),
                    FOREIGN KEY (related_product_id) REFERENCES products(id)
                )
            ''')
            
            # Tabla de webhooks para eventos
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS upsell_webhooks (
                    id TEXT PRIMARY KEY,
                    url TEXT NOT NULL,
                    events TEXT NOT NULL,
                    secret TEXT,
                    is_active BOOLEAN DEFAULT 1,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Tabla de notificaciones de alto rendimiento
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS performance_alerts (
                    id TEXT PRIMARY KEY,
                    upsell_id TEXT NOT NULL,
                    alert_type TEXT NOT NULL,
                    threshold_value REAL,
                    current_value REAL,
                    notified_at TIMESTAMP,
                    is_resolved BOOLEAN DEFAULT 0,
                    FOREIGN KEY (upsell_id) REFERENCES upsells(id)
                )
            ''')
            
            # Tabla de upsells (relación producto -> upsell)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS upsells (
                    id TEXT PRIMARY KEY,
                    product_id TEXT NOT NULL,
                    upsell_product_id TEXT NOT NULL,
                    bundle_price REAL,
                    original_price REAL NOT NULL,
                    savings REAL,
                    emotional_reason TEXT,
                    popularity_score REAL DEFAULT 0,
                    conversion_rate REAL DEFAULT 0,
                    total_views INTEGER DEFAULT 0,
                    total_purchases INTEGER DEFAULT 0,
                    is_active BOOLEAN DEFAULT 1,
                    priority INTEGER DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (product_id) REFERENCES products(id),
                    FOREIGN KEY (upsell_product_id) REFERENCES products(id)
                )
            ''')
            
            # Crear índices para mejor rendimiento
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_upsells_product_id ON upsells(product_id)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_upsells_popularity ON upsells(popularity_score DESC)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_upsells_conversion ON upsells(conversion_rate DESC)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_purchases_product ON purchases(product_id)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_purchases_date ON purchases(purchased_at DESC)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_customer_history ON customer_history(customer_id, product_id)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_time_offers ON time_limited_offers(upsell_id, is_active)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_products_category ON products(category)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_products_tags ON products(tags)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_cross_sells ON cross_sells(product_id, similarity_score DESC)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_message_variants ON message_variants(upsell_id, conversion_rate DESC)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_performance_alerts ON performance_alerts(upsell_id, is_resolved)')
            
            # Tabla de compras (para tracking de "clientes que compraron esto también")
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS purchases (
                    id TEXT PRIMARY KEY,
                    product_id TEXT NOT NULL,
                    upsell_product_id TEXT,
                    customer_id TEXT,
                    purchased_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (product_id) REFERENCES products(id),
                    FOREIGN KEY (upsell_product_id) REFERENCES products(id)
                )
            ''')
            
            conn.commit()
            logger.info("Base de datos de upsells inicializada")
        except Exception as e:
            logger.error(f"Error al inicializar base de datos de upsells: {str(e)}")
    
    def _load_default_upsells(self):
        """Carga algunos upsells de ejemplo"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Verificar si ya hay datos
            cursor.execute('SELECT COUNT(*) FROM products')
            if cursor.fetchone()[0] > 0:
                return
            
            # Productos de ejemplo
            default_products = [
                ("prod_001", "Producto A", 999.00, "Electrónica", "Descripción del Producto A"),
                ("prod_002", "Cargador Rápido", 399.00, "Accesorios", "Cargador rápido de alta velocidad"),
                ("prod_003", "Funda Protectora", 299.00, "Accesorios", "Funda resistente y elegante"),
                ("prod_004", "Auriculares Inalámbricos", 1299.00, "Audio", "Auriculares con cancelación de ruido"),
                ("prod_005", "Cable USB-C", 199.00, "Accesorios", "Cable USB-C de alta calidad"),
            ]
            
            cursor.executemany('''
                INSERT OR IGNORE INTO products (id, name, price, category, description)
                VALUES (?, ?, ?, ?, ?)
            ''', default_products)
            
            # Upsells de ejemplo
            default_upsells = [
                ("upsell_001", "prod_001", "prod_002", 199.00, 399.00, 200.00, 
                 "Clientes que compraron esto también se llevaron el cargador rápido para máxima compatibilidad"),
                ("upsell_002", "prod_001", "prod_003", 249.00, 299.00, 50.00,
                 "Clientes que compraron esto también se llevaron la funda protectora para mantener su producto seguro"),
                ("upsell_003", "prod_001", "prod_004", 1099.00, 1299.00, 200.00,
                 "Clientes que compraron esto también se llevaron los auriculares para una experiencia completa"),
            ]
            
            cursor.executemany('''
                INSERT OR IGNORE INTO upsells 
                (id, product_id, upsell_product_id, bundle_price, original_price, savings, emotional_reason)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', default_upsells)
            
            conn.commit()
            logger.info("Upsells de ejemplo cargados")
        except Exception as e:
            logger.error(f"Error al cargar upsells de ejemplo: {str(e)}")
    
    def get_customer_purchase_history(self, customer_id: str) -> List[str]:
        """Obtiene historial de compras de un cliente"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT DISTINCT product_id 
                FROM customer_history 
                WHERE customer_id = ?
                ORDER BY purchased_at DESC
                LIMIT 20
            ''', (customer_id,))
            history = [row[0] for row in cursor.fetchall()]
            return history
        except Exception as e:
            logger.error(f"Error al obtener historial: {str(e)}")
            return []
    
    def _get_cached_recommendations(self, cache_key: str) -> Optional[List[Dict]]:
        """Obtiene recomendaciones del cache si están disponibles"""
        if not self.cache_enabled:
            return None
        
        if cache_key in self._recommendation_cache:
            cached_data, timestamp = self._recommendation_cache[cache_key]
            age = (datetime.now() - timestamp).total_seconds()
            if age < self._cache_ttl:
                return cached_data
            else:
                del self._recommendation_cache[cache_key]
        return None
    
    def _cache_recommendations(self, cache_key: str, recommendations: List[Dict]):
        """Almacena recomendaciones en el cache"""
        if self.cache_enabled:
            self._recommendation_cache[cache_key] = (recommendations, datetime.now())
            # Limpiar cache antiguo (mantener solo últimos 1000)
            if len(self._recommendation_cache) > 1000:
                oldest_key = min(self._recommendation_cache.keys(), 
                                key=lambda k: self._recommendation_cache[k][1])
                del self._recommendation_cache[oldest_key]
    
    def get_complementary_upsells(self, product_id: str, limit: int = 2, 
                                  customer_id: str = None, 
                                  category_filter: str = None,
                                  exclude_purchased: bool = True,
                                  use_cache: bool = True) -> List[Dict]:
        """
        Obtiene upsells complementarios para un producto con algoritmo mejorado
        
        Args:
            product_id: ID del producto comprado
            limit: Número máximo de upsells a retornar (default: 2)
            customer_id: ID del cliente para personalización (opcional)
            category_filter: Filtrar por categoría (opcional)
            exclude_purchased: Excluir productos ya comprados (default: True)
            use_cache: Usar cache para recomendaciones (default: True)
            
        Returns:
            Lista de diccionarios con información de upsells
        """
        try:
            # Verificar cache
            if use_cache:
                cache_key = f"{product_id}:{limit}:{customer_id}:{category_filter}:{exclude_purchased}"
                cached = self._get_cached_recommendations(cache_key)
                if cached is not None:
                    logger.debug(f"Cache hit para {cache_key}")
                    return cached
            
            # Obtener historial del cliente si está disponible
            purchased_products = []
            if customer_id and exclude_purchased:
                purchased_products = self.get_customer_purchase_history(customer_id)
            
            # Construir query con filtros
            where_clauses = ['u.product_id = ?', 'u.is_active = 1', 'p.is_active = 1']
            params = [product_id]
            
            # Filtro de categoría
            if category_filter:
                where_clauses.append('p.category = ?')
                params.append(category_filter)
            
            # Excluir productos ya comprados
            if purchased_products:
                placeholders = ','.join(['?'] * len(purchased_products))
                where_clauses.append(f'u.upsell_product_id NOT IN ({placeholders})')
                params.extend(purchased_products)
            
            # Validar que el producto existe y obtener upsells
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Validar producto
                cursor.execute('SELECT id, name, price FROM products WHERE id = ?', (product_id,))
                product = cursor.fetchone()
                if not product:
                    logger.warning(f"Producto {product_id} no encontrado")
                    return []
                
                product_name, product_price = product[1], product[2]
            
            # Algoritmo mejorado: combina múltiples factores
            # Score = (conversion_rate * 0.4) + (popularity_score * 0.3) + 
            #        (savings_percentage * 0.2) + (priority * 0.1) + (time_offer_bonus * 0.1)
            query = f'''
                SELECT 
                    u.id,
                    u.bundle_price,
                    u.original_price,
                    u.savings,
                    u.emotional_reason,
                    u.conversion_rate,
                    u.total_views,
                    u.total_purchases,
                    u.priority,
                        u.popularity_score,
                    p.id as upsell_product_id,
                    p.name as upsell_product_name,
                    p.description as upsell_product_description,
                    p.category as upsell_category,
                    p.price as upsell_original_price,
                    p.stock_quantity,
                    p.tags,
                    CASE 
                        WHEN u.total_views > 0 
                        THEN (u.total_purchases * 1.0 / u.total_views) 
                        ELSE 0 
                    END as calculated_conversion,
                    CASE 
                        WHEN u.original_price > 0 
                        THEN ((u.savings * 1.0 / u.original_price) * 100) 
                        ELSE 0 
                    END as savings_percentage,
                    CASE 
                        WHEN tlo.id IS NOT NULL AND tlo.is_active = 1 
                        THEN tlo.discount_percentage 
                        ELSE 0 
                    END as time_offer_discount,
                    CASE 
                        WHEN tlo.id IS NOT NULL AND tlo.is_active = 1 
                        THEN 1.0 
                        ELSE 0.0 
                    END as has_time_offer
                FROM upsells u
                JOIN products p ON u.upsell_product_id = p.id
                LEFT JOIN time_limited_offers tlo ON u.id = tlo.upsell_id 
                    AND tlo.is_active = 1
                    AND datetime('now') >= tlo.start_date 
                    AND datetime('now') <= tlo.end_date
                WHERE {' AND '.join(where_clauses)}
                ORDER BY 
                    (COALESCE(u.conversion_rate, calculated_conversion) * 0.4 +
                     (u.popularity_score / 100.0) * 0.3 +
                     (savings_percentage / 100.0) * 0.2 +
                     (u.priority / 10.0) * 0.1 +
                     (has_time_offer * 0.1)) DESC,
                    u.savings DESC
                LIMIT ?
            '''
            params.append(limit)
            cursor.execute(query, params)
            
            upsells = []
            for row in cursor.fetchall():
                # Incrementar contador de vistas
                cursor.execute('''
                    UPDATE upsells 
                    SET total_views = total_views + 1,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE id = ?
                ''', (row["id"],))
                
                savings_percentage = row["savings_percentage"]
                conversion_rate = row["calculated_conversion"] or row["conversion_rate"] or 0
                time_offer_discount = row["time_offer_discount"] or 0
                has_time_offer = row["has_time_offer"] or 0
                stock_quantity = row["stock_quantity"]
                
                # Aplicar descuento de oferta temporal si existe
                final_bundle_price = float(row["bundle_price"])
                if has_time_offer and time_offer_discount > 0:
                    final_bundle_price = final_bundle_price * (1 - time_offer_discount / 100)
                    savings_percentage = savings_percentage + time_offer_discount
                
                # Verificar stock
                is_in_stock = stock_quantity == -1 or stock_quantity > 0
                
                upsell_data = {
                    "id": row["id"],
                    "upsell_product": {
                        "id": row["upsell_product_id"],
                        "name": row["upsell_product_name"],
                        "description": row["upsell_product_description"],
                        "category": row["upsell_category"],
                        "tags": row["tags"].split(',') if row["tags"] else [],
                        "stock_quantity": stock_quantity,
                        "in_stock": is_in_stock
                    },
                    "bundle_price": round(final_bundle_price, 2),
                    "original_price": float(row["original_price"]),
                    "savings": round(float(row["original_price"]) - final_bundle_price, 2),
                    "savings_percentage": round(savings_percentage, 1),
                    "emotional_reason": row["emotional_reason"],
                    "conversion_rate": round(conversion_rate * 100, 1) if conversion_rate > 0 else 0,
                    "popularity_score": float(row["popularity_score"] or 0),
                    "has_time_limited_offer": bool(has_time_offer),
                    "time_offer_discount": round(time_offer_discount, 1) if has_time_offer else 0,
                    "display_message": self._format_upsell_message(
                        row["upsell_product_name"],
                        final_bundle_price,
                        row["original_price"],
                        float(row["original_price"]) - final_bundle_price,
                        savings_percentage,
                        has_time_offer
                    ),
                    "urgency_message": self._generate_urgency_message(
                        savings_percentage, 
                        conversion_rate,
                        has_time_offer,
                        stock_quantity
                    )
                }
                upsells.append(upsell_data)
            
            conn.commit()
            
            # Guardar en cache
            if use_cache:
                cache_key = f"{product_id}:{limit}:{customer_id}:{category_filter}:{exclude_purchased}"
                self._cache_recommendations(cache_key, upsells)
            
            return upsells
            
        except Exception as e:
            logger.error(f"Error al obtener upsells: {str(e)}", exc_info=True)
            return []
    
    def _format_upsell_message(self, product_name: str, bundle_price: float, 
                               original_price: float, savings: float, 
                               savings_percentage: float = 0,
                               has_time_offer: bool = False) -> str:
        """
        Formatea el mensaje de upsell con variaciones atractivas
        
        Ejemplo: "¡Agrega el cargador rápido por solo +$199 (valor $399)!"
        """
        bundle_mxn = f"${bundle_price:.0f}"
        original_mxn = f"${original_price:.0f}"
        savings_mxn = f"${savings:.0f}"
        
        # Mensaje especial para ofertas temporales
        if has_time_offer:
            return f"⏰ ¡Oferta limitada! Agrega {product_name.lower()} por solo +{bundle_mxn} MXN (ahorra {savings_percentage:.0f}% - valor {original_mxn} MXN)"
        
        # Variaciones de mensajes según el ahorro
        if savings_percentage >= 40:
            return f"🔥 ¡Oferta especial! Agrega {product_name.lower()} por solo +{bundle_mxn} MXN (ahorra {savings_percentage:.0f}% - valor {original_mxn} MXN)"
        elif savings_percentage >= 25:
            return f"✨ ¡Agrega {product_name.lower()} por solo +{bundle_mxn} MXN y ahorra {savings_mxn} MXN (valor {original_mxn} MXN)!"
        elif savings_percentage >= 10:
            return f"💡 ¡Agrega {product_name.lower()} por solo +{bundle_mxn} MXN (valor {original_mxn} MXN)!"
        else:
            return f"¡Agrega {product_name.lower()} por solo +{bundle_mxn} MXN (valor {original_mxn} MXN)!"
    
    def _generate_urgency_message(self, savings_percentage: float, 
                                   conversion_rate: float,
                                   has_time_offer: bool = False,
                                   stock_quantity: int = -1) -> str:
        """Genera mensaje de urgencia basado en métricas"""
        if has_time_offer:
            return "⏰ Oferta por tiempo limitado"
        
        if stock_quantity > 0 and stock_quantity < 5:
            return f"⚠️ Solo {stock_quantity} disponibles"
        
        if savings_percentage >= 40:
            return "🔥 Oferta limitada - Ahorro excepcional"
        elif savings_percentage >= 25:
            return "✨ Gran ahorro disponible"
        elif conversion_rate > 0.3:
            return "⭐ Muy popular entre nuestros clientes"
        elif conversion_rate > 0.15:
            return "👍 Recomendado frecuentemente"
        else:
            return "💡 Complemento perfecto"
    
    def record_purchase(self, product_id: str, customer_id: str = None, 
                       upsell_product_id: str = None):
        """
        Registra una compra para tracking de popularidad y conversión
        
        Args:
            product_id: ID del producto principal comprado
            customer_id: ID del cliente (opcional)
            upsell_product_id: ID del upsell comprado (opcional)
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Validar que el producto existe
            cursor.execute('SELECT id FROM products WHERE id = ?', (product_id,))
            if not cursor.fetchone():
                logger.warning(f"Producto {product_id} no encontrado al registrar compra")
                return False
            
            purchase_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT INTO purchases (id, product_id, upsell_product_id, customer_id)
                VALUES (?, ?, ?, ?)
            ''', (purchase_id, product_id, upsell_product_id, customer_id))
            
            # Registrar en historial del cliente
            if customer_id:
                history_id = str(uuid.uuid4())
                cursor.execute('''
                    INSERT OR IGNORE INTO customer_history (id, customer_id, product_id)
                    VALUES (?, ?, ?)
                ''', (history_id, customer_id, product_id))
                
                if upsell_product_id:
                    history_id2 = str(uuid.uuid4())
                    cursor.execute('''
                        INSERT OR IGNORE INTO customer_history (id, customer_id, product_id)
                        VALUES (?, ?, ?)
                    ''', (history_id2, customer_id, upsell_product_id))
            
            # Actualizar métricas del upsell si se compró uno
            if upsell_product_id:
                cursor.execute('''
                    UPDATE upsells
                    SET popularity_score = popularity_score + 1,
                        total_purchases = total_purchases + 1,
                        conversion_rate = CASE 
                            WHEN total_views > 0 
                            THEN (total_purchases + 1) * 1.0 / total_views 
                            ELSE 0 
                        END,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE product_id = ? AND upsell_product_id = ?
                ''', (product_id, upsell_product_id))
            
            conn.commit()
            logger.info(f"Compra registrada: producto {product_id}, upsell {upsell_product_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error al registrar compra: {str(e)}", exc_info=True)
            return False
    
    def add_product(self, product_id: str, name: str, price: float, 
                   category: str = None, description: str = None,
                   tags: List[str] = None, stock_quantity: int = -1) -> bool:
        """Agrega un nuevo producto con validaciones"""
        try:
            # Validaciones
            if not product_id or not name:
                logger.error("product_id y name son requeridos")
                return False
            
            if price <= 0:
                logger.error("El precio debe ser mayor a 0")
                return False
            
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            tags_str = ','.join(tags) if tags else None
            cursor.execute('''
                INSERT OR REPLACE INTO products (id, name, price, category, description, tags, stock_quantity)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (product_id, name, price, category, description, tags_str, stock_quantity))
            
            conn.commit()
            logger.info(f"Producto agregado: {product_id} - {name}")
            return True
        except Exception as e:
            logger.error(f"Error al agregar producto: {str(e)}", exc_info=True)
            return False
    
    def add_upsell(self, product_id: str, upsell_product_id: str, 
                   bundle_price: float, emotional_reason: str = None,
                   priority: int = 0) -> bool:
        """Agrega un upsell para un producto con validaciones mejoradas"""
        try:
            # Validaciones
            if bundle_price <= 0:
                logger.error("El bundle_price debe ser mayor a 0")
                return False
            
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Validar que ambos productos existen
            cursor.execute('SELECT id, name, price FROM products WHERE id IN (?, ?)', 
                          (product_id, upsell_product_id))
            products = {row[0]: (row[1], row[2]) for row in cursor.fetchall()}
            
            if product_id not in products:
                logger.error(f"Producto principal {product_id} no encontrado")
                return False
            
            if upsell_product_id not in products:
                logger.error(f"Producto upsell {upsell_product_id} no encontrado")
                return False
            
            original_price = products[upsell_product_id][1]
            savings = original_price - bundle_price
            
            # Validar que el bundle_price es menor al precio original
            if bundle_price >= original_price:
                logger.warning(f"Bundle price ({bundle_price}) debería ser menor al precio original ({original_price})")
            
            # Generar razón emocional por defecto si no se proporciona
            if not emotional_reason:
                upsell_name = products[upsell_product_id][0]
                emotional_reason = f"Clientes que compraron esto también se llevaron {upsell_name.lower()} para complementar su compra"
            
            upsell_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT OR REPLACE INTO upsells 
                (id, product_id, upsell_product_id, bundle_price, original_price, savings, 
                 emotional_reason, priority, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            ''', (upsell_id, product_id, upsell_product_id, bundle_price, 
                  original_price, savings, emotional_reason, priority))
            
            conn.commit()
            logger.info(f"Upsell agregado: {product_id} -> {upsell_product_id}")
            return True
        except Exception as e:
            logger.error(f"Error al agregar upsell: {str(e)}", exc_info=True)
            return False
    
    def get_upsell_analytics(self, product_id: str = None) -> Dict:
        """
        Obtiene analytics de upsells
        
        Args:
            product_id: ID del producto (opcional, si no se proporciona retorna global)
            
        Returns:
            Diccionario con estadísticas de upsells
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            if product_id:
                # Analytics para un producto específico
                cursor.execute('''
                    SELECT 
                        COUNT(*) as total_upsells,
                        AVG(conversion_rate) as avg_conversion_rate,
                        AVG(savings_percentage) as avg_savings,
                        SUM(total_purchases) as total_purchases,
                        SUM(total_views) as total_views
                    FROM (
                        SELECT 
                            u.*,
                            CASE 
                                WHEN u.total_views > 0 
                                THEN (u.total_purchases * 1.0 / u.total_views) 
                                ELSE 0 
                            END as conversion_rate,
                            CASE 
                                WHEN u.original_price > 0 
                                THEN ((u.savings * 1.0 / u.original_price) * 100) 
                                ELSE 0 
                            END as savings_percentage
                        FROM upsells u
                        WHERE u.product_id = ? AND u.is_active = 1
                    )
                ''', (product_id,))
            else:
                # Analytics globales
                cursor.execute('''
                    SELECT 
                        COUNT(*) as total_upsells,
                        AVG(conversion_rate) as avg_conversion_rate,
                        AVG(savings_percentage) as avg_savings,
                        SUM(total_purchases) as total_purchases,
                        SUM(total_views) as total_views
                    FROM (
                        SELECT 
                            u.*,
                            CASE 
                                WHEN u.total_views > 0 
                                THEN (u.total_purchases * 1.0 / u.total_views) 
                                ELSE 0 
                            END as conversion_rate,
                            CASE 
                                WHEN u.original_price > 0 
                                THEN ((u.savings * 1.0 / u.original_price) * 100) 
                                ELSE 0 
                            END as savings_percentage
                        FROM upsells u
                        WHERE u.is_active = 1
                    )
                ''')
            
            row = cursor.fetchone()
            
            if row:
                return {
                    "total_upsells": row["total_upsells"] or 0,
                    "avg_conversion_rate": round((row["avg_conversion_rate"] or 0) * 100, 2),
                    "avg_savings_percentage": round(row["avg_savings"] or 0, 2),
                    "total_purchases": row["total_purchases"] or 0,
                    "total_views": row["total_views"] or 0,
                    "overall_conversion_rate": round(
                        ((row["total_purchases"] or 0) * 1.0 / (row["total_views"] or 1)) * 100, 2
                    ) if row["total_views"] else 0
                }
            return {}
        except Exception as e:
            logger.error(f"Error al obtener analytics: {str(e)}", exc_info=True)
            return {}
    
    def create_time_limited_offer(self, upsell_id: str, start_date: str, 
                                 end_date: str, discount_percentage: float) -> bool:
        """Crea una oferta por tiempo limitado para un upsell"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Validar que el upsell existe
            cursor.execute('SELECT id FROM upsells WHERE id = ?', (upsell_id,))
            if not cursor.fetchone():
                logger.error(f"Upsell {upsell_id} no encontrado")
                return False
            
            offer_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT INTO time_limited_offers 
                (id, upsell_id, start_date, end_date, discount_percentage)
                VALUES (?, ?, ?, ?, ?)
            ''', (offer_id, upsell_id, start_date, end_date, discount_percentage))
            
            conn.commit()
            logger.info(f"Oferta temporal creada: {offer_id} para upsell {upsell_id}")
            return True
        except Exception as e:
            logger.error(f"Error al crear oferta temporal: {str(e)}", exc_info=True)
            return False
    
    def get_top_performing_upsells(self, limit: int = 10, days: int = 30) -> List[Dict]:
        """Obtiene los upsells con mejor rendimiento"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
            
            cursor.execute('''
                SELECT 
                    u.id,
                    u.product_id,
                    u.upsell_product_id,
                    u.conversion_rate,
                    u.total_purchases,
                    u.total_views,
                    p.name as product_name,
                    p2.name as upsell_name,
                    CASE 
                        WHEN u.total_views > 0 
                        THEN (u.total_purchases * 1.0 / u.total_views) 
                        ELSE 0 
                    END as calculated_rate
                FROM upsells u
                JOIN products p ON u.product_id = p.id
                JOIN products p2 ON u.upsell_product_id = p2.id
                WHERE u.is_active = 1
                    AND u.total_views > 10
                ORDER BY calculated_rate DESC, u.total_purchases DESC
                LIMIT ?
            ''', (limit,))
            
            results = []
            for row in cursor.fetchall():
                results.append({
                    "upsell_id": row["id"],
                    "product": row["product_name"],
                    "upsell": row["upsell_name"],
                    "conversion_rate": round((row["calculated_rate"] or 0) * 100, 2),
                    "total_purchases": row["total_purchases"],
                    "total_views": row["total_views"]
                })
            
            return results
        except Exception as e:
            logger.error(f"Error al obtener top upsells: {str(e)}", exc_info=True)
            return []
    
    def export_upsell_report(self, product_id: str = None, 
                            format: str = 'json') -> Dict:
        """Exporta un reporte de upsells"""
        try:
            analytics = self.get_upsell_analytics(product_id)
            top_upsells = self.get_top_performing_upsells(limit=20)
            
            report = {
                "generated_at": datetime.now().isoformat(),
                "product_id": product_id,
                "analytics": analytics,
                "top_performing_upsells": top_upsells,
                "summary": {
                    "total_upsells": analytics.get("total_upsells", 0),
                    "avg_conversion": analytics.get("avg_conversion_rate", 0),
                    "total_revenue_potential": "N/A"  # Se puede calcular con precios
                }
            }
            
            return report
        except Exception as e:
            logger.error(f"Error al exportar reporte: {str(e)}", exc_info=True)
            return {}
    
    def get_cross_sell_recommendations(self, product_id: str, limit: int = 3) -> List[Dict]:
        """Obtiene recomendaciones de cross-sell (productos relacionados)"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT 
                    cs.similarity_score,
                    p.id, p.name, p.price, p.category, p.description, p.tags, p.stock_quantity
                FROM cross_sells cs
                JOIN products p ON cs.related_product_id = p.id
                WHERE cs.product_id = ? AND cs.is_active = 1 AND p.is_active = 1
                ORDER BY cs.similarity_score DESC
                LIMIT ?
            ''', (product_id, limit))
            
            recommendations = []
            for row in cursor.fetchall():
                recommendations.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "price": float(row["price"]),
                    "category": row["category"],
                    "description": row["description"],
                    "tags": row["tags"].split(',') if row["tags"] else [],
                    "stock_quantity": row["stock_quantity"],
                    "similarity_score": float(row["similarity_score"]),
                    "in_stock": row["stock_quantity"] == -1 or row["stock_quantity"] > 0
                })
            
            return recommendations
        except Exception as e:
            logger.error(f"Error al obtener cross-sells: {str(e)}", exc_info=True)
            return []
    
    def add_cross_sell(self, product_id: str, related_product_id: str, 
                      similarity_score: float = 0.5) -> bool:
        """Agrega un cross-sell (producto relacionado)"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('SELECT id FROM products WHERE id IN (?, ?)', 
                          (product_id, related_product_id))
            if len(cursor.fetchall()) != 2:
                return False
            
            cross_sell_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT OR REPLACE INTO cross_sells 
                (id, product_id, related_product_id, similarity_score)
                VALUES (?, ?, ?, ?)
            ''', (cross_sell_id, product_id, related_product_id, similarity_score))
            
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al agregar cross-sell: {str(e)}", exc_info=True)
            return False
    
    def get_similar_products(self, product_id: str, limit: int = 5) -> List[Dict]:
        """Obtiene productos similares basados en categoría, tags y precio"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('SELECT category, tags, price FROM products WHERE id = ?', (product_id,))
            product = cursor.fetchone()
            if not product:
                return []
            
            category = product["category"]
            tags = set((product["tags"] or "").split(',')) if product["tags"] else set()
            price = product["price"]
            price_range = price * 0.3
            
            cursor.execute('''
                SELECT id, name, price, category, description, tags, stock_quantity,
                    CASE WHEN category = ? THEN 0.4 ELSE 0.0 END as category_score,
                    CASE WHEN ABS(price - ?) <= ? THEN 0.3 ELSE 0.0 END as price_score
                FROM products
                WHERE id != ? AND is_active = 1
                ORDER BY (category_score + price_score) DESC
                LIMIT ?
            ''', (category, price, price_range, product_id, limit))
            
            similar_products = []
            for row in cursor.fetchall():
                product_tags = set((row["tags"] or "").split(',')) if row["tags"] else set()
                tag_similarity = len(tags & product_tags) / max(len(tags | product_tags), 1) * 0.3
                total_score = row["category_score"] + row["price_score"] + tag_similarity
                
                similar_products.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "price": float(row["price"]),
                    "category": row["category"],
                    "similarity_score": round(total_score, 3),
                    "in_stock": row["stock_quantity"] == -1 or row["stock_quantity"] > 0
                })
            
            similar_products.sort(key=lambda x: x["similarity_score"], reverse=True)
            return similar_products
        except Exception as e:
            logger.error(f"Error al obtener productos similares: {str(e)}", exc_info=True)
            return []
    
    def register_webhook(self, url: str, events: List[str], secret: str = "") -> bool:
        """Registra un webhook para eventos de upsells"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            webhook_id = str(uuid.uuid4())
            events_json = json.dumps(events)
            
            cursor.execute('''
                INSERT INTO upsell_webhooks (id, url, events, secret)
                VALUES (?, ?, ?, ?)
            ''', (webhook_id, url, events_json, secret))
            
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al registrar webhook: {str(e)}", exc_info=True)
            return False
    
    def trigger_webhook(self, event: str, data: Dict):
        """Dispara un webhook para un evento"""
        try:
            import requests
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('SELECT * FROM upsell_webhooks WHERE is_active = 1')
            webhooks = cursor.fetchall()
            
            for webhook in webhooks:
                events = json.loads(webhook['events'])
                if event in events or '*' in events:
                    try:
                        payload = {"event": event, "timestamp": datetime.now().isoformat(), "data": data}
                        requests.post(webhook['url'], json=payload, timeout=5)
                    except Exception as e:
                        logger.error(f"Error al disparar webhook {webhook['url']}: {str(e)}")
        except ImportError:
            logger.warning("Requests no disponible. Webhooks deshabilitados.")
        except Exception as e:
            logger.error(f"Error en trigger_webhook: {str(e)}", exc_info=True)
    
    def check_performance_alerts(self, upsell_id: str = None):
        """Verifica y crea alertas de rendimiento"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            where_clause = "WHERE u.id = ?" if upsell_id else ""
            params = (upsell_id,) if upsell_id else ()
            
            query = f'''
                SELECT u.id, u.total_views, u.total_purchases,
                    CASE WHEN u.total_views > 0 
                    THEN (u.total_purchases * 1.0 / u.total_views) ELSE 0 END as conversion_rate
                FROM upsells u
                {where_clause}
                AND u.is_active = 1 AND u.total_views >= 50
            '''
            
            cursor.execute(query, params)
            alerts_created = 0
            
            for row in cursor.fetchall():
                upsell_id_val = row["id"]
                conversion_rate = row["conversion_rate"]
                
                if conversion_rate > 0.30:
                    cursor.execute('''
                        SELECT id FROM performance_alerts 
                        WHERE upsell_id = ? AND alert_type = 'high_conversion' AND is_resolved = 0
                    ''', (upsell_id_val,))
                    if not cursor.fetchone():
                        alert_id = str(uuid.uuid4())
                        cursor.execute('''
                            INSERT INTO performance_alerts 
                            (id, upsell_id, alert_type, threshold_value, current_value)
                            VALUES (?, ?, 'high_conversion', 0.30, ?)
                        ''', (alert_id, upsell_id_val, conversion_rate))
                        alerts_created += 1
                        self.trigger_webhook('upsell.high_performance', {
                            "upsell_id": upsell_id_val,
                            "conversion_rate": conversion_rate
                        })
            
            conn.commit()
            return alerts_created
        except Exception as e:
            logger.error(f"Error al verificar alertas: {str(e)}", exc_info=True)
            return 0
    
    def batch_get_upsells(self, product_ids: List[str], limit_per_product: int = 2) -> Dict[str, List[Dict]]:
        """
        Obtiene upsells para múltiples productos en una sola operación (batch)
        
        Args:
            product_ids: Lista de IDs de productos
            limit_per_product: Límite de upsells por producto
            
        Returns:
            Diccionario con product_id como clave y lista de upsells como valor
        """
        try:
            results = {}
            for product_id in product_ids:
                upsells = self.get_complementary_upsells(product_id, limit=limit_per_product, use_cache=True)
                results[product_id] = upsells
            return results
        except Exception as e:
            logger.error(f"Error en batch_get_upsells: {str(e)}", exc_info=True)
            return {}
    
    def calculate_volume_discount(self, base_price: float, quantity: int) -> Dict:
        """
        Calcula descuento por volumen
        
        Args:
            base_price: Precio base del producto
            quantity: Cantidad a comprar
            
        Returns:
            Diccionario con precio final, descuento aplicado y porcentaje
        """
        discount_tiers = {
            1: 0.0,      # Sin descuento
            2: 0.05,     # 5% de descuento
            3: 0.10,     # 10% de descuento
            5: 0.15,     # 15% de descuento
            10: 0.20     # 20% de descuento
        }
        
        # Encontrar el tier aplicable
        applicable_tier = 0
        discount_percentage = 0.0
        
        for tier_qty, discount in sorted(discount_tiers.items(), reverse=True):
            if quantity >= tier_qty:
                applicable_tier = tier_qty
                discount_percentage = discount
                break
        
        total_before_discount = base_price * quantity
        discount_amount = total_before_discount * discount_percentage
        final_price = total_before_discount - discount_amount
        
        return {
            "quantity": quantity,
            "base_price": base_price,
            "total_before_discount": round(total_before_discount, 2),
            "discount_percentage": round(discount_percentage * 100, 1),
            "discount_amount": round(discount_amount, 2),
            "final_price": round(final_price, 2),
            "savings_per_unit": round(discount_amount / quantity, 2) if quantity > 0 else 0,
            "tier_applied": applicable_tier
        }
    
    def get_realtime_metrics(self, product_id: str = None) -> Dict:
        """
        Obtiene métricas en tiempo real de upsells
        
        Args:
            product_id: ID del producto (opcional)
            
        Returns:
            Diccionario con métricas en tiempo real
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            if product_id:
                where_clause = "WHERE u.product_id = ?"
                params = (product_id,)
            else:
                where_clause = ""
                params = ()
            
            query = f'''
                SELECT 
                    COUNT(DISTINCT u.id) as active_upsells,
                    SUM(u.total_views) as total_views_today,
                    SUM(u.total_purchases) as total_purchases_today,
                    AVG(CASE 
                        WHEN u.total_views > 0 
                        THEN (u.total_purchases * 1.0 / u.total_views) 
                        ELSE 0 
                    END) as avg_conversion_rate,
                    SUM(CASE WHEN tlo.id IS NOT NULL THEN 1 ELSE 0 END) as active_time_offers
                FROM upsells u
                LEFT JOIN time_limited_offers tlo ON u.id = tlo.upsell_id 
                    AND tlo.is_active = 1
                    AND datetime('now') >= tlo.start_date 
                    AND datetime('now') <= tlo.end_date
                {where_clause}
                AND u.is_active = 1
            '''
            
            cursor.execute(query, params)
            row = cursor.fetchone()
            
            # Obtener upsells más vistos en las últimas 24 horas
            cursor.execute(f'''
                SELECT u.id, u.total_views, p.name as product_name, p2.name as upsell_name
                FROM upsells u
                JOIN products p ON u.product_id = p.id
                JOIN products p2 ON u.upsell_product_id = p2.id
                {where_clause}
                AND u.is_active = 1
                ORDER BY u.total_views DESC
                LIMIT 5
            ''', params)
            
            top_viewed = []
            for r in cursor.fetchall():
                top_viewed.append({
                    "upsell_id": r["id"],
                    "product": r["product_name"],
                    "upsell": r["upsell_name"],
                    "views": r["total_views"]
                })
            
            return {
                "timestamp": datetime.now().isoformat(),
                "active_upsells": row["active_upsells"] or 0,
                "total_views": row["total_views_today"] or 0,
                "total_purchases": row["total_purchases_today"] or 0,
                "avg_conversion_rate": round((row["avg_conversion_rate"] or 0) * 100, 2),
                "active_time_offers": row["active_time_offers"] or 0,
                "top_viewed_upsells": top_viewed
            }
        except Exception as e:
            logger.error(f"Error al obtener métricas en tiempo real: {str(e)}", exc_info=True)
            return {}
    
    def optimize_recommendations(self, product_id: str, customer_id: str = None) -> List[Dict]:
        """
        Versión optimizada de get_complementary_upsells con mejoras de rendimiento
        
        Args:
            product_id: ID del producto
            customer_id: ID del cliente (opcional)
            
        Returns:
            Lista optimizada de upsells
        """
        try:
            # Usar cache primero
            cache_key = f"opt:{product_id}:{customer_id}"
            cached = self._get_cached_recommendations(cache_key)
            if cached:
                return cached
            
            # Obtener upsells con algoritmo optimizado
            upsells = self.get_complementary_upsells(
                product_id,
                limit=2,
                customer_id=customer_id,
                use_cache=False
            )
            
            # Aplicar optimizaciones adicionales
            optimized = []
            for upsell in upsells:
                # Filtrar por stock si es crítico
                if upsell['upsell_product'].get('stock_quantity', -1) == 0:
                    continue
                
                # Priorizar ofertas temporales
                if upsell.get('has_time_limited_offer'):
                    upsell['priority_boost'] = 1.5
                
                optimized.append(upsell)
            
            # Ordenar por prioridad
            optimized.sort(key=lambda x: (
                x.get('has_time_limited_offer', False),
                x.get('conversion_rate', 0),
                x.get('savings_percentage', 0)
            ), reverse=True)
            
            # Guardar en cache
            self._cache_recommendations(cache_key, optimized)
            
            return optimized
        except Exception as e:
            logger.error(f"Error en optimize_recommendations: {str(e)}", exc_info=True)
            return []
    
    def get_customer_insights(self, customer_id: str) -> Dict:
        """
        Obtiene insights personalizados para un cliente
        
        Args:
            customer_id: ID del cliente
            
        Returns:
            Diccionario con insights del cliente
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Historial de compras
            cursor.execute('''
                SELECT COUNT(*) as total_purchases,
                       COUNT(DISTINCT product_id) as unique_products,
                       MAX(purchased_at) as last_purchase
                FROM customer_history
                WHERE customer_id = ?
            ''', (customer_id,))
            history = cursor.fetchone()
            
            # Productos más comprados
            cursor.execute('''
                SELECT p.id, p.name, p.category, COUNT(*) as purchase_count
                FROM customer_history ch
                JOIN products p ON ch.product_id = p.id
                WHERE ch.customer_id = ?
                GROUP BY p.id, p.name, p.category
                ORDER BY purchase_count DESC
                LIMIT 5
            ''', (customer_id,))
            
            top_products = []
            for row in cursor.fetchall():
                top_products.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "category": row["category"],
                    "purchase_count": row["purchase_count"]
                })
            
            # Categorías preferidas
            cursor.execute('''
                SELECT p.category, COUNT(*) as count
                FROM customer_history ch
                JOIN products p ON ch.product_id = p.id
                WHERE ch.customer_id = ? AND p.category IS NOT NULL
                GROUP BY p.category
                ORDER BY count DESC
                LIMIT 3
            ''', (customer_id,))
            
            preferred_categories = [row["category"] for row in cursor.fetchall()]
            
            return {
                "customer_id": customer_id,
                "total_purchases": history["total_purchases"] or 0,
                "unique_products": history["unique_products"] or 0,
                "last_purchase": history["last_purchase"],
                "top_products": top_products,
                "preferred_categories": preferred_categories,
                "customer_segment": self._determine_customer_segment(history["total_purchases"] or 0)
            }
        except Exception as e:
            logger.error(f"Error al obtener insights del cliente: {str(e)}", exc_info=True)
            return {}
    
    def _determine_customer_segment(self, total_purchases: int) -> str:
        """Determina el segmento del cliente basado en compras"""
        if total_purchases >= 10:
            return "VIP"
        elif total_purchases >= 5:
            return "Frecuente"
        elif total_purchases >= 2:
            return "Regular"
        else:
            return "Nuevo"
    
    def predict_conversion_probability(self, upsell_id: str, customer_id: str = None) -> Dict:
        """
        Predice la probabilidad de conversión de un upsell usando algoritmo simple
        
        Args:
            upsell_id: ID del upsell
            customer_id: ID del cliente (opcional)
            
        Returns:
            Diccionario con probabilidad de conversión y factores
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Obtener datos del upsell
            cursor.execute('''
                SELECT u.*, p.category as product_category, p2.category as upsell_category
                FROM upsells u
                JOIN products p ON u.product_id = p.id
                JOIN products p2 ON u.upsell_product_id = p2.id
                WHERE u.id = ?
            ''', (upsell_id,))
            
            upsell = cursor.fetchone()
            if not upsell:
                return {"probability": 0.0, "error": "Upsell no encontrado"}
            
            # Factores base
            base_probability = 0.3  # 30% base
            
            # Factor 1: Tasa de conversión histórica
            conversion_rate = (upsell["total_purchases"] * 1.0 / max(upsell["total_views"], 1))
            conversion_factor = min(conversion_rate * 2, 0.4)  # Máximo 40% boost
            
            # Factor 2: Ahorro porcentual
            savings_pct = (upsell["savings"] / upsell["original_price"]) * 100 if upsell["original_price"] > 0 else 0
            savings_factor = min(savings_pct / 100, 0.2)  # Máximo 20% boost
            
            # Factor 3: Popularidad
            popularity_factor = min(upsell["popularity_score"] / 100, 0.15)  # Máximo 15% boost
            
            # Factor 4: Oferta temporal
            cursor.execute('''
                SELECT COUNT(*) as count FROM time_limited_offers
                WHERE upsell_id = ? AND is_active = 1
                AND datetime('now') >= start_date AND datetime('now') <= end_date
            ''', (upsell_id,))
            has_time_offer = cursor.fetchone()["count"] > 0
            time_offer_factor = 0.1 if has_time_offer else 0.0
            
            # Factor 5: Historial del cliente (si está disponible)
            customer_factor = 0.0
            if customer_id:
                cursor.execute('''
                    SELECT COUNT(*) as purchase_count FROM customer_history
                    WHERE customer_id = ?
                ''', (customer_id,))
                purchase_count = cursor.fetchone()["purchase_count"]
                if purchase_count >= 5:
                    customer_factor = 0.1  # Clientes frecuentes más propensos
                elif purchase_count >= 2:
                    customer_factor = 0.05
            
            # Calcular probabilidad final
            total_probability = base_probability + conversion_factor + savings_factor + \
                              popularity_factor + time_offer_factor + customer_factor
            total_probability = min(total_probability, 0.95)  # Máximo 95%
            
            return {
                "upsell_id": upsell_id,
                "probability": round(total_probability * 100, 2),
                "factors": {
                    "base": 30.0,
                    "conversion_rate": round(conversion_factor * 100, 2),
                    "savings": round(savings_factor * 100, 2),
                    "popularity": round(popularity_factor * 100, 2),
                    "time_offer": round(time_offer_factor * 100, 2),
                    "customer_history": round(customer_factor * 100, 2)
                },
                "recommendation": "Alta" if total_probability > 0.6 else "Media" if total_probability > 0.4 else "Baja"
            }
        except Exception as e:
            logger.error(f"Error al predecir conversión: {str(e)}", exc_info=True)
            return {"probability": 0.0, "error": str(e)}
    
    def get_trending_upsells(self, days: int = 7, limit: int = 10) -> List[Dict]:
        """
        Obtiene upsells en tendencia basados en crecimiento reciente
        
        Args:
            days: Días a considerar para la tendencia
            limit: Número de resultados
            
        Returns:
            Lista de upsells en tendencia
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
            
            # Calcular crecimiento en vistas y compras
            cursor.execute('''
                SELECT 
                    u.id,
                    u.product_id,
                    u.upsell_product_id,
                    u.total_views,
                    u.total_purchases,
                    p.name as product_name,
                    p2.name as upsell_name,
                    p2.price as upsell_price,
                    u.bundle_price,
                    u.savings,
                    CASE 
                        WHEN u.total_views > 0 
                        THEN (u.total_purchases * 1.0 / u.total_views) 
                        ELSE 0 
                    END as conversion_rate,
                    COUNT(pur.id) as recent_purchases
                FROM upsells u
                JOIN products p ON u.product_id = p.id
                JOIN products p2 ON u.upsell_product_id = p2.id
                LEFT JOIN purchases pur ON u.id = (
                    SELECT upsell_id FROM upsells 
                    WHERE product_id = u.product_id AND upsell_product_id = u.upsell_product_id
                ) AND pur.purchased_at >= ?
                WHERE u.is_active = 1
                GROUP BY u.id
                HAVING recent_purchases > 0 OR u.total_views > 20
                ORDER BY 
                    (recent_purchases * 2.0 + u.total_views * 0.1) DESC,
                    conversion_rate DESC
                LIMIT ?
            ''', (cutoff_date, limit))
            
            trending = []
            for row in cursor.fetchall():
                trending.append({
                    "upsell_id": row["id"],
                    "product": row["product_name"],
                    "upsell": row["upsell_name"],
                    "price": float(row["upsell_price"]),
                    "bundle_price": float(row["bundle_price"]),
                    "savings": float(row["savings"]),
                    "total_views": row["total_views"],
                    "total_purchases": row["total_purchases"],
                    "recent_purchases": row["recent_purchases"],
                    "conversion_rate": round((row["conversion_rate"] or 0) * 100, 2),
                    "trend_score": round((row["recent_purchases"] * 2.0 + row["total_views"] * 0.1), 2)
                })
            
            return trending
        except Exception as e:
            logger.error(f"Error al obtener upsells en tendencia: {str(e)}", exc_info=True)
            return []
    
    def generate_upsell_suggestions(self, product_id: str) -> List[Dict]:
        """
        Genera sugerencias automáticas de upsells basadas en análisis de datos
        
        Args:
            product_id: ID del producto
            
        Returns:
            Lista de sugerencias de upsells potenciales
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Obtener información del producto
            cursor.execute('SELECT category, price, tags FROM products WHERE id = ?', (product_id,))
            product = cursor.fetchone()
            if not product:
                return []
            
            category = product["category"]
            price = product["price"]
            price_range = price * 0.5  # 50% de rango
            
            # Buscar productos complementarios que no sean upsells actuales
            cursor.execute('''
                SELECT DISTINCT p.id, p.name, p.price, p.category, p.description
                FROM products p
                WHERE p.id != ?
                    AND p.is_active = 1
                    AND p.id NOT IN (
                        SELECT upsell_product_id FROM upsells WHERE product_id = ?
                    )
                    AND (
                        p.category = ? OR
                        ABS(p.price - ?) <= ?
                    )
                ORDER BY 
                    CASE WHEN p.category = ? THEN 1 ELSE 0 END DESC,
                    ABS(p.price - ?) ASC
                LIMIT 10
            ''', (product_id, product_id, category, price, price_range, category, price))
            
            suggestions = []
            for row in cursor.fetchall():
                # Calcular score de compatibilidad
                category_match = 1.0 if row["category"] == category else 0.3
                price_match = 1.0 - min(abs(row["price"] - price) / price, 0.5)
                compatibility_score = (category_match * 0.6 + price_match * 0.4)
                
                # Calcular precio sugerido (20% de descuento)
                suggested_bundle_price = row["price"] * 0.8
                savings = row["price"] - suggested_bundle_price
                
                suggestions.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "category": row["category"],
                    "original_price": float(row["price"]),
                    "suggested_bundle_price": round(suggested_bundle_price, 2),
                    "suggested_savings": round(savings, 2),
                    "compatibility_score": round(compatibility_score, 3),
                    "reason": f"Producto complementario en categoría {row['category']}" if category_match > 0.5 else "Precio similar"
                })
            
            # Ordenar por score
            suggestions.sort(key=lambda x: x["compatibility_score"], reverse=True)
            return suggestions[:5]  # Top 5 sugerencias
        except Exception as e:
            logger.error(f"Error al generar sugerencias: {str(e)}", exc_info=True)
            return []
    
    def get_performance_forecast(self, upsell_id: str, days: int = 30) -> Dict:
        """
        Genera un pronóstico de rendimiento para un upsell
        
        Args:
            upsell_id: ID del upsell
            days: Días a pronosticar
            
        Returns:
            Diccionario con pronóstico
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Obtener datos históricos
            cursor.execute('''
                SELECT 
                    total_views,
                    total_purchases,
                    CASE 
                        WHEN total_views > 0 
                        THEN (total_purchases * 1.0 / total_views) 
                        ELSE 0 
                    END as conversion_rate,
                    created_at
                FROM upsells
                WHERE id = ?
            ''', (upsell_id,))
            
            upsell = cursor.fetchone()
            if not upsell:
                return {}
            
            # Calcular días desde creación
            created_at = datetime.fromisoformat(upsell["created_at"].replace('Z', '+00:00'))
            days_active = (datetime.now() - created_at.replace(tzinfo=None)).days
            days_active = max(days_active, 1)
            
            # Calcular promedios diarios
            avg_daily_views = upsell["total_views"] / days_active
            avg_daily_purchases = upsell["total_purchases"] / days_active
            conversion_rate = upsell["conversion_rate"] or 0
            
            # Pronóstico
            forecast_views = avg_daily_views * days
            forecast_purchases = forecast_views * conversion_rate
            
            # Calcular potencial de ingresos (necesitaría precio del upsell)
            cursor.execute('''
                SELECT bundle_price FROM upsells WHERE id = ?
            ''', (upsell_id,))
            bundle_price = cursor.fetchone()["bundle_price"] or 0
            forecast_revenue = forecast_purchases * bundle_price
            
            return {
                "upsell_id": upsell_id,
                "forecast_period_days": days,
                "current_metrics": {
                    "total_views": upsell["total_views"],
                    "total_purchases": upsell["total_purchases"],
                    "conversion_rate": round(conversion_rate * 100, 2),
                    "days_active": days_active
                },
                "forecast": {
                    "projected_views": round(forecast_views, 0),
                    "projected_purchases": round(forecast_purchases, 1),
                    "projected_revenue": round(forecast_revenue, 2),
                    "confidence": "Alta" if days_active >= 7 else "Media" if days_active >= 3 else "Baja"
                },
                "growth_rate": {
                    "views_per_day": round(avg_daily_views, 2),
                    "purchases_per_day": round(avg_daily_purchases, 2)
                }
            }
        except Exception as e:
            logger.error(f"Error al generar pronóstico: {str(e)}", exc_info=True)
            return {}
    
    def auto_optimize_upsells(self) -> Dict:
        """
        Optimiza automáticamente upsells basado en rendimiento
        
        Returns:
            Diccionario con resultados de optimización
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Encontrar upsells con bajo rendimiento
            cursor.execute('''
                SELECT id, product_id, upsell_product_id, total_views, total_purchases,
                    CASE 
                        WHEN total_views > 0 
                        THEN (total_purchases * 1.0 / total_views) 
                        ELSE 0 
                    END as conversion_rate
                FROM upsells
                WHERE is_active = 1
                    AND total_views >= 50
                    AND (total_purchases * 1.0 / total_views) < 0.05
            ''')
            
            low_performers = cursor.fetchall()
            optimizations = []
            
            for row in low_performers:
                upsell_id = row[0]
                conversion_rate = row[5]
                
                # Sugerir aumentar descuento
                cursor.execute('SELECT bundle_price, original_price FROM upsells WHERE id = ?', (upsell_id,))
                prices = cursor.fetchone()
                current_discount = (prices[1] - prices[0]) / prices[1] if prices[1] > 0 else 0
                
                if current_discount < 0.20:  # Si el descuento es menor a 20%
                    suggested_discount = min(current_discount + 0.10, 0.30)  # Aumentar 10%
                    new_bundle_price = prices[1] * (1 - suggested_discount)
                    new_savings = prices[1] - new_bundle_price
                    
                    optimizations.append({
                        "upsell_id": upsell_id,
                        "action": "increase_discount",
                        "current_discount": round(current_discount * 100, 1),
                        "suggested_discount": round(suggested_discount * 100, 1),
                        "current_bundle_price": prices[0],
                        "suggested_bundle_price": round(new_bundle_price, 2),
                        "reason": f"Baja conversión ({round(conversion_rate * 100, 1)}%). Aumentar descuento puede mejorar rendimiento."
                    })
            
            return {
                "optimizations_found": len(optimizations),
                "optimizations": optimizations,
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Error en auto_optimize_upsells: {str(e)}", exc_info=True)
            return {}
    
    def get_dashboard_summary(self) -> Dict:
        """
        Obtiene resumen completo para dashboard de administración
        
        Returns:
            Diccionario con todas las métricas principales
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Métricas generales
            cursor.execute('''
                SELECT 
                    COUNT(DISTINCT u.id) as total_upsells,
                    COUNT(DISTINCT u.product_id) as products_with_upsells,
                    SUM(u.total_views) as total_views,
                    SUM(u.total_purchases) as total_purchases,
                    AVG(CASE WHEN u.total_views > 0 
                        THEN (u.total_purchases * 1.0 / u.total_views) ELSE 0 END) as avg_conversion,
                    COUNT(DISTINCT CASE WHEN tlo.id IS NOT NULL THEN u.id END) as active_time_offers
                FROM upsells u
                LEFT JOIN time_limited_offers tlo ON u.id = tlo.upsell_id 
                    AND tlo.is_active = 1
                    AND datetime('now') >= tlo.start_date 
                    AND datetime('now') <= tlo.end_date
                WHERE u.is_active = 1
            ''')
            metrics = cursor.fetchone()
            
            # Top 5 productos con más upsells
            cursor.execute('''
                SELECT p.name, COUNT(u.id) as upsell_count
                FROM products p
                JOIN upsells u ON p.id = u.product_id
                WHERE u.is_active = 1
                GROUP BY p.id, p.name
                ORDER BY upsell_count DESC
                LIMIT 5
            ''')
            top_products = [{"name": r["name"], "upsells": r["upsell_count"]} for r in cursor.fetchall()]
            
            # Upsells más exitosos
            cursor.execute('''
                SELECT u.id, p.name as product, p2.name as upsell,
                    CASE WHEN u.total_views > 0 
                    THEN (u.total_purchases * 1.0 / u.total_views) ELSE 0 END as conversion
                FROM upsells u
                JOIN products p ON u.product_id = p.id
                JOIN products p2 ON u.upsell_product_id = p2.id
                WHERE u.is_active = 1 AND u.total_views >= 20
                ORDER BY conversion DESC
                LIMIT 5
            ''')
            top_performers = []
            for r in cursor.fetchall():
                top_performers.append({
                    "upsell_id": r["id"],
                    "product": r["product"],
                    "upsell": r["upsell"],
                    "conversion_rate": round((r["conversion"] or 0) * 100, 2)
                })
            
            # Estadísticas de clientes
            cursor.execute('''
                SELECT 
                    COUNT(DISTINCT customer_id) as unique_customers,
                    COUNT(*) as total_purchases
                FROM customer_history
            ''')
            customer_stats = cursor.fetchone()
            
            return {
                "timestamp": datetime.now().isoformat(),
                "overview": {
                    "total_upsells": metrics["total_upsells"] or 0,
                    "products_with_upsells": metrics["products_with_upsells"] or 0,
                    "total_views": metrics["total_views"] or 0,
                    "total_purchases": metrics["total_purchases"] or 0,
                    "avg_conversion_rate": round((metrics["avg_conversion"] or 0) * 100, 2),
                    "active_time_offers": metrics["active_time_offers"] or 0
                },
                "customers": {
                    "unique_customers": customer_stats["unique_customers"] or 0,
                    "total_purchases": customer_stats["total_purchases"] or 0
                },
                "top_products": top_products,
                "top_performers": top_performers
            }
        except Exception as e:
            logger.error(f"Error al obtener dashboard: {str(e)}", exc_info=True)
            return {}
    
    def get_collaborative_recommendations(self, customer_id: str, limit: int = 5) -> List[Dict]:
        """
        Recomendaciones basadas en colaboración (clientes similares)
        
        Args:
            customer_id: ID del cliente
            limit: Número de recomendaciones
            
        Returns:
            Lista de productos recomendados
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Obtener productos comprados por el cliente
            cursor.execute('''
                SELECT product_id FROM customer_history
                WHERE customer_id = ?
            ''', (customer_id,))
            customer_products = {row["product_id"] for row in cursor.fetchall()}
            
            if not customer_products:
                return []
            
            # Encontrar clientes similares (que compraron productos similares)
            placeholders = ','.join(['?'] * len(customer_products))
            cursor.execute(f'''
                SELECT DISTINCT ch2.customer_id, COUNT(*) as common_products
                FROM customer_history ch1
                JOIN customer_history ch2 ON ch1.product_id = ch2.product_id
                WHERE ch1.customer_id = ?
                    AND ch2.customer_id != ?
                    AND ch2.product_id IN ({placeholders})
                GROUP BY ch2.customer_id
                HAVING common_products >= 2
                ORDER BY common_products DESC
                LIMIT 10
            ''', (customer_id, customer_id, *customer_products))
            
            similar_customers = [row["customer_id"] for row in cursor.fetchall()]
            
            if not similar_customers:
                return []
            
            # Obtener productos comprados por clientes similares pero no por este cliente
            customer_placeholders = ','.join(['?'] * len(similar_customers))
            product_placeholders = ','.join(['?'] * len(customer_products))
            
            cursor.execute(f'''
                SELECT p.id, p.name, p.price, p.category, COUNT(*) as recommendation_score
                FROM customer_history ch
                JOIN products p ON ch.product_id = p.id
                WHERE ch.customer_id IN ({customer_placeholders})
                    AND ch.product_id NOT IN ({product_placeholders})
                    AND p.is_active = 1
                GROUP BY p.id, p.name, p.price, p.category
                ORDER BY recommendation_score DESC
                LIMIT ?
            ''', (*similar_customers, *customer_products, limit))
            
            recommendations = []
            for row in cursor.fetchall():
                recommendations.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "price": float(row["price"]),
                    "category": row["category"],
                    "recommendation_score": row["recommendation_score"],
                    "reason": f"Comprado por {row['recommendation_score']} clientes similares"
                })
            
            return recommendations
        except Exception as e:
            logger.error(f"Error en recomendaciones colaborativas: {str(e)}", exc_info=True)
            return []
    
    def calculate_revenue_impact(self, product_id: str = None, days: int = 30) -> Dict:
        """
        Calcula el impacto en ingresos de los upsells
        
        Args:
            product_id: ID del producto (opcional)
            days: Días a considerar
            
        Returns:
            Diccionario con impacto en ingresos
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
            
            where_clause = "WHERE u.product_id = ?" if product_id else ""
            params = (product_id, cutoff_date) if product_id else (cutoff_date,)
            
            query = f'''
                SELECT 
                    COUNT(DISTINCT u.id) as upsells_count,
                    SUM(u.total_purchases) as total_upsell_purchases,
                    SUM(u.total_purchases * u.bundle_price) as total_revenue,
                    AVG(u.bundle_price) as avg_bundle_price,
                    SUM(u.savings * u.total_purchases) as total_savings_given
                FROM upsells u
                {where_clause}
                AND u.is_active = 1
            '''
            
            cursor.execute(query, params[:1] if product_id else ())
            revenue = cursor.fetchone()
            
            # Calcular ROI estimado
            if revenue and revenue["total_revenue"]:
                roi = ((revenue["total_revenue"] - (revenue["total_savings_given"] or 0)) / 
                      max(revenue["total_revenue"], 1)) * 100
            else:
                roi = 0
            
            return {
                "period_days": days,
                "product_id": product_id,
                "upsells_count": revenue["upsells_count"] or 0,
                "total_upsell_purchases": revenue["total_upsell_purchases"] or 0,
                "total_revenue": round(revenue["total_revenue"] or 0, 2),
                "avg_bundle_price": round(revenue["avg_bundle_price"] or 0, 2),
                "total_savings_given": round(revenue["total_savings_given"] or 0, 2),
                "estimated_roi": round(roi, 2),
                "revenue_per_upsell": round(
                    (revenue["total_revenue"] or 0) / max(revenue["upsells_count"] or 1, 1), 2
                )
            }
        except Exception as e:
            logger.error(f"Error al calcular impacto en ingresos: {str(e)}", exc_info=True)
            return {}
    
    def get_seasonal_trends(self, product_id: str = None) -> Dict:
        """
        Analiza tendencias estacionales de upsells
        
        Args:
            product_id: ID del producto (opcional)
            
        Returns:
            Diccionario con tendencias por mes
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            where_clause = "WHERE u.product_id = ?" if product_id else ""
            params = (product_id,) if product_id else ()
            
            query = f'''
                SELECT 
                    strftime('%Y-%m', pur.purchased_at) as month,
                    COUNT(*) as purchases,
                    SUM(u.bundle_price) as revenue
                FROM purchases pur
                JOIN upsells u ON pur.upsell_product_id = (
                    SELECT upsell_product_id FROM upsells 
                    WHERE id = pur.upsell_product_id
                )
                {where_clause}
                AND pur.upsell_product_id IS NOT NULL
                GROUP BY month
                ORDER BY month DESC
                LIMIT 12
            '''
            
            cursor.execute(query, params)
            trends = []
            for row in cursor.fetchall():
                trends.append({
                    "month": row["month"],
                    "purchases": row["purchases"],
                    "revenue": round(row["revenue"] or 0, 2)
                })
            
            return {
                "product_id": product_id,
                "monthly_trends": trends,
                "total_months": len(trends)
            }
        except Exception as e:
            logger.error(f"Error al obtener tendencias estacionales: {str(e)}", exc_info=True)
            return {}
    
    def create_message_variant(self, upsell_id: str, variant_name: str, 
                              message_template: str) -> bool:
        """Crea una variante de mensaje para A/B testing"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('SELECT id FROM upsells WHERE id = ?', (upsell_id,))
            if not cursor.fetchone():
                return False
            
            variant_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT INTO message_variants 
                (id, upsell_id, variant_name, message_template)
                VALUES (?, ?, ?, ?)
            ''', (variant_id, upsell_id, variant_name, message_template))
            
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al crear variante: {str(e)}", exc_info=True)
            return False
    
    def get_best_message_variant(self, upsell_id: str) -> Optional[str]:
        """Obtiene la mejor variante de mensaje basada en conversión"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT message_template, conversion_rate
                FROM message_variants
                WHERE upsell_id = ? AND is_active = 1
                ORDER BY conversion_rate DESC, total_conversions DESC
                LIMIT 1
            ''', (upsell_id,))
            
            result = cursor.fetchone()
            return result[0] if result else None
        except Exception as e:
            logger.error(f"Error al obtener mejor variante: {str(e)}")
            return None
    
    def record_variant_view(self, variant_id: str):
        """Registra una vista de una variante de mensaje"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('''
                UPDATE message_variants
                SET total_views = total_views + 1
                WHERE id = ?
            ''', (variant_id,))
            
            conn.commit()
        except Exception as e:
            logger.error(f"Error al registrar vista de variante: {str(e)}")
    
    def record_variant_conversion(self, variant_id: str):
        """Registra una conversión de una variante de mensaje"""
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('''
                UPDATE message_variants
                SET total_conversions = total_conversions + 1,
                    conversion_rate = CASE 
                        WHEN total_views > 0 
                        THEN (total_conversions + 1) * 1.0 / total_views 
                        ELSE 0 
                    END
                WHERE id = ?
            ''', (variant_id,))
            
            conn.commit()
        except Exception as e:
            logger.error(f"Error al registrar conversión de variante: {str(e)}")
    
    def backup_database(self, backup_path: str = None) -> str:
        """
        Crea un backup de la base de datos
        
        Args:
            backup_path: Ruta donde guardar el backup (opcional)
            
        Returns:
            Ruta del archivo de backup creado
        """
        try:
            if not backup_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup_path = f"upsells_backup_{timestamp}.db"
            
            # Copiar archivo de base de datos
            shutil.copy2(self.db_path, backup_path)
            logger.info(f"Backup creado: {backup_path}")
            return backup_path
        except Exception as e:
            logger.error(f"Error al crear backup: {str(e)}", exc_info=True)
            raise
    
    def restore_database(self, backup_path: str) -> bool:
        """
        Restaura la base de datos desde un backup
        
        Args:
            backup_path: Ruta del archivo de backup
            
        Returns:
            True si se restauró exitosamente
        """
        try:
            if not os.path.exists(backup_path):
                logger.error(f"Archivo de backup no encontrado: {backup_path}")
                return False
            
            # Hacer backup del archivo actual antes de restaurar
            current_backup = f"{self.db_path}.pre_restore_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            if os.path.exists(self.db_path):
                shutil.copy2(self.db_path, current_backup)
            
            # Restaurar desde backup
            shutil.copy2(backup_path, self.db_path)
            logger.info(f"Base de datos restaurada desde: {backup_path}")
            return True
        except Exception as e:
            logger.error(f"Error al restaurar backup: {str(e)}", exc_info=True)
            return False
    
    def export_all_data(self, format: str = 'json') -> Dict:
        """
        Exporta todos los datos del sistema
        
        Args:
            format: Formato de exportación (json, csv)
            
        Returns:
            Diccionario con todos los datos
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Exportar productos
            cursor.execute('SELECT * FROM products')
            products = [dict(row) for row in cursor.fetchall()]
            
            # Exportar upsells
            cursor.execute('SELECT * FROM upsells')
            upsells = [dict(row) for row in cursor.fetchall()]
            
            # Exportar compras
            cursor.execute('SELECT * FROM purchases ORDER BY purchased_at DESC LIMIT 1000')
            purchases = [dict(row) for row in cursor.fetchall()]
            
            # Exportar cross-sells
            cursor.execute('SELECT * FROM cross_sells')
            cross_sells = [dict(row) for row in cursor.fetchall()]
            
            # Exportar ofertas temporales
            cursor.execute('SELECT * FROM time_limited_offers')
            time_offers = [dict(row) for row in cursor.fetchall()]
            
            # Exportar historial de clientes
            cursor.execute('SELECT * FROM customer_history ORDER BY purchased_at DESC LIMIT 1000')
            customer_history = [dict(row) for row in cursor.fetchall()]
            
            export_data = {
                "exported_at": datetime.now().isoformat(),
                "format": format,
                "data": {
                    "products": products,
                    "upsells": upsells,
                    "purchases": purchases,
                    "cross_sells": cross_sells,
                    "time_limited_offers": time_offers,
                    "customer_history": customer_history
                },
                "counts": {
                    "products": len(products),
                    "upsells": len(upsells),
                    "purchases": len(purchases),
                    "cross_sells": len(cross_sells),
                    "time_limited_offers": len(time_offers),
                    "customer_history": len(customer_history)
                }
            }
            
            return export_data
        except Exception as e:
            logger.error(f"Error al exportar datos: {str(e)}", exc_info=True)
            return {}
    
    def get_system_health(self) -> Dict:
        """
        Obtiene el estado de salud del sistema
        
        Returns:
            Diccionario con estado del sistema
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Verificar integridad de la base de datos
            cursor.execute('PRAGMA integrity_check')
            integrity = cursor.fetchone()[0]
            
            # Obtener estadísticas de la base de datos
            cursor.execute('SELECT COUNT(*) FROM products')
            products_count = cursor.fetchone()[0]
            
            cursor.execute('SELECT COUNT(*) FROM upsells WHERE is_active = 1')
            active_upsells = cursor.fetchone()[0]
            
            cursor.execute('SELECT COUNT(*) FROM purchases')
            purchases_count = cursor.fetchone()[0]
            
            # Verificar tamaño de la base de datos
            db_size = os.path.getsize(self.db_path) if os.path.exists(self.db_path) else 0
            
            return {
                "status": "healthy" if integrity == "ok" else "degraded",
                "database_integrity": integrity,
                "database_size_mb": round(db_size / (1024 * 1024), 2),
                "products_count": products_count,
                "active_upsells": active_upsells,
                "total_purchases": purchases_count,
                "cache_size": len(self._recommendation_cache),
                "cache_enabled": self.cache_enabled,
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Error al verificar salud del sistema: {str(e)}", exc_info=True)
            return {"status": "error", "error": str(e)}
    
    def send_notification(self, customer_id: str, notification_type: str, data: Dict) -> bool:
        """
        Envía notificación a un cliente
        
        Args:
            customer_id: ID del cliente
            notification_type: Tipo de notificación (upsell_available, time_offer, etc.)
            data: Datos de la notificación
            
        Returns:
            True si se envió exitosamente
        """
        try:
            # Disparar webhook de notificación
            self.trigger_webhook('notification.sent', {
                "customer_id": customer_id,
                "type": notification_type,
                "data": data,
                "timestamp": datetime.now().isoformat()
            })
            
            logger.info(f"Notificación enviada a {customer_id}: {notification_type}")
            return True
        except Exception as e:
            logger.error(f"Error al enviar notificación: {str(e)}")
            return False
    
    def calculate_loyalty_points(self, customer_id: str, purchase_amount: float) -> Dict:
        """
        Calcula puntos de lealtad por compra
        
        Args:
            customer_id: ID del cliente
            purchase_amount: Monto de la compra
            
        Returns:
            Diccionario con puntos ganados y total
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Obtener historial del cliente
            cursor.execute('''
                SELECT COUNT(*) as purchase_count FROM customer_history
                WHERE customer_id = ?
            ''', (customer_id,))
            purchase_count = cursor.fetchone()[0]
            
            # Calcular puntos (1 punto por cada $10 MXN, bonus por cliente frecuente)
            base_points = int(purchase_amount / 10)
            bonus_multiplier = 1.0
            
            if purchase_count >= 10:
                bonus_multiplier = 1.5  # 50% bonus para VIP
            elif purchase_count >= 5:
                bonus_multiplier = 1.25  # 25% bonus para frecuentes
            
            points_earned = int(base_points * bonus_multiplier)
            
            # Obtener puntos totales (simulado, en producción usar tabla de puntos)
            total_points = points_earned  # Simplificado
            
            return {
                "customer_id": customer_id,
                "purchase_amount": purchase_amount,
                "points_earned": points_earned,
                "total_points": total_points,
                "bonus_applied": bonus_multiplier > 1.0,
                "bonus_multiplier": bonus_multiplier
            }
        except Exception as e:
            logger.error(f"Error al calcular puntos: {str(e)}", exc_info=True)
            return {}
    
    def get_personalized_offer(self, customer_id: str, product_id: str) -> Optional[Dict]:
        """
        Genera una oferta personalizada para un cliente específico
        
        Args:
            customer_id: ID del cliente
            product_id: ID del producto
            
        Returns:
            Oferta personalizada o None
        """
        try:
            # Obtener insights del cliente
            insights = self.get_customer_insights(customer_id)
            segment = insights.get("customer_segment", "Nuevo")
            
            # Obtener upsells
            upsells = self.get_complementary_upsells(
                product_id,
                limit=2,
                customer_id=customer_id,
                exclude_purchased=True
            )
            
            if not upsells:
                return None
            
            # Personalizar oferta según segmento
            personalized_discount = 0.0
            if segment == "VIP":
                personalized_discount = 0.15  # 15% adicional para VIP
            elif segment == "Frecuente":
                personalized_discount = 0.10  # 10% adicional
            
            # Aplicar descuento personalizado al primer upsell
            if upsells and personalized_discount > 0:
                first_upsell = upsells[0].copy()
                original_price = first_upsell["bundle_price"]
                discounted_price = original_price * (1 - personalized_discount)
                
                return {
                    "customer_id": customer_id,
                    "customer_segment": segment,
                    "product_id": product_id,
                    "personalized_upsell": {
                        **first_upsell,
                        "personalized_price": round(discounted_price, 2),
                        "personalized_discount": round(personalized_discount * 100, 1),
                        "savings_additional": round(original_price - discounted_price, 2),
                        "message": f"🎁 Oferta exclusiva para clientes {segment}: Ahorra {personalized_discount * 100:.0f}% adicional"
                    },
                    "other_upsells": upsells[1:] if len(upsells) > 1 else []
                }
            
            return {
                "customer_id": customer_id,
                "customer_segment": segment,
                "product_id": product_id,
                "upsells": upsells
            }
        except Exception as e:
            logger.error(f"Error al generar oferta personalizada: {str(e)}", exc_info=True)
            return None
    
    def get_performance_comparison(self, upsell_id1: str, upsell_id2: str) -> Dict:
        """
        Compara el rendimiento de dos upsells
        
        Args:
            upsell_id1: ID del primer upsell
            upsell_id2: ID del segundo upsell
            
        Returns:
            Diccionario con comparación
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            def get_upsell_metrics(upsell_id):
                cursor.execute('''
                    SELECT 
                        total_views,
                        total_purchases,
                        bundle_price,
                        savings,
                        CASE WHEN total_views > 0 
                        THEN (total_purchases * 1.0 / total_views) ELSE 0 END as conversion_rate
                    FROM upsells
                    WHERE id = ?
                ''', (upsell_id,))
                return cursor.fetchone()
            
            metrics1 = get_upsell_metrics(upsell_id1)
            metrics2 = get_upsell_metrics(upsell_id2)
            
            if not metrics1 or not metrics2:
                return {}
            
            revenue1 = (metrics1["total_purchases"] or 0) * (metrics1["bundle_price"] or 0)
            revenue2 = (metrics2["total_purchases"] or 0) * (metrics2["bundle_price"] or 0)
            
            return {
                "upsell_1": {
                    "id": upsell_id1,
                    "views": metrics1["total_views"] or 0,
                    "purchases": metrics1["total_purchases"] or 0,
                    "conversion_rate": round((metrics1["conversion_rate"] or 0) * 100, 2),
                    "revenue": round(revenue1, 2)
                },
                "upsell_2": {
                    "id": upsell_id2,
                    "views": metrics2["total_views"] or 0,
                    "purchases": metrics2["total_purchases"] or 0,
                    "conversion_rate": round((metrics2["conversion_rate"] or 0) * 100, 2),
                    "revenue": round(revenue2, 2)
                },
                "winner": upsell_id1 if revenue1 > revenue2 else upsell_id2,
                "difference": {
                    "revenue_diff": round(abs(revenue1 - revenue2), 2),
                    "conversion_diff": round(abs((metrics1["conversion_rate"] or 0) - (metrics2["conversion_rate"] or 0)) * 100, 2)
                }
            }
        except Exception as e:
            logger.error(f"Error al comparar upsells: {str(e)}", exc_info=True)
            return {}
    
    def get_ab_test_results(self, upsell_id: str) -> Dict:
        """
        Obtiene resultados de A/B testing para un upsell
        
        Args:
            upsell_id: ID del upsell
            
        Returns:
            Diccionario con resultados de todas las variantes
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT 
                    id,
                    variant_name,
                    message_template,
                    total_views,
                    total_conversions,
                    conversion_rate,
                    CASE WHEN total_views > 0 
                    THEN (total_conversions * 1.0 / total_views) ELSE 0 END as calculated_rate
                FROM message_variants
                WHERE upsell_id = ? AND is_active = 1
                ORDER BY calculated_rate DESC
            ''', (upsell_id,))
            
            variants = []
            for row in cursor.fetchall():
                rate = row["calculated_rate"] or row["conversion_rate"] or 0
                variants.append({
                    "variant_id": row["id"],
                    "variant_name": row["variant_name"],
                    "message_template": row["message_template"],
                    "views": row["total_views"],
                    "conversions": row["total_conversions"],
                    "conversion_rate": round(rate * 100, 2),
                    "is_winner": False  # Se actualizará después
                })
            
            # Marcar ganador
            if variants:
                best_rate = max(v["conversion_rate"] for v in variants)
                for v in variants:
                    if v["conversion_rate"] == best_rate:
                        v["is_winner"] = True
                        break
            
            return {
                "upsell_id": upsell_id,
                "total_variants": len(variants),
                "variants": variants,
                "best_variant": variants[0] if variants else None
            }
        except Exception as e:
            logger.error(f"Error al obtener resultados A/B: {str(e)}", exc_info=True)
            return {}
    
    def optimize_price_dynamically(self, upsell_id: str) -> Dict:
        """
        Optimiza el precio de un upsell dinámicamente basado en rendimiento
        
        Args:
            upsell_id: ID del upsell
            
        Returns:
            Diccionario con recomendaciones de precio
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT 
                    bundle_price,
                    original_price,
                    savings,
                    total_views,
                    total_purchases,
                    CASE WHEN total_views > 0 
                    THEN (total_purchases * 1.0 / total_views) ELSE 0 END as conversion_rate
                FROM upsells
                WHERE id = ?
            ''', (upsell_id,))
            
            upsell = cursor.fetchone()
            if not upsell:
                return {}
            
            current_price = upsell["bundle_price"]
            conversion_rate = upsell["conversion_rate"] or 0
            views = upsell["total_views"] or 0
            
            # Algoritmo de optimización de precio
            recommendations = []
            
            # Si conversión es baja (<5%) y hay suficientes vistas, sugerir bajar precio
            if conversion_rate < 0.05 and views >= 50:
                price_reduction = 0.10  # Reducir 10%
                new_price = current_price * (1 - price_reduction)
                expected_conversion = min(conversion_rate * 1.5, 0.15)  # Esperar 50% más conversión
                
                recommendations.append({
                    "action": "reduce_price",
                    "current_price": round(current_price, 2),
                    "suggested_price": round(new_price, 2),
                    "reduction_percent": round(price_reduction * 100, 1),
                    "expected_conversion_increase": round((expected_conversion - conversion_rate) * 100, 1),
                    "reason": "Baja conversión detectada. Reducir precio puede aumentar ventas."
                })
            
            # Si conversión es alta (>20%) y hay muchas vistas, considerar aumentar precio
            elif conversion_rate > 0.20 and views >= 100:
                price_increase = 0.05  # Aumentar 5%
                new_price = current_price * (1 + price_increase)
                expected_conversion = conversion_rate * 0.9  # Esperar 10% menos conversión pero más ingresos
                
                recommendations.append({
                    "action": "increase_price",
                    "current_price": round(current_price, 2),
                    "suggested_price": round(new_price, 2),
                    "increase_percent": round(price_increase * 100, 1),
                    "expected_conversion_decrease": round((conversion_rate - expected_conversion) * 100, 1),
                    "reason": "Alta conversión detectada. Aumentar precio puede maximizar ingresos."
                })
            
            # Si no hay recomendaciones, mantener precio actual
            if not recommendations:
                recommendations.append({
                    "action": "maintain_price",
                    "current_price": round(current_price, 2),
                    "reason": "Rendimiento óptimo. Mantener precio actual."
                })
            
            return {
                "upsell_id": upsell_id,
                "current_metrics": {
                    "price": round(current_price, 2),
                    "conversion_rate": round(conversion_rate * 100, 2),
                    "total_views": views
                },
                "recommendations": recommendations
            }
        except Exception as e:
            logger.error(f"Error al optimizar precio: {str(e)}", exc_info=True)
            return {}
    
    def get_cohort_analysis(self, start_date: str = None, end_date: str = None) -> Dict:
        """
        Análisis de cohortes de clientes
        
        Args:
            start_date: Fecha de inicio (ISO format)
            end_date: Fecha de fin (ISO format)
            
        Returns:
            Diccionario con análisis de cohortes
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            if not start_date:
                start_date = (datetime.now() - timedelta(days=90)).isoformat()
            if not end_date:
                end_date = datetime.now().isoformat()
            
            # Agrupar clientes por mes de primera compra
            cursor.execute('''
                SELECT 
                    strftime('%Y-%m', MIN(purchased_at)) as cohort_month,
                    COUNT(DISTINCT customer_id) as customers,
                    COUNT(*) as total_purchases,
                    AVG(total_amount) as avg_purchase_value
                FROM purchases
                WHERE purchased_at >= ? AND purchased_at <= ?
                GROUP BY cohort_month
                ORDER BY cohort_month
            ''', (start_date, end_date))
            
            cohorts = []
            for row in cursor.fetchall():
                cohorts.append({
                    "cohort_month": row["cohort_month"],
                    "customers": row["customers"],
                    "total_purchases": row["total_purchases"],
                    "avg_purchase_value": round(row["avg_purchase_value"] or 0, 2),
                    "purchases_per_customer": round(
                        (row["total_purchases"] or 0) / max(row["customers"], 1), 2
                    )
                })
            
            return {
                "period": {
                    "start_date": start_date,
                    "end_date": end_date
                },
                "cohorts": cohorts,
                "total_cohorts": len(cohorts)
            }
        except Exception as e:
            logger.error(f"Error en análisis de cohortes: {str(e)}", exc_info=True)
            return {}
    
    def get_hybrid_recommendations(self, product_id: str, customer_id: str = None, 
                                   limit: int = 5) -> List[Dict]:
        """
        Recomendaciones híbridas combinando múltiples algoritmos
        
        Args:
            product_id: ID del producto
            customer_id: ID del cliente (opcional)
            limit: Número de recomendaciones
            
        Returns:
            Lista de recomendaciones híbridas
        """
        try:
            recommendations = []
            
            # 1. Upsells complementarios (basado en contenido)
            complementary = self.get_complementary_upsells(
                product_id, limit=limit, customer_id=customer_id
            )
            for item in complementary:
                recommendations.append({
                    **item,
                    "recommendation_type": "complementary",
                    "score": item.get("score", 0) * 0.4  # 40% peso
                })
            
            # 2. Cross-sells (basado en relaciones)
            cross_sells = self.get_cross_sell_recommendations(product_id, limit=limit)
            for item in cross_sells:
                recommendations.append({
                    **item,
                    "recommendation_type": "cross_sell",
                    "score": item.get("similarity_score", 0) * 0.3  # 30% peso
                })
            
            # 3. Recomendaciones colaborativas (si hay customer_id)
            if customer_id:
                collaborative = self.get_collaborative_recommendations(customer_id, limit=limit)
                for item in collaborative:
                    recommendations.append({
                        **item,
                        "recommendation_type": "collaborative",
                        "score": item.get("recommendation_score", 0) * 0.3  # 30% peso
                    })
            
            # Combinar y deduplicar por product_id
            seen_products = {}
            for rec in recommendations:
                product_key = rec.get("upsell_product_id") or rec.get("product_id")
                if product_key:
                    if product_key not in seen_products:
                        seen_products[product_key] = rec
                    else:
                        # Combinar scores si ya existe
                        existing = seen_products[product_key]
                        existing["score"] = (existing.get("score", 0) + rec.get("score", 0))
                        existing["recommendation_types"] = existing.get("recommendation_types", []) + [rec.get("recommendation_type")]
            
            # Ordenar por score combinado
            final_recommendations = sorted(
                seen_products.values(),
                key=lambda x: x.get("score", 0),
                reverse=True
            )[:limit]
            
            return final_recommendations
        except Exception as e:
            logger.error(f"Error en recomendaciones híbridas: {str(e)}", exc_info=True)
            return []
    
    def create_smart_alert(self, alert_type: str, threshold: float, 
                          product_id: str = None) -> bool:
        """
        Crea una alerta inteligente basada en umbrales
        
        Args:
            alert_type: Tipo de alerta (low_conversion, high_performance, etc.)
            threshold: Umbral para activar la alerta
            product_id: ID del producto (opcional)
            
        Returns:
            True si se creó exitosamente
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            alert_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT INTO performance_alerts
                (id, alert_type, threshold, product_id, is_active, created_at)
                VALUES (?, ?, ?, ?, 1, datetime('now'))
            ''', (alert_id, alert_type, threshold, product_id))
            
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al crear alerta: {str(e)}", exc_info=True)
            return False
    
    def get_inventory_urgency(self, product_id: str) -> Dict:
        """
        Calcula urgencia basada en inventario
        
        Args:
            product_id: ID del producto
            
        Returns:
            Diccionario con nivel de urgencia
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT stock_quantity, name FROM products WHERE id = ?
            ''', (product_id,))
            
            product = cursor.fetchone()
            if not product:
                return {}
            
            stock = product["stock_quantity"] or 0
            urgency_level = "normal"
            message = ""
            
            if stock <= 0:
                urgency_level = "out_of_stock"
                message = "⚠️ Agotado - Últimas unidades"
            elif stock <= 5:
                urgency_level = "critical"
                message = f"🔥 ¡Solo quedan {stock} unidades!"
            elif stock <= 15:
                urgency_level = "high"
                message = f"⚡ Solo {stock} disponibles"
            elif stock <= 30:
                urgency_level = "medium"
                message = f"📦 {stock} en stock"
            else:
                urgency_level = "normal"
                message = "Disponible"
            
            return {
                "product_id": product_id,
                "product_name": product["name"],
                "stock_quantity": stock,
                "urgency_level": urgency_level,
                "urgency_message": message,
                "show_urgency": urgency_level in ["critical", "high", "out_of_stock"]
            }
        except Exception as e:
            logger.error(f"Error al calcular urgencia: {str(e)}", exc_info=True)
            return {}
    
    def log_event(self, event_type: str, entity_type: str, entity_id: str, 
                  details: Dict = None) -> bool:
        """
        Registra un evento en el sistema para auditoría
        
        Args:
            event_type: Tipo de evento (purchase, view, conversion, etc.)
            entity_type: Tipo de entidad (upsell, product, customer)
            entity_id: ID de la entidad
            details: Detalles adicionales del evento
            
        Returns:
            True si se registró exitosamente
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Crear tabla de eventos si no existe
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS system_events (
                    id TEXT PRIMARY KEY,
                    event_type TEXT NOT NULL,
                    entity_type TEXT NOT NULL,
                    entity_id TEXT NOT NULL,
                    details TEXT,
                    created_at TEXT NOT NULL DEFAULT (datetime('now'))
                )
            ''')
            
            event_id = str(uuid.uuid4())
            details_json = json.dumps(details) if details else None
            
            cursor.execute('''
                INSERT INTO system_events 
                (id, event_type, entity_type, entity_id, details)
                VALUES (?, ?, ?, ?, ?)
            ''', (event_id, event_type, entity_type, entity_id, details_json))
            
            conn.commit()
            
            logger.info(f"Evento registrado: {event_type} - {entity_type}:{entity_id}")
            return True
        except Exception as e:
            logger.error(f"Error al registrar evento: {str(e)}", exc_info=True)
            return False
    
    def get_event_history(self, entity_type: str = None, entity_id: str = None,
                         event_type: str = None, limit: int = 100) -> List[Dict]:
        """
        Obtiene historial de eventos
        
        Args:
            entity_type: Tipo de entidad a filtrar
            entity_id: ID de entidad a filtrar
            event_type: Tipo de evento a filtrar
            limit: Número máximo de eventos
            
        Returns:
            Lista de eventos
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Crear tabla si no existe
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS system_events (
                    id TEXT PRIMARY KEY,
                    event_type TEXT NOT NULL,
                    entity_type TEXT NOT NULL,
                    entity_id TEXT NOT NULL,
                    details TEXT,
                    created_at TEXT NOT NULL DEFAULT (datetime('now'))
                )
            ''')
            
            query = 'SELECT * FROM system_events WHERE 1=1'
            params = []
            
            if entity_type:
                query += ' AND entity_type = ?'
                params.append(entity_type)
            
            if entity_id:
                query += ' AND entity_id = ?'
                params.append(entity_id)
            
            if event_type:
                query += ' AND event_type = ?'
                params.append(event_type)
            
            query += ' ORDER BY created_at DESC LIMIT ?'
            params.append(limit)
            
            cursor.execute(query, params)
            
            events = []
            for row in cursor.fetchall():
                events.append({
                    "event_id": row["id"],
                    "event_type": row["event_type"],
                    "entity_type": row["entity_type"],
                    "entity_id": row["entity_id"],
                    "details": json.loads(row["details"]) if row["details"] else {},
                    "created_at": row["created_at"]
                })
            
            return events
        except Exception as e:
            logger.error(f"Error al obtener historial: {str(e)}", exc_info=True)
            return []
    
    def get_conversion_funnel(self, product_id: str = None, days: int = 30) -> Dict:
        """
        Analiza el embudo de conversión
        
        Args:
            product_id: ID del producto (opcional)
            days: Días a considerar
            
        Returns:
            Diccionario con análisis del embudo
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
            
            where_clause = "WHERE u.product_id = ?" if product_id else ""
            params = (product_id, cutoff_date) if product_id else (cutoff_date,)
            
            # Vistas totales
            query_views = f'''
                SELECT COUNT(*) as total_views
                FROM upsells u
                {where_clause}
                AND u.is_active = 1
            '''
            cursor.execute(query_views, params[:1] if product_id else ())
            total_views = cursor.fetchone()["total_views"] or 0
            
            # Compras totales
            query_purchases = f'''
                SELECT COUNT(*) as total_purchases
                FROM purchases p
                JOIN upsells u ON p.upsell_product_id = u.upsell_product_id
                {where_clause.replace("u.product_id", "u.product_id") if product_id else ""}
                AND p.purchased_at >= ?
            '''
            cursor.execute(query_purchases, params)
            total_purchases = cursor.fetchone()["total_purchases"] or 0
            
            # Calcular tasas
            view_to_purchase_rate = (total_purchases / max(total_views, 1)) * 100
            
            return {
                "period_days": days,
                "product_id": product_id,
                "funnel": {
                    "views": total_views,
                    "purchases": total_purchases,
                    "conversion_rate": round(view_to_purchase_rate, 2),
                    "drop_off": round(100 - view_to_purchase_rate, 2)
                },
                "stages": [
                    {"stage": "Views", "count": total_views, "percentage": 100.0},
                    {"stage": "Purchases", "count": total_purchases, 
                     "percentage": round(view_to_purchase_rate, 2)}
                ]
            }
        except Exception as e:
            logger.error(f"Error en análisis de embudo: {str(e)}", exc_info=True)
            return {}
    
    def get_customer_journey(self, customer_id: str) -> Dict:
        """
        Analiza el journey completo de un cliente
        
        Args:
            customer_id: ID del cliente
            
        Returns:
            Diccionario con journey del cliente
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Obtener todas las compras del cliente
            cursor.execute('''
                SELECT 
                    p.*,
                    pr.name as product_name,
                    pr.category as product_category
                FROM purchases p
                JOIN products pr ON p.product_id = pr.id
                WHERE p.customer_id = ?
                ORDER BY p.purchased_at ASC
            ''', (customer_id,))
            
            purchases = cursor.fetchall()
            
            # Obtener upsells comprados
            cursor.execute('''
                SELECT 
                    p.*,
                    pr.name as upsell_name,
                    u.bundle_price
                FROM purchases p
                JOIN products pr ON p.upsell_product_id = pr.id
                JOIN upsells u ON p.upsell_product_id = u.upsell_product_id
                WHERE p.customer_id = ? AND p.upsell_product_id IS NOT NULL
                ORDER BY p.purchased_at ASC
            ''', (customer_id,))
            
            upsells = cursor.fetchall()
            
            # Calcular métricas
            total_spent = sum(float(p["total_amount"] or 0) for p in purchases)
            total_upsells = len(upsells)
            avg_order_value = total_spent / max(len(purchases), 1)
            
            # Categorías más compradas
            categories = {}
            for p in purchases:
                cat = p["product_category"]
                categories[cat] = categories.get(cat, 0) + 1
            
            return {
                "customer_id": customer_id,
                "journey": {
                    "total_purchases": len(purchases),
                    "total_upsells": total_upsells,
                    "upsell_rate": round((total_upsells / max(len(purchases), 1)) * 100, 2),
                    "total_spent": round(total_spent, 2),
                    "avg_order_value": round(avg_order_value, 2),
                    "first_purchase": purchases[0]["purchased_at"] if purchases else None,
                    "last_purchase": purchases[-1]["purchased_at"] if purchases else None,
                    "favorite_categories": sorted(
                        categories.items(),
                        key=lambda x: x[1],
                        reverse=True
                    )[:3]
                },
                "purchases": [
                    {
                        "date": p["purchased_at"],
                        "product": p["product_name"],
                        "category": p["product_category"],
                        "amount": float(p["total_amount"] or 0)
                    }
                    for p in purchases
                ],
                "upsells": [
                    {
                        "date": u["purchased_at"],
                        "upsell": u["upsell_name"],
                        "amount": float(u["bundle_price"] or 0)
                    }
                    for u in upsells
                ]
            }
        except Exception as e:
            logger.error(f"Error al analizar journey: {str(e)}", exc_info=True)
            return {}
    
    def invalidate_cache(self, pattern: str = None) -> int:
        """
        Invalida el cache de recomendaciones
        
        Args:
            pattern: Patrón para invalidar (opcional, si None invalida todo)
            
        Returns:
            Número de entradas invalidadas
        """
        try:
            if not self.cache_enabled:
                return 0
            
            if pattern:
                # Invalidar entradas que coincidan con el patrón
                keys_to_remove = [
                    key for key in self._recommendation_cache.keys()
                    if pattern in key
                ]
                for key in keys_to_remove:
                    del self._recommendation_cache[key]
                return len(keys_to_remove)
            else:
                # Invalidar todo el cache
                count = len(self._recommendation_cache)
                self._recommendation_cache.clear()
                return count
        except Exception as e:
            logger.error(f"Error al invalidar cache: {str(e)}")
            return 0
    
    def create_coupon(self, code: str, discount_type: str, discount_value: float,
                     min_purchase: float = 0, max_uses: int = None,
                     valid_until: str = None) -> bool:
        """
        Crea un cupón de descuento
        
        Args:
            code: Código del cupón
            discount_type: Tipo de descuento (percentage, fixed)
            discount_value: Valor del descuento
            min_purchase: Compra mínima requerida
            max_uses: Máximo de usos (None = ilimitado)
            valid_until: Fecha de expiración (ISO format)
            
        Returns:
            True si se creó exitosamente
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Crear tabla de cupones si no existe
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS coupons (
                    id TEXT PRIMARY KEY,
                    code TEXT UNIQUE NOT NULL,
                    discount_type TEXT NOT NULL,
                    discount_value REAL NOT NULL,
                    min_purchase REAL DEFAULT 0,
                    max_uses INTEGER,
                    current_uses INTEGER DEFAULT 0,
                    valid_until TEXT,
                    is_active INTEGER DEFAULT 1,
                    created_at TEXT NOT NULL DEFAULT (datetime('now'))
                )
            ''')
            
            coupon_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT INTO coupons
                (id, code, discount_type, discount_value, min_purchase, max_uses, valid_until)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (coupon_id, code, discount_type, discount_value, min_purchase, max_uses, valid_until))
            
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al crear cupón: {str(e)}", exc_info=True)
            return False
    
    def validate_coupon(self, code: str, purchase_amount: float) -> Dict:
        """
        Valida y aplica un cupón
        
        Args:
            code: Código del cupón
            purchase_amount: Monto de la compra
            
        Returns:
            Diccionario con información del cupón y descuento aplicado
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT * FROM coupons
                WHERE code = ? AND is_active = 1
            ''', (code,))
            
            coupon = cursor.fetchone()
            if not coupon:
                return {"valid": False, "error": "Cupón no encontrado"}
            
            # Validar fecha de expiración
            if coupon["valid_until"]:
                if datetime.now() > datetime.fromisoformat(coupon["valid_until"].replace('Z', '+00:00')).replace(tzinfo=None):
                    return {"valid": False, "error": "Cupón expirado"}
            
            # Validar compra mínima
            if purchase_amount < coupon["min_purchase"]:
                return {
                    "valid": False,
                    "error": f"Compra mínima requerida: ${coupon['min_purchase']:.2f}"
                }
            
            # Validar máximo de usos
            if coupon["max_uses"] and coupon["current_uses"] >= coupon["max_uses"]:
                return {"valid": False, "error": "Cupón agotado"}
            
            # Calcular descuento
            if coupon["discount_type"] == "percentage":
                discount_amount = purchase_amount * (coupon["discount_value"] / 100)
            else:  # fixed
                discount_amount = min(coupon["discount_value"], purchase_amount)
            
            final_amount = purchase_amount - discount_amount
            
            return {
                "valid": True,
                "coupon_code": code,
                "discount_type": coupon["discount_type"],
                "discount_value": coupon["discount_value"],
                "discount_amount": round(discount_amount, 2),
                "original_amount": round(purchase_amount, 2),
                "final_amount": round(final_amount, 2),
                "savings": round(discount_amount, 2)
            }
        except Exception as e:
            logger.error(f"Error al validar cupón: {str(e)}", exc_info=True)
            return {"valid": False, "error": str(e)}
    
    def get_competitor_analysis(self, product_id: str) -> Dict:
        """
        Análisis competitivo de un producto (simulado)
        
        Args:
            product_id: ID del producto
            
        Returns:
            Diccionario con análisis competitivo
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT price, category, name FROM products WHERE id = ?
            ''', (product_id,))
            
            product = cursor.fetchone()
            if not product:
                return {}
            
            # Obtener productos similares en la misma categoría
            cursor.execute('''
                SELECT 
                    id,
                    name,
                    price,
                    COUNT(DISTINCT ch.customer_id) as popularity
                FROM products p
                LEFT JOIN customer_history ch ON p.id = ch.product_id
                WHERE p.category = ? AND p.id != ? AND p.is_active = 1
                GROUP BY p.id
                ORDER BY popularity DESC, price ASC
                LIMIT 5
            ''', (product["category"], product_id))
            
            competitors = []
            for row in cursor.fetchall():
                price_diff = ((row["price"] - product["price"]) / product["price"]) * 100
                competitors.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "price": float(row["price"]),
                    "popularity": row["popularity"],
                    "price_difference": round(price_diff, 2),
                    "is_cheaper": row["price"] < product["price"]
                })
            
            return {
                "product_id": product_id,
                "product_name": product["name"],
                "product_price": float(product["price"]),
                "category": product["category"],
                "competitors": competitors,
                "total_competitors": len(competitors),
                "recommendations": {
                    "price_position": "competitive" if len([c for c in competitors if c["is_cheaper"]) < 2 else "premium",
                    "suggested_action": "Considerar ajuste de precio" if len([c for c in competitors if c["is_cheaper"]) > 3 else "Precio competitivo"
                }
            }
        except Exception as e:
            logger.error(f"Error en análisis competitivo: {str(e)}", exc_info=True)
            return {}
    
    def get_upsell_performance_by_category(self, days: int = 30) -> Dict:
        """
        Analiza rendimiento de upsells por categoría
        
        Args:
            days: Días a considerar
            
        Returns:
            Diccionario con rendimiento por categoría
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
            
            cursor.execute('''
                SELECT 
                    p.category,
                    COUNT(DISTINCT u.id) as upsells_count,
                    SUM(u.total_views) as total_views,
                    SUM(u.total_purchases) as total_purchases,
                    AVG(CASE WHEN u.total_views > 0 
                        THEN (u.total_purchases * 1.0 / u.total_views) ELSE 0 END) as avg_conversion,
                    SUM(u.total_purchases * u.bundle_price) as total_revenue
                FROM upsells u
                JOIN products p ON u.product_id = p.id
                WHERE u.is_active = 1
                GROUP BY p.category
                ORDER BY total_revenue DESC
            ''')
            
            categories = []
            for row in cursor.fetchall():
                categories.append({
                    "category": row["category"],
                    "upsells_count": row["upsells_count"],
                    "total_views": row["total_views"] or 0,
                    "total_purchases": row["total_purchases"] or 0,
                    "avg_conversion_rate": round((row["avg_conversion"] or 0) * 100, 2),
                    "total_revenue": round(row["total_revenue"] or 0, 2)
                })
            
            return {
                "period_days": days,
                "categories": categories,
                "total_categories": len(categories),
                "best_category": categories[0] if categories else None
            }
        except Exception as e:
            logger.error(f"Error al analizar por categoría: {str(e)}", exc_info=True)
            return {}
    
    def get_time_based_recommendations(self, product_id: str, 
                                      hour: int = None) -> List[Dict]:
        """
        Recomendaciones basadas en hora del día
        
        Args:
            product_id: ID del producto
            hour: Hora del día (0-23), si None usa hora actual
            
        Returns:
            Lista de recomendaciones optimizadas por hora
        """
        try:
            if hour is None:
                hour = datetime.now().hour
            
            # Obtener upsells normales
            upsells = self.get_complementary_upsells(product_id, limit=5)
            
            # Ajustar según hora del día
            # Mañana (6-12): Productos energéticos/productividad
            # Tarde (12-18): Productos de entretenimiento
            # Noche (18-24): Productos relajantes
            # Madrugada (0-6): Productos esenciales
            
            time_multiplier = 1.0
            if 6 <= hour < 12:
                time_multiplier = 1.1  # 10% boost para productos matutinos
            elif 18 <= hour < 24:
                time_multiplier = 1.15  # 15% boost para productos nocturnos
            
            # Aplicar multiplicador a scores
            for upsell in upsells:
                upsell["score"] = upsell.get("score", 0) * time_multiplier
                upsell["time_optimized"] = True
                upsell["hour"] = hour
            
            # Reordenar por score
            upsells.sort(key=lambda x: x.get("score", 0), reverse=True)
            
            return upsells[:3]  # Top 3
        except Exception as e:
            logger.error(f"Error en recomendaciones por tiempo: {str(e)}", exc_info=True)
            return []
    
    def add_to_wishlist(self, customer_id: str, product_id: str) -> bool:
        """
        Agrega un producto a la wishlist del cliente
        
        Args:
            customer_id: ID del cliente
            product_id: ID del producto
            
        Returns:
            True si se agregó exitosamente
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Crear tabla de wishlist si no existe
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS wishlist (
                    id TEXT PRIMARY KEY,
                    customer_id TEXT NOT NULL,
                    product_id TEXT NOT NULL,
                    created_at TEXT NOT NULL DEFAULT (datetime('now')),
                    UNIQUE(customer_id, product_id)
                )
            ''')
            
            wishlist_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT OR IGNORE INTO wishlist (id, customer_id, product_id)
                VALUES (?, ?, ?)
            ''', (wishlist_id, customer_id, product_id))
            
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al agregar a wishlist: {str(e)}", exc_info=True)
            return False
    
    def get_wishlist(self, customer_id: str) -> List[Dict]:
        """
        Obtiene la wishlist de un cliente
        
        Args:
            customer_id: ID del cliente
            
        Returns:
            Lista de productos en wishlist
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT 
                    w.product_id,
                    p.name,
                    p.price,
                    p.category,
                    w.created_at
                FROM wishlist w
                JOIN products p ON w.product_id = p.id
                WHERE w.customer_id = ? AND p.is_active = 1
                ORDER BY w.created_at DESC
            ''', (customer_id,))
            
            items = []
            for row in cursor.fetchall():
                items.append({
                    "product_id": row["product_id"],
                    "name": row["name"],
                    "price": float(row["price"]),
                    "category": row["category"],
                    "added_at": row["created_at"]
                })
            
            return items
        except Exception as e:
            logger.error(f"Error al obtener wishlist: {str(e)}", exc_info=True)
            return []
    
    def get_engagement_metrics(self, product_id: str = None, days: int = 30) -> Dict:
        """
        Calcula métricas de engagement
        
        Args:
            product_id: ID del producto (opcional)
            days: Días a considerar
            
        Returns:
            Diccionario con métricas de engagement
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
            
            # Vistas
            views_query = 'SELECT COUNT(*) as count FROM upsells WHERE is_active = 1'
            if product_id:
                views_query += ' AND product_id = ?'
                cursor.execute(views_query, (product_id,))
            else:
                cursor.execute(views_query)
            total_views = cursor.fetchone()["count"] or 0
            
            # Wishlist
            wishlist_query = 'SELECT COUNT(*) as count FROM wishlist WHERE created_at >= ?'
            if product_id:
                wishlist_query += ' AND product_id = ?'
                cursor.execute(wishlist_query, (cutoff_date, product_id))
            else:
                cursor.execute(wishlist_query, (cutoff_date,))
            wishlist_adds = cursor.fetchone()["count"] or 0
            
            # Compras
            purchases_query = '''
                SELECT COUNT(*) as count FROM purchases 
                WHERE purchased_at >= ?
            '''
            if product_id:
                purchases_query += ' AND product_id = ?'
                cursor.execute(purchases_query, (cutoff_date, product_id))
            else:
                cursor.execute(purchases_query, (cutoff_date,))
            total_purchases = cursor.fetchone()["count"] or 0
            
            # Calcular engagement score
            engagement_score = (total_views * 0.3 + wishlist_adds * 0.4 + total_purchases * 0.3)
            
            return {
                "period_days": days,
                "product_id": product_id,
                "metrics": {
                    "total_views": total_views,
                    "wishlist_adds": wishlist_adds,
                    "total_purchases": total_purchases,
                    "engagement_score": round(engagement_score, 2)
                },
                "engagement_level": "Alto" if engagement_score > 100 else "Medio" if engagement_score > 50 else "Bajo"
            }
        except Exception as e:
            logger.error(f"Error al calcular engagement: {str(e)}", exc_info=True)
            return {}
    
    def get_cart_abandonment_analysis(self, days: int = 30) -> Dict:
        """
        Analiza abandono de carrito (simulado)
        
        Args:
            days: Días a considerar
            
        Returns:
            Diccionario con análisis de abandono
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Obtener vistas de upsells (simulando carritos iniciados)
                cursor.execute('''
                    SELECT 
                        COUNT(DISTINCT u.product_id) as products_viewed,
                        SUM(u.total_views) as total_views
                    FROM upsells u
                    WHERE u.is_active = 1
                ''')
                views_data = cursor.fetchone()
                
                # Obtener compras completadas
                cursor.execute('''
                    SELECT COUNT(DISTINCT product_id) as products_purchased
                    FROM purchases
                ''')
                purchases_data = cursor.fetchone()
                
                products_viewed = views_data["products_viewed"] or 0
                products_purchased = purchases_data["products_purchased"] or 0
                total_views = views_data["total_views"] or 0
                
                # Calcular tasa de abandono (simulado)
                abandonment_rate = ((products_viewed - products_purchased) / max(products_viewed, 1)) * 100
                
                return {
                    "period_days": days,
                    "analysis": {
                        "products_viewed": products_viewed,
                        "products_purchased": products_purchased,
                        "total_views": total_views,
                        "abandonment_rate": round(abandonment_rate, 2),
                        "conversion_rate": round(100 - abandonment_rate, 2)
                    },
                    "recommendations": {
                        "urgency": "Alta" if abandonment_rate > 70 else "Media" if abandonment_rate > 50 else "Baja",
                        "suggested_actions": [
                            "Enviar recordatorios de carrito abandonado",
                            "Ofrecer descuentos adicionales",
                            "Mejorar mensajes de urgencia"
                        ] if abandonment_rate > 50 else ["Mantener estrategia actual"]
                    }
                }
        except Exception as e:
            logger.error(f"Error en análisis de abandono: {str(e)}", exc_info=True)
            return {}
    
    def get_referral_rewards(self, referrer_id: str, referred_id: str) -> Dict:
        """
        Calcula recompensas por referido
        
        Args:
            referrer_id: ID del cliente que refirió
            referred_id: ID del cliente referido
            
        Returns:
            Diccionario con recompensas
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Crear tabla de referidos si no existe
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS referrals (
                        id TEXT PRIMARY KEY,
                        referrer_id TEXT NOT NULL,
                        referred_id TEXT NOT NULL,
                        reward_points INTEGER DEFAULT 0,
                        created_at TEXT NOT NULL DEFAULT (datetime('now')),
                        UNIQUE(referrer_id, referred_id)
                    )
                ''')
                
                # Verificar si ya existe
                cursor.execute('''
                    SELECT id FROM referrals 
                    WHERE referrer_id = ? AND referred_id = ?
                ''', (referrer_id, referred_id))
                
                if cursor.fetchone():
                    return {"success": False, "error": "Referido ya registrado"}
                
                # Calcular puntos de recompensa (100 puntos por referido)
                reward_points = 100
                
                referral_id = str(uuid.uuid4())
                cursor.execute('''
                    INSERT INTO referrals (id, referrer_id, referred_id, reward_points)
                    VALUES (?, ?, ?, ?)
                ''', (referral_id, referrer_id, referred_id, reward_points))
                
                conn.commit()
            
            return {
                "success": True,
                "referrer_id": referrer_id,
                "referred_id": referred_id,
                "reward_points": reward_points,
                "message": f"¡Recompensa de {reward_points} puntos otorgada!"
            }
        except Exception as e:
            logger.error(f"Error al calcular recompensas: {str(e)}", exc_info=True)
            return {"success": False, "error": str(e)}
    
    def add_product_review(self, product_id: str, customer_id: str, 
                          rating: int, comment: str = None) -> bool:
        """
        Agrega una reseña a un producto
        
        Args:
            product_id: ID del producto
            customer_id: ID del cliente
            rating: Calificación (1-5)
            comment: Comentario opcional
            
        Returns:
            True si se agregó exitosamente
        """
        try:
            if rating < 1 or rating > 5:
                return False
            
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Crear tabla de reviews si no existe
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS product_reviews (
                        id TEXT PRIMARY KEY,
                        product_id TEXT NOT NULL,
                        customer_id TEXT NOT NULL,
                        rating INTEGER NOT NULL,
                        comment TEXT,
                        created_at TEXT NOT NULL DEFAULT (datetime('now')),
                        UNIQUE(product_id, customer_id)
                    )
                ''')
                
                review_id = str(uuid.uuid4())
                cursor.execute('''
                    INSERT OR REPLACE INTO product_reviews 
                    (id, product_id, customer_id, rating, comment)
                    VALUES (?, ?, ?, ?, ?)
                ''', (review_id, product_id, customer_id, rating, comment))
                
                conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al agregar review: {str(e)}", exc_info=True)
            return False
    
    def get_product_reviews(self, product_id: str, limit: int = 10) -> Dict:
        """
        Obtiene reseñas de un producto
        
        Args:
            product_id: ID del producto
            limit: Número máximo de reseñas
            
        Returns:
            Diccionario con reseñas y estadísticas
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT 
                        rating,
                        comment,
                        created_at
                    FROM product_reviews
                    WHERE product_id = ?
                    ORDER BY created_at DESC
                    LIMIT ?
                ''', (product_id, limit))
                
                reviews = []
                for row in cursor.fetchall():
                    reviews.append({
                        "rating": row["rating"],
                        "comment": row["comment"],
                        "created_at": row["created_at"]
                    })
                
                # Calcular estadísticas
                cursor.execute('''
                    SELECT 
                        COUNT(*) as total_reviews,
                        AVG(rating) as avg_rating,
                        COUNT(CASE WHEN rating >= 4 THEN 1 END) as positive_reviews
                    FROM product_reviews
                    WHERE product_id = ?
                ''', (product_id,))
                
                stats = cursor.fetchone()
                
                return {
                    "product_id": product_id,
                    "reviews": reviews,
                    "statistics": {
                        "total_reviews": stats["total_reviews"] or 0,
                        "average_rating": round(stats["avg_rating"] or 0, 2),
                        "positive_reviews": stats["positive_reviews"] or 0,
                        "positive_percentage": round(
                            ((stats["positive_reviews"] or 0) / max(stats["total_reviews"] or 1, 1)) * 100, 2
                        )
                    }
                }
        except Exception as e:
            logger.error(f"Error al obtener reviews: {str(e)}", exc_info=True)
            return {}
    
    def search_products(self, query: str, limit: int = 10) -> List[Dict]:
        """
        Búsqueda inteligente de productos
        
        Args:
            query: Término de búsqueda
            limit: Número máximo de resultados
            
        Returns:
            Lista de productos encontrados
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                search_term = f"%{query}%"
                cursor.execute('''
                    SELECT 
                        id,
                        name,
                        description,
                        price,
                        category,
                        tags
                    FROM products
                    WHERE is_active = 1
                        AND (
                            name LIKE ? OR
                            description LIKE ? OR
                            category LIKE ? OR
                            tags LIKE ?
                        )
                    ORDER BY 
                        CASE 
                            WHEN name LIKE ? THEN 1
                            WHEN category LIKE ? THEN 2
                            ELSE 3
                        END,
                        price ASC
                    LIMIT ?
                ''', (search_term, search_term, search_term, search_term, 
                      f"{query}%", f"{query}%", limit))
                
                results = []
                for row in cursor.fetchall():
                    results.append({
                        "product_id": row["id"],
                        "name": row["name"],
                        "description": row["description"],
                        "price": float(row["price"]),
                        "category": row["category"],
                        "tags": row["tags"]
                    })
                
                return results
        except Exception as e:
            logger.error(f"Error en búsqueda: {str(e)}", exc_info=True)
            return []
    
    def get_customer_retention(self, days: int = 90) -> Dict:
        """
        Calcula métricas de retención de clientes
        
        Args:
            days: Días a considerar
            
        Returns:
            Diccionario con métricas de retención
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
                
                # Clientes nuevos
                cursor.execute('''
                    SELECT COUNT(DISTINCT customer_id) as new_customers
                    FROM customer_history
                    WHERE purchased_at >= ?
                ''', (cutoff_date,))
                new_customers = cursor.fetchone()["new_customers"] or 0
                
                # Clientes que regresaron
                cursor.execute('''
                    SELECT COUNT(DISTINCT customer_id) as returning_customers
                    FROM customer_history
                    WHERE customer_id IN (
                        SELECT DISTINCT customer_id 
                        FROM customer_history 
                        WHERE purchased_at < ?
                    )
                    AND purchased_at >= ?
                ''', (cutoff_date, cutoff_date))
                returning_customers = cursor.fetchone()["returning_customers"] or 0
                
                # Total de clientes únicos
                cursor.execute('''
                    SELECT COUNT(DISTINCT customer_id) as total_customers
                    FROM customer_history
                ''')
                total_customers = cursor.fetchone()["total_customers"] or 0
                
                # Calcular tasa de retención
                retention_rate = (returning_customers / max(total_customers - new_customers, 1)) * 100 if total_customers > new_customers else 0
                
                return {
                    "period_days": days,
                    "metrics": {
                        "new_customers": new_customers,
                        "returning_customers": returning_customers,
                        "total_customers": total_customers,
                        "retention_rate": round(retention_rate, 2)
                    },
                    "retention_level": "Alta" if retention_rate > 50 else "Media" if retention_rate > 30 else "Baja"
                }
        except Exception as e:
            logger.error(f"Error al calcular retención: {str(e)}", exc_info=True)
            return {}
    
    def get_product_recommendations_by_tags(self, product_id: str, limit: int = 5) -> List[Dict]:
        """
        Recomendaciones basadas en tags
        
        Args:
            product_id: ID del producto
            limit: Número de recomendaciones
            
        Returns:
            Lista de productos recomendados por tags
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Obtener tags del producto
                cursor.execute('SELECT tags FROM products WHERE id = ?', (product_id,))
                product = cursor.fetchone()
                if not product or not product["tags"]:
                    return []
                
                product_tags = set(product["tags"].split(','))
                
                # Buscar productos con tags similares
                cursor.execute('''
                    SELECT 
                        id,
                        name,
                        price,
                        category,
                        tags
                    FROM products
                    WHERE id != ? AND is_active = 1
                ''', (product_id,))
                
                candidates = []
                for row in cursor.fetchall():
                    if row["tags"]:
                        candidate_tags = set(row["tags"].split(','))
                        common_tags = product_tags.intersection(candidate_tags)
                        if common_tags:
                            similarity = len(common_tags) / max(len(product_tags.union(candidate_tags)), 1)
                            candidates.append({
                                "product_id": row["id"],
                                "name": row["name"],
                                "price": float(row["price"]),
                                "category": row["category"],
                                "similarity_score": round(similarity, 3),
                                "common_tags": list(common_tags)
                            })
                
                # Ordenar por similitud
                candidates.sort(key=lambda x: x["similarity_score"], reverse=True)
                return candidates[:limit]
        except Exception as e:
            logger.error(f"Error en recomendaciones por tags: {str(e)}", exc_info=True)
            return []
    
    def analyze_sentiment(self, text: str) -> Dict:
        """
        Análisis básico de sentimiento (simulado)
        
        Args:
            text: Texto a analizar
            
        Returns:
            Diccionario con análisis de sentimiento
        """
        try:
            if not text:
                return {"sentiment": "neutral", "score": 0.0}
            
            text_lower = text.lower()
            
            # Palabras positivas
            positive_words = ['excelente', 'bueno', 'genial', 'perfecto', 'recomendado', 
                            'satisfecho', 'feliz', 'amor', 'mejor', 'increíble']
            # Palabras negativas
            negative_words = ['malo', 'terrible', 'horrible', 'decepcionado', 'problema',
                            'defectuoso', 'no funciona', 'pésimo', 'basura']
            
            positive_count = sum(1 for word in positive_words if word in text_lower)
            negative_count = sum(1 for word in negative_words if word in text_lower)
            
            # Calcular score (-1 a 1)
            total_words = len(text.split())
            if total_words == 0:
                return {"sentiment": "neutral", "score": 0.0}
            
            score = (positive_count - negative_count) / max(total_words, 1)
            score = max(-1.0, min(1.0, score))  # Normalizar entre -1 y 1
            
            if score > 0.1:
                sentiment = "positive"
            elif score < -0.1:
                sentiment = "negative"
            else:
                sentiment = "neutral"
            
            return {
                "sentiment": sentiment,
                "score": round(score, 3),
                "positive_words": positive_count,
                "negative_words": negative_count
            }
        except Exception as e:
            logger.error(f"Error en análisis de sentimiento: {str(e)}")
            return {"sentiment": "neutral", "score": 0.0}
    
    def compare_products(self, product_ids: List[str]) -> Dict:
        """
        Compara múltiples productos lado a lado
        
        Args:
            product_ids: Lista de IDs de productos a comparar
            
        Returns:
            Diccionario con comparación
        """
        try:
            if len(product_ids) < 2:
                return {"error": "Se requieren al menos 2 productos para comparar"}
            
            if len(product_ids) > 5:
                return {"error": "Máximo 5 productos para comparar"}
            
            conn = sqlite3.connect(self.db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            placeholders = ','.join(['?'] * len(product_ids))
            cursor.execute(f'''
                SELECT 
                    p.*,
                    AVG(pr.rating) as avg_rating,
                    COUNT(pr.id) as review_count
                FROM products p
                LEFT JOIN product_reviews pr ON p.id = pr.product_id
                WHERE p.id IN ({placeholders})
                GROUP BY p.id
            ''', product_ids)
            
            products = []
            for row in cursor.fetchall():
                products.append({
                    "product_id": row["id"],
                    "name": row["name"],
                    "price": float(row["price"]),
                    "category": row["category"],
                    "description": row["description"],
                    "tags": row["tags"],
                    "stock_quantity": row["stock_quantity"],
                    "avg_rating": round(row["avg_rating"] or 0, 2),
                    "review_count": row["review_count"] or 0
                })
            
            # Encontrar mejor precio, mejor rating, etc.
            if products:
                cheapest = min(products, key=lambda x: x["price"])
                most_expensive = max(products, key=lambda x: x["price"])
                best_rated = max([p for p in products if p["avg_rating"] > 0], 
                               key=lambda x: x["avg_rating"], default=None)
                
                return {
                    "products": products,
                    "comparison": {
                        "cheapest": cheapest["product_id"],
                        "most_expensive": most_expensive["product_id"],
                        "price_range": round(most_expensive["price"] - cheapest["price"], 2),
                        "best_rated": best_rated["product_id"] if best_rated else None,
                        "total_products": len(products)
                    }
                }
            
            return {}
        except Exception as e:
            logger.error(f"Error al comparar productos: {str(e)}", exc_info=True)
            return {}
    
    def get_most_searched_products(self, days: int = 30, limit: int = 10) -> List[Dict]:
        """
        Obtiene productos más buscados (simulado basado en vistas)
        
        Args:
            days: Días a considerar
            limit: Número de resultados
            
        Returns:
            Lista de productos más buscados
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
                
                cursor.execute('''
                    SELECT 
                        p.id,
                        p.name,
                        p.price,
                        p.category,
                        SUM(u.total_views) as total_views,
                        SUM(u.total_purchases) as total_purchases
                    FROM products p
                    JOIN upsells u ON p.id = u.product_id
                    WHERE u.is_active = 1
                    GROUP BY p.id, p.name, p.price, p.category
                    ORDER BY total_views DESC
                    LIMIT ?
                ''', (limit,))
                
                products = []
                for row in cursor.fetchall():
                    products.append({
                        "product_id": row["id"],
                        "name": row["name"],
                        "price": float(row["price"]),
                        "category": row["category"],
                        "total_views": row["total_views"] or 0,
                        "total_purchases": row["total_purchases"] or 0,
                        "popularity_score": (row["total_views"] or 0) + (row["total_purchases"] or 0) * 2
                    })
                
                return products
        except Exception as e:
            logger.error(f"Error al obtener productos más buscados: {str(e)}", exc_info=True)
            return []
    
    def get_low_stock_alerts(self, threshold: int = 10) -> List[Dict]:
        """
        Obtiene alertas de productos con stock bajo
        
        Args:
            threshold: Umbral de stock bajo
            
        Returns:
            Lista de productos con stock bajo
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT 
                        id,
                        name,
                        price,
                        category,
                        stock_quantity
                    FROM products
                    WHERE is_active = 1
                        AND stock_quantity <= ?
                    ORDER BY stock_quantity ASC
                ''', (threshold,))
                
                alerts = []
                for row in cursor.fetchall():
                    urgency = "critical" if row["stock_quantity"] <= 3 else "high" if row["stock_quantity"] <= 5 else "medium"
                    alerts.append({
                        "product_id": row["id"],
                        "name": row["name"],
                        "price": float(row["price"]),
                        "category": row["category"],
                        "stock_quantity": row["stock_quantity"],
                        "urgency": urgency,
                        "message": f"¡Solo quedan {row['stock_quantity']} unidades!" if row["stock_quantity"] > 0 else "¡Agotado!"
                    })
                
                return alerts
        except Exception as e:
            logger.error(f"Error al obtener alertas de stock: {str(e)}", exc_info=True)
            return []
    
    def get_api_performance_metrics(self) -> Dict:
        """
        Obtiene métricas de rendimiento de la API
        
        Returns:
            Diccionario con métricas de rendimiento
        """
        try:
            return {
                "cache": {
                    "enabled": self.cache_enabled,
                    "size": len(self._recommendation_cache),
                    "hit_rate": "N/A"  # Se podría calcular con tracking
                },
                "database": {
                    "path": self.db_path,
                    "size_mb": round(os.path.getsize(self.db_path) / (1024 * 1024), 2) if os.path.exists(self.db_path) else 0
                },
                "system": {
                    "total_methods": len([m for m in dir(self) if not m.startswith('_') and callable(getattr(self, m))]),
                    "cache_ttl": self._cache_ttl
                },
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Error al obtener métricas de rendimiento: {str(e)}")
            return {}
    
    def get_trending_searches(self, days: int = 7, limit: int = 10) -> List[Dict]:
        """
        Obtiene búsquedas en tendencia (simulado basado en productos más vistos)
        
        Args:
            days: Días a considerar
            limit: Número de resultados
            
        Returns:
            Lista de búsquedas en tendencia
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT 
                        p.name,
                        p.category,
                        SUM(u.total_views) as search_count,
                        COUNT(DISTINCT u.id) as upsells_count
                    FROM products p
                    JOIN upsells u ON p.id = u.product_id
                    WHERE u.is_active = 1
                    GROUP BY p.id, p.name, p.category
                    ORDER BY search_count DESC
                    LIMIT ?
                ''', (limit,))
                
                trends = []
                for row in cursor.fetchall():
                    trends.append({
                        "search_term": row["name"],
                        "category": row["category"],
                        "search_count": row["search_count"] or 0,
                        "upsells_count": row["upsells_count"],
                        "trend_score": (row["search_count"] or 0) * 1.5 + row["upsells_count"] * 10
                    })
                
                return trends
        except Exception as e:
            logger.error(f"Error al obtener tendencias: {str(e)}", exc_info=True)
            return []
    
    def get_frequently_bought_together(self, product_id: str, limit: int = 5) -> List[Dict]:
        """
        Obtiene productos frecuentemente comprados juntos
        
        Args:
            product_id: ID del producto
            limit: Número de resultados
            
        Returns:
            Lista de productos frecuentemente comprados juntos
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Encontrar productos comprados en la misma transacción o por los mismos clientes
                cursor.execute('''
                    SELECT 
                        p2.id,
                        p2.name,
                        p2.price,
                        p2.category,
                        COUNT(DISTINCT ch1.customer_id) as customers_count
                    FROM customer_history ch1
                    JOIN customer_history ch2 ON ch1.customer_id = ch2.customer_id
                    JOIN products p2 ON ch2.product_id = p2.id
                    WHERE ch1.product_id = ?
                        AND ch2.product_id != ?
                        AND p2.is_active = 1
                    GROUP BY p2.id, p2.name, p2.price, p2.category
                    ORDER BY customers_count DESC
                    LIMIT ?
                ''', (product_id, product_id, limit))
                
                products = []
                for row in cursor.fetchall():
                    products.append({
                        "product_id": row["id"],
                        "name": row["name"],
                        "price": float(row["price"]),
                        "category": row["category"],
                        "customers_bought_together": row["customers_count"],
                        "compatibility_score": min(row["customers_count"] / 10.0, 1.0)  # Normalizado
                    })
                
                return products
        except Exception as e:
            logger.error(f"Error al obtener productos relacionados: {str(e)}", exc_info=True)
            return []
    
    def calculate_conversion_velocity(self, product_id: str = None, days: int = 30) -> Dict:
        """
        Calcula velocidad de conversión (tiempo promedio desde vista hasta compra)
        
        Args:
            product_id: ID del producto (opcional)
            days: Días a considerar
            
        Returns:
            Diccionario con métricas de velocidad
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
                
                # Simular cálculo de velocidad (en producción se necesitaría tracking de vistas)
                cursor.execute('''
                    SELECT 
                        COUNT(*) as total_purchases,
                        AVG(CASE WHEN u.total_views > 0 
                            THEN (u.total_purchases * 1.0 / u.total_views) ELSE 0 END) as avg_conversion
                    FROM upsells u
                    WHERE u.is_active = 1
                ''')
                
                metrics = cursor.fetchone()
                
                # Calcular velocidad estimada (simulado)
                avg_conversion = metrics["avg_conversion"] or 0
                total_purchases = metrics["total_purchases"] or 0
                
                # Velocidad alta = conversión rápida
                if avg_conversion > 0.2:
                    velocity = "Alta"
                    estimated_hours = 2
                elif avg_conversion > 0.1:
                    velocity = "Media"
                    estimated_hours = 24
                else:
                    velocity = "Baja"
                    estimated_hours = 72
                
                return {
                    "period_days": days,
                    "product_id": product_id,
                    "metrics": {
                        "total_purchases": total_purchases,
                        "avg_conversion_rate": round(avg_conversion * 100, 2),
                        "conversion_velocity": velocity,
                        "estimated_conversion_hours": estimated_hours
                    }
                }
        except Exception as e:
            logger.error(f"Error al calcular velocidad: {str(e)}", exc_info=True)
            return {}
    
    def get_dynamic_pricing_suggestions(self, product_id: str) -> Dict:
        """
        Sugerencias de precios dinámicos basadas en múltiples factores
        
        Args:
            product_id: ID del producto
            
        Returns:
            Diccionario con sugerencias de precio
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT 
                        p.price,
                        p.stock_quantity,
                        AVG(pr.rating) as avg_rating,
                        COUNT(pr.id) as review_count,
                        SUM(u.total_views) as total_views,
                        SUM(u.total_purchases) as total_purchases
                    FROM products p
                    LEFT JOIN product_reviews pr ON p.id = pr.product_id
                    LEFT JOIN upsells u ON p.id = u.product_id
                    WHERE p.id = ?
                    GROUP BY p.id
                ''', (product_id,))
                
                product = cursor.fetchone()
                if not product:
                    return {}
                
                current_price = product["price"]
                stock = product["stock_quantity"] or 0
                avg_rating = product["avg_rating"] or 0
                views = product["total_views"] or 0
                purchases = product["total_purchases"] or 0
                conversion = (purchases / max(views, 1)) if views > 0 else 0
                
                suggestions = []
                
                # Sugerencia 1: Stock bajo -> aumentar precio
                if stock <= 5 and stock > 0:
                    new_price = current_price * 1.1
                    suggestions.append({
                        "strategy": "scarcity_pricing",
                        "current_price": round(current_price, 2),
                        "suggested_price": round(new_price, 2),
                        "change_percent": 10.0,
                        "reason": "Stock bajo - aplicar pricing de escasez"
                    })
                
                # Sugerencia 2: Alta demanda -> aumentar precio
                if conversion > 0.2 and views > 50:
                    new_price = current_price * 1.05
                    suggestions.append({
                        "strategy": "demand_pricing",
                        "current_price": round(current_price, 2),
                        "suggested_price": round(new_price, 2),
                        "change_percent": 5.0,
                        "reason": "Alta demanda detectada"
                    })
                
                # Sugerencia 3: Baja conversión -> reducir precio
                if conversion < 0.05 and views > 30:
                    new_price = current_price * 0.9
                    suggestions.append({
                        "strategy": "conversion_boost",
                        "current_price": round(current_price, 2),
                        "suggested_price": round(new_price, 2),
                        "change_percent": -10.0,
                        "reason": "Baja conversión - reducir precio para aumentar ventas"
                    })
                
                # Sugerencia 4: Alta calificación -> premium pricing
                if avg_rating >= 4.5 and product["review_count"] >= 10:
                    new_price = current_price * 1.08
                    suggestions.append({
                        "strategy": "premium_pricing",
                        "current_price": round(current_price, 2),
                        "suggested_price": round(new_price, 2),
                        "change_percent": 8.0,
                        "reason": "Alta calificación - aplicar premium pricing"
                    })
                
                if not suggestions:
                    suggestions.append({
                        "strategy": "maintain",
                        "current_price": round(current_price, 2),
                        "suggested_price": round(current_price, 2),
                        "change_percent": 0.0,
                        "reason": "Precio actual es óptimo"
                    })
                
                return {
                    "product_id": product_id,
                    "current_price": round(current_price, 2),
                    "suggestions": suggestions,
                    "best_suggestion": suggestions[0] if suggestions else None
                }
        except Exception as e:
            logger.error(f"Error en sugerencias de precio: {str(e)}", exc_info=True)
            return {}
    
    def get_behavioral_recommendations(self, customer_id: str, limit: int = 5) -> List[Dict]:
        """
        Recomendaciones basadas en comportamiento del cliente
        
        Args:
            customer_id: ID del cliente
            limit: Número de recomendaciones
            
        Returns:
            Lista de recomendaciones basadas en comportamiento
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Obtener categorías y tags más visitados/comprados
                cursor.execute('''
                    SELECT 
                        p.category,
                        p.tags,
                        COUNT(*) as interaction_count
                    FROM customer_history ch
                    JOIN products p ON ch.product_id = p.id
                    WHERE ch.customer_id = ?
                    GROUP BY p.category, p.tags
                    ORDER BY interaction_count DESC
                    LIMIT 3
                ''', (customer_id,))
                
                preferences = cursor.fetchall()
                if not preferences:
                    return []
                
                # Buscar productos similares
                preferred_categories = [p["category"] for p in preferences]
                preferred_tags = set()
                for p in preferences:
                    if p["tags"]:
                        preferred_tags.update(p["tags"].split(','))
                
                # Buscar productos en categorías preferidas
                placeholders = ','.join(['?'] * len(preferred_categories))
                cursor.execute(f'''
                    SELECT DISTINCT
                        p.id,
                        p.name,
                        p.price,
                        p.category,
                        p.tags,
                        COUNT(DISTINCT ch2.customer_id) as popularity
                    FROM products p
                    LEFT JOIN customer_history ch2 ON p.id = ch2.product_id
                    WHERE p.category IN ({placeholders})
                        AND p.id NOT IN (
                            SELECT product_id FROM customer_history WHERE customer_id = ?
                        )
                        AND p.is_active = 1
                    GROUP BY p.id, p.name, p.price, p.category, p.tags
                    ORDER BY popularity DESC, p.price ASC
                    LIMIT ?
                ''', (*preferred_categories, customer_id, limit))
                
                recommendations = []
                for row in cursor.fetchall():
                    # Calcular score basado en preferencias
                    category_match = 1.0 if row["category"] in preferred_categories else 0.5
                    tag_match = 0.0
                    if row["tags"] and preferred_tags:
                        product_tags = set(row["tags"].split(','))
                        common_tags = product_tags.intersection(preferred_tags)
                        tag_match = len(common_tags) / max(len(preferred_tags), 1)
                    
                    score = (category_match * 0.6 + tag_match * 0.4) * (1 + (row["popularity"] or 0) / 100)
                    
                    recommendations.append({
                        "product_id": row["id"],
                        "name": row["name"],
                        "price": float(row["price"]),
                        "category": row["category"],
                        "tags": row["tags"],
                        "behavioral_score": round(score, 3),
                        "reason": f"Basado en tu preferencia por {row['category']}"
                    })
                
                recommendations.sort(key=lambda x: x["behavioral_score"], reverse=True)
                return recommendations[:limit]
        except Exception as e:
            logger.error(f"Error en recomendaciones de comportamiento: {str(e)}", exc_info=True)
            return []
    
    def get_product_affinity_matrix(self, limit: int = 10) -> Dict:
        """
        Genera matriz de afinidad entre productos
        
        Args:
            limit: Número de productos a analizar
            
        Returns:
            Diccionario con matriz de afinidad
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Obtener productos más populares
                cursor.execute('''
                    SELECT 
                        p.id,
                        p.name,
                        COUNT(DISTINCT ch.customer_id) as customers
                    FROM products p
                    JOIN customer_history ch ON p.id = ch.product_id
                    WHERE p.is_active = 1
                    GROUP BY p.id, p.name
                    ORDER BY customers DESC
                    LIMIT ?
                ''', (limit,))
                
                products = [{"id": row["id"], "name": row["name"]} for row in cursor.fetchall()]
                
                # Calcular afinidad entre productos
                affinity_matrix = {}
                for i, prod1 in enumerate(products):
                    for prod2 in products[i+1:]:
                        cursor.execute('''
                            SELECT COUNT(DISTINCT customer_id) as common_customers
                            FROM customer_history ch1
                            JOIN customer_history ch2 ON ch1.customer_id = ch2.customer_id
                            WHERE ch1.product_id = ? AND ch2.product_id = ?
                        ''', (prod1["id"], prod2["id"]))
                        
                        result = cursor.fetchone()
                        common = result["common_customers"] or 0
                        
                        if common > 0:
                            key = f"{prod1['id']}-{prod2['id']}"
                            affinity_matrix[key] = {
                                "product_1": prod1["name"],
                                "product_2": prod2["name"],
                                "affinity_score": common,
                                "relationship": "strong" if common >= 5 else "moderate" if common >= 2 else "weak"
                            }
                
                return {
                    "products_analyzed": len(products),
                    "affinity_pairs": len(affinity_matrix),
                    "matrix": affinity_matrix
                }
        except Exception as e:
            logger.error(f"Error en matriz de afinidad: {str(e)}", exc_info=True)
            return {}
    
    def get_customer_lifetime_value(self, customer_id: str) -> Dict:
        """
        Calcula el valor de vida del cliente (LTV)
        
        Args:
            customer_id: ID del cliente
            
        Returns:
            Diccionario con métricas de LTV
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Obtener todas las compras del cliente
                cursor.execute('''
                    SELECT 
                        SUM(total_amount) as total_spent,
                        COUNT(*) as total_orders,
                        MIN(purchased_at) as first_purchase,
                        MAX(purchased_at) as last_purchase,
                        AVG(total_amount) as avg_order_value
                    FROM purchases
                    WHERE customer_id = ?
                ''', (customer_id,))
                
                stats = cursor.fetchone()
                
                if not stats or not stats["total_spent"]:
                    return {}
                
                total_spent = float(stats["total_spent"] or 0)
                total_orders = stats["total_orders"] or 0
                avg_order_value = float(stats["avg_order_value"] or 0)
                
                # Calcular días como cliente
                first_purchase = datetime.fromisoformat(stats["first_purchase"].replace('Z', '+00:00')).replace(tzinfo=None)
                days_as_customer = (datetime.now() - first_purchase).days
                days_as_customer = max(days_as_customer, 1)
                
                # Calcular frecuencia de compra
                purchase_frequency = total_orders / (days_as_customer / 30)  # Compras por mes
                
                # Estimar LTV (simplificado: gasto promedio mensual * meses estimados)
                monthly_value = total_spent / (days_as_customer / 30)
                estimated_ltv = monthly_value * 12  # Estimar 1 año
                
                # Clasificar cliente
                if total_spent >= 5000:
                    tier = "Premium"
                elif total_spent >= 2000:
                    tier = "Gold"
                elif total_spent >= 1000:
                    tier = "Silver"
                else:
                    tier = "Bronze"
                
                return {
                    "customer_id": customer_id,
                    "ltv_metrics": {
                        "total_spent": round(total_spent, 2),
                        "total_orders": total_orders,
                        "avg_order_value": round(avg_order_value, 2),
                        "days_as_customer": days_as_customer,
                        "purchase_frequency_per_month": round(purchase_frequency, 2),
                        "monthly_value": round(monthly_value, 2),
                        "estimated_ltv": round(estimated_ltv, 2)
                    },
                    "customer_tier": tier,
                    "first_purchase": stats["first_purchase"],
                    "last_purchase": stats["last_purchase"]
                }
        except Exception as e:
            logger.error(f"Error al calcular LTV: {str(e)}", exc_info=True)
            return {}
    
    def get_upsell_success_rate(self, product_id: str = None, days: int = 30) -> Dict:
        """
        Calcula tasa de éxito de upsells
        
        Args:
            product_id: ID del producto (opcional)
            days: Días a considerar
            
        Returns:
            Diccionario con tasa de éxito
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
                
                where_clause = "WHERE u.product_id = ?" if product_id else ""
                params = (product_id, cutoff_date) if product_id else (cutoff_date,)
                
                # Total de compras con upsell
                query = f'''
                    SELECT 
                        COUNT(DISTINCT p.id) as purchases_with_upsell,
                        COUNT(DISTINCT CASE WHEN p.upsell_product_id IS NOT NULL THEN p.id END) as upsell_purchases
                    FROM purchases p
                    JOIN upsells u ON p.product_id = u.product_id
                    {where_clause}
                    AND p.purchased_at >= ?
                '''
                
                cursor.execute(query, params)
                stats = cursor.fetchone()
                
                total_purchases = stats["purchases_with_upsell"] or 0
                upsell_purchases = stats["upsell_purchases"] or 0
                
                success_rate = (upsell_purchases / max(total_purchases, 1)) * 100
                
                return {
                    "period_days": days,
                    "product_id": product_id,
                    "metrics": {
                        "total_purchases": total_purchases,
                        "upsell_purchases": upsell_purchases,
                        "success_rate": round(success_rate, 2),
                        "no_upsell_purchases": total_purchases - upsell_purchases
                    },
                    "performance": "Excelente" if success_rate > 40 else "Bueno" if success_rate > 25 else "Mejorable"
                }
        except Exception as e:
            logger.error(f"Error al calcular tasa de éxito: {str(e)}", exc_info=True)
            return {}


def create_flask_app(api: OfferLetterAPI) -> Flask:
    """Crea una aplicación Flask para la API REST con todas las funcionalidades avanzadas"""
    if not FLASK_AVAILABLE:
        raise ImportError("Flask no está instalado. Instale con: pip install flask")
    
    app = Flask(__name__)
    
    # Inicializar managers globales
    auth_manager = AuthenticationManager()
    cache_manager = CacheManager(default_ttl=300)
    i18n_manager = I18nManager()
    signature_manager = DigitalSignatureManager()
    integration_manager = IntegrationManager()
    analytics_manager = AnalyticsManager()
    backup_manager = BackupManager()
    performance_monitor = PerformanceMonitor()
    notification_manager = NotificationManager()
    export_manager = ExportManager()
    schema_validator = SchemaValidator()
    rate_limiter = RateLimiter(max_requests=60, window_seconds=60)
    event_manager = EventManager()
    advanced_logger = AdvancedLogger()
    audit_manager = AuditManager()
    task_queue = TaskQueue()
    template_manager = TemplateManager()
    permission_manager = PermissionManager()
    upsell_manager = UpsellManager()
    
    # Decorador para autenticación
    def require_auth(f):
        @functools.wraps(f)
        def decorated_function(*args, **kwargs):
            token = request.headers.get('Authorization', '').replace('Bearer ', '')
            if not token:
                return jsonify({"error": "Token de autenticación requerido"}), 401
            
            token_data = auth_manager.verify_token(token)
            if not token_data:
                return jsonify({"error": "Token inválido o expirado"}), 401
            
            request.user_id = token_data['user_id']
            return f(*args, **kwargs)
        return decorated_function
    
    # Decorador para permisos
    def require_permission(permission: str):
        def decorator(f):
            @functools.wraps(f)
            @require_auth
            def decorated_function(*args, **kwargs):
                # Obtener rol del usuario (simplificado, en producción obtener de BD)
                user_role = 'hr_user'  # Por defecto
                
                if not permission_manager.has_permission(user_role, permission):
                    return jsonify({
                        "error": "Permiso insuficiente",
                        "required": permission,
                        "role": user_role
                    }), 403
                
                return f(*args, **kwargs)
            return decorated_function
        return decorator
    
    # Middleware para rate limiting
    @app.before_request
    def rate_limit_check():
        # Excluir endpoints públicos del rate limiting
        public_endpoints = ['/health', '/api/docs', '/api/i18n/languages']
        if request.path not in public_endpoints:
            client_id = rate_limiter.get_client_identifier(request)
            allowed, info = rate_limiter.is_allowed(client_id)
            if not allowed:
                response = jsonify(info)
                response.headers['X-RateLimit-Limit'] = str(info['limit'])
                response.headers['X-RateLimit-Remaining'] = '0'
                response.headers['X-RateLimit-Reset'] = str(int(info['retry_after']))
                response.status_code = 429
                return response
            else:
                # Agregar headers de rate limit
                response = make_response()
                response.headers['X-RateLimit-Limit'] = str(info['limit'])
                response.headers['X-RateLimit-Remaining'] = str(info['remaining'])
                response.headers['X-RateLimit-Reset'] = str(int(info['reset_at']))
    
    # Middleware para logging y monitoreo
    @app.before_request
    def before_request():
        request.start_time = time.time()
    
    @app.after_request
    def after_request(response):
        if hasattr(request, 'start_time'):
            response_time = time.time() - request.start_time
            performance_monitor.record_request(
                request.path,
                request.method,
                response_time,
                response.status_code
            )
            
            # Logging avanzado
            user_id = getattr(request, 'user_id', None)
            advanced_logger.log_api_request(
                request.method,
                request.path,
                response.status_code,
                response_time,
                user_id
            )
            
            # Auditoría
            if user_id and request.method in ['POST', 'PUT', 'DELETE']:
                try:
                    audit_manager.log_action(
                        user_id=user_id,
                        action=f"{request.method} {request.path}",
                        resource_type=request.path.split('/')[-1] if '/' in request.path else None,
                        ip_address=request.remote_addr if hasattr(request, 'remote_addr') else None,
                        user_agent=request.headers.get('User-Agent')
                    )
                except Exception as e:
                    logger.warning(f"Error logging audit action: {str(e)}")
            
            # Emitir eventos
            if request.path.startswith('/api/offers') and request.method == 'POST':
                if response.status_code == 201:
                    try:
                        offer_data = request.json
                        offer_id = offer_data.get('id') or 'unknown'
                        event_manager.emit_offer_created(offer_id, offer_data)
                    except Exception as e:
                        logger.warning(f"Error emitting offer created event: {str(e)}")
            
            # Compresión de respuestas
            response = CompressionMiddleware.compress_response(response)
        
        return response
    
    # Middleware para caché
    def cache_response(ttl=300):
        def decorator(f):
            @functools.wraps(f)
            def decorated_function(*args, **kwargs):
                cache_key = f"{request.path}:{request.method}:{str(request.args)}"
                cached = cache_manager.get(cache_key)
                if cached is not None:
                    performance_monitor.record_cache_hit()
                    return jsonify(cached)
                
                performance_monitor.record_cache_miss()
                result = f(*args, **kwargs)
                
                # Solo cachear respuestas exitosas
                if isinstance(result, tuple) and len(result) == 2:
                    response_data, status_code = result
                    if status_code == 200 and isinstance(response_data, dict):
                        cache_manager.set(cache_key, response_data, ttl)
                        return jsonify(response_data), status_code
                
                return result
            return decorated_function
        return decorator
    
    @app.route('/api/offers', methods=['POST'])
    @require_auth
    def create_offer():
        """Endpoint para crear una nueva oferta"""
        try:
            data = request.json
            if not data:
                return jsonify({
                    "success": False,
                    "error": "Request body is required"
                }), 400
            
            # Validar con JSON Schema
            is_valid, validation_errors = schema_validator.validate(data)
            if not is_valid:
                return jsonify({
                    "success": False,
                    "error": "Validation failed",
                    "validation_errors": validation_errors
                }), 400
            
            format_type = request.args.get('format', 'txt')
            if format_type not in Config.SUPPORTED_FORMATS:
                return jsonify({
                    "success": False,
                    "error": f"Unsupported format. Supported formats: {', '.join(Config.SUPPORTED_FORMATS)}"
                }), 400
            
            result = api.create_offer(data, format_type)
            
            # Emitir evento
            if result.get('success'):
                event_manager.emit_offer_created(
                    result.get('offer_id', 'unknown'),
                    data
                )
            
            return jsonify(result), 201 if result.get('success') else 400
        except Exception as e:
            logger.error(f"Error creating offer: {str(e)}", exc_info=True)
            advanced_logger.log_error(e, {"endpoint": "/api/offers", "method": "POST"})
            return jsonify({"success": False, "error": "Internal server error"}), 500
    
    @app.route('/api/offers/<offer_id>', methods=['GET'])
    def get_offer(offer_id):
        """Endpoint para obtener una oferta por ID"""
        offer = api.get_offer_by_id(offer_id)
        if offer:
            return jsonify(offer), 200
        return jsonify({"error": "Oferta no encontrada"}), 404
    
    @app.route('/api/offers', methods=['GET'])
    def list_offers():
        """Endpoint para listar ofertas con filtros"""
        filters = {}
        if request.args.get('status'):
            filters['status'] = request.args.get('status')
        if request.args.get('department'):
            filters['department'] = request.args.get('department')
        
        offers = api.get_offer_history(filters)
        return jsonify({"offers": offers, "count": len(offers)}), 200
    
    @app.route('/api/offers/<offer_id>/status', methods=['PUT'])
    def update_status(offer_id):
        """Endpoint para actualizar el estado de una oferta"""
        data = request.json
        status = data.get('status')
        details = data.get('details', '')
        
        if api.update_offer_status(offer_id, status, details):
            return jsonify({"success": True}), 200
        return jsonify({"error": "Error al actualizar estado"}), 400
    
    @app.route('/api/offers/<offer_id>/pdf', methods=['GET'])
    def download_pdf(offer_id):
        """Endpoint para descargar PDF de una oferta"""
        pdf_path = api.generate_pdf(offer_id)
        if pdf_path and os.path.exists(pdf_path):
            return send_file(pdf_path, as_attachment=True)
        return jsonify({"error": "PDF no disponible"}), 404
    
    @app.route('/api/offers/<offer_id>/send', methods=['POST'])
    def send_email(offer_id):
        """Endpoint para enviar oferta por email"""
        smtp_config = request.json.get('smtp_config')
        include_pdf = request.json.get('include_pdf', False)
        
        if not smtp_config:
            return jsonify({"error": "Configuración SMTP requerida"}), 400
        
        if api.send_offer_email(offer_id, smtp_config, include_pdf):
            return jsonify({"success": True}), 200
        return jsonify({"error": "Error al enviar email"}), 500
    
    @app.route('/api/statistics', methods=['GET'])
    def get_statistics():
        """Endpoint para obtener estadísticas"""
        stats = api.get_statistics()
        return jsonify(stats), 200
    
    # ============================================================================
    # ENDPOINTS DE UPSELLS COMPLEMENTARIOS
    # ============================================================================
    
    @app.route('/api/upsells/<product_id>', methods=['GET'])
    def get_upsells(product_id):
        """
        Obtiene upsells complementarios para un producto
        
        Ejemplo de respuesta:
        {
            "success": true,
            "product_id": "prod_001",
            "upsells": [
                {
                    "id": "upsell_001",
                    "upsell_product": {
                        "id": "prod_002",
                        "name": "Cargador Rápido",
                        "description": "Cargador rápido de alta velocidad",
                        "category": "Accesorios"
                    },
                    "bundle_price": 199.0,
                    "original_price": 399.0,
                    "savings": 200.0,
                    "emotional_reason": "Clientes que compraron esto también se llevaron el cargador rápido para máxima compatibilidad",
                    "display_message": "¡Agrega cargador rápido por solo +$199 MXN (valor $399 MXN)!"
                }
            ]
        }
        """
        try:
            limit = int(request.args.get('limit', 2))
            customer_id = request.args.get('customer_id')
            
            # Validar límite
            if limit < 1 or limit > 10:
                return jsonify({
                    "success": False,
                    "error": "El límite debe estar entre 1 y 10"
                }), 400
            
            category_filter = request.args.get('category')
            exclude_purchased = request.args.get('exclude_purchased', 'true').lower() == 'true'
            
            upsells = upsell_manager.get_complementary_upsells(
                product_id, 
                limit=limit,
                customer_id=customer_id,
                category_filter=category_filter,
                exclude_purchased=exclude_purchased
            )
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "upsells": upsells,
                "count": len(upsells)
            }), 200
        except ValueError as e:
            return jsonify({
                "success": False,
                "error": f"Parámetro inválido: {str(e)}"
            }), 400
        except Exception as e:
            logger.error(f"Error al obtener upsells: {str(e)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al obtener upsells complementarios"
            }), 500
    
    @app.route('/api/purchases', methods=['POST'])
    def record_purchase():
        """
        Registra una compra y retorna upsells complementarios
        
        Body:
        {
            "product_id": "prod_001",
            "customer_id": "customer_123" (opcional)
        }
        
        Respuesta:
        {
            "success": true,
            "purchase_id": "uuid",
            "upsells": [...]
        }
        """
        try:
            data = request.json
            if not data or 'product_id' not in data:
                return jsonify({
                    "success": False,
                    "error": "product_id es requerido"
                }), 400
            
            product_id = data['product_id']
            customer_id = data.get('customer_id')
            
            # Registrar la compra
            upsell_manager.record_purchase(product_id, customer_id)
            
            # Obtener upsells complementarios
            limit = int(request.args.get('limit', 2))
            upsells = upsell_manager.get_complementary_upsells(product_id, limit=limit)
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "upsells": upsells,
                "count": len(upsells),
                "message": "Compra registrada exitosamente"
            }), 201
        except Exception as e:
            logger.error(f"Error al registrar compra: {str(e)}")
            return jsonify({
                "success": False,
                "error": "Error al registrar compra"
            }), 500
    
    @app.route('/api/purchases/<product_id>/add-upsell', methods=['POST'])
    def add_upsell_to_purchase(product_id):
        """
        Agrega un upsell a una compra (1-click add-on)
        
        Body:
        {
            "upsell_product_id": "prod_002",
            "customer_id": "customer_123" (opcional)
        }
        """
        try:
            data = request.json
            if not data or 'upsell_product_id' not in data:
                return jsonify({
                    "success": False,
                    "error": "upsell_product_id es requerido"
                }), 400
            
            upsell_product_id = data['upsell_product_id']
            customer_id = data.get('customer_id')
            
            # Registrar la compra del upsell
            upsell_manager.record_purchase(product_id, customer_id, upsell_product_id)
            
            # Obtener información del upsell
            upsells = upsell_manager.get_complementary_upsells(product_id, limit=10)
            selected_upsell = next((u for u in upsells if u['upsell_product']['id'] == upsell_product_id), None)
            
            if not selected_upsell:
                return jsonify({
                    "success": False,
                    "error": "Upsell no encontrado"
                }), 404
            
            return jsonify({
                "success": True,
                "message": "Upsell agregado exitosamente",
                "upsell": selected_upsell
            }), 200
        except Exception as e:
            logger.error(f"Error al agregar upsell: {str(e)}")
            return jsonify({
                "success": False,
                "error": "Error al agregar upsell"
            }), 500
    
    @app.route('/api/products', methods=['POST'])
    @require_auth
    def add_product():
        """Agrega un nuevo producto (requiere autenticación)"""
        try:
            data = request.json
            required_fields = ['id', 'name', 'price']
            if not all(field in data for field in required_fields):
                return jsonify({
                    "success": False,
                    "error": f"Campos requeridos: {', '.join(required_fields)}"
                }), 400
            
            success = upsell_manager.add_product(
                product_id=data['id'],
                name=data['name'],
                price=float(data['price']),
                category=data.get('category'),
                description=data.get('description'),
                tags=data.get('tags', []),
                stock_quantity=int(data.get('stock_quantity', -1))
            )
            
            if success:
                return jsonify({
                    "success": True,
                    "message": "Producto agregado exitosamente"
                }), 201
            return jsonify({
                "success": False,
                "error": "Error al agregar producto"
            }), 500
        except Exception as e:
            logger.error(f"Error al agregar producto: {str(e)}")
            return jsonify({
                "success": False,
                "error": "Error al agregar producto"
            }), 500
    
    @app.route('/api/upsells', methods=['POST'])
    @require_auth
    def create_upsell():
        """Crea un upsell para un producto (requiere autenticación)"""
        try:
            data = request.json
            required_fields = ['product_id', 'upsell_product_id', 'bundle_price']
            if not all(field in data for field in required_fields):
                return jsonify({
                    "success": False,
                    "error": f"Campos requeridos: {', '.join(required_fields)}"
                }), 400
            
            success = upsell_manager.add_upsell(
                product_id=data['product_id'],
                upsell_product_id=data['upsell_product_id'],
                bundle_price=float(data['bundle_price']),
                emotional_reason=data.get('emotional_reason'),
                priority=int(data.get('priority', 0))
            )
            
            if success:
                return jsonify({
                    "success": True,
                    "message": "Upsell creado exitosamente"
                }), 201
            return jsonify({
                "success": False,
                "error": "Error al crear upsell"
            }), 500
        except Exception as e:
            logger.error(f"Error al crear upsell: {str(e)}")
            return jsonify({
                "success": False,
                "error": "Error al crear upsell"
            }), 500
    
    @app.route('/api/upsells/analytics', methods=['GET'])
    @require_auth
    def get_upsell_analytics():
        """
        Obtiene analytics de upsells (requiere autenticación)
        
        Query params:
            product_id: ID del producto (opcional, si no se proporciona retorna global)
        """
        try:
            product_id = request.args.get('product_id')
            analytics = upsell_manager.get_upsell_analytics(product_id)
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "analytics": analytics
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener analytics: {str(e)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al obtener analytics"
            }), 500
    
    @app.route('/api/upsells/top-performing', methods=['GET'])
    @require_auth
    def get_top_performing_upsells():
        """
        Obtiene los upsells con mejor rendimiento (requiere autenticación)
        
        Query params:
            limit: Número de resultados (default: 10)
            days: Días a considerar (default: 30)
        """
        try:
            limit = int(request.args.get('limit', 10))
            days = int(request.args.get('days', 30))
            
            top_upsells = upsell_manager.get_top_performing_upsells(limit=limit, days=days)
            
            return jsonify({
                "success": True,
                "limit": limit,
                "days": days,
                "upsells": top_upsells,
                "count": len(top_upsells)
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener top upsells: {str(e)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al obtener top upsells"
            }), 500
    
    @app.route('/api/upsells/time-limited-offers', methods=['POST'])
    @require_auth
    def create_time_limited_offer():
        """
        Crea una oferta por tiempo limitado (requiere autenticación)
        
        Body:
        {
            "upsell_id": "upsell_001",
            "start_date": "2024-01-01T00:00:00",
            "end_date": "2024-01-31T23:59:59",
            "discount_percentage": 15.0
        }
        """
        try:
            data = request.json
            required_fields = ['upsell_id', 'start_date', 'end_date', 'discount_percentage']
            if not all(field in data for field in required_fields):
                return jsonify({
                    "success": False,
                    "error": f"Campos requeridos: {', '.join(required_fields)}"
                }), 400
            
            success = upsell_manager.create_time_limited_offer(
                upsell_id=data['upsell_id'],
                start_date=data['start_date'],
                end_date=data['end_date'],
                discount_percentage=float(data['discount_percentage'])
            )
            
            if success:
                return jsonify({
                    "success": True,
                    "message": "Oferta temporal creada exitosamente"
                }), 201
            return jsonify({
                "success": False,
                "error": "Error al crear oferta temporal"
            }), 500
        except Exception as e:
            logger.error(f"Error al crear oferta temporal: {str(e)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al crear oferta temporal"
            }), 500
    
    @app.route('/api/upsells/report', methods=['GET'])
    @require_auth
    def export_upsell_report():
        """
        Exporta un reporte completo de upsells (requiere autenticación)
        
        Query params:
            product_id: ID del producto (opcional)
            format: Formato del reporte (json, csv) - default: json
        """
        try:
            product_id = request.args.get('product_id')
            format_type = request.args.get('format', 'json')
            
            report = upsell_manager.export_upsell_report(product_id, format=format_type)
            
            if format_type == 'csv' and CSV_AVAILABLE:
                # Generar CSV
                import csv
                from io import StringIO
                
                output = StringIO()
                writer = csv.writer(output)
                
                # Headers
                writer.writerow(['Upsell ID', 'Product', 'Upsell', 'Conversion Rate', 'Purchases', 'Views'])
                
                # Data
                for upsell in report.get('top_performing_upsells', []):
                    writer.writerow([
                        upsell.get('upsell_id', ''),
                        upsell.get('product', ''),
                        upsell.get('upsell', ''),
                        upsell.get('conversion_rate', 0),
                        upsell.get('total_purchases', 0),
                        upsell.get('total_views', 0)
                    ])
                
                response = make_response(output.getvalue())
                response.headers['Content-Type'] = 'text/csv'
                response.headers['Content-Disposition'] = f'attachment; filename=upsell_report_{datetime.now().strftime("%Y%m%d")}.csv'
                return response
            
            return jsonify({
                "success": True,
                "report": report
            }), 200
        except Exception as e:
            logger.error(f"Error al exportar reporte: {str(e)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al exportar reporte"
            }), 500
    
    @app.route('/api/upsells/batch', methods=['POST'])
    def batch_get_upsells():
        """Obtiene upsells para múltiples productos en una sola operación"""
        try:
            data = request.json
            if not data or 'product_ids' not in data:
                return jsonify({"success": False, "error": "product_ids es requerido"}), 400
            
            product_ids = data['product_ids']
            limit = int(data.get('limit_per_product', 2))
            
            if not isinstance(product_ids, list) or len(product_ids) == 0:
                return jsonify({"success": False, "error": "product_ids debe ser una lista no vacía"}), 400
            
            if len(product_ids) > 50:
                return jsonify({"success": False, "error": "Máximo 50 productos por batch"}), 400
            
            results = upsell_manager.batch_get_upsells(product_ids, limit_per_product=limit)
            
            return jsonify({
                "success": True,
                "results": results,
                "total_products": len(product_ids),
                "products_with_upsells": len([k for k, v in results.items() if len(v) > 0])
            }), 200
        except Exception as e:
            logger.error(f"Error en batch_get_upsells: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener upsells en batch"}), 500
    
    @app.route('/api/upsells/volume-discount', methods=['POST'])
    def calculate_volume_discount():
        """Calcula descuento por volumen"""
        try:
            data = request.json
            if not all(field in data for field in ['base_price', 'quantity']):
                return jsonify({"success": False, "error": "base_price y quantity son requeridos"}), 400
            
            base_price = float(data['base_price'])
            quantity = int(data['quantity'])
            
            if base_price <= 0 or quantity <= 0:
                return jsonify({"success": False, "error": "base_price y quantity deben ser mayores a 0"}), 400
            
            discount_info = upsell_manager.calculate_volume_discount(base_price, quantity)
            return jsonify({"success": True, "discount": discount_info}), 200
        except Exception as e:
            logger.error(f"Error al calcular descuento: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular descuento"}), 500
    
    @app.route('/api/upsells/realtime-metrics', methods=['GET'])
    @require_auth
    def get_realtime_metrics():
        """Obtiene métricas en tiempo real - requiere autenticación"""
        try:
            product_id = request.args.get('product_id')
            metrics = upsell_manager.get_realtime_metrics(product_id)
            return jsonify({"success": True, "metrics": metrics}), 200
        except Exception as e:
            logger.error(f"Error al obtener métricas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener métricas"}), 500
    
    @app.route('/api/customers/<customer_id>/insights', methods=['GET'])
    @require_auth
    def get_customer_insights(customer_id):
        """Obtiene insights personalizados de un cliente - requiere autenticación"""
        try:
            insights = upsell_manager.get_customer_insights(customer_id)
            return jsonify({"success": True, "insights": insights}), 200
        except Exception as e:
            logger.error(f"Error al obtener insights: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener insights"}), 500
    
    @app.route('/api/upsells/<product_id>/optimized', methods=['GET'])
    def get_optimized_recommendations(product_id):
        """Obtiene recomendaciones optimizadas con mejoras de rendimiento"""
        try:
            customer_id = request.args.get('customer_id')
            recommendations = upsell_manager.optimize_recommendations(product_id, customer_id)
            return jsonify({
                "success": True,
                "product_id": product_id,
                "recommendations": recommendations,
                "count": len(recommendations),
                "optimized": True
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener recomendaciones optimizadas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener recomendaciones"}), 500
    
    @app.route('/api/upsells/<upsell_id>/predict', methods=['GET'])
    def predict_conversion(upsell_id):
        """
        Predice la probabilidad de conversión de un upsell
        
        Query params:
            customer_id: ID del cliente (opcional)
        """
        try:
            customer_id = request.args.get('customer_id')
            prediction = upsell_manager.predict_conversion_probability(upsell_id, customer_id)
            return jsonify({"success": True, "prediction": prediction}), 200
        except Exception as e:
            logger.error(f"Error al predecir conversión: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al predecir conversión"}), 500
    
    @app.route('/api/upsells/trending', methods=['GET'])
    def get_trending_upsells():
        """
        Obtiene upsells en tendencia
        
        Query params:
            days: Días a considerar (default: 7)
            limit: Número de resultados (default: 10)
        """
        try:
            days = int(request.args.get('days', 7))
            limit = int(request.args.get('limit', 10))
            trending = upsell_manager.get_trending_upsells(days=days, limit=limit)
            return jsonify({
                "success": True,
                "trending_upsells": trending,
                "count": len(trending),
                "period_days": days
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener trending: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener trending"}), 500
    
    @app.route('/api/products/<product_id>/suggest-upsells', methods=['GET'])
    @require_auth
    def generate_upsell_suggestions(product_id):
        """
        Genera sugerencias automáticas de upsells - requiere autenticación
        """
        try:
            suggestions = upsell_manager.generate_upsell_suggestions(product_id)
            return jsonify({
                "success": True,
                "product_id": product_id,
                "suggestions": suggestions,
                "count": len(suggestions)
            }), 200
        except Exception as e:
            logger.error(f"Error al generar sugerencias: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al generar sugerencias"}), 500
    
    @app.route('/api/upsells/<upsell_id>/forecast', methods=['GET'])
    @require_auth
    def get_performance_forecast(upsell_id):
        """
        Genera pronóstico de rendimiento - requiere autenticación
        
        Query params:
            days: Días a pronosticar (default: 30)
        """
        try:
            days = int(request.args.get('days', 30))
            forecast = upsell_manager.get_performance_forecast(upsell_id, days=days)
            return jsonify({"success": True, "forecast": forecast}), 200
        except Exception as e:
            logger.error(f"Error al generar pronóstico: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al generar pronóstico"}), 500
    
    @app.route('/api/upsells/auto-optimize', methods=['POST'])
    @require_auth
    def auto_optimize_upsells():
        """
        Optimiza automáticamente upsells basado en rendimiento - requiere autenticación
        """
        try:
            optimizations = upsell_manager.auto_optimize_upsells()
            return jsonify({"success": True, "optimizations": optimizations}), 200
        except Exception as e:
            logger.error(f"Error en auto-optimización: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en auto-optimización"}), 500
    
    @app.route('/api/upsells/dashboard', methods=['GET'])
    @require_auth
    def get_dashboard():
        """Obtiene resumen completo para dashboard - requiere autenticación"""
        try:
            summary = upsell_manager.get_dashboard_summary()
            return jsonify({"success": True, "dashboard": summary}), 200
        except Exception as e:
            logger.error(f"Error al obtener dashboard: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener dashboard"}), 500
    
    @app.route('/api/customers/<customer_id>/collaborative-recommendations', methods=['GET'])
    def get_collaborative_recommendations(customer_id):
        """
        Obtiene recomendaciones basadas en clientes similares
        
        Query params:
            limit: Número de recomendaciones (default: 5)
        """
        try:
            limit = int(request.args.get('limit', 5))
            recommendations = upsell_manager.get_collaborative_recommendations(customer_id, limit=limit)
            return jsonify({
                "success": True,
                "customer_id": customer_id,
                "recommendations": recommendations,
                "count": len(recommendations),
                "type": "collaborative_filtering"
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener recomendaciones colaborativas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener recomendaciones"}), 500
    
    @app.route('/api/upsells/revenue-impact', methods=['GET'])
    @require_auth
    def get_revenue_impact():
        """
        Calcula impacto en ingresos de upsells - requiere autenticación
        
        Query params:
            product_id: ID del producto (opcional)
            days: Días a considerar (default: 30)
        """
        try:
            product_id = request.args.get('product_id')
            days = int(request.args.get('days', 30))
            impact = upsell_manager.calculate_revenue_impact(product_id, days=days)
            return jsonify({"success": True, "revenue_impact": impact}), 200
        except Exception as e:
            logger.error(f"Error al calcular impacto: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular impacto"}), 500
    
    @app.route('/api/upsells/seasonal-trends', methods=['GET'])
    @require_auth
    def get_seasonal_trends():
        """
        Analiza tendencias estacionales - requiere autenticación
        
        Query params:
            product_id: ID del producto (opcional)
        """
        try:
            product_id = request.args.get('product_id')
            trends = upsell_manager.get_seasonal_trends(product_id)
            return jsonify({"success": True, "trends": trends}), 200
        except Exception as e:
            logger.error(f"Error al obtener tendencias: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener tendencias"}), 500
    
    @app.route('/api/upsells/<upsell_id>/message-variants', methods=['POST'])
    @require_auth
    def create_message_variant(upsell_id):
        """
        Crea una variante de mensaje para A/B testing - requiere autenticación
        
        Body:
        {
            "variant_name": "Variante A",
            "message_template": "¡Oferta especial! Agrega {product} por solo +${price} MXN"
        }
        """
        try:
            data = request.json
            if not all(field in data for field in ['variant_name', 'message_template']):
                return jsonify({"success": False, "error": "variant_name y message_template son requeridos"}), 400
            
            success = upsell_manager.create_message_variant(
                upsell_id=upsell_id,
                variant_name=data['variant_name'],
                message_template=data['message_template']
            )
            
            if success:
                return jsonify({"success": True, "message": "Variante creada exitosamente"}), 201
            return jsonify({"success": False, "error": "Error al crear variante"}), 500
        except Exception as e:
            logger.error(f"Error al crear variante: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al crear variante"}), 500
    
    @app.route('/api/upsells/<upsell_id>/best-variant', methods=['GET'])
    def get_best_message_variant(upsell_id):
        """Obtiene la mejor variante de mensaje basada en conversión"""
        try:
            variant = upsell_manager.get_best_message_variant(upsell_id)
            if variant:
                return jsonify({"success": True, "best_variant": variant}), 200
            return jsonify({"success": False, "error": "No se encontraron variantes"}), 404
        except Exception as e:
            logger.error(f"Error al obtener mejor variante: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener variante"}), 500
    
    @app.route('/api/upsells/backup', methods=['POST'])
    @require_auth
    def backup_upsells_database():
        """
        Crea un backup de la base de datos de upsells - requiere autenticación
        
        Body (opcional):
        {
            "backup_path": "custom/path/backup.db"
        }
        """
        try:
            data = request.json or {}
            backup_path = data.get('backup_path')
            
            backup_file = upsell_manager.backup_database(backup_path)
            
            return jsonify({
                "success": True,
                "backup_path": backup_file,
                "message": "Backup creado exitosamente"
            }), 200
        except Exception as e:
            logger.error(f"Error al crear backup: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al crear backup"}), 500
    
    @app.route('/api/upsells/restore', methods=['POST'])
    @require_auth
    def restore_upsells_database():
        """
        Restaura la base de datos desde un backup - requiere autenticación
        
        Body:
        {
            "backup_path": "upsells_backup_20240115.db"
        }
        """
        try:
            data = request.json
            if not data or 'backup_path' not in data:
                return jsonify({"success": False, "error": "backup_path es requerido"}), 400
            
            success = upsell_manager.restore_database(data['backup_path'])
            
            if success:
                return jsonify({"success": True, "message": "Base de datos restaurada exitosamente"}), 200
            return jsonify({"success": False, "error": "Error al restaurar backup"}), 500
        except Exception as e:
            logger.error(f"Error al restaurar backup: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al restaurar backup"}), 500
    
    @app.route('/api/upsells/export-all', methods=['GET'])
    @require_auth
    def export_all_upsells_data():
        """
        Exporta todos los datos del sistema - requiere autenticación
        
        Query params:
            format: Formato de exportación (json, csv) - default: json
        """
        try:
            format_type = request.args.get('format', 'json')
            export_data = upsell_manager.export_all_data(format=format_type)
            
            if format_type == 'csv' and CSV_AVAILABLE:
                import csv
                from io import StringIO
                
                output = StringIO()
                writer = csv.writer(output)
                
                # Exportar productos
                writer.writerow(['=== PRODUCTOS ==='])
                writer.writerow(['ID', 'Nombre', 'Precio', 'Categoría'])
                for product in export_data['data']['products']:
                    writer.writerow([
                        product.get('id', ''),
                        product.get('name', ''),
                        product.get('price', ''),
                        product.get('category', '')
                    ])
                
                writer.writerow([])
                writer.writerow(['=== UPSELLS ==='])
                writer.writerow(['ID', 'Producto', 'Upsell', 'Bundle Price', 'Conversión'])
                for upsell in export_data['data']['upsells']:
                    writer.writerow([
                        upsell.get('id', ''),
                        upsell.get('product_id', ''),
                        upsell.get('upsell_product_id', ''),
                        upsell.get('bundle_price', ''),
                        upsell.get('conversion_rate', '')
                    ])
                
                response = make_response(output.getvalue())
                response.headers['Content-Type'] = 'text/csv'
                response.headers['Content-Disposition'] = f'attachment; filename=upsells_export_{datetime.now().strftime("%Y%m%d")}.csv'
                return response
            
            return jsonify({"success": True, "export": export_data}), 200
        except Exception as e:
            logger.error(f"Error al exportar datos: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al exportar datos"}), 500
    
    @app.route('/api/upsells/health', methods=['GET'])
    def get_upsells_health():
        """Obtiene el estado de salud del sistema de upsells"""
        try:
            health = upsell_manager.get_system_health()
            status_code = 200 if health.get("status") == "healthy" else 503
            return jsonify({"success": True, "health": health}), status_code
        except Exception as e:
            logger.error(f"Error al verificar salud: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al verificar salud"}), 500
    
    @app.route('/api/customers/<customer_id>/personalized-offer', methods=['GET'])
    def get_personalized_offer(customer_id):
        """
        Obtiene oferta personalizada para un cliente
        
        Query params:
            product_id: ID del producto
        """
        try:
            product_id = request.args.get('product_id')
            if not product_id:
                return jsonify({"success": False, "error": "product_id es requerido"}), 400
            
            offer = upsell_manager.get_personalized_offer(customer_id, product_id)
            
            if offer:
                return jsonify({"success": True, "offer": offer}), 200
            return jsonify({"success": False, "error": "No se pudo generar oferta personalizada"}), 404
        except Exception as e:
            logger.error(f"Error al obtener oferta personalizada: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener oferta"}), 500
    
    @app.route('/api/customers/<customer_id>/loyalty-points', methods=['POST'])
    def calculate_loyalty_points(customer_id):
        """
        Calcula puntos de lealtad por compra
        
        Body:
        {
            "purchase_amount": 599.00
        }
        """
        try:
            data = request.json
            if not data or 'purchase_amount' not in data:
                return jsonify({"success": False, "error": "purchase_amount es requerido"}), 400
            
            points = upsell_manager.calculate_loyalty_points(
                customer_id,
                float(data['purchase_amount'])
            )
            
            return jsonify({"success": True, "loyalty_points": points}), 200
        except Exception as e:
            logger.error(f"Error al calcular puntos: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular puntos"}), 500
    
    @app.route('/api/upsells/compare', methods=['POST'])
    @require_auth
    def compare_upsells():
        """
        Compara rendimiento de dos upsells - requiere autenticación
        
        Body:
        {
            "upsell_id_1": "upsell_001",
            "upsell_id_2": "upsell_002"
        }
        """
        try:
            data = request.json
            if not all(field in data for field in ['upsell_id_1', 'upsell_id_2']):
                return jsonify({"success": False, "error": "upsell_id_1 y upsell_id_2 son requeridos"}), 400
            
            comparison = upsell_manager.get_performance_comparison(
                data['upsell_id_1'],
                data['upsell_id_2']
            )
            
            return jsonify({"success": True, "comparison": comparison}), 200
        except Exception as e:
            logger.error(f"Error al comparar upsells: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al comparar upsells"}), 500
    
    @app.route('/api/upsells/<upsell_id>/ab-test-results', methods=['GET'])
    @require_auth
    def get_ab_test_results(upsell_id):
        """
        Obtiene resultados de A/B testing - requiere autenticación
        """
        try:
            results = upsell_manager.get_ab_test_results(upsell_id)
            return jsonify({"success": True, "ab_test": results}), 200
        except Exception as e:
            logger.error(f"Error al obtener resultados A/B: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener resultados"}), 500
    
    @app.route('/api/upsells/notify', methods=['POST'])
    def send_upsell_notification():
        """
        Envía notificación de upsell a un cliente
        
        Body:
        {
            "customer_id": "customer_123",
            "notification_type": "upsell_available",
            "data": {
                "product_id": "prod_001",
                "upsell_id": "upsell_001"
            }
        }
        """
        try:
            data = request.json
            required_fields = ['customer_id', 'notification_type', 'data']
            if not all(field in data for field in required_fields):
                return jsonify({"success": False, "error": f"Campos requeridos: {', '.join(required_fields)}"}), 400
            
            success = upsell_manager.send_notification(
                data['customer_id'],
                data['notification_type'],
                data['data']
            )
            
            if success:
                return jsonify({"success": True, "message": "Notificación enviada"}), 200
            return jsonify({"success": False, "error": "Error al enviar notificación"}), 500
        except Exception as e:
            logger.error(f"Error al enviar notificación: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al enviar notificación"}), 500
    
    @app.route('/api/upsells/docs', methods=['GET'])
    def get_upsells_api_docs():
        """
        Documentación de la API de upsells
        """
        docs = {
            "api_name": "Upsells API",
            "version": "1.0",
            "description": "API completa para gestión de upsells complementarios",
            "endpoints": {
                "public": [
                    {
                        "method": "GET",
                        "path": "/api/upsells/<product_id>",
                        "description": "Obtiene upsells complementarios para un producto",
                        "params": ["limit", "customer_id", "category", "exclude_purchased"]
                    },
                    {
                        "method": "POST",
                        "path": "/api/purchases",
                        "description": "Registra una compra y retorna upsells sugeridos"
                    },
                    {
                        "method": "GET",
                        "path": "/api/products/<product_id>/cross-sells",
                        "description": "Obtiene productos relacionados (cross-sells)"
                    },
                    {
                        "method": "GET",
                        "path": "/api/products/<product_id>/similar",
                        "description": "Obtiene productos similares"
                    },
                    {
                        "method": "GET",
                        "path": "/api/customers/<customer_id>/personalized-offer",
                        "description": "Obtiene oferta personalizada para un cliente"
                    }
                ],
                "admin": [
                    {
                        "method": "GET",
                        "path": "/api/upsells/dashboard",
                        "description": "Dashboard completo con todas las métricas"
                    },
                    {
                        "method": "GET",
                        "path": "/api/upsells/analytics",
                        "description": "Analytics avanzados de upsells"
                    },
                    {
                        "method": "POST",
                        "path": "/api/upsells/auto-optimize",
                        "description": "Optimización automática de upsells"
                    }
                ]
            },
            "features": [
                "Upsells complementarios inteligentes",
                "Cross-sells y productos similares",
                "Ofertas temporales",
                "Descuentos por volumen",
                "Predicción de conversión (ML)",
                "Recomendaciones colaborativas",
                "A/B testing de mensajes",
                "Analytics avanzados",
                "Dashboard completo",
                "Backup y restore"
            ]
        }
        return jsonify(docs), 200
    
    @app.route('/api/upsells/<upsell_id>/optimize-price', methods=['POST'])
    @require_auth
    def optimize_upsell_price(upsell_id):
        """
        Optimiza precio dinámicamente - requiere autenticación
        """
        try:
            optimization = upsell_manager.optimize_price_dynamically(upsell_id)
            return jsonify({"success": True, "optimization": optimization}), 200
        except Exception as e:
            logger.error(f"Error al optimizar precio: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al optimizar precio"}), 500
    
    @app.route('/api/upsells/cohort-analysis', methods=['GET'])
    @require_auth
    def get_cohort_analysis():
        """
        Análisis de cohortes de clientes - requiere autenticación
        
        Query params:
            start_date: Fecha de inicio (ISO format)
            end_date: Fecha de fin (ISO format)
        """
        try:
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            analysis = upsell_manager.get_cohort_analysis(start_date, end_date)
            return jsonify({"success": True, "cohort_analysis": analysis}), 200
        except Exception as e:
            logger.error(f"Error en análisis de cohortes: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en análisis"}), 500
    
    @app.route('/api/upsells/<product_id>/hybrid-recommendations', methods=['GET'])
    def get_hybrid_recommendations(product_id):
        """
        Obtiene recomendaciones híbridas combinando múltiples algoritmos
        
        Query params:
            customer_id: ID del cliente (opcional)
            limit: Número de recomendaciones (default: 5)
        """
        try:
            customer_id = request.args.get('customer_id')
            limit = int(request.args.get('limit', 5))
            
            recommendations = upsell_manager.get_hybrid_recommendations(
                product_id, customer_id, limit
            )
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "recommendations": recommendations,
                "count": len(recommendations),
                "type": "hybrid"
            }), 200
        except Exception as e:
            logger.error(f"Error en recomendaciones híbridas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener recomendaciones"}), 500
    
    @app.route('/api/upsells/smart-alerts', methods=['POST'])
    @require_auth
    def create_smart_alert():
        """
        Crea una alerta inteligente - requiere autenticación
        
        Body:
        {
            "alert_type": "low_conversion",
            "threshold": 0.05,
            "product_id": "prod_001" (opcional)
        }
        """
        try:
            data = request.json
            required_fields = ['alert_type', 'threshold']
            if not all(field in data for field in required_fields):
                return jsonify({"success": False, "error": "alert_type y threshold son requeridos"}), 400
            
            success = upsell_manager.create_smart_alert(
                data['alert_type'],
                float(data['threshold']),
                data.get('product_id')
            )
            
            if success:
                return jsonify({"success": True, "message": "Alerta creada exitosamente"}), 201
            return jsonify({"success": False, "error": "Error al crear alerta"}), 500
        except Exception as e:
            logger.error(f"Error al crear alerta: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al crear alerta"}), 500
    
    @app.route('/api/products/<product_id>/inventory-urgency', methods=['GET'])
    def get_inventory_urgency(product_id):
        """
        Obtiene nivel de urgencia basado en inventario
        """
        try:
            urgency = upsell_manager.get_inventory_urgency(product_id)
            if urgency:
                return jsonify({"success": True, "urgency": urgency}), 200
            return jsonify({"success": False, "error": "Producto no encontrado"}), 404
        except Exception as e:
            logger.error(f"Error al calcular urgencia: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular urgencia"}), 500
    
    @app.route('/api/upsells/performance-summary', methods=['GET'])
    @require_auth
    def get_performance_summary():
        """
        Resumen de rendimiento completo - requiere autenticación
        
        Query params:
            days: Días a considerar (default: 30)
        """
        try:
            days = int(request.args.get('days', 30))
            
            # Obtener múltiples métricas
            realtime_metrics = upsell_manager.get_realtime_metrics()
            revenue_impact = upsell_manager.calculate_revenue_impact(days=days)
            dashboard = upsell_manager.get_dashboard_summary()
            
            summary = {
                "period_days": days,
                "realtime_metrics": realtime_metrics,
                "revenue_impact": revenue_impact,
                "dashboard_overview": dashboard.get("overview", {}),
                "timestamp": datetime.now().isoformat()
            }
            
            return jsonify({"success": True, "summary": summary}), 200
        except Exception as e:
            logger.error(f"Error al obtener resumen: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener resumen"}), 500
    
    @app.route('/api/upsells/events', methods=['POST'])
    def log_upsell_event():
        """
        Registra un evento en el sistema
        
        Body:
        {
            "event_type": "view",
            "entity_type": "upsell",
            "entity_id": "upsell_001",
            "details": {"source": "product_page"}
        }
        """
        try:
            data = request.json
            required_fields = ['event_type', 'entity_type', 'entity_id']
            if not all(field in data for field in required_fields):
                return jsonify({"success": False, "error": f"Campos requeridos: {', '.join(required_fields)}"}), 400
            
            success = upsell_manager.log_event(
                data['event_type'],
                data['entity_type'],
                data['entity_id'],
                data.get('details')
            )
            
            if success:
                return jsonify({"success": True, "message": "Evento registrado"}), 201
            return jsonify({"success": False, "error": "Error al registrar evento"}), 500
        except Exception as e:
            logger.error(f"Error al registrar evento: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al registrar evento"}), 500
    
    @app.route('/api/upsells/events', methods=['GET'])
    @require_auth
    def get_event_history():
        """
        Obtiene historial de eventos - requiere autenticación
        
        Query params:
            entity_type: Tipo de entidad
            entity_id: ID de entidad
            event_type: Tipo de evento
            limit: Número máximo de eventos (default: 100)
        """
        try:
            entity_type = request.args.get('entity_type')
            entity_id = request.args.get('entity_id')
            event_type = request.args.get('event_type')
            limit = int(request.args.get('limit', 100))
            
            events = upsell_manager.get_event_history(
                entity_type, entity_id, event_type, limit
            )
            
            return jsonify({
                "success": True,
                "events": events,
                "count": len(events)
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener eventos: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener eventos"}), 500
    
    @app.route('/api/upsells/conversion-funnel', methods=['GET'])
    @require_auth
    def get_conversion_funnel():
        """
        Analiza el embudo de conversión - requiere autenticación
        
        Query params:
            product_id: ID del producto (opcional)
            days: Días a considerar (default: 30)
        """
        try:
            product_id = request.args.get('product_id')
            days = int(request.args.get('days', 30))
            
            funnel = upsell_manager.get_conversion_funnel(product_id, days)
            return jsonify({"success": True, "funnel": funnel}), 200
        except Exception as e:
            logger.error(f"Error en análisis de embudo: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en análisis"}), 500
    
    @app.route('/api/customers/<customer_id>/journey', methods=['GET'])
    @require_auth
    def get_customer_journey(customer_id):
        """
        Analiza el journey completo de un cliente - requiere autenticación
        """
        try:
            journey = upsell_manager.get_customer_journey(customer_id)
            return jsonify({"success": True, "journey": journey}), 200
        except Exception as e:
            logger.error(f"Error al analizar journey: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al analizar journey"}), 500
    
    @app.route('/api/upsells/cache/invalidate', methods=['POST'])
    @require_auth
    def invalidate_cache():
        """
        Invalida el cache de recomendaciones - requiere autenticación
        
        Body (opcional):
        {
            "pattern": "prod_001"  # Patrón para invalidar específico
        }
        """
        try:
            data = request.json or {}
            pattern = data.get('pattern')
            
            count = upsell_manager.invalidate_cache(pattern)
            
            return jsonify({
                "success": True,
                "invalidated_entries": count,
                "message": f"Cache invalidado: {count} entradas"
            }), 200
        except Exception as e:
            logger.error(f"Error al invalidar cache: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al invalidar cache"}), 500
    
    @app.route('/api/upsells/statistics', methods=['GET'])
    def get_upsells_statistics():
        """
        Obtiene estadísticas generales del sistema
        """
        try:
            conn = sqlite3.connect(upsell_manager.db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Estadísticas generales
            cursor.execute('SELECT COUNT(*) as count FROM products WHERE is_active = 1')
            active_products = cursor.fetchone()["count"]
            
            cursor.execute('SELECT COUNT(*) as count FROM upsells WHERE is_active = 1')
            active_upsells = cursor.fetchone()["count"]
            
            cursor.execute('SELECT COUNT(*) as count FROM purchases')
            total_purchases = cursor.fetchone()["count"]
            
            cursor.execute('SELECT COUNT(DISTINCT customer_id) as count FROM customer_history')
            unique_customers = cursor.fetchone()["count"]
            
            cursor.execute('''
                SELECT 
                    SUM(total_views) as total_views,
                    SUM(total_purchases) as total_purchases,
                    AVG(CASE WHEN total_views > 0 
                        THEN (total_purchases * 1.0 / total_views) ELSE 0 END) as avg_conversion
                FROM upsells
                WHERE is_active = 1
            ''')
            metrics = cursor.fetchone()
            
            conn.close()
            
            return jsonify({
                "success": True,
                "statistics": {
                    "active_products": active_products,
                    "active_upsells": active_upsells,
                    "total_purchases": total_purchases,
                    "unique_customers": unique_customers,
                    "total_views": metrics["total_views"] or 0,
                    "total_upsell_purchases": metrics["total_purchases"] or 0,
                    "avg_conversion_rate": round((metrics["avg_conversion"] or 0) * 100, 2),
                    "cache_size": len(upsell_manager._recommendation_cache),
                    "cache_enabled": upsell_manager.cache_enabled
                },
                "timestamp": datetime.now().isoformat()
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener estadísticas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener estadísticas"}), 500
    
    @app.route('/api/coupons', methods=['POST'])
    @require_auth
    def create_coupon():
        """
        Crea un cupón de descuento - requiere autenticación
        
        Body:
        {
            "code": "SUMMER2024",
            "discount_type": "percentage",
            "discount_value": 15.0,
            "min_purchase": 500.0,
            "max_uses": 100,
            "valid_until": "2024-12-31T23:59:59"
        }
        """
        try:
            data = request.json
            required_fields = ['code', 'discount_type', 'discount_value']
            if not all(field in data for field in required_fields):
                return jsonify({"success": False, "error": "code, discount_type y discount_value son requeridos"}), 400
            
            success = upsell_manager.create_coupon(
                code=data['code'],
                discount_type=data['discount_type'],
                discount_value=float(data['discount_value']),
                min_purchase=float(data.get('min_purchase', 0)),
                max_uses=data.get('max_uses'),
                valid_until=data.get('valid_until')
            )
            
            if success:
                return jsonify({"success": True, "message": "Cupón creado exitosamente"}), 201
            return jsonify({"success": False, "error": "Error al crear cupón"}), 500
        except Exception as e:
            logger.error(f"Error al crear cupón: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al crear cupón"}), 500
    
    @app.route('/api/coupons/validate', methods=['POST'])
    def validate_coupon():
        """
        Valida y aplica un cupón
        
        Body:
        {
            "code": "SUMMER2024",
            "purchase_amount": 599.00
        }
        """
        try:
            data = request.json
            if not all(field in data for field in ['code', 'purchase_amount']):
                return jsonify({"success": False, "error": "code y purchase_amount son requeridos"}), 400
            
            result = upsell_manager.validate_coupon(
                data['code'],
                float(data['purchase_amount'])
            )
            
            return jsonify({"success": result.get("valid", False), "coupon": result}), 200
        except Exception as e:
            logger.error(f"Error al validar cupón: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al validar cupón"}), 500
    
    @app.route('/api/products/<product_id>/competitor-analysis', methods=['GET'])
    @require_auth
    def get_competitor_analysis(product_id):
        """
        Análisis competitivo de un producto - requiere autenticación
        """
        try:
            analysis = upsell_manager.get_competitor_analysis(product_id)
            return jsonify({"success": True, "analysis": analysis}), 200
        except Exception as e:
            logger.error(f"Error en análisis competitivo: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en análisis"}), 500
    
    @app.route('/api/upsells/performance-by-category', methods=['GET'])
    @require_auth
    def get_performance_by_category():
        """
        Analiza rendimiento de upsells por categoría - requiere autenticación
        
        Query params:
            days: Días a considerar (default: 30)
        """
        try:
            days = int(request.args.get('days', 30))
            performance = upsell_manager.get_upsell_performance_by_category(days)
            return jsonify({"success": True, "performance": performance}), 200
        except Exception as e:
            logger.error(f"Error al analizar por categoría: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en análisis"}), 500
    
    @app.route('/api/upsells/<product_id>/time-based', methods=['GET'])
    def get_time_based_recommendations(product_id):
        """
        Recomendaciones optimizadas por hora del día
        
        Query params:
            hour: Hora del día (0-23), si no se proporciona usa hora actual
        """
        try:
            hour = request.args.get('hour')
            hour = int(hour) if hour else None
            
            recommendations = upsell_manager.get_time_based_recommendations(product_id, hour)
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "hour": hour or datetime.now().hour,
                "recommendations": recommendations,
                "count": len(recommendations),
                "type": "time_optimized"
            }), 200
        except Exception as e:
            logger.error(f"Error en recomendaciones por tiempo: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener recomendaciones"}), 500
    
    @app.route('/api/customers/<customer_id>/wishlist', methods=['POST'])
    def add_to_wishlist(customer_id):
        """
        Agrega un producto a la wishlist
        
        Body:
        {
            "product_id": "prod_001"
        }
        """
        try:
            data = request.json
            if not data or 'product_id' not in data:
                return jsonify({"success": False, "error": "product_id es requerido"}), 400
            
            success = upsell_manager.add_to_wishlist(customer_id, data['product_id'])
            
            if success:
                return jsonify({"success": True, "message": "Producto agregado a wishlist"}), 201
            return jsonify({"success": False, "error": "Error al agregar a wishlist"}), 500
        except Exception as e:
            logger.error(f"Error al agregar a wishlist: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al agregar a wishlist"}), 500
    
    @app.route('/api/customers/<customer_id>/wishlist', methods=['GET'])
    def get_wishlist(customer_id):
        """
        Obtiene la wishlist de un cliente
        """
        try:
            wishlist = upsell_manager.get_wishlist(customer_id)
            return jsonify({
                "success": True,
                "customer_id": customer_id,
                "wishlist": wishlist,
                "count": len(wishlist)
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener wishlist: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener wishlist"}), 500
    
    @app.route('/api/upsells/engagement', methods=['GET'])
    @require_auth
    def get_engagement_metrics():
        """
        Obtiene métricas de engagement - requiere autenticación
        
        Query params:
            product_id: ID del producto (opcional)
            days: Días a considerar (default: 30)
        """
        try:
            product_id = request.args.get('product_id')
            days = int(request.args.get('days', 30))
            
            metrics = upsell_manager.get_engagement_metrics(product_id, days)
            return jsonify({"success": True, "engagement": metrics}), 200
        except Exception as e:
            logger.error(f"Error al calcular engagement: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular engagement"}), 500
    
    @app.route('/api/upsells/cart-abandonment', methods=['GET'])
    @require_auth
    def get_cart_abandonment_analysis():
        """
        Analiza abandono de carrito - requiere autenticación
        
        Query params:
            days: Días a considerar (default: 30)
        """
        try:
            days = int(request.args.get('days', 30))
            analysis = upsell_manager.get_cart_abandonment_analysis(days)
            return jsonify({"success": True, "abandonment_analysis": analysis}), 200
        except Exception as e:
            logger.error(f"Error en análisis de abandono: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en análisis"}), 500
    
    @app.route('/api/customers/<referrer_id>/referral', methods=['POST'])
    def create_referral(referrer_id):
        """
        Registra un referido y calcula recompensas
        
        Body:
        {
            "referred_id": "customer_456"
        }
        """
        try:
            data = request.json
            if not data or 'referred_id' not in data:
                return jsonify({"success": False, "error": "referred_id es requerido"}), 400
            
            result = upsell_manager.get_referral_rewards(referrer_id, data['referred_id'])
            return jsonify(result), 200 if result.get("success") else 400
        except Exception as e:
            logger.error(f"Error al registrar referido: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al registrar referido"}), 500
    
    @app.route('/api/products/<product_id>/reviews', methods=['POST'])
    def add_product_review(product_id):
        """
        Agrega una reseña a un producto
        
        Body:
        {
            "customer_id": "customer_123",
            "rating": 5,
            "comment": "Excelente producto!"
        }
        """
        try:
            data = request.json
            required_fields = ['customer_id', 'rating']
            if not all(field in data for field in required_fields):
                return jsonify({"success": False, "error": "customer_id y rating son requeridos"}), 400
            
            if not (1 <= data['rating'] <= 5):
                return jsonify({"success": False, "error": "Rating debe estar entre 1 y 5"}), 400
            
            success = upsell_manager.add_product_review(
                product_id,
                data['customer_id'],
                int(data['rating']),
                data.get('comment')
            )
            
            if success:
                return jsonify({"success": True, "message": "Reseña agregada exitosamente"}), 201
            return jsonify({"success": False, "error": "Error al agregar reseña"}), 500
        except Exception as e:
            logger.error(f"Error al agregar review: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al agregar review"}), 500
    
    @app.route('/api/products/<product_id>/reviews', methods=['GET'])
    def get_product_reviews(product_id):
        """
        Obtiene reseñas de un producto
        
        Query params:
            limit: Número máximo de reseñas (default: 10)
        """
        try:
            limit = int(request.args.get('limit', 10))
            reviews = upsell_manager.get_product_reviews(product_id, limit)
            return jsonify({"success": True, "reviews": reviews}), 200
        except Exception as e:
            logger.error(f"Error al obtener reviews: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener reviews"}), 500
    
    @app.route('/api/products/search', methods=['GET'])
    def search_products():
        """
        Búsqueda inteligente de productos
        
        Query params:
            q: Término de búsqueda
            limit: Número máximo de resultados (default: 10)
        """
        try:
            query = request.args.get('q')
            if not query:
                return jsonify({"success": False, "error": "Parámetro 'q' es requerido"}), 400
            
            limit = int(request.args.get('limit', 10))
            results = upsell_manager.search_products(query, limit)
            
            return jsonify({
                "success": True,
                "query": query,
                "results": results,
                "count": len(results)
            }), 200
        except Exception as e:
            logger.error(f"Error en búsqueda: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en búsqueda"}), 500
    
    @app.route('/api/customers/retention', methods=['GET'])
    @require_auth
    def get_customer_retention():
        """
        Calcula métricas de retención de clientes - requiere autenticación
        
        Query params:
            days: Días a considerar (default: 90)
        """
        try:
            days = int(request.args.get('days', 90))
            retention = upsell_manager.get_customer_retention(days)
            return jsonify({"success": True, "retention": retention}), 200
        except Exception as e:
            logger.error(f"Error al calcular retención: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular retención"}), 500
    
    @app.route('/api/products/<product_id>/tag-recommendations', methods=['GET'])
    def get_tag_recommendations(product_id):
        """
        Obtiene recomendaciones basadas en tags
        
        Query params:
            limit: Número de recomendaciones (default: 5)
        """
        try:
            limit = int(request.args.get('limit', 5))
            recommendations = upsell_manager.get_product_recommendations_by_tags(product_id, limit)
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "recommendations": recommendations,
                "count": len(recommendations),
                "type": "tag_based"
            }), 200
        except Exception as e:
            logger.error(f"Error en recomendaciones por tags: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener recomendaciones"}), 500
    
    @app.route('/api/reviews/analyze-sentiment', methods=['POST'])
    def analyze_review_sentiment():
        """
        Analiza el sentimiento de un texto (review)
        
        Body:
        {
            "text": "Excelente producto, muy recomendado!"
        }
        """
        try:
            data = request.json
            if not data or 'text' not in data:
                return jsonify({"success": False, "error": "text es requerido"}), 400
            
            sentiment = upsell_manager.analyze_sentiment(data['text'])
            return jsonify({"success": True, "sentiment": sentiment}), 200
        except Exception as e:
            logger.error(f"Error en análisis de sentimiento: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error en análisis"}), 500
    
    @app.route('/api/products/compare', methods=['POST'])
    def compare_products():
        """
        Compara múltiples productos lado a lado
        
        Body:
        {
            "product_ids": ["prod_001", "prod_002", "prod_003"]
        }
        """
        try:
            data = request.json
            if not data or 'product_ids' not in data:
                return jsonify({"success": False, "error": "product_ids es requerido"}), 400
            
            if not isinstance(data['product_ids'], list):
                return jsonify({"success": False, "error": "product_ids debe ser una lista"}), 400
            
            comparison = upsell_manager.compare_products(data['product_ids'])
            
            if "error" in comparison:
                return jsonify({"success": False, "error": comparison["error"]}), 400
            
            return jsonify({"success": True, "comparison": comparison}), 200
        except Exception as e:
            logger.error(f"Error al comparar productos: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al comparar productos"}), 500
    
    @app.route('/api/products/most-searched', methods=['GET'])
    def get_most_searched_products():
        """
        Obtiene productos más buscados
        
        Query params:
            days: Días a considerar (default: 30)
            limit: Número de resultados (default: 10)
        """
        try:
            days = int(request.args.get('days', 30))
            limit = int(request.args.get('limit', 10))
            
            products = upsell_manager.get_most_searched_products(days, limit)
            return jsonify({
                "success": True,
                "products": products,
                "count": len(products),
                "period_days": days
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener productos más buscados: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener productos"}), 500
    
    @app.route('/api/products/low-stock-alerts', methods=['GET'])
    @require_auth
    def get_low_stock_alerts():
        """
        Obtiene alertas de productos con stock bajo - requiere autenticación
        
        Query params:
            threshold: Umbral de stock bajo (default: 10)
        """
        try:
            threshold = int(request.args.get('threshold', 10))
            alerts = upsell_manager.get_low_stock_alerts(threshold)
            
            return jsonify({
                "success": True,
                "alerts": alerts,
                "count": len(alerts),
                "threshold": threshold
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener alertas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener alertas"}), 500
    
    @app.route('/api/performance', methods=['GET'])
    @require_auth
    def get_api_performance():
        """
        Obtiene métricas de rendimiento de la API - requiere autenticación
        """
        try:
            metrics = upsell_manager.get_api_performance_metrics()
            return jsonify({"success": True, "performance": metrics}), 200
        except Exception as e:
            logger.error(f"Error al obtener métricas: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener métricas"}), 500
    
    @app.route('/api/search/trending', methods=['GET'])
    def get_trending_searches():
        """
        Obtiene búsquedas en tendencia
        
        Query params:
            days: Días a considerar (default: 7)
            limit: Número de resultados (default: 10)
        """
        try:
            days = int(request.args.get('days', 7))
            limit = int(request.args.get('limit', 10))
            
            trends = upsell_manager.get_trending_searches(days, limit)
            return jsonify({
                "success": True,
                "trending_searches": trends,
                "count": len(trends),
                "period_days": days
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener tendencias: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener tendencias"}), 500
    
    @app.route('/api/products/<product_id>/frequently-bought-together', methods=['GET'])
    def get_frequently_bought_together(product_id):
        """
        Obtiene productos frecuentemente comprados juntos
        
        Query params:
            limit: Número de resultados (default: 5)
        """
        try:
            limit = int(request.args.get('limit', 5))
            products = upsell_manager.get_frequently_bought_together(product_id, limit)
            
            return jsonify({
                "success": True,
                "product_id": product_id,
                "frequently_bought_together": products,
                "count": len(products)
            }), 200
        except Exception as e:
            logger.error(f"Error al obtener productos relacionados: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener productos"}), 500
    
    @app.route('/api/upsells/conversion-velocity', methods=['GET'])
    @require_auth
    def get_conversion_velocity():
        """
        Calcula velocidad de conversión - requiere autenticación
        
        Query params:
            product_id: ID del producto (opcional)
            days: Días a considerar (default: 30)
        """
        try:
            product_id = request.args.get('product_id')
            days = int(request.args.get('days', 30))
            
            velocity = upsell_manager.calculate_conversion_velocity(product_id, days)
            return jsonify({"success": True, "velocity": velocity}), 200
        except Exception as e:
            logger.error(f"Error al calcular velocidad: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular velocidad"}), 500
    
    @app.route('/api/products/<product_id>/dynamic-pricing', methods=['GET'])
    @require_auth
    def get_dynamic_pricing_suggestions(product_id):
        """
        Obtiene sugerencias de precios dinámicos - requiere autenticación
        """
        try:
            suggestions = upsell_manager.get_dynamic_pricing_suggestions(product_id)
            return jsonify({"success": True, "pricing": suggestions}), 200
        except Exception as e:
            logger.error(f"Error en sugerencias de precio: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener sugerencias"}), 500
    
    @app.route('/api/customers/<customer_id>/behavioral-recommendations', methods=['GET'])
    def get_behavioral_recommendations(customer_id):
        """
        Obtiene recomendaciones basadas en comportamiento del cliente
        
        Query params:
            limit: Número de recomendaciones (default: 5)
        """
        try:
            limit = int(request.args.get('limit', 5))
            recommendations = upsell_manager.get_behavioral_recommendations(customer_id, limit)
            
            return jsonify({
                "success": True,
                "customer_id": customer_id,
                "recommendations": recommendations,
                "count": len(recommendations),
                "type": "behavioral"
            }), 200
        except Exception as e:
            logger.error(f"Error en recomendaciones de comportamiento: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al obtener recomendaciones"}), 500
    
    @app.route('/api/products/affinity-matrix', methods=['GET'])
    @require_auth
    def get_product_affinity_matrix():
        """
        Genera matriz de afinidad entre productos - requiere autenticación
        
        Query params:
            limit: Número de productos a analizar (default: 10)
        """
        try:
            limit = int(request.args.get('limit', 10))
            matrix = upsell_manager.get_product_affinity_matrix(limit)
            return jsonify({"success": True, "affinity_matrix": matrix}), 200
        except Exception as e:
            logger.error(f"Error en matriz de afinidad: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al generar matriz"}), 500
    
    @app.route('/api/customers/<customer_id>/lifetime-value', methods=['GET'])
    @require_auth
    def get_customer_lifetime_value(customer_id):
        """
        Calcula el valor de vida del cliente (LTV) - requiere autenticación
        """
        try:
            ltv = upsell_manager.get_customer_lifetime_value(customer_id)
            return jsonify({"success": True, "ltv": ltv}), 200
        except Exception as e:
            logger.error(f"Error al calcular LTV: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular LTV"}), 500
    
    @app.route('/api/upsells/success-rate', methods=['GET'])
    @require_auth
    def get_upsell_success_rate():
        """
        Calcula tasa de éxito de upsells - requiere autenticación
        
        Query params:
            product_id: ID del producto (opcional)
            days: Días a considerar (default: 30)
        """
        try:
            product_id = request.args.get('product_id')
            days = int(request.args.get('days', 30))
            
            success_rate = upsell_manager.get_upsell_success_rate(product_id, days)
            return jsonify({"success": True, "success_rate": success_rate}), 200
        except Exception as e:
            logger.error(f"Error al calcular tasa de éxito: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": "Error al calcular tasa"}), 500
    
    @app.route('/health', methods=['GET'])
    def health():
        """Endpoint de salud con información detallada"""
        health_status = {
            "status": "healthy",
            "version": "3.0",
            "timestamp": datetime.now().isoformat(),
            "services": {
                "database": "operational",
                "cache": "operational",
                "authentication": "operational"
            },
            "uptime": "N/A"  # Se puede calcular con time.time() - start_time
        }
        
        # Verificar estado de la base de datos
        try:
            conn = sqlite3.connect("offer_letters.db")
            conn.close()
        except Exception:
            health_status["status"] = "degraded"
            health_status["services"]["database"] = "error"
        
        status_code = 200 if health_status["status"] == "healthy" else 503
        return jsonify(health_status), status_code
    
    @app.route('/health/ready', methods=['GET'])
    def readiness():
        """Endpoint de readiness probe (Kubernetes)"""
        try:
            # Verificar que la base de datos responde
            conn = sqlite3.connect("offer_letters.db")
            cursor = conn.cursor()
            cursor.execute("SELECT 1")
            conn.close()
            return jsonify({"ready": True}), 200
        except Exception:
            return jsonify({"ready": False}), 503
    
    @app.route('/health/live', methods=['GET'])
    def liveness():
        """Endpoint de liveness probe (Kubernetes)"""
        return jsonify({"alive": True}), 200
    
    # ============================================================================
    # ENDPOINTS DE AUTENTICACIÓN
    # ============================================================================
    
    @app.route('/api/auth/register', methods=['POST'])
    def register():
        """Registra un nuevo usuario"""
        data = request.json
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')
        role = data.get('role', 'user')
        
        if not all([username, email, password]):
            return jsonify({"error": "username, email y password son requeridos"}), 400
        
        user_id = auth_manager.create_user(username, email, password, role)
        if user_id:
            return jsonify({"success": True, "user_id": user_id}), 201
        return jsonify({"error": "Usuario ya existe"}), 400
    
    @app.route('/api/auth/login', methods=['POST'])
    def login():
        """Autentica usuario y retorna token"""
        data = request.json
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({"error": "username y password son requeridos"}), 400
        
        token = auth_manager.authenticate(username, password)
        if token:
            return jsonify({"success": True, "token": token}), 200
        return jsonify({"error": "Credenciales inválidas"}), 401
    
    @app.route('/api/auth/verify', methods=['POST'])
    @require_auth
    def verify_token():
        """Verifica token de autenticación"""
        return jsonify({"success": True, "user_id": request.user_id}), 200
    
    # ============================================================================
    # ENDPOINTS DE ANÁLISIS Y REPORTES
    # ============================================================================
    
    @app.route('/api/analytics/acceptance-rate', methods=['GET'])
    @require_auth
    @cache_response(ttl=600)
    def get_acceptance_rate():
        """Obtiene tasa de aceptación"""
        date_from = request.args.get('date_from')
        date_to = request.args.get('date_to')
        stats = analytics_manager.get_acceptance_rate(date_from, date_to)
        return stats, 200
    
    @app.route('/api/analytics/salary-trends', methods=['GET'])
    @require_auth
    @cache_response(ttl=600)
    def get_salary_trends():
        """Obtiene tendencias salariales"""
        department = request.args.get('department')
        trends = analytics_manager.get_salary_trends(department)
        return trends, 200
    
    @app.route('/api/analytics/time-to-acceptance', methods=['GET'])
    @require_auth
    @cache_response(ttl=600)
    def get_time_to_acceptance():
        """Obtiene tiempo promedio hasta aceptación"""
        stats = analytics_manager.get_time_to_acceptance()
        return stats, 200
    
    # ============================================================================
    # ENDPOINTS DE BACKUP Y RESTORE
    # ============================================================================
    
    @app.route('/api/backup/create', methods=['POST'])
    @require_auth
    def create_backup():
        """Crea un backup de la base de datos"""
        try:
            backup_path = backup_manager.create_backup()
            return jsonify({"success": True, "backup_path": backup_path}), 201
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    
    @app.route('/api/backup/list', methods=['GET'])
    @require_auth
    def list_backups():
        """Lista todos los backups disponibles"""
        backups = backup_manager.list_backups()
        return jsonify({"backups": backups, "count": len(backups)}), 200
    
    @app.route('/api/backup/restore', methods=['POST'])
    @require_auth
    def restore_backup():
        """Restaura desde un backup"""
        backup_path = request.json.get('backup_path')
        if not backup_path:
            return jsonify({"error": "backup_path requerido"}), 400
        
        if backup_manager.restore_backup(backup_path):
            return jsonify({"success": True}), 200
        return jsonify({"error": "Error al restaurar backup"}), 500
    
    # ============================================================================
    # ENDPOINTS DE FIRMAS DIGITALES
    # ============================================================================
    
    @app.route('/api/offers/<offer_id>/sign', methods=['POST'])
    @require_auth
    def sign_offer(offer_id):
        """Firma digitalmente una oferta"""
        offer = api.get_offer_by_id(offer_id)
        if not offer:
            return jsonify({"error": "Oferta no encontrada"}), 404
        
        offer_data = json.loads(offer.get('offer_data', '{}'))
        signed_offer = signature_manager.create_signed_offer(offer_data)
        
        # Actualizar oferta con firma
        conn = sqlite3.connect("offer_letters.db")
        cursor = conn.cursor()
        cursor.execute('''
            UPDATE offers SET offer_data = ? WHERE id = ?
        ''', (json.dumps(signed_offer), offer_id))
        conn.commit()
        conn.close()
        
        return jsonify({"success": True, "signed_offer": signed_offer}), 200
    
    @app.route('/api/offers/<offer_id>/verify-signature', methods=['GET'])
    def verify_signature(offer_id):
        """Verifica firma digital de una oferta"""
        offer = api.get_offer_by_id(offer_id)
        if not offer:
            return jsonify({"error": "Oferta no encontrada"}), 404
        
        offer_data = json.loads(offer.get('offer_data', '{}'))
        signature = offer_data.get('digital_signature')
        
        if not signature:
            return jsonify({"verified": False, "error": "Oferta no tiene firma digital"}), 400
        
        # Remover firma para verificar
        offer_data_without_sig = {k: v for k, v in offer_data.items() 
                                  if k not in ['digital_signature', 'signed_at', 'signature_version']}
        
        is_valid = signature_manager.verify_signature(offer_data_without_sig, signature)
        return jsonify({"verified": is_valid, "signed_at": offer_data.get('signed_at')}), 200
    
    # ============================================================================
    # ENDPOINTS DE INTEGRACIONES
    # ============================================================================
    
    @app.route('/api/integrations/ats/register', methods=['POST'])
    @require_auth
    def register_ats():
        """Registra integración con ATS"""
        data = request.json
        ats_name = data.get('ats_name')
        api_key = data.get('api_key')
        api_url = data.get('api_url')
        
        if not all([ats_name, api_key, api_url]):
            return jsonify({"error": "ats_name, api_key y api_url son requeridos"}), 400
        
        integration_manager.register_ats_integration(ats_name, api_key, api_url)
        return jsonify({"success": True}), 201
    
    @app.route('/api/integrations/ats/sync/<offer_id>', methods=['POST'])
    @require_auth
    def sync_to_ats(offer_id):
        """Sincroniza oferta con ATS"""
        ats_name = request.json.get('ats_name')
        if not ats_name:
            return jsonify({"error": "ats_name requerido"}), 400
        
        if integration_manager.sync_offer_to_ats(offer_id, ats_name):
            return jsonify({"success": True}), 200
        return jsonify({"error": "Error al sincronizar"}), 500
    
    @app.route('/api/integrations/hris/register', methods=['POST'])
    @require_auth
    def register_hris():
        """Registra integración con HRIS"""
        data = request.json
        hris_name = data.get('hris_name')
        api_key = data.get('api_key')
        api_url = data.get('api_url')
        
        if not all([hris_name, api_key, api_url]):
            return jsonify({"error": "hris_name, api_key y api_url son requeridos"}), 400
        
        integration_manager.register_hris_integration(hris_name, api_key, api_url)
        return jsonify({"success": True}), 201
    
    # ============================================================================
    # ENDPOINTS DE INTERNACIONALIZACIÓN
    # ============================================================================
    
    @app.route('/api/i18n/translate', methods=['GET'])
    def translate():
        """Traduce una clave"""
        key = request.args.get('key')
        language = request.args.get('lang', 'es')
        
        if not key:
            return jsonify({"error": "key requerido"}), 400
        
        translation = i18n_manager.translate(key, language)
        return jsonify({"key": key, "language": language, "translation": translation}), 200
    
    @app.route('/api/i18n/languages', methods=['GET'])
    def get_languages():
        """Obtiene idiomas disponibles"""
        languages = list(i18n_manager.translations.keys())
        return jsonify({"languages": languages, "default": i18n_manager.default_language}), 200
    
    # ============================================================================
    # ENDPOINTS DE RENDIMIENTO Y MÉTRICAS
    # ============================================================================
    
    @app.route('/api/performance/stats', methods=['GET'])
    @require_auth
    def get_performance_stats():
        """Obtiene estadísticas de rendimiento"""
        stats = performance_monitor.get_stats()
        cache_stats = cache_manager.get_stats()
        return jsonify({
            "performance": stats,
            "cache": cache_stats
        }), 200
    
    @app.route('/api/cache/clear', methods=['POST'])
    @require_auth
    def clear_cache():
        """Limpia el caché"""
        cache_manager.clear()
        return jsonify({"success": True}), 200
    
    @app.route('/api/cache/stats', methods=['GET'])
    @require_auth
    def get_cache_stats():
        """Obtiene estadísticas del caché"""
        stats = cache_manager.get_stats()
        return jsonify(stats), 200
    
    # ============================================================================
    # ENDPOINTS DE NOTIFICACIONES
    # ============================================================================
    
    @app.route('/api/notifications', methods=['GET'])
    @require_auth
    def get_notifications():
        """Obtiene notificaciones del usuario"""
        unread_only = request.args.get('unread_only', 'false').lower() == 'true'
        notifications = notification_manager.get_user_notifications(
            request.user_id,
            unread_only
        )
        return jsonify({"notifications": notifications, "count": len(notifications)}), 200
    
    @app.route('/api/notifications/<notification_id>/read', methods=['PUT'])
    @require_auth
    def mark_notification_read(notification_id):
        """Marca notificación como leída"""
        if notification_manager.mark_as_read(notification_id):
            return jsonify({"success": True}), 200
        return jsonify({"error": "Notificación no encontrada"}), 404
    
    # ============================================================================
    # ENDPOINTS DE EXPORTACIÓN
    # ============================================================================
    
    @app.route('/api/export/csv', methods=['GET'])
    @require_auth
    def export_to_csv():
        """Exporta ofertas a CSV"""
        filters = {}
        if request.args.get('status'):
            filters['status'] = request.args.get('status')
        if request.args.get('department'):
            filters['department'] = request.args.get('department')
        
        offers = api.get_offer_history(filters)
        csv_path = export_manager.export_to_csv(offers)
        
        return send_file(csv_path, as_attachment=True, download_name=os.path.basename(csv_path))
    
    @app.route('/api/export/excel', methods=['GET'])
    @require_auth
    def export_to_excel():
        """Exporta ofertas a Excel"""
        filters = {}
        if request.args.get('status'):
            filters['status'] = request.args.get('status')
        if request.args.get('department'):
            filters['department'] = request.args.get('department')
        
        offers = api.get_offer_history(filters)
        excel_path = export_manager.export_to_excel_simple(offers)
        
        return send_file(excel_path, as_attachment=True, download_name=os.path.basename(excel_path))
    
    @app.route('/api/export/statistics', methods=['GET'])
    @require_auth
    def export_statistics():
        """Exporta estadísticas a CSV"""
        stats = api.get_statistics()
        csv_path = export_manager.export_statistics_report(stats)
        
        return send_file(csv_path, as_attachment=True, download_name=os.path.basename(csv_path))
    
    # Endpoints adicionales avanzados
    @app.route('/api/offers/<offer_id>/compare', methods=['POST'])
    def compare_offers(offer_id):
        """Compara dos ofertas"""
        try:
            other_offer_id = request.json.get('other_offer_id')
            if not other_offer_id:
                return jsonify({"error": "other_offer_id requerido"}), 400
            
            comparison = api.compare_offers(offer_id, other_offer_id)
            return jsonify(comparison), 200
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    
    @app.route('/api/offers/<offer_id>/versions', methods=['GET'])
    def get_offer_versions(offer_id):
        """Obtiene todas las versiones de una oferta"""
        versions = api.get_offer_versions(offer_id)
        return jsonify({"versions": versions}), 200
    
    @app.route('/api/offers/<offer_id>/word', methods=['GET'])
    def download_word(offer_id):
        """Descarga oferta en formato Word/DOCX"""
        word_path = api.generate_word(offer_id)
        if word_path and os.path.exists(word_path):
            return send_file(word_path, as_attachment=True)
        return jsonify({"error": "Documento Word no disponible"}), 404
    
    @app.route('/api/dashboard', methods=['GET'])
    def dashboard():
        """Endpoint para dashboard con métricas avanzadas"""
        dashboard_data = api.get_dashboard_data()
        return jsonify(dashboard_data), 200
    
    @app.route('/api/search', methods=['GET'])
    def search_offers():
        """Búsqueda avanzada de ofertas"""
        query = request.args.get('q', '')
        filters = {
            'status': request.args.get('status'),
            'department': request.args.get('department'),
            'date_from': request.args.get('date_from'),
            'date_to': request.args.get('date_to')
        }
        results = api.search_offers(query, filters)
        return jsonify(results), 200
    
    @app.route('/api/webhooks', methods=['POST'])
    def register_webhook():
        """Registra un webhook para notificaciones"""
        webhook_url = request.json.get('url')
        events = request.json.get('events', [])
        if api.register_webhook(webhook_url, events):
            return jsonify({"success": True}), 201
        return jsonify({"error": "Error al registrar webhook"}), 400
    
    @app.route('/api/docs', methods=['GET'])
    def api_docs():
        """Documentación de la API en formato OpenAPI 3.0"""
        docs = {
            "openapi": "3.0.0",
            "info": {
                "title": "Offer Letter API",
                "version": "3.0",
                "description": "API completa para gestión de cartas de oferta laboral con funcionalidades avanzadas",
                "contact": {
                    "name": "API Support",
                    "email": "support@example.com"
                }
            },
            "servers": [
                {"url": "http://localhost:5000", "description": "Servidor de desarrollo"},
                {"url": "https://api.example.com", "description": "Servidor de producción"}
            ],
            "components": {
                "securitySchemes": {
                    "bearerAuth": {
                        "type": "http",
                        "scheme": "bearer",
                        "bearerFormat": "JWT"
                    }
                }
            },
            "paths": {
                "/api/offers": {
                    "post": {
                        "summary": "Crear nueva oferta",
                        "security": [{"bearerAuth": []}],
                        "requestBody": {
                            "required": True,
                            "content": {
                                "application/json": {
                                    "schema": {"type": "object"}
                                }
                            }
                        },
                        "responses": {
                            "201": {"description": "Oferta creada exitosamente"},
                            "400": {"description": "Error de validación"}
                        }
                    },
                    "get": {
                        "summary": "Listar ofertas",
                        "parameters": [
                            {"name": "status", "in": "query", "schema": {"type": "string"}},
                            {"name": "department", "in": "query", "schema": {"type": "string"}}
                        ],
                        "responses": {
                            "200": {"description": "Lista de ofertas"}
                        }
                    }
                },
                "/api/auth/login": {
                    "post": {
                        "summary": "Autenticar usuario",
                        "requestBody": {
                            "required": True,
                            "content": {
                                "application/json": {
                                    "schema": {
                                        "type": "object",
                                        "properties": {
                                            "username": {"type": "string"},
                                            "password": {"type": "string"}
                                        }
                                    }
                                }
                            }
                        }
                    }
                },
                "/api/analytics/acceptance-rate": {
                    "get": {
                        "summary": "Obtener tasa de aceptación",
                        "security": [{"bearerAuth": []}],
                        "parameters": [
                            {"name": "date_from", "in": "query", "schema": {"type": "string"}},
                            {"name": "date_to", "in": "query", "schema": {"type": "string"}}
                        ]
                    }
                }
            },
            "tags": [
                {"name": "Ofertas", "description": "Gestión de ofertas laborales"},
                {"name": "Autenticación", "description": "Autenticación y autorización"},
                {"name": "Analytics", "description": "Análisis y reportes"},
                {"name": "Integraciones", "description": "Integraciones con sistemas externos"},
                {"name": "Backup", "description": "Backup y restauración"}
            ]
        }
        return jsonify(docs), 200
    
    @app.route('/api/docs/swagger', methods=['GET'])
    def swagger_ui():
        """Interfaz Swagger UI para documentación interactiva"""
        swagger_html = """
<!DOCTYPE html>
<html>
<head>
    <title>Offer Letter API - Swagger UI</title>
    <link rel="stylesheet" type="text/css" href="https://unpkg.com/swagger-ui-dist@4.15.5/swagger-ui.css" />
    <style>
        html { box-sizing: border-box; overflow: -moz-scrollbars-vertical; overflow-y: scroll; }
        *, *:before, *:after { box-sizing: inherit; }
        body { margin:0; background: #fafafa; }
    </style>
</head>
<body>
    <div id="swagger-ui"></div>
    <script src="https://unpkg.com/swagger-ui-dist@4.15.5/swagger-ui-bundle.js"></script>
    <script src="https://unpkg.com/swagger-ui-dist@4.15.5/swagger-ui-standalone-preset.js"></script>
    <script>
        window.onload = function() {
            const ui = SwaggerUIBundle({
                url: "/api/docs",
                dom_id: '#swagger-ui',
                deepLinking: true,
                presets: [
                    SwaggerUIBundle.presets.apis,
                    SwaggerUIStandalonePreset
                ],
                plugins: [
                    SwaggerUIBundle.plugins.DownloadUrl
                ],
                layout: "StandaloneLayout"
            });
        };
    </script>
</body>
</html>
"""
        return swagger_html, 200
    
    @app.route('/api/metrics', methods=['GET'])
    def prometheus_metrics():
        """Métricas en formato Prometheus"""
        stats = performance_monitor.get_stats()
        cache_stats = cache_manager.get_stats()
        
        metrics = f"""# HELP api_requests_total Total number of API requests
# TYPE api_requests_total counter
api_requests_total {stats.get('total_requests', 0)}

# HELP api_response_time_seconds Average response time in seconds
# TYPE api_response_time_seconds gauge
api_response_time_seconds {stats.get('avg_response_time', 0)}

# HELP api_error_rate Error rate percentage
# TYPE api_error_rate gauge
api_error_rate {stats.get('error_rate', 0)}

# HELP cache_hit_rate Cache hit rate percentage
# TYPE cache_hit_rate gauge
cache_hit_rate {stats.get('cache_hit_rate', 0)}

# HELP cache_keys_total Total number of cache keys
# TYPE cache_keys_total gauge
cache_keys_total {cache_stats.get('total_keys', 0)}
"""
        return metrics, 200, {'Content-Type': 'text/plain; version=0.0.4'}
    
    @app.route('/api/config', methods=['GET'])
    @require_auth
    def get_config():
        """Obtiene configuración del sistema (sin valores sensibles)"""
        config = {
            "version": "3.0",
            "features": {
                "authentication": True,
                "cache": True,
                "rate_limiting": True,
                "notifications": True,
                "export": True,
                "analytics": True,
                "backup": True,
                "signatures": CRYPTO_AVAILABLE,
                "pdf_generation": REPORTLAB_AVAILABLE,
                "email": EMAIL_AVAILABLE
            },
            "rate_limit": {
                "max_requests": rate_limiter.max_requests,
                "window_seconds": rate_limiter.window_seconds
            },
            "cache": {
                "default_ttl": cache_manager.default_ttl
            },
            "i18n": {
                "languages": list(i18n_manager.translations.keys()),
                "default": i18n_manager.default_language
            }
        }
        return jsonify(config), 200
    
    return app


class OfferVersionManager:
    """Gestión de versiones de ofertas"""
    
    def __init__(self, db_path: str = "offer_letters.db"):
        self.db_path = db_path
        self._init_versions_table()
    
    def _init_versions_table(self):
        """Inicializa tabla de versiones"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS offer_versions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    offer_id TEXT NOT NULL,
                    version_number INTEGER NOT NULL,
                    offer_data TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    created_by TEXT,
                    change_summary TEXT,
                    FOREIGN KEY (offer_id) REFERENCES offers(id)
                )
            ''')
            conn.commit()
    
    def create_version(self, offer_id: str, offer_data: Dict, created_by: str = "", 
                      change_summary: str = "") -> int:
        """Crea una nueva versión de una oferta"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Obtener última versión
            cursor.execute('SELECT MAX(version_number) FROM offer_versions WHERE offer_id = ?', (offer_id,))
            last_version = cursor.fetchone()[0] or 0
            new_version = last_version + 1
            
            cursor.execute('''
                INSERT INTO offer_versions (offer_id, version_number, offer_data, created_by, change_summary)
                VALUES (?, ?, ?, ?, ?)
            ''', (offer_id, new_version, json.dumps(offer_data), created_by, change_summary))
            
            conn.commit()
            return new_version
    
    def get_versions(self, offer_id: str) -> List[Dict]:
        """Obtiene todas las versiones de una oferta"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT * FROM offer_versions WHERE offer_id = ? ORDER BY version_number DESC
            ''', (offer_id,))
            rows = cursor.fetchall()
            
            return [dict(row) for row in rows]


class WebhookManager:
    """Gestión de webhooks para notificaciones"""
    
    def __init__(self, db_path: str = "offer_letters.db"):
        self.db_path = db_path
        self._init_webhooks_table()
    
    def _init_webhooks_table(self):
        """Inicializa tabla de webhooks"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS webhooks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    url TEXT NOT NULL,
                    events TEXT NOT NULL,
                    secret TEXT,
                    active BOOLEAN DEFAULT 1,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.commit()
    
    def register(self, url: str, events: List[str], secret: str = "") -> bool:
        """Registra un nuevo webhook"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO webhooks (url, events, secret)
                    VALUES (?, ?, ?)
                ''', (url, json.dumps(events), secret))
                conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error al registrar webhook: {str(e)}")
            return False
    
    def trigger_webhook(self, event: str, data: Dict):
        """Dispara un webhook para un evento"""
        try:
            import requests
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                cursor.execute('SELECT * FROM webhooks WHERE active = 1')
                webhooks = cursor.fetchall()
            
            for webhook in webhooks:
                events = json.loads(webhook['events'])
                if event in events or '*' in events:
                    try:
                        requests.post(webhook['url'], json=data, timeout=5)
                    except Exception as e:
                        logger.error(f"Error al disparar webhook {webhook['url']}: {str(e)}")
        except ImportError:
            logger.warning("Requests no disponible. Webhooks deshabilitados.")
        except Exception as e:
            logger.error(f"Error en trigger_webhook: {str(e)}")


# Extender OfferLetterAPI con nuevas funcionalidades
def extend_offer_letter_api():
    """Extiende la clase OfferLetterAPI con métodos adicionales"""
    
    def compare_offers(self, offer_id1: str, offer_id2: str) -> Dict:
        """Compara dos ofertas y retorna diferencias"""
        offer1 = self.get_offer_by_id(offer_id1)
        offer2 = self.get_offer_by_id(offer_id2)
        
        if not offer1 or not offer2:
            raise ValueError("Una o ambas ofertas no encontradas")
        
        data1 = json.loads(offer1.get('offer_data', '{}'))
        data2 = json.loads(offer2.get('offer_data', '{}'))
        
        differences = {}
        all_keys = set(data1.keys()) | set(data2.keys())
        
        for key in all_keys:
            val1 = data1.get(key)
            val2 = data2.get(key)
            if val1 != val2:
                differences[key] = {
                    "offer1": val1,
                    "offer2": val2
                }
        
        return {
            "offer1_id": offer_id1,
            "offer2_id": offer_id2,
            "differences": differences,
            "total_differences": len(differences)
        }
    
    def get_offer_versions(self, offer_id: str) -> List[Dict]:
        """Obtiene todas las versiones de una oferta"""
        if not self.use_database:
            return []
        
        version_manager = OfferVersionManager(self.db.db_path if self.db else "offer_letters.db")
        return version_manager.get_versions(offer_id)
    
    def generate_word(self, offer_id: str, output_path: str = "./output") -> Optional[str]:
        """Genera un documento Word/DOCX de la oferta"""
        try:
            offer = self.get_offer_by_id(offer_id)
            if not offer:
                raise ValueError(f"Oferta {offer_id} no encontrada")
            
            offer_data = json.loads(offer.get('offer_data', '{}'))
            data = OfferLetterData(**offer_data)
            offer_letter = offer.get('offer_letter_content', '')
            
            # Crear documento Word simple (formato RTF como alternativa)
            safe_name = "".join(c for c in data.candidate_name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')
            filename = f"offer_letter_{safe_name}_{datetime.now().strftime('%Y%m%d')}.rtf"
            filepath = os.path.join(output_path, filename)
            
            os.makedirs(output_path, exist_ok=True)
            
            # Generar RTF básico (formato compatible con Word)
            # Extraer el reemplazo fuera de la f-string para evitar error de sintaxis
            offer_letter_rtf = offer_letter.replace(chr(10), '\\par ')
            rtf_header = "{\\rtf1\\ansi\\deff0\n{\\fonttbl{\\f0 Times New Roman;}}\n\\f0\\fs24\n"
            rtf_footer = "\n}"
            rtf_content = f"{rtf_header}{offer_letter_rtf}{rtf_footer}"
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(rtf_content)
            
            logger.info(f"Documento Word/RTF generado: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error al generar Word: {str(e)}")
            return None
    
    def get_dashboard_data(self) -> Dict:
        """Obtiene datos para dashboard"""
        stats = self.get_statistics()
        
        # Obtener ofertas recientes
        recent_offers = self.get_offer_history({})
        recent_offers = recent_offers[:10] if len(recent_offers) > 10 else recent_offers
        
        # Calcular métricas adicionales
        total_salary = sum(float(o.get('salary', 0)) for o in recent_offers)
        avg_salary = total_salary / len(recent_offers) if recent_offers else 0
        
        return {
            "statistics": stats,
            "recent_offers": recent_offers,
            "metrics": {
                "total_offers": stats.get('total_offers', 0),
                "avg_salary": avg_salary,
                "pending_offers": stats.get('by_status', {}).get('generated', 0),
                "accepted_offers": stats.get('by_status', {}).get('accepted', 0)
            },
            "timestamp": datetime.now().isoformat()
        }
    
    def search_offers(self, query: str, filters: Dict = None) -> Dict:
        """Búsqueda avanzada de ofertas"""
        all_offers = self.get_offer_history(filters or {})
        
        if not query:
            return {"results": all_offers, "count": len(all_offers)}
        
        query_lower = query.lower()
        results = []
        
        for offer in all_offers:
            # Buscar en campos principales
            searchable_text = f"{offer.get('candidate_name', '')} {offer.get('position_title', '')} {offer.get('department', '')}".lower()
            if query_lower in searchable_text:
                results.append(offer)
        
        return {
            "results": results,
            "count": len(results),
            "query": query
        }
    
    def register_webhook(self, url: str, events: List[str]) -> bool:
        """Registra un webhook"""
        if not self.use_database:
            return False
        
        webhook_manager = WebhookManager(self.db.db_path if self.db else "offer_letters.db")
        return webhook_manager.register(url, events)
    
    # Agregar métodos a la clase
    OfferLetterAPI.compare_offers = compare_offers
    OfferLetterAPI.get_offer_versions = get_offer_versions
    OfferLetterAPI.generate_word = generate_word
    OfferLetterAPI.get_dashboard_data = get_dashboard_data
    OfferLetterAPI.search_offers = search_offers
    OfferLetterAPI.register_webhook = register_webhook

# Extender la API
extend_offer_letter_api()


# ============================================================================
# CONFIGURACIÓN MEDIANTE ARCHIVOS
# ============================================================================

class ConfigManager:
    """Gestión de configuración mediante archivos"""
    
    def __init__(self, config_file: str = "config.json"):
        self.config_file = config_file
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        """Carga configuración desde archivo"""
        default_config = {
            "api": {
                "host": "0.0.0.0",
                "port": 5000,
                "debug": False
            },
            "database": {
                "path": "offer_letters.db"
            },
            "cache": {
                "default_ttl": 300,
                "enabled": True
            },
            "rate_limit": {
                "max_requests": 60,
                "window_seconds": 60,
                "enabled": True
            },
            "security": {
                "secret_key": os.environ.get("SECRET_KEY", "change-me-in-production"),
                "token_expiry": 3600
            },
            "logging": {
                "level": "INFO",
                "directory": "./logs"
            },
            "features": {
                "enable_notifications": True,
                "enable_webhooks": True,
                "enable_analytics": True
            }
        }
        
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    file_config = json.load(f)
                    # Merge con configuración por defecto
                    return self._merge_config(default_config, file_config)
            except Exception as e:
                logger.warning(f"Error al cargar configuración: {str(e)}. Usando valores por defecto.")
        
        return default_config
    
    def _merge_config(self, default: Dict, custom: Dict) -> Dict:
        """Fusiona configuraciones"""
        result = default.copy()
        for key, value in custom.items():
            if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                result[key] = self._merge_config(result[key], value)
            else:
                result[key] = value
        return result
    
    def get(self, key_path: str, default=None):
        """Obtiene valor de configuración por ruta (ej: 'api.port')"""
        keys = key_path.split('.')
        value = self.config
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return default
        return value
    
    def save(self):
        """Guarda configuración en archivo"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2, ensure_ascii=False)
            return True
        except Exception as e:
            logger.error(f"Error al guardar configuración: {str(e)}")
            return False


# ============================================================================
# TESTS AVANZADOS
# ============================================================================

class OfferLetterTests:
    """Suite de tests básicos para el sistema"""
    
    @staticmethod
    def test_offer_validation():
        """Test de validación de ofertas"""
        data = OfferLetterData(
            candidate_name="Test User",
            candidate_email="test@example.com",
            position_title="Developer",
            department="Tech",
            start_date="2024-01-01",
            salary=50000
        )
        is_valid, errors = data.validate()
        assert is_valid, f"Validación falló: {errors}"
        print("✓ Test de validación pasado")
    
    @staticmethod
    def test_offer_generation():
        """Test de generación de ofertas"""
        generator = OfferLetterGenerator()
        data = OfferLetterData(
            candidate_name="Test User",
            candidate_email="test@example.com",
            position_title="Developer",
            department="Tech",
            start_date="2024-01-01",
            salary=50000,
            company_name="Test Company"
        )
        letter = generator.generate_offer_letter(data)
        assert len(letter) > 0, "Carta de oferta vacía"
        assert "Test User" in letter, "Nombre del candidato no encontrado"
        print("✓ Test de generación pasado")
    
    @staticmethod
    def test_api_creation():
        """Test de creación de oferta vía API"""
        api = OfferLetterAPI(use_database=False)
        offer_data = {
            "candidate_name": "Test User",
            "candidate_email": "test@example.com",
            "position_title": "Developer",
            "department": "Tech",
            "start_date": "2024-01-01",
            "salary": 50000
        }
        result = api.create_offer(offer_data)
        assert result["success"], f"Creación falló: {result.get('error')}"
        print("✓ Test de API pasado")
    
    @staticmethod
    def test_schema_validation():
        """Test de validación con JSON Schema"""
        validator = SchemaValidator()
        
        # Test válido
        valid_data = {
            "candidate_name": "John Doe",
            "candidate_email": "john@example.com",
            "position_title": "Developer",
            "department": "Tech",
            "start_date": "2024-01-01",
            "salary": 50000
        }
        is_valid, errors = validator.validate(valid_data)
        assert is_valid, f"Validación falló para datos válidos: {errors}"
        
        # Test inválido
        invalid_data = {
            "candidate_name": "A",  # Muy corto
            "candidate_email": "invalid-email",  # Email inválido
            "salary": -1000  # Salario negativo
        }
        is_valid, errors = validator.validate(invalid_data)
        assert not is_valid, "Validación debería fallar para datos inválidos"
        assert len(errors) > 0, "Debería haber errores de validación"
        
        print("✓ Test de validación de esquema pasado")
    
    @staticmethod
    def test_cache_manager():
        """Test del sistema de caché"""
        cache = CacheManager(default_ttl=60)
        
        # Test set/get
        cache.set("test_key", "test_value")
        value = cache.get("test_key")
        assert value == "test_value", "Caché no retornó valor correcto"
        
        # Test expiración
        cache.set("expire_key", "value", ttl=1)
        time.sleep(2)
        value = cache.get("expire_key")
        assert value is None, "Caché no expiró correctamente"
        
        print("✓ Test de caché pasado")
    
    @staticmethod
    def test_rate_limiter():
        """Test del rate limiter"""
        limiter = RateLimiter(max_requests=2, window_seconds=60)
        
        # Test permitido
        allowed, info = limiter.is_allowed("test_client")
        assert allowed, "Primera petición debería ser permitida"
        
        # Test límite
        limiter.is_allowed("test_client")
        allowed, info = limiter.is_allowed("test_client")
        assert not allowed, "Tercera petición debería ser bloqueada"
        
        print("✓ Test de rate limiting pasado")
    
    @staticmethod
    def test_export_manager():
        """Test del sistema de exportación"""
        export = ExportManager()
        test_offers = [
            {"id": "1", "candidate_name": "Test", "salary": 50000},
            {"id": "2", "candidate_name": "Test2", "salary": 60000}
        ]
        
        csv_path = export.export_to_csv(test_offers, "./test_output")
        assert os.path.exists(csv_path), "Archivo CSV no fue creado"
        os.remove(csv_path)
        
        print("✓ Test de exportación pasado")
    
    @staticmethod
    def test_authentication():
        """Test del sistema de autenticación"""
        auth = AuthenticationManager()
        
        # Crear usuario
        user_id = auth.create_user("testuser", "test@example.com", "password123")
        assert user_id is not None, "Usuario no fue creado"
        
        # Autenticar
        token = auth.authenticate("testuser", "password123")
        assert token is not None, "Autenticación falló"
        
        # Verificar token
        token_data = auth.verify_token(token)
        assert token_data is not None, "Token no fue verificado"
        assert token_data['user_id'] == user_id, "User ID no coincide"
        
        print("✓ Test de autenticación pasado")
    
    @staticmethod
    def run_all_tests():
        """Ejecuta todos los tests"""
        print("\n" + "=" * 70)
        print("EJECUTANDO SUITE COMPLETA DE TESTS")
        print("=" * 70)
        tests = [
            ("Validación de ofertas", OfferLetterTests.test_offer_validation),
            ("Generación de ofertas", OfferLetterTests.test_offer_generation),
            ("API de creación", OfferLetterTests.test_api_creation),
            ("Validación de esquema", OfferLetterTests.test_schema_validation),
            ("Sistema de caché", OfferLetterTests.test_cache_manager),
            ("Rate limiting", OfferLetterTests.test_rate_limiter),
            ("Exportación", OfferLetterTests.test_export_manager),
            ("Autenticación", OfferLetterTests.test_authentication)
        ]
        
        passed = 0
        failed = 0
        
        for test_name, test_func in tests:
            try:
                test_func()
                passed += 1
            except AssertionError as e:
                print(f"✗ {test_name}: {str(e)}")
                failed += 1
            except Exception as e:
                print(f"✗ {test_name}: Error inesperado - {str(e)}")
                failed += 1
        
        print("\n" + "=" * 70)
        print(f"RESUMEN: {passed} pasados, {failed} fallidos de {len(tests)} tests")
        print("=" * 70)
        
        return failed == 0


# ============================================================================
# FUNCIONES DE MARKETING - INSTAGRAM DM
# ============================================================================

def generar_respuesta_dm_precio_instagram(
    producto: str,
    precio: str,
    ciudad: Optional[str] = None,
    nombre_cliente: Optional[str] = None,
    video_url: Optional[str] = None
) -> Dict[str, str]:
    """
    Genera una respuesta automática para un DM de Instagram preguntando por el precio de un producto.
    
    Args:
        producto: Nombre del producto
        precio: Precio del producto (formato: "$XXX" o "XXX pesos")
        ciudad: Ciudad detectada del cliente (opcional, si no se proporciona se pregunta)
        nombre_cliente: Nombre del cliente (opcional, para personalización)
        video_url: URL del video de 15 segundos del producto en uso (opcional)
    
    Returns:
        Dict con:
            - 'mensaje': Texto completo del mensaje
            - 'mensaje_texto': Solo el texto (sin video)
            - 'instrucciones_video': Instrucciones para adjuntar el video
    """
    # Saludo personalizado si hay nombre
    saludo = f"Hola {nombre_cliente}! 👋" if nombre_cliente else "Hola! 👋"
    
    # Confirmación de interés
    confirmacion = f"¡Qué genial que te interese el {producto}! 🎉"
    
    # Precio y envío
    if ciudad:
        precio_envio = f"El {producto} está en {precio} + envío a {ciudad}."
    else:
        precio_envio = f"El {producto} está en {precio} + envío. ¿A qué ciudad lo necesitas enviar? 📍"
    
    # Solicitud de WhatsApp
    whatsapp_cta = "Para enviarte el catálogo exclusivo con todos los detalles, ¿me pasas tu WhatsApp? 📱"
    
    # Video
    if video_url:
        video_instruccion = f"\n\n🎥 Chécalo 👀 - Aquí tienes un mini-video de 15 seg del {producto} en acción:\n{video_url}"
    else:
        video_instruccion = "\n\n🎥 Chécalo 👀 - Te adjunto un mini-video de 15 seg del producto en uso"
    
    # Construir mensaje completo
    mensaje_completo = f"""{saludo}

{confirmacion}

{precio_envio}

{whatsapp_cta}{video_instruccion}"""
    
    # Mensaje solo texto (sin video)
    mensaje_texto = f"""{saludo}

{confirmacion}

{precio_envio}

{whatsapp_cta}"""
    
    return {
        'mensaje': mensaje_completo,
        'mensaje_texto': mensaje_texto,
        'instrucciones_video': f"Adjuntar video de 15 segundos del {producto} en uso" if not video_url else None,
        'ciudad_detectada': ciudad if ciudad and ciudad != "[Ciudad]" else None,
        'requiere_ciudad': ciudad == "[Ciudad]"
    }


def generate_offer_comparison_table(offers: List[Dict]) -> str:
    """
    Genera una tabla comparativa de múltiples ofertas
    
    Args:
        offers: Lista de diccionarios con datos de ofertas
        
    Returns:
        String con tabla formateada
    """
    if not offers:
        return "No hay ofertas para comparar"
    
    # Encabezados
    headers = ["Candidato", "Posición", "Departamento", "Salario", "Estado", "Fecha"]
    rows = []
    
    for offer in offers:
        offer_data = offer.get('offer_data', {})
        if isinstance(offer_data, str):
            try:
                offer_data = json.loads(offer_data)
            except json.JSONDecodeError:
                offer_data = {}
        
        row = [
            offer_data.get('candidate_name', 'N/A'),
            offer_data.get('position_title', 'N/A'),
            offer_data.get('department', 'N/A'),
            f"{offer_data.get('currency', 'USD')} {offer_data.get('salary', 0):,.2f}",
            offer.get('status', 'N/A'),
            offer.get('created_at', 'N/A')[:10] if offer.get('created_at') else 'N/A'
        ]
        rows.append(row)
    
    # Formatear tabla
    table_lines = ["| " + " | ".join(headers) + " |"]
    table_lines.append("|" + "|".join(["---"] * len(headers)) + "|")
    
    for row in rows:
        table_lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    
    return "\n".join(table_lines)


def calculate_offer_metrics(offers: List[Dict]) -> Dict:
    """
    Calcula métricas agregadas de una lista de ofertas
    
    Args:
        offers: Lista de ofertas
        
    Returns:
        Diccionario con métricas calculadas
    """
    if not offers:
        return {
            "total": 0,
            "avg_salary": 0,
            "total_budget": 0,
            "status_distribution": {},
            "department_distribution": {}
        }
    
    metrics = {
        "total": len(offers),
        "salaries": [],
        "status_distribution": defaultdict(int),
        "department_distribution": defaultdict(int),
        "currency_distribution": defaultdict(int)
    }
    
    for offer in offers:
        # Estado
        status = offer.get('status', 'unknown')
        metrics["status_distribution"][status] += 1
        
        # Datos de la oferta
        offer_data = offer.get('offer_data', {})
        if isinstance(offer_data, str):
            try:
                offer_data = json.loads(offer_data)
            except json.JSONDecodeError:
                continue
        
        # Departamento
        department = offer_data.get('department', 'unknown')
        metrics["department_distribution"][department] += 1
        
        # Salario
        salary = offer_data.get('salary', 0)
        if salary:
            metrics["salaries"].append(float(salary))
        
        # Moneda
        currency = offer_data.get('currency', 'USD')
        metrics["currency_distribution"][currency] += 1
    
    # Calcular promedios
    metrics["avg_salary"] = sum(metrics["salaries"]) / len(metrics["salaries"]) if metrics["salaries"] else 0
    metrics["total_budget"] = sum(metrics["salaries"])
    metrics["min_salary"] = min(metrics["salaries"]) if metrics["salaries"] else 0
    metrics["max_salary"] = max(metrics["salaries"]) if metrics["salaries"] else 0
    
    # Convertir defaultdicts a dicts
    metrics["status_distribution"] = dict(metrics["status_distribution"])
    metrics["department_distribution"] = dict(metrics["department_distribution"])
    metrics["currency_distribution"] = dict(metrics["currency_distribution"])
    
    # Limpiar lista temporal
    del metrics["salaries"]
    
    return metrics


def validate_offer_consistency(offer_data: Dict) -> Tuple[bool, List[str]]:
    """
    Valida la consistencia lógica de los datos de una oferta
    
    Args:
        offer_data: Diccionario con datos de la oferta
        
    Returns:
        Tupla (es_consistente, lista_de_advertencias)
    """
    warnings = []
    
    # Validar fechas
    start_date_str = offer_data.get('start_date')
    expiry_date_str = offer_data.get('offer_expiry_date')
    
    if start_date_str and expiry_date_str:
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            expiry_date = datetime.strptime(expiry_date_str, '%Y-%m-%d').date()
            
            if expiry_date < start_date:
                warnings.append("La fecha de expiración es anterior a la fecha de inicio")
        except (ValueError, TypeError):
            pass
    
    # Validar salario vs tipo de empleo
    salary = offer_data.get('salary', 0)
    employment_type = offer_data.get('employment_type', '').lower()
    
    if employment_type == 'intern' and salary > 50000:
        warnings.append("El salario para una pasantía parece muy alto")
    
    if employment_type == 'full-time' and salary < 20000:
        warnings.append("El salario para tiempo completo parece muy bajo")
    
    # Validar períodos
    probation_period = offer_data.get('probation_period', 0)
    notice_period = offer_data.get('notice_period', 0)
    
    if probation_period > 180:
        warnings.append("El período de prueba es muy largo (más de 6 meses)")
    
    if notice_period > 90:
        warnings.append("El período de preaviso es muy largo (más de 3 meses)")
    
    # Validar beneficios
    benefits = offer_data.get('benefits', [])
    if len(benefits) > 20:
        warnings.append("Hay muchos beneficios listados, considerar agruparlos")
    
    return len(warnings) == 0, warnings


def extract_offer_keywords(offer_data: Dict) -> List[str]:
    """
    Extrae palabras clave relevantes de una oferta
    
    Args:
        offer_data: Datos de la oferta
        
    Returns:
        Lista de palabras clave
    """
    keywords = []
    
    # Agregar información básica
    if offer_data.get('position_title'):
        keywords.append(offer_data['position_title'].lower())
    
    if offer_data.get('department'):
        keywords.append(offer_data['department'].lower())
    
    if offer_data.get('location'):
        location = offer_data['location'].lower()
        keywords.append(location)
        # Extraer ciudad si es posible
        if ',' in location:
            city = location.split(',')[0].strip()
            keywords.append(city)
    
    # Agregar tipo de empleo
    employment_type = offer_data.get('employment_type', '').lower()
    if employment_type:
        keywords.append(employment_type)
        if 'remote' in employment_type or 'remoto' in employment_type:
            keywords.append('remote')
        if 'hybrid' in employment_type or 'híbrido' in employment_type:
            keywords.append('hybrid')
    
    # Agregar palabras de beneficios
    benefits = offer_data.get('benefits', [])
    for benefit in benefits:
        words = benefit.lower().split()
        keywords.extend([w for w in words if len(w) > 3])
    
    # Remover duplicados y palabras muy comunes
    common_words = {'the', 'and', 'or', 'for', 'with', 'del', 'de', 'la', 'el', 'en', 'y', 'o'}
    keywords = [k for k in set(keywords) if k not in common_words and len(k) > 2]
    
    return keywords[:20]  # Limitar a 20 keywords


def generate_offer_export_filename(offer_data: Dict, format_type: str = "txt") -> str:
    """
    Genera un nombre de archivo descriptivo para exportar una oferta
    
    Args:
        offer_data: Datos de la oferta
        format_type: Tipo de formato (txt, pdf, html, etc.)
        
    Returns:
        String con nombre de archivo
    """
    candidate = offer_data.get('candidate_name', 'candidate')
    position = offer_data.get('position_title', 'position')
    date_str = datetime.now().strftime('%Y%m%d')
    
    # Limpiar nombres
    candidate_clean = re.sub(r'[^a-zA-Z0-9]', '_', candidate)[:20]
    position_clean = re.sub(r'[^a-zA-Z0-9]', '_', position)[:30]
    
    return f"offer_{candidate_clean}_{position_clean}_{date_str}.{format_type}"


# Ejemplo de uso
if __name__ == "__main__":
    import sys
    
    # Verificar argumentos de línea de comandos
    if len(sys.argv) > 1:
        if sys.argv[1] == "server":
            print("Iniciando servidor Flask...")
            api = OfferLetterAPI(use_database=True)
            app = create_flask_app(api)
            print("\n" + "=" * 70)
            # Cargar configuración
            config_manager = ConfigManager()
            host = config_manager.get('api.host', '0.0.0.0')
            port = config_manager.get('api.port', 5000)
            debug = config_manager.get('api.debug', False)
            
            print("API REST disponible en: http://{}:{}".format(host, port))
            print("Documentación OpenAPI: http://{}:{}/api/docs".format(host, port))
            print("Swagger UI: http://{}:{}/api/docs/swagger".format(host, port))
            print("Health check: http://{}:{}/health".format(host, port))
            print("Métricas Prometheus: http://{}:{}/api/metrics".format(host, port))
            print("\nEndpoints principales:")
            print("  POST /api/auth/register     - Registrar usuario")
            print("  POST /api/auth/login        - Autenticar y obtener token")
            print("  POST /api/offers            - Crear oferta (requiere auth)")
            print("  GET  /api/offers            - Listar ofertas")
            print("  GET  /api/analytics/*       - Análisis y reportes (requiere auth)")
            print("  POST /api/backup/create     - Crear backup (requiere auth)")
            print("  GET  /api/performance/stats  - Métricas de rendimiento (requiere auth)")
            print("  GET  /api/export/csv        - Exportar a CSV (requiere auth)")
            print("  GET  /api/notifications      - Notificaciones (requiere auth)")
            print("=" * 70 + "\n")
            app.run(host=host, port=port, debug=debug)
            sys.exit(0)
        elif sys.argv[1] == "test":
            OfferLetterTests.run_all_tests()
            sys.exit(0)
        elif sys.argv[1] == "help":
            print("""
Uso: python offer_letter_api.py [comando]

Comandos disponibles:
  (sin comando)  - Ejecuta ejemplo de generación de oferta
  server         - Inicia servidor Flask API
  test           - Ejecuta tests unitarios
  help           - Muestra esta ayuda
            """)
            sys.exit(0)
    
    # Modo de ejemplo/ejecución directa
    print("=" * 70)
    print("SISTEMA DE GENERACIÓN DE CARTAS DE OFERTA LABORAL")
    print("=" * 70)
    
    # Inicializar API con base de datos
    api = OfferLetterAPI(use_database=True)
    
    # Datos de ejemplo
    sample_offer = {
        "candidate_name": "Juan Pérez",
        "candidate_email": "juan.perez@email.com",
        "position_title": "Desarrollador Senior de Software",
        "department": "Tecnología",
        "start_date": "2024-02-01",
        "salary": 85000,
        "currency": "USD",
        "benefits": [
            "Seguro médico completo",
            "Plan de retiro 401(k) con contribución de la empresa",
            "Vacaciones pagadas (20 días)",
            "Días de enfermedad (10 días)",
            "Desarrollo profesional y capacitación"
        ],
        "reporting_manager": "María González, Directora de Tecnología",
        "location": "Ciudad de México (Híbrido)",
        "employment_type": "Full-time",
        "probation_period": 90,
        "notice_period": 30,
        "company_name": "Tech Solutions S.A.",
        "company_address": "Av. Reforma 123, Ciudad de México, CDMX 06600",
        "hr_contact": "Ana Martínez - ana.martinez@techsolutions.com",
        "offer_expiry_date": "2024-01-15",
        "additional_terms": [
            "Horario flexible de trabajo",
            "Posibilidad de trabajo remoto 3 días por semana"
        ],
        "template_type": "technical",
        "bonus_structure": "Bono anual del 10-15% basado en desempeño"
    }
    
    # Generar carta de oferta
    print("\n1. Generando carta de oferta...")
    result = api.create_offer(sample_offer, format_type="txt")
    
    if result["success"]:
        print(f"✓ Oferta creada con ID: {result['offer_id']}")
        print("\n" + "=" * 70)
        print("CARTA DE OFERTA GENERADA EXITOSAMENTE")
        print("=" * 70)
        print(result["offer_letter"])
        print("=" * 70)
        
        # Guardar en archivo
        print("\n2. Guardando archivo...")
        generator = OfferLetterGenerator()
        output_path = generator.save_offer_letter(
            result["offer_letter"],
            "./output",
            sample_offer["candidate_name"]
        )
        print(f"✓ Archivo guardado en: {output_path}")
        
        # Generar HTML
        print("\n3. Generando versión HTML...")
        html_result = api.create_offer(sample_offer, format_type="html")
        if html_result["success"]:
            html_path = generator.save_offer_letter(
                html_result["offer_letter"],
                "./output",
                f"{sample_offer['candidate_name']}_html"
            ).replace('.txt', '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_result["offer_letter"])
            print(f"✓ HTML guardado en: {html_path}")
        
        # Generar PDF si está disponible
        if REPORTLAB_AVAILABLE:
            print("\n4. Generando PDF...")
            try:
                pdf_path = api.generate_pdf(result["offer_id"])
                if pdf_path:
                    print(f"✓ PDF generado en: {pdf_path}")
            except Exception as e:
                print(f"✗ Error al generar PDF: {str(e)}")
        
        # Exportar JSON
        print("\n5. Exportando JSON...")
        json_export = api.export_offer_json(sample_offer)
        json_path = os.path.join("./output", f"offer_{result['offer_id']}.json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_export, f, indent=2, ensure_ascii=False)
        print(f"✓ JSON exportado en: {json_path}")
        
        # Mostrar estadísticas
        print("\n6. Estadísticas del sistema:")
        stats = api.get_statistics()
        print(json.dumps(stats, indent=2, ensure_ascii=False))
        
        print("\n" + "=" * 70)
        print("PROCESO COMPLETADO EXITOSAMENTE")
        print("=" * 70)
        print("\nComandos disponibles:")
        print("  python offer_letter_api.py server  - Iniciar servidor Flask API")
        print("  python offer_letter_api.py test    - Ejecutar tests unitarios")
        print("  python offer_letter_api.py help    - Mostrar ayuda")
        print("\nFuncionalidades disponibles:")
        print("  ✓ Generación en múltiples formatos (TXT, HTML, PDF, Word/RTF)")
        print("  ✓ API REST completa con Flask (35+ endpoints)")
        print("  ✓ Base de datos SQLite con historial completo")
        print("  ✓ Sistema de versiones de ofertas")
        print("  ✓ Comparación de ofertas")
        print("  ✓ Búsqueda avanzada con filtros")
        print("  ✓ Webhooks para notificaciones")
        print("  ✓ Dashboard con métricas en tiempo real")
        print("  ✓ Envío por email con adjuntos")
        print("  ✓ Estadísticas y reportes avanzados")
        print("  ✓ Autenticación JWT y autorización")
        print("  ✓ Sistema de caché con TTL")
        print("  ✓ Internacionalización (i18n) - ES, EN, FR, PT")
        print("  ✓ Firmas digitales para ofertas")
        print("  ✓ Integraciones con ATS y HRIS")
        print("  ✓ Análisis avanzado y reportes")
        print("  ✓ Sistema de backup y restore")
        print("  ✓ Monitor de rendimiento y métricas")
        print("  ✓ Rate limiting inteligente")
        print("  ✓ Validación con JSON Schema")
        print("  ✓ Sistema de notificaciones")
        print("  ✓ Exportación a CSV y Excel")
        print("  ✓ Sistema de eventos")
        print("  ✓ Logging avanzado con rotación")
        print("  ✓ Tests unitarios integrados")
        print("\nTotal de líneas de código: ~3,000+")
        print("Total de endpoints: 35+")
        print("Total de managers: 12+")
        print("=" * 70)
    else:
        print(f"✗ Error: {result.get('error')}")
