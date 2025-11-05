#!/usr/bin/env python3
"""
Script completo para crear el documento Word del Entregable 2 con todos los ejercicios
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime

def create_complete_word_document():
    """Crear el documento Word completo del Entregable 2"""
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # PORTADA
    title = doc.add_heading('ENTREGABLE 2: ECUACIONES DIFERENCIALES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información de la universidad
    doc.add_paragraph('UNIVERSIDAD [NOMBRE DE LA UNIVERSIDAD]')
    doc.add_paragraph('FACULTAD DE [NOMBRE DE LA FACULTAD]')
    doc.add_paragraph('DEPARTAMENTO DE MATEMÁTICAS')
    
    # Línea separadora
    doc.add_paragraph('─' * 50)
    
    # Información del curso
    doc.add_heading('Reporte de Investigación: Aplicaciones de las Ecuaciones Diferenciales', level=1)
    doc.add_paragraph('─' * 50)
    
    # Datos del estudiante
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run('Materia: ').bold = True
    info_paragraph.add_run('Ecuaciones Diferenciales\n')
    
    info_paragraph.add_run('Profesor: ').bold = True
    info_paragraph.add_run('[Nombre del Profesor]\n')
    
    info_paragraph.add_run('Alumno: ').bold = True
    info_paragraph.add_run('[Tu Nombre Completo]\n')
    
    info_paragraph.add_run('Matrícula: ').bold = True
    info_paragraph.add_run('[Tu Matrícula]\n')
    
    info_paragraph.add_run('Carrera: ').bold = True
    info_paragraph.add_run('[Tu Carrera]\n')
    
    info_paragraph.add_run('Fecha: ').bold = True
    info_paragraph.add_run(f'{datetime.now().strftime("%d de %B de %Y")}')
    
    # Salto de página
    doc.add_page_break()
    
    # PARTE A: REPORTE DE INVESTIGACIÓN
    doc.add_heading('PARTE A: REPORTE DE INVESTIGACIÓN', level=1)
    
    # Introducción
    doc.add_heading('INTRODUCCIÓN', level=2)
    intro_text = """
Las ecuaciones diferenciales constituyen una herramienta matemática fundamental en el análisis y modelado de fenómenos dinámicos que involucran cambios continuos en el tiempo o en el espacio. Estas ecuaciones relacionan una función desconocida con sus derivadas, permitiendo describir matemáticamente procesos naturales, físicos, químicos, biológicos y tecnológicos.

En el contexto de la ingeniería, las ecuaciones diferenciales son indispensables para el diseño, análisis y optimización de sistemas complejos. Desde el modelado de circuitos eléctricos hasta el análisis de estructuras mecánicas, pasando por el control de procesos industriales y la simulación de sistemas dinámicos, estas ecuaciones proporcionan el marco matemático necesario para comprender y predecir el comportamiento de sistemas reales.

La importancia del estudio de las ecuaciones diferenciales radica en su capacidad para transformar problemas prácticos complejos en modelos matemáticos manejables, facilitando así la toma de decisiones técnicas fundamentadas y el desarrollo de soluciones innovadoras a los desafíos de la ingeniería moderna.
    """
    doc.add_paragraph(intro_text.strip())
    
    # Aplicaciones
    doc.add_heading('APLICACIONES A LA INGENIERÍA EN GENERAL Y A [TU CARRERA] EN PARTICULAR', level=2)
    
    doc.add_heading('Aplicaciones Generales en Ingeniería', level=3)
    
    # Ingeniería Mecánica
    doc.add_heading('1. Ingeniería Mecánica', level=4)
    mech_text = """
• Vibraciones y Oscilaciones: Las ecuaciones diferenciales de segundo orden modelan sistemas vibratorios como amortiguadores, suspensiones de vehículos y estructuras sometidas a cargas dinámicas.
• Transferencia de Calor: La ecuación del calor (ecuación diferencial parcial) describe la distribución de temperatura en sólidos, esencial para el diseño de intercambiadores de calor y sistemas de refrigeración.
• Dinámica de Fluidos: Las ecuaciones de Navier-Stokes modelan el comportamiento de fluidos, fundamentales en el diseño de turbinas, bombas y sistemas de ventilación.
    """
    doc.add_paragraph(mech_text.strip())
    
    # Ingeniería Eléctrica
    doc.add_heading('2. Ingeniería Eléctrica', level=4)
    elec_text = """
• Circuitos RLC: Los circuitos eléctricos con resistencias, inductores y capacitores se modelan mediante ecuaciones diferenciales lineales de segundo orden.
• Sistemas de Control: La teoría de control utiliza ecuaciones diferenciales para diseñar sistemas de retroalimentación que mantengan la estabilidad y el rendimiento deseado.
• Transmisión de Señales: Las ecuaciones de onda describen la propagación de señales electromagnéticas en cables y sistemas de comunicación.
    """
    doc.add_paragraph(elec_text.strip())
    
    # Ingeniería Química
    doc.add_heading('3. Ingeniería Química', level=4)
    chem_text = """
• Cinética de Reacciones: Las ecuaciones diferenciales modelan la velocidad de reacciones químicas y la concentración de reactivos y productos en el tiempo.
• Transferencia de Masa: Los procesos de difusión y convección se describen mediante ecuaciones diferenciales parciales.
• Reactores Químicos: El diseño y optimización de reactores requiere el modelado matemático de procesos de mezclado y reacción.
    """
    doc.add_paragraph(chem_text.strip())
    
    # Ingeniería Civil
    doc.add_heading('4. Ingeniería Civil', level=4)
    civil_text = """
• Análisis Estructural: Las ecuaciones diferenciales modelan la respuesta dinámica de estructuras sometidas a cargas variables como sismos y viento.
• Mecánica de Suelos: El comportamiento de suelos bajo carga se modela mediante ecuaciones diferenciales que consideran la consolidación y la deformación.
• Hidráulica: El flujo de agua en canales y tuberías se describe mediante ecuaciones diferenciales que consideran la conservación de masa y momentum.
    """
    doc.add_paragraph(civil_text.strip())
    
    # Aplicaciones específicas
    doc.add_heading('Aplicaciones Específicas en [TU CARRERA]', level=3)
    specific_text = """
Si estudias Ingeniería Mecánica:
• Modelado de sistemas de suspensión automotriz
• Análisis de vibraciones en máquinas rotativas
• Diseño de sistemas de control de temperatura en motores
• Simulación de procesos de manufactura

Si estudias Ingeniería Eléctrica:
• Diseño de filtros analógicos y digitales
• Análisis de estabilidad en sistemas de potencia
• Modelado de convertidores de energía
• Sistemas de comunicación inalámbrica

Si estudias Ingeniería Química:
• Optimización de procesos de separación
• Control de reactores de polimerización
• Modelado de sistemas de intercambio de calor
• Análisis de procesos de fermentación

Si estudias Ingeniería Civil:
• Análisis sísmico de estructuras
• Modelado de flujo de tráfico
• Diseño de sistemas de drenaje
• Análisis de estabilidad de taludes
    """
    doc.add_paragraph(specific_text.strip())
    
    # Conclusiones
    doc.add_heading('CONCLUSIONES', level=2)
    conclusions_text = """
El estudio de las ecuaciones diferenciales representa un pilar fundamental en la formación de cualquier ingeniero, ya que proporciona las herramientas matemáticas necesarias para abordar problemas complejos del mundo real. A través de este análisis, he comprendido que estas ecuaciones no son meros ejercicios académicos, sino herramientas prácticas que permiten modelar, analizar y optimizar sistemas tecnológicos.

La importancia de dominar los diferentes métodos de solución (separación de variables, ecuaciones homogéneas, exactas, lineales, transformada de Laplace, etc.) radica en la versatilidad que proporcionan para abordar diversos tipos de problemas. Cada método tiene su ámbito de aplicación específico, y la elección del método adecuado puede significar la diferencia entre una solución elegante y eficiente o un proceso tedioso y propenso a errores.

En mi carrera específica, las ecuaciones diferenciales me permitirán abordar desafíos técnicos con una perspectiva matemática sólida, facilitando el diseño de soluciones innovadoras y la optimización de procesos existentes. El conocimiento de estas herramientas matemáticas me prepara para enfrentar los retos de la ingeniería moderna, donde la modelación matemática y la simulación computacional son cada vez más importantes.

Finalmente, reconozco que el dominio de las ecuaciones diferenciales no solo es importante para el éxito académico, sino que constituye una competencia profesional esencial que me permitirá contribuir de manera significativa al desarrollo tecnológico y la resolución de problemas de ingeniería en mi campo de especialización.
    """
    doc.add_paragraph(conclusions_text.strip())
    
    # Salto de página
    doc.add_page_break()
    
    # PARTE B: EJERCICIOS
    doc.add_heading('PARTE B: EJERCICIOS DE REPASO', level=1)
    
    doc.add_heading('EJERCICIOS DE MÉTODOS ANTERIORES', level=2)
    
    # Ejercicio 1
    doc.add_heading('Ejercicio 1: Separación de Variables', level=3)
    doc.add_paragraph('Resolver: y\' = -2^(x-y)')
    doc.add_paragraph('Solución:')
    sol1_text = """
Separando variables:
dy/dx = -2^(x-y) = -2^x/2^y

2^y dy = -2^x dx

Integrando ambos lados:
∫ 2^y dy = -∫ 2^x dx

2^y/ln(2) = -2^x/ln(2) + C

2^y = -2^x + C ln(2)

y = log₂(-2^x + C ln(2))
    """
    doc.add_paragraph(sol1_text.strip())
    
    # Ejercicio 2
    doc.add_heading('Ejercicio 2: Ecuaciones Diferenciales Homogéneas', level=3)
    doc.add_paragraph('Resolver: (x² + y²)dx - 2xy dy = 0')
    doc.add_paragraph('Solución:')
    sol2_text = """
Esta es una ecuación homogénea. Hacemos la sustitución y = vx, entonces dy = v dx + x dv.

Sustituyendo:
(x² + v²x²)dx - 2x(vx)(v dx + x dv) = 0
x²(1 + v²)dx - 2x²v(v dx + x dv) = 0
(1 + v²)dx - 2v(v dx + x dv) = 0
(1 + v²)dx - 2v² dx - 2vx dv = 0
(1 - v²)dx - 2vx dv = 0

Separando variables:
dx/x = 2v/(1 - v²) dv

Integrando:
ln|x| = -ln|1 - v²| + C
ln|x| + ln|1 - v²| = C
ln|x(1 - v²)| = C
x(1 - v²) = C₁

Sustituyendo v = y/x:
x(1 - y²/x²) = C₁
x - y²/x = C₁
x² - y² = C₁x
    """
    doc.add_paragraph(sol2_text.strip())
    
    # Ejercicio 3
    doc.add_heading('Ejercicio 3: Ecuaciones Diferenciales Exactas y Factor Integrante', level=3)
    doc.add_paragraph('Resolver: (2xy + 3)dx + (x² - 1)dy = 0')
    doc.add_paragraph('Solución:')
    sol3_text = """
Verificamos si es exacta:
∂M/∂y = ∂/∂y(2xy + 3) = 2x
∂N/∂x = ∂/∂x(x² - 1) = 2x

Como ∂M/∂y = ∂N/∂x, la ecuación es exacta.

Existe una función F(x,y) tal que:
∂F/∂x = 2xy + 3
∂F/∂y = x² - 1

Integrando la primera ecuación respecto a x:
F(x,y) = ∫ (2xy + 3)dx = x²y + 3x + g(y)

Derivando respecto a y:
∂F/∂y = x² + g'(y) = x² - 1

Por tanto: g'(y) = -1, entonces g(y) = -y + C

La solución general es:
x²y + 3x - y = C
    """
    doc.add_paragraph(sol3_text.strip())
    
    # Ejercicio 4
    doc.add_heading('Ejercicio 4: Ecuaciones Diferenciales Lineales de Primer Orden', level=3)
    doc.add_paragraph('Resolver: y\' + y/x = x²')
    doc.add_paragraph('Solución:')
    sol4_text = """
Esta es una ecuación lineal de primer orden de la forma y' + P(x)y = Q(x) donde P(x) = 1/x y Q(x) = x².

El factor integrante es:
μ(x) = e^(∫ P(x)dx) = e^(∫ 1/x dx) = e^(ln|x|) = |x| = x (para x > 0)

Multiplicando la ecuación por el factor integrante:
xy' + y = x³

El lado izquierdo es d/dx(xy), entonces:
d/dx(xy) = x³

Integrando:
xy = ∫ x³ dx = x⁴/4 + C

Por tanto:
y = x³/4 + C/x
    """
    doc.add_paragraph(sol4_text.strip())
    
    # Ejercicio 5
    doc.add_heading('Ejercicio 5: Lineales de Orden Superior Homogéneas', level=3)
    doc.add_paragraph('Resolver: y\'\' - 3y\' + 2y = 0')
    doc.add_paragraph('Solución:')
    sol5_text = """
La ecuación característica es:
r² - 3r + 2 = 0
(r - 1)(r - 2) = 0
r = 1, 2

Como las raíces son reales y distintas, la solución general es:
y = C₁ e^x + C₂ e^(2x)
    """
    doc.add_paragraph(sol5_text.strip())
    
    # Ejercicio 6
    doc.add_heading('Ejercicio 6: Lineales de Orden Superior No Homogéneas', level=3)
    doc.add_paragraph('Resolver: y\'\' - 3y\' + 2y = 4e^(3x)')
    doc.add_paragraph('Solución:')
    sol6_text = """
Primero resolvemos la ecuación homogénea (del ejercicio anterior):
y_h = C₁ e^x + C₂ e^(2x)

Para la solución particular, usamos el método de coeficientes indeterminados. Como 4e^(3x) no es solución de la ecuación homogénea, proponemos:
y_p = Ae^(3x)

Derivando:
y_p' = 3Ae^(3x)
y_p'' = 9Ae^(3x)

Sustituyendo en la ecuación:
9Ae^(3x) - 3(3Ae^(3x)) + 2(Ae^(3x)) = 4e^(3x)
9Ae^(3x) - 9Ae^(3x) + 2Ae^(3x) = 4e^(3x)
2Ae^(3x) = 4e^(3x)

Por tanto: 2A = 4, entonces A = 2

La solución particular es: y_p = 2e^(3x)

La solución general es:
y = y_h + y_p = C₁ e^x + C₂ e^(2x) + 2e^(3x)
    """
    doc.add_paragraph(sol6_text.strip())
    
    # Ejercicio 7
    doc.add_heading('Ejercicio 7: Solución de EDO usando Transformada de Laplace', level=3)
    doc.add_paragraph('Resolver: y\'\' + 4y\' + 4y = e^(-2t) con y(0) = 1 e y\'(0) = 0')
    doc.add_paragraph('Solución:')
    sol7_text = """
Aplicando la transformada de Laplace a ambos lados:
L{y''} + 4L{y'} + 4L{y} = L{e^(-2t)}

s²Y(s) - sy(0) - y'(0) + 4(sY(s) - y(0)) + 4Y(s) = 1/(s+2)

s²Y(s) - s - 0 + 4sY(s) - 4 + 4Y(s) = 1/(s+2)

Y(s)(s² + 4s + 4) - s - 4 = 1/(s+2)

Y(s)(s+2)² = 1/(s+2) + s + 4 = (1 + s(s+2) + 4(s+2))/(s+2)

Y(s)(s+2)² = (1 + s² + 2s + 4s + 8)/(s+2) = (s² + 6s + 9)/(s+2)

Y(s) = (s² + 6s + 9)/(s+2)³ = (s+3)²/(s+2)³

Usando fracciones parciales:
(s+3)²/(s+2)³ = A/(s+2) + B/(s+2)² + C/(s+2)³

Resolviendo: A = 1, B = 2, C = 1

Por tanto:
Y(s) = 1/(s+2) + 2/(s+2)² + 1/(s+2)³

Aplicando la transformada inversa:
y(t) = e^(-2t) + 2te^(-2t) + (t²/2)e^(-2t) = e^(-2t)(1 + 2t + t²/2)
    """
    doc.add_paragraph(sol7_text.strip())
    
    # Salto de página
    doc.add_page_break()
    
    # EJERCICIOS ADICIONALES
    doc.add_heading('EJERCICIOS ADICIONALES DE ESTA SEMANA', level=2)
    
    # Ejercicio 8 - Serie de Fourier
    doc.add_heading('Ejercicio 8: Serie de Fourier', level=3)
    doc.add_paragraph('Desarrollar en serie de Fourier la función periódica de período 2π dada por:')
    doc.add_paragraph('f(x) = {0 si -π ≤ x < 0')
    doc.add_paragraph('       {x si 0 ≤ x < π')
    
    doc.add_paragraph('Solución:')
    fourier_text = """
Para obtener la serie de Fourier, calculamos los coeficientes:

a₀ = (1/π) ∫₋π^π f(x) dx = (1/π) ∫₀^π x dx = (1/π) [x²/2]₀^π = π/2

aₙ = (1/π) ∫₋π^π f(x) cos(nx) dx = (1/π) ∫₀^π x cos(nx) dx
    = (1/π) [x sin(nx)/n + cos(nx)/n²]₀^π = (1/π) [(cos(nπ) - 1)/n²]
    = ((-1)ⁿ - 1)/(π n²) = {0 si n es par
                           {-2/(π n²) si n es impar

bₙ = (1/π) ∫₋π^π f(x) sin(nx) dx = (1/π) ∫₀^π x sin(nx) dx
    = (1/π) [-x cos(nx)/n + sin(nx)/n²]₀^π = (1/π) [-π cos(nπ)/n]
    = -(-1)ⁿ/n = (-1)^(n+1)/n

Por lo tanto, la serie de Fourier es:

f(x) = π/4 + Σₙ₌₁^∞ [((-1)ⁿ - 1)/(π n²) cos(nx) + (-1)^(n+1)/n sin(nx)]

Código para SageMath Cell:
f = piecewise([((-pi,0), 0), ((0,pi), x)])
n = 1  # Cambiar este valor de 1 a 10
s = f.fourier_series_partial_sum(n)
plot(f,(-2*pi,2*pi), thickness=3) + plot(s,(x,-2*pi,2*pi), color='red', thickness=1)
    """
    doc.add_paragraph(fourier_text.strip())
    
    # Ejercicio 9
    doc.add_heading('Ejercicio 9: Ecuación Diferencial con Condiciones Iniciales', level=3)
    doc.add_paragraph('Resolver: x\'\'(t) + 4x(t) = f(t) donde x(1) = 0 e x\'(1) = n')
    doc.add_paragraph('Solución:')
    sol9_text = """
Esta es una ecuación diferencial lineal de segundo orden no homogénea. Primero resolvemos la ecuación homogénea:

x''(t) + 4x(t) = 0

La ecuación característica es: r² + 4 = 0, entonces r = ±2i

La solución general de la ecuación homogénea es:
x_h(t) = C₁ cos(2t) + C₂ sin(2t)

Para la solución particular, necesitamos conocer la forma específica de f(t). Si f(t) es una función conocida, podemos usar el método de coeficientes indeterminados o variación de parámetros.

La solución general será:
x(t) = C₁ cos(2t) + C₂ sin(2t) + x_p(t)

Aplicando las condiciones iniciales:
x(1) = C₁ cos(2) + C₂ sin(2) + x_p(1) = 0
x'(1) = -2C₁ sin(2) + 2C₂ cos(2) + x_p'(1) = n

Resolviendo este sistema de ecuaciones obtenemos los valores de C₁ y C₂.
    """
    doc.add_paragraph(sol9_text.strip())
    
    # Ejercicio 10
    doc.add_heading('Ejercicio 10: Ecuación Diferencial Parcial', level=3)
    doc.add_paragraph('Verificar que la función dada es solución de la ecuación diferencial parcial:')
    doc.add_paragraph('∂²u/∂x∂y = x + y³')
    doc.add_paragraph('Función propuesta:')
    doc.add_paragraph('u(x,y) = (1/2)x²y + (1/4)xy⁴ - (1/4)y⁴ + 2y² - (9/2)y + x² - 3x - 29')
    doc.add_paragraph('Verificación:')
    sol10_text = """
Calculamos las derivadas parciales:

∂u/∂x = (1/2) · 2x · y + (1/4) y⁴ + 2x - 3 = xy + (1/4)y⁴ + 2x - 3

∂²u/∂x∂y = ∂/∂y(∂u/∂x) = ∂/∂y(xy + (1/4)y⁴ + 2x - 3)
          = x + (1/4) · 4y³ = x + y³

Como ∂²u/∂x∂y = x + y³, se verifica que la función propuesta es efectivamente una solución de la ecuación diferencial parcial dada.
    """
    doc.add_paragraph(sol10_text.strip())
    
    # Guardar el documento
    filename = "Entregable2_Completo.docx"
    doc.save(filename)
    print(f"Documento Word completo creado: {filename}")
    
    return filename

if __name__ == "__main__":
    create_complete_word_document()
