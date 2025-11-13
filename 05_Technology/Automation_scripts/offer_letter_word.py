#!/usr/bin/env python3
"""
Exportación a Word para Cartas de Oferta
Convierte cartas de oferta a formato Word (.docx)
"""

import sys
from typing import Optional


def generate_word_from_text(text_file: str, word_file: Optional[str] = None) -> bool:
    """Genera archivo Word desde archivo de texto."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if not word_file:
            word_file = text_file.replace('.txt', '.docx')
        
        print(f"🔄 Generando Word desde {text_file}...")
        
        # Leer archivo de texto
        with open(text_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Procesar contenido
        lines = content.split('\n')
        in_section = False
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Título principal
            if 'OFFER OF EMPLOYMENT' in line and line.startswith('='):
                p = doc.add_paragraph('OFFER OF EMPLOYMENT')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.bold = True
                doc.add_paragraph()
                continue
            
            # Encabezados de sección
            if line.isupper() and len(line) < 50 and not line.startswith('•') and '---' not in line:
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = None  # Color por defecto
                in_section = True
                continue
            
            # Separadores
            if line.startswith('=') or line.startswith('-'):
                continue
            
            # Texto normal
            p = doc.add_paragraph(line)
            if in_section:
                p.paragraph_format.left_indent = Inches(0.25)
        
        # Guardar
        doc.save(word_file)
        
        print(f"✅ Word generado exitosamente: {word_file}")
        return True
        
    except ImportError:
        print("❌ Error: python-docx no está instalado", file=sys.stderr)
        print("   Instálalo con: pip install python-docx", file=sys.stderr)
        return False
    except Exception as e:
        print(f"❌ Error generando Word: {e}", file=sys.stderr)
        return False


def generate_word_from_html(html_file: str, word_file: Optional[str] = None) -> bool:
    """Genera archivo Word desde HTML."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        if not word_file:
            word_file = html_file.replace('.html', '.docx')
        
        print(f"🔄 Generando Word desde {html_file}...")
        
        # Leer HTML
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Crear documento Word
        doc = Document()
        
        # Extraer texto del HTML (simple, sin tags)
        text = re.sub(r'<[^>]+>', '', html_content)
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        
        # Procesar líneas
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()
        
        # Guardar
        doc.save(word_file)
        
        print(f"✅ Word generado exitosamente: {word_file}")
        return True
        
    except ImportError:
        print("❌ Error: python-docx no está instalado", file=sys.stderr)
        print("   Instálalo con: pip install python-docx", file=sys.stderr)
        return False
    except Exception as e:
        print(f"❌ Error generando Word: {e}", file=sys.stderr)
        return False


def main():
    """Función principal."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Genera archivos Word de cartas de oferta')
    parser.add_argument('input_file',
                       help='Archivo de entrada (TXT o HTML)')
    parser.add_argument('--output', '-o',
                       help='Archivo Word de salida')
    parser.add_argument('--from-html', action='store_true',
                       help='Forzar entrada como HTML')
    
    args = parser.parse_args()
    
    # Determinar tipo
    if args.from_html or args.input_file.endswith('.html'):
        success = generate_word_from_html(args.input_file, args.output)
    else:
        success = generate_word_from_text(args.input_file, args.output)
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()





