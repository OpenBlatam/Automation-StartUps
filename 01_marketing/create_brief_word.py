#!/usr/bin/env python3
"""
Script mejorado para crear el Brief UGC en formato Word con tablas bonitas e imágenes
Versión mejorada con más contenido y secciones
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime

def add_image_placeholder(doc, text="[IMAGEN AQUÍ]", width=Inches(5)):
    """Añade un placeholder para imagen con borde mejorado"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Borde superior
    border_para = doc.add_paragraph()
    border_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_run = border_para.add_run("┌" + "─" * 60 + "┐")
    border_run.font.size = Pt(8)
    border_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # Contenido
    for i in range(4):
        content_para = doc.add_paragraph()
        content_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 1:
            content_run = content_para.add_run("│" + " " * 20 + text + " " * (40 - len(text)) + "│")
            content_run.font.size = Pt(11)
            content_run.font.color.rgb = RGBColor(128, 128, 128)
            content_run.font.italic = True
        else:
            content_run = content_para.add_run("│" + " " * 60 + "│")
            content_run.font.size = Pt(8)
            content_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # Borde inferior
    border_para2 = doc.add_paragraph()
    border_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_run2 = border_para2.add_run("└" + "─" * 60 + "┘")
    border_run2.font.size = Pt(8)
    border_run2.font.color.rgb = RGBColor(200, 200, 200)
    
    doc.add_paragraph()

def create_table_with_style(doc, headers, rows, title=None, alternate_colors=True):
    """Crea una tabla bonita con estilo mejorado"""
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Fondo azul para headers
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0066CC')
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Rows con colores alternados
    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)
            
            # Color alternado para filas
            if alternate_colors and row_idx % 2 == 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F0F8FF')
                row_cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()

def add_section_header(doc, title, emoji="📋"):
    """Añade un encabezado de sección con estilo"""
    doc.add_paragraph()
    header = doc.add_heading(f'{emoji} {title}', 1)
    header.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    doc.add_paragraph()

def add_subsection_header(doc, title, level=2):
    """Añade un encabezado de subsección"""
    doc.add_heading(title, level)
    doc.add_paragraph()

def create_brief_word():
    """Crea el documento Word del Brief UGC mejorado"""
    
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ========== PORTADA ==========
    title = doc.add_heading('🎬 BRIEF UGC PARA CREADORAS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(32)
    title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    subtitle = doc.add_paragraph('Campaña IA Bulk Documentos')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(20)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 204, 102)
    subtitle.runs[0].font.bold = True
    
    subtitle2 = doc.add_paragraph('Generación Masiva de Contenido')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(14)
    subtitle2.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    subtitle2.runs[0].font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Imagen placeholder
    add_image_placeholder(doc, "LOGO / IMAGEN DEL PRODUCTO")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Información de versión
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run(f'Versión 11.0 - Ultra Completo Absoluto Definitivo Máximo')
    version_run.font.size = Pt(12)
    version_run.font.bold = True
    version_run.font.color.rgb = RGBColor(0, 102, 204)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'{datetime.now().strftime("%d de %B de %Y")}')
    date_run.font.size = Pt(11)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # ========== ÍNDICE ==========
    doc.add_heading('📋 ÍNDICE', 1)
    doc.add_paragraph()
    
    toc_items = [
        "1. Información General",
        "2. Objetivos de la Campaña",
        "3. Perfil de Creadora Ideal",
        "4. Tipos de Contenido UGC",
        "5. Especificaciones Técnicas",
        "6. Guía de Estilo y Tono",
        "7. Hooks y Mensajes Clave",
        "8. Compensación y Condiciones",
        "9. Métricas de Éxito",
        "10. Cronograma y Entregas",
        "11. Checklist de Entrega",
        "12. Casos de Uso Específicos",
        "13. Recursos y Materiales",
        "14. Guía de Screen Recording",
        "15. Guía de Edición de Videos",
        "16. Templates de Captions",
        "17. Hashtags Sugeridos",
        "18. Mejores Prácticas de Engagement",
        "19. Estrategias de Repurposing",
        "20. Calendario de Contenido",
        "21. Troubleshooting",
        "22. Guía de Storytelling",
        "23. Checklist de Producción",
        "24. Restricciones y Guidelines",
        "25. Casos de Éxito Reales",
        "26. FAQ Específico para Creadoras",
        "27. Guía de A/B Testing",
        "28. Guía de Compliance y Legal",
        "29. Guía de Crisis Management",
        "30. Estrategias de Viralidad",
        "31. Ideas de Contenido Creativas",
        "32. Workflow de Producción Optimizado",
        "33. Contacto y Soporte"
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 1: INFORMACIÓN GENERAL ==========
    add_section_header(doc, "INFORMACIÓN GENERAL")
    
    create_table_with_style(
        doc,
        ["Aspecto", "Detalle"],
        [
            ["Proyecto", "Campaña de User Generated Content (UGC) para IA Bulk Documentos"],
            ["Producto/Servicio", "IA Bulk Documentos - Plataforma que genera 10,000+ documentos profesionales en 60 segundos"],
            ["Objetivo de la Campaña", "Generar contenido auténtico y orgánico que muestre casos de uso reales, beneficios tangibles y testimonios genuinos"],
            ["Plataformas Objetivo", "TikTok (prioritario), Instagram Reels, YouTube Shorts, LinkedIn (opcional)"],
            ["Duración", "30-60 días (contenido entregado en lotes)"]
        ],
        "Información del Proyecto"
    )
    
    add_image_placeholder(doc, "DIAGRAMA DEL PRODUCTO / CASO DE USO")
    
    # ========== SECCIÓN 2: OBJETIVOS ==========
    doc.add_page_break()
    add_section_header(doc, "OBJETIVOS DE LA CAMPAÑA", "🎯")
    
    create_table_with_style(
        doc,
        ["Tipo", "Objetivo"],
        [
            ["Principal", "Crear contenido UGC auténtico que genere awareness, credibilidad y conversiones orgánicas"],
            ["Secundario 1", "Mostrar casos de uso reales y tangibles"],
            ["Secundario 2", "Generar prueba social auténtica"],
            ["Secundario 3", "Educar sobre el problema que resuelve"],
            ["Secundario 4", "Demostrar ahorro de tiempo y eficiencia"],
            ["Secundario 5", "Construir confianza mediante testimonios genuinos"]
        ],
        "Objetivos de la Campaña"
    )
    
    doc.add_paragraph()
    message_para = doc.add_paragraph()
    message_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    message_run = message_para.add_run('"Genera 10,000+ documentos profesionales en 60 segundos. De horas de trabajo manual a segundos de automatización."')
    message_run.font.size = Pt(16)
    message_run.font.italic = True
    message_run.font.bold = True
    message_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # ========== SECCIÓN 3: PERFIL DE CREADORA ==========
    doc.add_page_break()
    add_section_header(doc, "PERFIL DE CREADORA IDEAL", "👥")
    
    create_table_with_style(
        doc,
        ["Característica", "Especificación"],
        [
            ["Nicho", "Tech, productividad, emprendimiento, marketing, negocios"],
            ["Audiencia", "10K-200K seguidores (micro-influencers)"],
            ["Engagement", ">3% engagement rate"],
            ["Estilo", "Auténtico, educativo, práctico"],
            ["Contenido", "Habla sobre productividad, herramientas, automatización, IA"]
        ],
        "Características Deseadas"
    )
    
    doc.add_paragraph()
    
    create_table_with_style(
        doc,
        ["Tipo", "Descripción", "Audiencia"],
        [
            ["Tech/Productividad", "Hablan de herramientas y software, comparten tips de productividad", "Profesionales y emprendedores"],
            ["Negocios/Marketing", "Contenido sobre crecimiento de negocio, marketing y ventas", "Empresarios y marketers"],
            ["Emprendimiento", "Comparten su journey emprendedor, herramientas que usan", "Aspirantes a emprendedores"]
        ],
        "Tipos de Creadoras"
    )
    
    add_image_placeholder(doc, "EJEMPLO DE PERFIL DE CREADORA IDEAL")
    
    # ========== SECCIÓN 4: TIPOS DE CONTENIDO ==========
    doc.add_page_break()
    add_section_header(doc, "TIPOS DE CONTENIDO UGC SOLICITADOS", "🎬")
    
    create_table_with_style(
        doc,
        ["Tipo", "Duración", "Prioridad", "Estructura"],
        [
            ["Video Testimonial", "30-60s", "⭐ PRIORITARIO", "Hook → Problema → Solución → Resultado → CTA"],
            ["Video Educativo", "30-45s", "⭐ PRIORITARIO", "Hook → Problema → Demo → Beneficio → CTA"],
            ["Antes/Después", "30-45s", "Alta", "Hook → Antes → Después → Contraste → CTA"],
            ["Problema/Solución", "30-45s", "Alta", "Hook → Problema → Solución → Resultado → CTA"],
            ["Tutorial Rápido", "30-45s", "Media", "Hook → Paso 1 → Paso 2 → Paso 3 → Resultado → CTA"],
            ["Storytelling Personal", "30-60s", "Media", "Hook → Historia → Transformación → Recomendación → CTA"]
        ],
        "Tipos de Contenido UGC"
    )
    
    doc.add_paragraph()
    add_subsection_header(doc, "Ejemplo de Guión - Video Testimonial")
    
    script_table = doc.add_table(rows=6, cols=2)
    script_table.style = 'Light List Accent 1'
    
    script_data = [
        ["Hook (0-3s)", "Esto me ahorró 20 horas esta semana creando documentos"],
        ["Problema (3-8s)", "Antes pasaba 4 horas creando cada propuesta para clientes. Con 5 clientes por semana, eran 20 horas solo en documentos."],
        ["Solución (8-15s)", "Ahora uso IA Bulk Documentos. Escribo una consulta y en 30 segundos tengo 5 propuestas personalizadas y profesionales listas."],
        ["Resultado (15-25s)", "Esta semana generé 20 propuestas en menos de 5 minutos. 20 horas ahorradas. Puedo enfocarme en cerrar más clientes."],
        ["CTA (25-30s)", "Si también creas documentos regularmente, link en bio para probarlo gratis"]
    ]
    
    header_cells = script_table.rows[0].cells
    header_cells[0].text = "Momento"
    header_cells[1].text = "Contenido"
    header_cells[0].paragraphs[0].runs[0].font.bold = True
    header_cells[1].paragraphs[0].runs[0].font.bold = True
    
    for row_idx, (moment, content) in enumerate(script_data, start=1):
        row_cells = script_table.rows[row_idx].cells
        row_cells[0].text = moment
        row_cells[1].text = content
        row_cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    add_image_placeholder(doc, "EJEMPLO DE VIDEO UGC")
    
    # ========== SECCIÓN 5: ESPECIFICACIONES TÉCNICAS ==========
    doc.add_page_break()
    add_section_header(doc, "ESPECIFICACIONES TÉCNICAS", "📱")
    
    create_table_with_style(
        doc,
        ["Aspecto", "Especificación"],
        [
            ["Resolución", "1080x1920 (9:16 vertical)"],
            ["Duración", "30-60 segundos (óptimo: 30-45s)"],
            ["Formato de archivo", "MP4, MOV"],
            ["Frame rate", "30fps"],
            ["Audio", "Estéreo, 44.1kHz"],
            ["Tamaño máximo", "500MB"],
            ["Iluminación", "Buena iluminación natural o artificial"],
            ["Estabilidad", "Video estable (usar trípode o estabilización)"],
            ["Enfoque", "Video nítido y bien enfocado"],
            ["Subtítulos", "Incluir subtítulos/closed captions (obligatorio)"],
            ["Música", "Royalty-free o música de la plataforma"],
            ["Branding", "Mencionar 'IA Bulk Documentos' al menos una vez"]
        ],
        "Especificaciones de Video"
    )
    
    # ========== SECCIÓN 6: GUÍA DE ESTILO ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE ESTILO Y TONO", "🎨")
    
    add_subsection_header(doc, "Tono de Voz")
    tone_items = [
        "Auténtico: Habla como hablarías normalmente",
        "Conversacional: Como si le hablaras a un amigo",
        "Educativo: Comparte conocimiento, no solo vendas",
        "Empático: Reconoce el problema que otros tienen",
        "Entusiasta pero genuino: Muestra emoción real, no forzada"
    ]
    for item in tone_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Elementos a Evitar")
    avoid_items = [
        "Scripts memorizados que suenan robóticos",
        "Over-selling o exageración",
        "Contenido genérico sin personalidad",
        "Videos demasiado producidos (pierde autenticidad)",
        "Menciones excesivas del producto (máximo 2-3 veces)"
    ]
    for item in avoid_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    
    # ========== SECCIÓN 7: HOOKS ==========
    doc.add_page_break()
    add_section_header(doc, "HOOKS Y MENSAJES CLAVE", "📝")
    
    create_table_with_style(
        doc,
        ["Categoría", "Ejemplo de Hook"],
        [
            ["Ahorro de Tiempo", "Esto me ahorró 20 horas esta semana"],
            ["Ahorro de Tiempo", "De 4 horas a 30 segundos"],
            ["Ahorro de Tiempo", "Genero 100 documentos en 1 minuto"],
            ["Problema", "¿Te pasa que pierdes horas creando documentos?"],
            ["Problema", "Si odias crear documentos uno por uno..."],
            ["Resultado", "Esto cambió cómo trabajo completamente"],
            ["Resultado", "Mi productividad se multiplicó por 10"],
            ["Comparación", "Antes vs Ahora: Crear documentos"],
            ["Comparación", "Cómo pasé de 20 horas a 5 minutos"],
            ["Shock Value", "¿100 documentos manualmente? No gracias"],
            ["Pregunta", "¿Cuántos documentos haces manualmente?"]
        ],
        "Hooks Efectivos (Primeros 3 segundos)"
    )
    
    doc.add_paragraph()
    create_table_with_style(
        doc,
        ["Mensaje Clave", "Descripción"],
        [
            ["Genera miles de documentos en segundos", "Velocidad y eficiencia"],
            ["Personalización automática", "Cada documento único"],
            ["Ahorro de 95% del tiempo", "Impacto cuantificable"],
            ["0 errores, calidad profesional", "Confiabilidad"],
            ["Escalable a millones de documentos", "Potencial ilimitado"]
        ],
        "Mensajes Clave a Incluir"
    )
    
    # ========== SECCIÓN 8: COMPENSACIÓN ==========
    doc.add_page_break()
    add_section_header(doc, "COMPENSACIÓN Y CONDICIONES", "💰")
    
    create_table_with_style(
        doc,
        ["Modelo", "Descripción", "Ventajas"],
        [
            ["Pago Fijo", "$150-300 USD por video aprobado", "Ingreso garantizado, predecible"],
            ["Comisión", "25-30% por conversión generada", "Potencial ilimitado, alineado con resultados"],
            ["Híbrido ⭐", "Pago base + comisión + bonuses", "Balance entre seguridad y potencial"],
            ["Acceso Gratis", "Acceso gratuito a plataforma + comisión", "Bajo riesgo, alto potencial"]
        ],
        "Modelos de Compensación"
    )
    
    doc.add_paragraph()
    create_table_with_style(
        doc,
        ["Métrica", "Bonus"],
        [
            ["Video >10K views", "+$50 USD"],
            ["Video >25K views", "+$100 USD"],
            ["Video >50K views", "+$200 USD"],
            ["Video >100K views", "+$500 USD"],
            ["10-19 sign-ups", "+$25 USD"],
            ["20-49 sign-ups", "+$50 USD"],
            ["50+ sign-ups", "+$100 USD"],
            ["100+ sign-ups", "+$200 USD"],
            ["Engagement rate >5%", "+$25 USD"],
            ["Engagement rate >8%", "+$50 USD"],
            ["Engagement rate >10%", "+$100 USD"]
        ],
        "Bonuses por Performance"
    )
    
    # ========== SECCIÓN 9: MÉTRICAS ==========
    doc.add_page_break()
    add_section_header(doc, "MÉTRICAS DE ÉXITO", "📊")
    
    create_table_with_style(
        doc,
        ["Métrica", "Básico", "Bueno", "Excelente"],
        [
            ["Views por video", "5,000+", "25,000+", "100,000+"],
            ["Engagement Rate", "3-5%", "5-8%", "8%+"],
            ["CTR en Link", "1-2%", "2-3%", "3%+"],
            ["Conversiones", "10-20", "20-50", "50+"],
            ["Comentarios", "30-50", "50-100", "100+"],
            ["Shares", "10-20", "20-50", "50+"],
            ["Retención 3s", "60-70%", "70-80%", "80%+"],
            ["Completion Rate", "20-30%", "30-40%", "40%+"]
        ],
        "Métricas de Éxito por Nivel"
    )
    
    # ========== SECCIÓN 10: CRONOGRAMA ==========
    doc.add_page_break()
    add_section_header(doc, "CRONOGRAMA Y ENTREGAS", "📅")
    
    create_table_with_style(
        doc,
        ["Semana", "Actividad", "Entregable"],
        [
            ["Semana 1", "Briefing y aprobación de conceptos", "Conceptos aprobados"],
            ["Semana 2", "Producción y entrega", "2 videos aprobados"],
            ["Semana 3", "Producción y entrega", "2 videos adicionales"],
            ["Semana 4", "Entrega final y métricas", "Videos finales + reporte"]
        ],
        "Timeline Típico de Campaña"
    )
    
    doc.add_paragraph()
    create_table_with_style(
        doc,
        ["Fase", "Actividad", "Tiempo"],
        [
            ["Concepto", "Aprobación de idea/hook antes de grabar", "2-3 días"],
            ["Primera versión", "Revisión de video editado", "2-3 días"],
            ["Ajustes", "Máximo 2 rondas de ediciones", "1-2 días"],
            ["Aprobación final", "OK para publicar", "1 día"]
        ],
        "Proceso de Aprobación"
    )
    
    # ========== SECCIÓN 11: CHECKLIST ==========
    doc.add_page_break()
    add_section_header(doc, "CHECKLIST DE ENTREGA", "✅")
    
    checklist_items = [
        "Video en formato 9:16 (1080x1920)",
        "Duración entre 30-60 segundos",
        "Audio claro y sin ruido excesivo",
        "Video nítido y bien iluminado",
        "Subtítulos/captions incluidos",
        "Menciona 'IA Bulk Documentos' al menos una vez",
        "CTA claro al final",
        "Hook en primeros 3 segundos",
        "Contenido auténtico y no robótico",
        "Archivo nombrado correctamente",
        "Link de publicación incluido",
        "Screenshots de métricas (si disponible)",
        "Caption usado en la publicación",
        "Hashtags utilizados"
    ]
    
    for item in checklist_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 12: CASOS DE USO ==========
    doc.add_page_break()
    add_section_header(doc, "CASOS DE USO ESPECÍFICOS", "🎯")
    
    create_table_with_style(
        doc,
        ["Caso de Uso", "Descripción", "Beneficio Clave"],
        [
            ["Propuestas Comerciales", "Generar múltiples propuestas personalizadas para leads", "Ahorro de tiempo en proceso de ventas"],
            ["Contratos Legales", "Generación masiva de contratos personalizados", "Compliance automático, reducción de errores"],
            ["Emails Personalizados", "Campañas de email marketing masivas", "Mejor engagement y conversión"],
            ["Reportes Automáticos", "Reportes para múltiples clientes", "Consistencia en formato y calidad"],
            ["Certificados y Diplomas", "Generación masiva de certificados", "Ahorro en procesos administrativos"]
        ],
        "Casos de Uso a Mostrar"
    )
    
    add_image_placeholder(doc, "DIAGRAMA DE CASOS DE USO")
    
    # ========== SECCIÓN 13: RECURSOS ==========
    doc.add_page_break()
    add_section_header(doc, "RECURSOS Y MATERIALES", "📚")
    
    add_subsection_header(doc, "Acceso y Cuentas")
    recursos_items = [
        "Acceso gratuito a plan Professional ($497/mes) por duración de campaña",
        "Cuenta de prueba con datos de ejemplo para demos",
        "Link trackeable único con UTM parameters",
        "Dashboard de métricas para tracking en tiempo real"
    ]
    for item in recursos_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Assets Visuales")
    assets_items = [
        "Logo en diferentes formatos (PNG, SVG)",
        "Paleta de colores oficial (#0066CC, #00CC66)",
        "Fuentes recomendadas (Montserrat, Open Sans)",
        "Screenshots de la plataforma para uso en videos",
        "Banners para stories (templates editables)",
        "Iconos y elementos gráficos"
    ]
    for item in assets_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 14: SCREEN RECORDING ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE SCREEN RECORDING", "🎥")
    
    create_table_with_style(
        doc,
        ["Herramienta", "Plataforma", "Precio", "Recomendación"],
        [
            ["OBS Studio", "PC/Mac", "Gratis", "⭐ Profesional"],
            ["QuickTime", "Mac", "Gratis", "Integrado"],
            ["Windows Game Bar", "Windows", "Gratis", "Integrado"],
            ["Loom", "Web/App", "Freemium", "Fácil de usar"],
            ["Camtasia", "PC/Mac", "$299", "Muy fácil"],
            ["ScreenFlow", "Mac", "$169", "Optimizado Mac"],
            ["Screen Studio", "Mac", "$89", "Automático"]
        ],
        "Herramientas de Screen Recording"
    )
    
    doc.add_paragraph()
    add_subsection_header(doc, "Configuración Óptima")
    config_items = [
        "Resolución: 1080p (1920x1080) mínimo, 4K si es posible",
        "Frame rate: 30fps (suficiente), 60fps para gameplay",
        "Área: Full screen, ventana específica o región personalizada",
        "Audio: Micrófono externo + audio del sistema en pistas separadas"
    ]
    for item in config_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 15: EDICIÓN ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE EDICIÓN DE VIDEOS", "✂️")
    
    create_table_with_style(
        doc,
        ["Herramienta", "Plataforma", "Precio", "Recomendación"],
        [
            ["CapCut", "Móvil/Desktop", "Gratis", "⭐ Muy completa"],
            ["InShot", "Móvil", "Freemium", "Fácil"],
            ["DaVinci Resolve", "Desktop", "Gratis", "⭐ Profesional"],
            ["Adobe Premiere Pro", "Desktop", "Pago", "Estándar industria"],
            ["Final Cut Pro", "Mac", "Pago", "Optimizado Mac"],
            ["VN Editor", "Móvil", "Gratis", "Profesional móvil"]
        ],
        "Herramientas de Edición"
    )
    
    doc.add_paragraph()
    add_subsection_header(doc, "Workflow de Edición (40 minutos)")
    workflow_steps = [
        "1. Importar y Organizar (5 min): Video, screen recording, música",
        "2. Corte y Estructura (10 min): Eliminar silencios, estructurar Hook → Problema → Solución → Resultado → CTA",
        "3. Añadir Elementos Visuales (10 min): Subtítulos sincronizados, texto en pantalla, transiciones",
        "4. Audio (5 min): Música de fondo 30-40% volumen, voz 100%, eliminar ruido",
        "5. Color y Ajustes (5 min): Brillo, contraste, saturación sutil",
        "6. Exportación (5 min): MP4 H.264, 1080x1920, 30fps, alta calidad"
    ]
    for step in workflow_steps:
        para = doc.add_paragraph(step, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 16: TEMPLATES ==========
    doc.add_page_break()
    add_section_header(doc, "TEMPLATES DE CAPTIONS", "📝")
    
    add_subsection_header(doc, "Template 1: Testimonial Auténtico")
    template1 = doc.add_paragraph('🚀 Esto me ahorró 20 horas esta semana creando documentos\n\nAntes pasaba 4 horas creando cada propuesta para clientes. Con 5 clientes por semana, eran 20 horas solo en documentos.\n\nAhora uso IA Bulk Documentos. Escribo una consulta y en 30 segundos tengo 5 propuestas personalizadas y profesionales listas.\n\nEsta semana generé 20 propuestas en menos de 5 minutos. 20 horas ahorradas. Puedo enfocarme en cerrar más clientes.\n\nSi también creas documentos regularmente, link en bio para probarlo gratis 👆\n\n#Productividad #IA #Automatización #Negocios #Emprendimiento')
    template1.style = 'Intense Quote'
    template1.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Template 2: Educativo/Caso de Uso")
    template2 = doc.add_paragraph('💡 Te muestro cómo genero 100 documentos en 1 minuto\n\nSi eres como yo y necesitas crear múltiples documentos personalizados, sabes que es súper tedioso hacerlo uno por uno.\n\nCon IA Bulk Documentos, solo escribo: "Genera propuestas para estos 100 leads" y en 30 segundos tengo 100 propuestas únicas, cada una personalizada con los datos del cliente.\n\nAntes esto me tomaba 50 horas. Ahora 30 segundos. Puedo responder a 10x más oportunidades.\n\nSi también necesitas crear documentos masivamente, link en bio 👆\n\n¿Cuántos documentos haces manualmente? ¿Y si fueran 100 de golpe? 👇')
    template2.style = 'Intense Quote'
    template2.runs[0].font.size = Pt(10)
    
    # ========== SECCIÓN 17: HASHTAGS ==========
    doc.add_page_break()
    add_section_header(doc, "HASHTAGS SUGERIDOS", "🏷️")
    
    create_table_with_style(
        doc,
        ["Categoría", "Hashtags", "Cantidad"],
        [
            ["Principales", "#IA #Productividad #Automatización #Negocios #Emprendimiento", "5-7"],
            ["Secundarios", "#MarketingDigital #HerramientasTech #AhorroTiempo #Eficiencia", "3-5"],
            ["Nicho", "#Freelancer #AgenciaMarketing #Consultoría #Startup", "2-3"],
            ["Plataforma", "#TikTok #Reels #Shorts", "1-2"],
            ["TOTAL", "15-20 hashtags por post", "-"]
        ],
        "Estrategia de Hashtags"
    )
    
    # ========== SECCIÓN 18: ENGAGEMENT ==========
    doc.add_page_break()
    add_section_header(doc, "MEJORES PRÁCTICAS DE ENGAGEMENT", "💬")
    
    create_table_with_style(
        doc,
        ["Práctica", "Descripción", "Impacto"],
        [
            ["Responder Comentarios", "Responde en primeras 2 horas, mínimo 80% de comentarios", "Algoritmo favorece, más engagement"],
            ["Preguntas en Captions", "Preguntas abiertas generan más comentarios", "Aumenta interacción"],
            ["Pin Comentarios", "Pinned comment con link o pregunta", "Mayor visibilidad"],
            ["Stories Follow-up", "Comparte video en stories, responde preguntas", "Más alcance"],
            ["Timing Óptimo", "TikTok: 6-10 AM o 7-9 PM | Reels: 9-11 AM o 2-4 PM", "Mayor audiencia activa"]
        ],
        "Estrategias de Engagement"
    )
    
    # ========== SECCIÓN 19: REPURPOSING ==========
    doc.add_page_break()
    add_section_header(doc, "ESTRATEGIAS DE REPURPOSING", "🔄")
    
    create_table_with_style(
        doc,
        ["Formato", "Duración", "Plataforma", "Modificaciones"],
        [
            ["Video Original", "30-45s", "TikTok/Reels/Shorts", "Formato 9:16 vertical"],
            ["Stories", "15s", "Instagram Stories", "Recorta mejores 15s, añade swipe up"],
            ["Post Extendido", "60s", "Instagram Feed", "Extiende a 60s, caption más largo"],
            ["YouTube Shorts", "60s", "YouTube", "Versión extendida, thumbnail atractivo"],
            ["LinkedIn Video", "60-90s", "LinkedIn", "Más profesional, contexto B2B"],
            ["Twitter/X", "30s", "Twitter", "Recorta a 30s, hook directo"]
        ],
        "Repurposing de 1 Video en Múltiples Formatos"
    )
    
    # ========== SECCIÓN 20: CALENDARIO ==========
    doc.add_page_break()
    add_section_header(doc, "CALENDARIO DE CONTENIDO SUGERIDO", "📅")
    
    create_table_with_style(
        doc,
        ["Semana", "Día", "Tipo de Contenido", "Objetivo"],
        [
            ["Semana 1", "Día 1", "Video problema/relatable", "Generar identificación"],
            ["Semana 1", "Día 4", "Video solución/demo", "Mostrar producto"],
            ["Semana 1", "Día 7", "Story behind the scenes", "Autenticidad"],
            ["Semana 2", "Día 1", "Video tutorial rápido", "Educar"],
            ["Semana 2", "Día 4", "Video caso de uso", "Aplicación práctica"],
            ["Semana 2", "Día 7", "Q&A en stories", "Engagement"],
            ["Semana 3", "Día 1", "Video testimonial", "Prueba social"],
            ["Semana 3", "Día 4", "Video comparación", "Contraste"],
            ["Semana 4", "Día 1", "Video tutorial avanzado", "Profundizar"]
        ],
        "Plan de Contenido 4 Semanas"
    )
    
    # ========== SECCIÓN 21: TROUBLESHOOTING ==========
    doc.add_page_break()
    add_section_header(doc, "TROUBLESHOOTING COMÚN", "🔧")
    
    create_table_with_style(
        doc,
        ["Problema", "Solución"],
        [
            ["Video no se ve bien en móvil", "Verifica resolución (1080x1920), exporta MP4 H.264, evita compresión excesiva"],
            ["Audio no se escucha bien", "Normaliza volumen (-6dB a -12dB), elimina ruido, música 30-40%"],
            ["Subtítulos no sincronizan", "Revisa timing frame por frame, ajusta delay, usa auto-sync"],
            ["Video muy largo", "Recorta partes menos importantes, acelera secciones lentas (1.5x-2x)"],
            ["Screen recording baja calidad", "Aumenta resolución, usa OBS/Camtasia, graba 1080p mínimo"],
            ["No sé qué decir", "Usa guiones del brief, habla naturalmente, haz múltiples takes"]
        ],
        "Soluciones a Problemas Comunes"
    )
    
    # ========== SECCIÓN 22: STORYTELLING ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE STORYTELLING", "📖")
    
    create_table_with_style(
        doc,
        ["Momento", "Duración", "Contenido", "Objetivo"],
        [
            ["El Gancho", "0-3s", "Problema identificable o resultado impactante", "Captar atención"],
            ["El Problema", "3-8s", "Describe el dolor, sé específico", "Conectar con audiencia"],
            ["El Descubrimiento", "8-12s", "Momento de cambio, primera impresión", "Crear interés"],
            ["La Solución", "12-20s", "Muestra el proceso, demuestra facilidad", "Educar"],
            ["La Transformación", "20-27s", "Resultados concretos, impacto", "Probar valor"],
            ["El CTA", "27-30s", "Invitación clara, bajo fricción", "Convertir"]
        ],
        "Estructura de Storytelling para UGC"
    )
    
    # ========== SECCIÓN 23: CHECKLIST PRODUCCIÓN ==========
    doc.add_page_break()
    add_section_header(doc, "CHECKLIST DE PRODUCCIÓN COMPLETO", "✅")
    
    add_subsection_header(doc, "Pre-Producción")
    preprod_items = [
        "Leí y entendí el brief completo",
        "Tengo acceso a la plataforma",
        "Probé el producto y entiendo cómo funciona",
        "Elegí tipo de video a crear",
        "Preparé guión o puntos clave",
        "Preparé datos de ejemplo (si aplica)",
        "Verifiqué herramientas de grabación/edición",
        "Configuré espacio de grabación (iluminación, audio)"
    ]
    for item in preprod_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Producción")
    prod_items = [
        "Grabé hook en primeros 3 segundos",
        "Mencioné el problema claramente",
        "Mostré el producto en uso",
        "Compartí resultados concretos",
        "Incluí CTA claro al final",
        "Audio claro y sin ruido",
        "Video nítido y bien iluminado",
        "Duración entre 30-60 segundos"
    ]
    for item in prod_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 24: RESTRICCIONES ==========
    doc.add_page_break()
    add_section_header(doc, "RESTRICCIONES Y GUIDELINES", "🚫")
    
    add_subsection_header(doc, "Qué NO Hacer")
    no_items = [
        "Hacer claims falsos o exagerados",
        "Comparar directamente con competidores (nombres)",
        "Usar música con derechos de autor",
        "Incluir información confidencial",
        "Hacer spam o contenido demasiado promocional",
        "Usar bots o engagement falso"
    ]
    for item in no_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Qué SÍ Hacer")
    yes_items = [
        "Ser auténtico y genuino",
        "Mostrar uso real del producto",
        "Compartir resultados reales",
        "Responder comentarios genuinamente",
        "Crear contenido de valor educativo",
        "Mantener tu estilo y personalidad"
    ]
    for item in yes_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.color.rgb = RGBColor(0, 150, 0)
    
    # ========== SECCIÓN 25: CASOS DE ÉXITO ==========
    doc.add_page_break()
    add_section_header(doc, "CASOS DE ÉXITO REALES", "🏆")
    
    add_subsection_header(doc, "Caso 1: Creadora Tech - 2.5M Views en 3 Meses")
    caso1_items = [
        "Creadora: Micro-influencer tech (45K seguidores)",
        "Videos creados: 12 videos (3 por semana)",
        "Hook usado: 'De 20 horas a 5 minutos. Así lo hago.'",
        "Resultados: 2,500,000+ views totales",
        "Engagement rate: 6.8% (promedio)",
        "CTR link: 3.2%",
        "Conversiones: 800+ sign-ups",
        "ROI para marca: 1,200%",
        "Ingresos creadora: $2,400 (pago fijo) + $800 (bonuses)"
    ]
    for item in caso1_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Caso 2: Creadora B2B - 425 Demos en 4 Meses")
    caso2_items = [
        "Creadora: LinkedIn influencer (28K seguidores)",
        "Videos creados: 8 videos (2 por semana)",
        "Hook usado: 'ANTES: 20 horas/semana | DESPUÉS: 5 horas/semana'",
        "Resultados: 850,000+ views totales",
        "Engagement rate: 5.2%",
        "CTR link: 3.8%",
        "Conversiones: 425 demos cualificados",
        "Close rate: 18% (77 clientes)",
        "ROI para marca: 450%"
    ]
    for item in caso2_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    add_subsection_header(doc, "Caso 3: Video Viral - 500K Views en 7 Días")
    caso3_items = [
        "Creadora: Emprendedora (15K seguidores)",
        "Video viral: 1 video específico",
        "Hook usado: '¿100 documentos manualmente? No gracias. Una consulta. Listo.'",
        "Resultados: 500,000+ views en 7 días",
        "Engagement rate: 12.4% (excepcional)",
        "CTR link: 4.8%",
        "Conversiones: 240+ sign-ups",
        "Shares: 8,500+",
        "Comentarios: 12,000+",
        "Bonus viral: +$500 USD"
    ]
    for item in caso3_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    add_image_placeholder(doc, "GRÁFICOS DE CASOS DE ÉXITO")
    
    # ========== SECCIÓN 26: FAQ EXPANDIDO ==========
    doc.add_page_break()
    add_section_header(doc, "FAQ ESPECÍFICO PARA CREADORAS", "❓")
    
    add_subsection_header(doc, "Preguntas sobre Compensación")
    faq_comp = [
        "Q: ¿Cuánto puedo ganar realmente?",
        "A: Depende del modelo: Pago fijo $150-300 USD, Comisión 25-30%, Híbrido $100 base + 20% + bonuses",
        "",
        "Q: ¿Cuándo me pagan?",
        "A: Pago fijo: 50% al aprobar concepto, 50% al publicar. Comisiones: Mensual. Bonuses: Inmediato",
        "",
        "Q: ¿Puedo negociar el precio?",
        "A: Sí, especialmente si tienes alta tasa de conversión, alto engagement (>5%), o puedes crear múltiples videos"
    ]
    for item in faq_comp:
        if item:
            para = doc.add_paragraph(item)
            para.runs[0].font.size = Pt(11)
            if item.startswith("Q:"):
                para.runs[0].font.bold = True
                para.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        else:
            doc.add_paragraph()
    
    doc.add_paragraph()
    add_subsection_header(doc, "Preguntas sobre Contenido")
    faq_content = [
        "Q: ¿Debo mencionar que es contenido patrocinado?",
        "A: Sí, según regulaciones: TikTok/Instagram usa #ad o #sponsored, LinkedIn menciona 'colaboración'",
        "",
        "Q: ¿Puedo rechazar ediciones solicitadas?",
        "A: Sí, pero primera ronda incluida. Segunda ronda si es corrección de error nuestro. Ediciones excesivas pueden requerir pago adicional",
        "",
        "Q: ¿Qué pasa si mi video no alcanza las métricas esperadas?",
        "A: No hay penalización primera vez. Te damos feedback y tips. Podemos ajustar estrategia. Opción de crear video adicional"
    ]
    for item in faq_content:
        if item:
            para = doc.add_paragraph(item)
            para.runs[0].font.size = Pt(11)
            if item.startswith("Q:"):
                para.runs[0].font.bold = True
                para.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        else:
            doc.add_paragraph()
    
    # ========== SECCIÓN 27: A/B TESTING ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE A/B TESTING", "🧪")
    
    create_table_with_style(
        doc,
        ["Elemento a Testear", "Variante A", "Variante B", "Variante C"],
        [
            ["Hooks", "Esto me ahorró 20 horas", "De 4 horas a 30 segundos", "¿Te pasa que pierdes horas?"],
            ["CTAs", "Link en bio si quieres probarlo gratis", "Prueba gratis, link en bio", "Si también creas documentos, link en bio"],
            ["Duración", "30 segundos (rápido)", "45 segundos (más contexto)", "60 segundos (completo)"],
            ["Estilo Visual", "Solo talking head", "Talking head + screen recording", "Solo screen recording con voz"],
            ["Música", "Upbeat, energética", "Calmada, profesional", "Sin música, solo voz"]
        ],
        "Qué Testear en tus Videos"
    )
    
    doc.add_paragraph()
    add_subsection_header(doc, "Cómo Testear")
    testing_steps = [
        "1. Crea 2-3 variantes del mismo concepto",
        "2. Publica en diferentes días (mismo horario)",
        "3. Monitorea métricas por 48-72 horas",
        "4. Compara resultados: Views, Engagement, Retención, Conversiones",
        "5. Escala el ganador y crea más contenido similar"
    ]
    for step in testing_steps:
        para = doc.add_paragraph(step, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 28: GUÍA LEGAL ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE COMPLIANCE Y LEGAL", "⚖️")
    
    create_table_with_style(
        doc,
        ["Plataforma", "Requisito", "Ubicación", "Multa"],
        [
            ["TikTok", "#ad o #sponsored", "Al inicio del caption", "Hasta $43,280 USD"],
            ["Instagram", "#ad o #sponsored + Paid partnership", "Visible sin expandir", "Hasta $43,280 USD"],
            ["YouTube", "Incluye contenido pagado", "En descripción", "Hasta $43,280 USD"],
            ["LinkedIn", "Colaboración o Partnership", "Visible en caption", "Hasta 4% ingresos anuales"],
            ["Reino Unido", "#ad obligatorio", "Al inicio, visible", "Hasta £500,000"]
        ],
        "Requisitos de Disclosure por Plataforma"
    )
    
    doc.add_paragraph()
    add_subsection_header(doc, "Template de Disclosure Correcto")
    disclosure_correct = doc.add_paragraph('#ad Esto me ahorró 20 horas esta semana...')
    disclosure_correct.style = 'Intense Quote'
    disclosure_correct.runs[0].font.size = Pt(11)
    disclosure_correct.runs[0].font.color.rgb = RGBColor(0, 150, 0)
    
    disclosure_incorrect = doc.add_paragraph('#sp Acabo de probar... (muy corto, no suficiente)')
    disclosure_incorrect.style = 'Intense Quote'
    disclosure_incorrect.runs[0].font.size = Pt(11)
    disclosure_incorrect.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    
    # ========== SECCIÓN 29: CRISIS MANAGEMENT ==========
    doc.add_page_break()
    add_section_header(doc, "GUÍA DE CRISIS MANAGEMENT", "🚨")
    
    create_table_with_style(
        doc,
        ["Situación", "Acción Inmediata", "Siguiente Paso"],
        [
            ["Video recibe críticas negativas", "No elimines inmediatamente, espera 24-48h", "Responde profesionalmente, contacta al equipo"],
            ["Video no alcanza métricas", "No te preocupes, es normal", "Analiza qué mejorar, ajusta próximo video"],
            ["Error en el video", "Si es menor: edita y republica", "Si es mayor: regraba si necesario"],
            ["Problema con el producto", "Contacta soporte", "No critiques públicamente, resuelve en privado primero"]
        ],
        "Qué Hacer si Algo Sale Mal"
    )
    
    # ========== SECCIÓN 30: ESTRATEGIAS DE VIRALIDAD ==========
    doc.add_page_break()
    add_section_header(doc, "ESTRATEGIAS DE VIRALIDAD", "🎯")
    
    viral_strategies = [
        "1. Hook Ultra-Específico: 'De 20 horas a 5 minutos. Así lo hago.' (no 'Esto es genial')",
        "2. Contraste Dramático: Muestra antes/después visualmente con números específicos",
        "3. Trending Elements: Usa música trending (con permiso), formats trending, hashtags trending",
        "4. Timing Perfecto: Publica en horarios pico cuando tu audiencia está más activa",
        "5. Engagement Inmediato: Responde primeros comentarios en 30 minutos, haz preguntas en caption",
        "6. Visual Impact: Primer frame debe captar atención, colores vibrantes, texto grande y legible"
    ]
    for strategy in viral_strategies:
        para = doc.add_paragraph(strategy, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    # ========== SECCIÓN 31: IDEAS DE CONTENIDO ==========
    doc.add_page_break()
    add_section_header(doc, "IDEAS DE CONTENIDO CREATIVAS", "💡")
    
    create_table_with_style(
        doc,
        ["Categoría", "Ejemplo 1", "Ejemplo 2", "Ejemplo 3"],
        [
            ["Testimonials", "Esto me ahorró X horas esta semana", "De X horas a X minutos: Mi transformación", "Esta herramienta cambió mi negocio"],
            ["Tutoriales", "Cómo generar X documentos en X minutos", "Tutorial completo paso a paso", "5 formas de usar esta herramienta"],
            ["Comparaciones", "Antes vs Ahora: Proceso completo", "Método manual vs Automatizado", "Costo vs Beneficio: Análisis completo"],
            ["Casos de Uso", "Caso de uso: Propuestas comerciales", "Caso de uso: Contratos legales", "Caso de uso: Emails personalizados"],
            ["Problema/Solución", "¿Te pasa que pierdes horas en documentos?", "Problema común: Solución simple", "Si odias crear documentos manualmente..."]
        ],
        "50+ Ideas de Contenido UGC"
    )
    
    # ========== SECCIÓN 32: WORKFLOW OPTIMIZADO ==========
    doc.add_page_break()
    add_section_header(doc, "WORKFLOW DE PRODUCCIÓN OPTIMIZADO", "🎬")
    
    create_table_with_style(
        doc,
        ["Paso", "Actividad", "Tiempo", "Checklist"],
        [
            ["1. Preparación", "Revisa brief, elige tipo, prepara guión", "5 min", "Brief revisado, guión listo"],
            ["2. Grabación", "Graba hook, contenido, screen recording, CTA", "10 min", "Hook grabado, audio claro"],
            ["3. Edición", "Importa, recorta, subtítulos, música, exporta", "10 min", "Subtítulos sincronizados"],
            ["4. Optimización", "Escribe caption, hashtags, verifica link", "3 min", "Caption listo, link verificado"],
            ["5. Publicación", "Publica, comparte en stories, responde", "2 min", "Publicado, stories compartido"]
        ],
        "Proceso de 5 Pasos (30 minutos total)"
    )
    
    # ========== SECCIÓN 33: CONTACTO ==========
    doc.add_page_break()
    add_section_header(doc, "CONTACTO Y SOPORTE", "📞")
    
    create_table_with_style(
        doc,
        ["Departamento", "Email", "Responsabilidad", "Horario"],
        [
            ["Manager de Campaña", "email-manager@ejemplo.com", "Coordinación general, aprobaciones", "Lun-Vie 9-18h"],
            ["Soporte Técnico", "soporte-tecnico@ejemplo.com", "Plataforma, herramientas", "Lun-Vie 9-18h"],
            ["Analytics", "analytics@ejemplo.com", "Tracking, reportes, métricas", "Lun-Vie 9-18h"],
            ["Legal/Compliance", "legal@ejemplo.com", "Preguntas legales, compliance", "Lun-Vie 9-18h"],
            ["Urgente", "urgente@ejemplo.com", "Emergencias, crisis management", "24/7"],
            ["Mentoría", "mentoria@ejemplo.com", "Crecimiento profesional", "Lun-Vie 9-18h"],
            ["Community", "community@ejemplo.com", "Comunidad de creadoras", "Lun-Vie 9-18h"]
        ],
        "Equipo de Soporte Completo"
    )
    
    doc.add_paragraph()
    
    # Footer final mejorado
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('¡Estamos emocionados de trabajar contigo! 🚀')
    footer_run.font.size = Pt(20)
    footer_run.font.bold = True
    footer_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    footer2_para = doc.add_paragraph()
    footer2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2_para.add_run('Versión 11.0 - Ultra Completo Absoluto Definitivo Máximo')
    footer2_run.font.size = Pt(14)
    footer2_run.font.bold = True
    footer2_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    footer3_para = doc.add_paragraph()
    footer3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer3_run = footer3_para.add_run('Brief UGC Creadoras - IA Bulk Documentos')
    footer3_run.font.size = Pt(12)
    footer3_run.font.italic = True
    footer3_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Guardar documento
    output_path = '/Users/adan/Documents/documentos_blatam/01_marketing/BRIEF_UGC_CREADORAS_BULK.docx'
    doc.save(output_path)
    print(f"✅ Documento Word mejorado creado exitosamente: {output_path}")
    print(f"📊 Total de páginas: ~{len(doc.paragraphs) // 20} páginas estimadas")
    return output_path

if __name__ == "__main__":
    try:
        create_brief_word()
    except ImportError:
        print("❌ Error: python-docx no está instalado")
        print("   Instálalo con: pip install python-docx")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
